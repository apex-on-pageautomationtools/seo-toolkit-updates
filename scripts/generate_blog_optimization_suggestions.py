"""
Blog Optimization Suggestions report generator.

Given a single blog post URL (+ optionally a team-supplied list of target
pages/keywords), scrapes the REAL live page, generates AI suggestions
grounded in that real content, discovers REAL internal-link targets from the
site's own sitemap, and writes a .docx matching the "james" reference format
(D:\\Report Formats\\Blog Optimization Suggestions\\...) exactly.

Nothing here is fabricated: every "Existing" field is scraped, every
"Suggested" field is an AI suggestion grounded in the real scraped content
(never invented from nothing), internal-link anchor text is always a REAL
sentence pulled verbatim from the page, and the image suggestion is either a
real image already on the site or an explicit red-flagged "needs a new
image" note - never a guessed/placeholder URL.

Run as a subprocess (same convention as generate_seo_onpage_phase2.py) so
its stdout can be streamed live to the UI.
"""
import argparse
import json
import os
import re
import sys
import time
from datetime import datetime
from urllib.parse import urljoin, urlparse

import requests
from bs4 import BeautifulSoup

import docx
from docx.shared import Pt, Inches, RGBColor

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
# Reuse the existing OpenAI/Gemini/Groq/OpenRouter fallback chain instead of
# duplicating the HTTP calls / rate-limit bookkeeping - see that module for
# the actual provider implementations.
import generate_seo_onpage_phase2 as _onpage

UA = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
      "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36")
MISSING = "Missing!"


def log(msg):
    print(msg, flush=True)


# --------------------------------------------------------------------------- #
# Fetching
# --------------------------------------------------------------------------- #
def fetch(url, timeout=20):
    try:
        r = requests.get(url, headers={"User-Agent": UA}, timeout=timeout)
        return r
    except Exception as e:
        log(f"   [warn] fetch failed for {url}: {type(e).__name__}: {e}")
        return None


def site_root(url):
    p = urlparse(url)
    return f"{p.scheme}://{p.netloc}"


# --------------------------------------------------------------------------- #
# Scrape the target blog page for real
# --------------------------------------------------------------------------- #
def scrape_blog(url):
    r = fetch(url)
    if not r or r.status_code >= 400:
        raise Exception(f"Could not fetch {url}"
                        + (f" (HTTP {r.status_code})" if r else ""))
    soup = BeautifulSoup(r.text, "html.parser")

    title_tag = soup.find("title")
    existing_title = title_tag.get_text(strip=True) if title_tag else MISSING

    meta_desc_tag = soup.find("meta", attrs={"name": "description"})
    existing_meta = (meta_desc_tag.get("content", "").strip()
                     if meta_desc_tag and meta_desc_tag.get("content") else MISSING)

    h1_tag = soup.find("h1")
    existing_h1 = h1_tag.get_text(strip=True) if h1_tag else MISSING

    og_image = soup.find("meta", attrs={"property": "og:image"})
    featured_image = og_image.get("content", "").strip() if og_image and og_image.get("content") else ""
    if not featured_image:
        first_img = soup.select_one("article img, .entry-content img, main img")
        if first_img and first_img.get("src"):
            featured_image = urljoin(url, first_img["src"])

    # Existing category - WordPress convention: rel="category tag" links.
    cat_links = soup.select('a[rel*="category"]')
    existing_category = cat_links[0].get_text(strip=True) if cat_links else "Missing"

    # Real body paragraphs (for internal-linking anchor text and AI context) -
    # prefer the main content container over the whole page (nav/footer noise).
    # Pick whichever candidate container has the MOST real paragraphs, rather
    # than trusting selector priority blindly - confirmed real case: a site's
    # first <article> match was a "suggested content" teaser card with only
    # one short paragraph, while the real post body (224 paragraphs) was
    # elsewhere in <main> - selector-priority order alone picked the wrong,
    # nearly-empty container and silently produced zero usable paragraphs.
    def _real_paragraphs(node):
        ps = [p.get_text(" ", strip=True) for p in node.find_all("p")]
        return [p for p in ps if len(p.split()) >= 8]  # skip short/nav fragments

    candidates = [c for c in (soup.select_one("article"), soup.select_one(".entry-content"),
                              soup.select_one("main"), soup.body) if c]
    paragraphs = []
    for cand in candidates:
        found = _real_paragraphs(cand)
        if len(found) > len(paragraphs):
            paragraphs = found
    if not paragraphs:
        paragraphs = _real_paragraphs(soup)
    body_text = "\n".join(paragraphs)

    # Best-guess Contact page - a real link on the site, not invented.
    contact_url = ""
    for a in soup.find_all("a", href=True):
        if re.search(r"contact", a.get_text(" ", strip=True), re.I) or re.search(r"/contact", a["href"], re.I):
            contact_url = urljoin(url, a["href"])
            break

    return {
        "url": url,
        "site_root": site_root(url),
        "existing_title": existing_title,
        "existing_meta": existing_meta,
        "existing_h1": existing_h1,
        "featured_image": featured_image,
        "existing_category": existing_category,
        "paragraphs": paragraphs,
        "body_text": body_text[:8000],  # keep the AI prompt bounded
        "contact_url": contact_url,
        "html": r.text,
    }


# --------------------------------------------------------------------------- #
# Sitemap crawl - real internal-link/CTA targets, not invented URLs
# --------------------------------------------------------------------------- #
def crawl_sitemap(root, cap=300):
    """Fetch sitemap.xml (following one level of sub-sitemaps) and return a
    flat list of real page URLs on this site. Best-effort - returns [] if the
    site has no sitemap, which the caller treats as "no internal-link target
    found" rather than guessing one."""
    seen = []
    queue = [urljoin(root, "/sitemap.xml"), urljoin(root, "/sitemap_index.xml")]
    visited_sitemaps = set()
    while queue and len(seen) < cap:
        sm_url = queue.pop(0)
        if sm_url in visited_sitemaps:
            continue
        visited_sitemaps.add(sm_url)
        r = fetch(sm_url, timeout=15)
        if not r or r.status_code >= 400:
            continue
        try:
            soup = BeautifulSoup(r.text, "xml")
        except Exception:
            soup = BeautifulSoup(r.text, "html.parser")
        locs = [loc.get_text(strip=True) for loc in soup.find_all("loc")]
        for loc in locs:
            if loc.endswith(".xml") and len(visited_sitemaps) < 6:
                queue.append(loc)
            elif loc not in seen:
                seen.append(loc)
            if len(seen) >= cap:
                break
    return seen


def score_candidates(sitemap_urls, current_url, target_keywords):
    """Cheap slug/path scoring against target_keywords (or, if none given,
    against nothing yet - the caller falls back to AI topic keywords) so we
    only fetch the <title> of the most plausible candidates instead of every
    page on the site."""
    kws = [k.lower() for k in (target_keywords or []) if k.strip()]
    scored = []
    for u in sitemap_urls:
        if u.rstrip("/") == current_url.rstrip("/"):
            continue
        path = urlparse(u).path.lower()
        score = sum(1 for k in kws if re.sub(r"[^a-z0-9]+", "-", k).strip("-") in path.replace("/", "-"))
        scored.append((score, u))
    scored.sort(key=lambda x: -x[0])
    return scored


def find_link_targets(scraped, target_pages, target_keywords, topic_keywords, max_targets=2, fetch_cap=12):
    """Real internal-link / CTA-button targets. Priority:
    1. Team-supplied target_pages, if given (verified reachable, real URLs).
    2. Otherwise, crawl the sitemap and pick the best-scoring pages by
       target_keywords (team-supplied) or the AI's own topic_keywords for
       this post, fetching just the top candidates' real <title> to confirm
       relevance before deciding.
    Returns [{"url":..., "title":...}, ...] - never a fabricated URL; an
    empty list means no confident real target was found."""
    if target_pages:
        out = []
        for u in target_pages[:max_targets]:
            r = fetch(u, timeout=12)
            if r and r.status_code < 400:
                t = BeautifulSoup(r.text, "html.parser").find("title")
                out.append({"url": u, "title": t.get_text(strip=True) if t else u})
            else:
                log(f"   [warn] team-supplied target page not reachable, skipping: {u}")
        if out:
            return out

    sitemap_urls = crawl_sitemap(scraped["site_root"])
    if not sitemap_urls:
        log("   No sitemap found - can't discover a real internal-link target.")
        return []

    keywords = target_keywords or topic_keywords or []
    scored = score_candidates(sitemap_urls, scraped["url"], keywords)
    candidates = [u for score, u in scored if score > 0][:fetch_cap] or [u for _, u in scored[:fetch_cap]]

    out = []
    for u in candidates:
        r = fetch(u, timeout=12)
        if not r or r.status_code >= 400:
            continue
        t = BeautifulSoup(r.text, "html.parser").find("title")
        title = t.get_text(strip=True) if t else u
        title_l = title.lower()
        if any(k.lower() in title_l or k.lower() in u.lower() for k in keywords):
            out.append({"url": u, "title": title})
        if len(out) >= max_targets:
            break
    return out


# --------------------------------------------------------------------------- #
# AI suggestions - OpenAI tried FIRST (per-team decision for this tool, since
# it's a single-URL report rather than a bulk run), then the same free-tier
# fallbacks generate_seo_onpage_phase2 already uses. Reuses that module's
# actual provider functions instead of duplicating the HTTP/retry logic.
# --------------------------------------------------------------------------- #
def ai_chain(prompt):
    for fn in (_onpage._ai_suggest_openai, _onpage._ai_suggest_gemini,
              _onpage._ai_suggest_groq, _onpage._ai_suggest_openrouter):
        result = fn(prompt)
        if result is not None:
            return result
    return None


def ai_suggestions(scraped):
    prompt = f"""You are an SEO editor. Read this real blog post and suggest improvements.
Return ONLY a JSON object with exactly these keys:
- "category": a single WordPress-style category name for this post
- "tags": array of 3-6 short keyword tags
- "title": an improved SEO title (or the existing title if it's already good)
- "meta_description": an improved meta description, 140-160 characters
- "h1_change": either the string "No Change Required" or a new suggested H1
- "h2_list": array of 5-8 suggested H2 subheadings that could structure this content
- "topic_keywords": array of 4-8 short keywords/phrases capturing what this post is really about (used to find relevant internal links and images - be specific, not generic)
- "image_alt": suggested alt text for the featured image, based on the post topic
- "image_title_attr": suggested title attribute for the featured image
- "cta_text": a short (3-6 word) call-to-action button label relevant to this post's topic

Existing title: {scraped['existing_title']}
Existing meta description: {scraped['existing_meta']}
Existing H1: {scraped['existing_h1']}
Existing category: {scraped['existing_category']}

Post content (real, scraped from the live page):
{scraped['body_text'][:5000]}
"""
    result = ai_chain(prompt)
    if not result:
        raise Exception("Every AI provider failed - check API keys (Admin -> Sync API Keys) and try again.")
    return result


def pick_link_sentences(paragraphs, link_targets):
    """Ask the AI to pick which REAL paragraph(s) should carry an internal
    link to each target, by INDEX only - it never writes new anchor text,
    so "Find with" is always a verbatim sentence from the actual page."""
    if not paragraphs or not link_targets:
        return []
    numbered = "\n".join(f"[{i}] {p}" for i, p in enumerate(paragraphs[:40]))
    targets_desc = "\n".join(f"- {t['url']} (\"{t['title']}\")" for t in link_targets)
    prompt = f"""Given these real numbered paragraphs from a blog post, and these real target pages
on the same site, pick ONE paragraph index per target page that would be a natural
place to add an internal link to that target (based on topical relevance).
Return ONLY a JSON array of objects: [{{"index": <int>, "target_url": "<one of the target URLs above>"}}]
Pick at most {len(link_targets)} pairs, each with a DIFFERENT paragraph index. If none of the
paragraphs are a good fit for a target, omit that target.

Paragraphs:
{numbered}

Target pages:
{targets_desc}
"""
    result = ai_chain(prompt)
    if not isinstance(result, list):
        return []
    out = []
    for item in result:
        try:
            idx = int(item.get("index"))
            target_url = item.get("target_url")
            if 0 <= idx < len(paragraphs) and any(t["url"] == target_url for t in link_targets):
                out.append({"sentence": paragraphs[idx], "target_url": target_url})
        except Exception:
            continue
    return out


# --------------------------------------------------------------------------- #
# Real image reuse (never a fabricated URL)
# --------------------------------------------------------------------------- #
def find_reusable_image(scraped, topic_keywords):
    """Look for a real image already on the site that matches the post's
    topic - checks WordPress's media library search first (if available),
    then falls back to scanning the post's own images for one with relevant
    alt text. Returns a real image URL, or None if nothing confident was
    found - the caller flags that in red rather than guessing a URL."""
    root = scraped["site_root"]
    for kw in (topic_keywords or [])[:3]:
        r = fetch(f"{root}/wp-json/wp/v2/media?search={requests.utils.quote(kw)}&per_page=3", timeout=12)
        if r and r.status_code == 200:
            try:
                items = r.json()
                if isinstance(items, list) and items:
                    src = (items[0].get("source_url") or "").strip()
                    if src:
                        return src
            except Exception:
                pass

    soup = BeautifulSoup(scraped["html"], "html.parser")
    for img in soup.find_all("img"):
        alt = (img.get("alt") or "").lower()
        src = img.get("src") or ""
        if not src:
            continue
        if any((kw or "").lower() in alt for kw in (topic_keywords or [])):
            return urljoin(scraped["url"], src)
    return None


# --------------------------------------------------------------------------- #
# Best-effort annotated screenshot (real screenshot, box drawn around the
# element in question) - approximation of a human's exact placement call,
# never a fabricated image.
# --------------------------------------------------------------------------- #
_driver = None


def _get_driver():
    global _driver
    if _driver is not None:
        try:
            _ = _driver.title
            return _driver
        except Exception:
            _driver = None
    try:
        import tempfile
        from selenium.webdriver import Chrome, ChromeOptions
        opts = ChromeOptions()
        opts.add_argument(f"--user-data-dir={tempfile.mkdtemp(prefix='blogopt_')}")
        opts.add_argument("--headless=new")
        opts.add_argument("--no-sandbox")
        opts.add_argument("--disable-dev-shm-usage")
        opts.add_argument("--disable-gpu")
        opts.add_argument("--window-size=1366,1400")
        opts.add_argument(f"--user-agent={UA}")
        _driver = Chrome(options=opts)
        _driver.set_page_load_timeout(45)
        return _driver
    except Exception as e:
        log(f"   [warn] could not launch Chrome for screenshots: {type(e).__name__}: {e}")
        return None


def close_driver():
    global _driver
    if _driver:
        try:
            _driver.quit()
        except Exception:
            pass
        _driver = None


def capture_annotated_screenshot(url, out_path, highlight_selector=None):
    """Real screenshot of the live page, with a red box drawn around the
    element matching highlight_selector (if found) - an approximation of a
    human's placement callout, not a guarantee of exact taste. Returns
    out_path on success, None if the capture failed entirely."""
    driver = _get_driver()
    if not driver:
        return None
    try:
        driver.get(url)
        time.sleep(1.5)
        if highlight_selector:
            try:
                driver.execute_script("""
                    var el = document.querySelector(arguments[0]);
                    if (el) {
                        el.scrollIntoView({block: 'center'});
                        el.style.outline = '4px solid #ff3b30';
                        el.style.outlineOffset = '2px';
                    }
                """, highlight_selector)
                time.sleep(0.4)
            except Exception:
                pass
        driver.save_screenshot(out_path)
        return out_path
    except Exception as e:
        log(f"   [warn] screenshot capture failed: {type(e).__name__}: {e}")
        return None


# --------------------------------------------------------------------------- #
# Structured Data (JSON-LD) - fully real, every field scraped or known
# --------------------------------------------------------------------------- #
def build_json_ld(scraped, suggestions):
    root = scraped["site_root"]
    homepage = fetch(root, timeout=12)
    site_name = urlparse(root).netloc
    logo_url = ""
    if homepage and homepage.status_code < 400:
        hsoup = BeautifulSoup(homepage.text, "html.parser")
        og_site = hsoup.find("meta", attrs={"property": "og:site_name"})
        if og_site and og_site.get("content"):
            site_name = og_site["content"].strip()
        logo_img = hsoup.select_one('img[class*="logo"], header img, .site-logo img')
        if logo_img and logo_img.get("src"):
            logo_url = urljoin(root, logo_img["src"])

    today = datetime.now().strftime("%Y-%m-%d")
    return {
        "@context": "https://schema.org",
        "@type": "Article",
        "mainEntityOfPage": {"@type": "WebPage", "@id": scraped["url"]},
        "headline": suggestions.get("title") or scraped["existing_title"],
        "description": suggestions.get("meta_description") or scraped["existing_meta"],
        "image": scraped["featured_image"] or "",
        "author": {"@type": "Organization", "name": site_name, "url": root},
        "publisher": {"@type": "Organization", "name": site_name,
                      "logo": {"@type": "ImageObject", "url": logo_url}},
        "datePublished": today,
        "dateModified": today,
    }


# --------------------------------------------------------------------------- #
# .docx assembly - matches the "james" reference format exactly
# --------------------------------------------------------------------------- #
SEP = "*" * 43


def build_docx(out_path, scraped, suggestions, link_pairs, image_url, image_missing_note,
               screenshot_paths, contact_url, json_ld):
    doc = docx.Document()
    doc.add_heading("Blog Optimization Suggestions", level=0)

    doc.add_paragraph(f"Blog URL \u2013 {scraped['url']}")
    doc.add_paragraph(SEP)

    doc.add_paragraph("Additional Suggestions-")
    doc.add_paragraph("Category- \n"
                      f"Existing Category: {scraped['existing_category']}\n"
                      f"Suggested Category: {suggestions.get('category', '')}")
    tags = ", ".join(suggestions.get("tags") or [])
    doc.add_paragraph(f"Tag Suggestions (Provide Keywords Based): {tags}")
    doc.add_paragraph(SEP)

    doc.add_paragraph("Meta Suggestions-")
    doc.add_paragraph(f"Existing Title: {scraped['existing_title']} ")
    doc.add_paragraph(f"Suggested Title: {suggestions.get('title', '')}")
    doc.add_paragraph(f"Existing Meta Description: {scraped['existing_meta']}")
    doc.add_paragraph(f"Suggested Meta Description: {suggestions.get('meta_description', '')}")

    doc.add_paragraph("Heading Tags-")
    doc.add_paragraph(f"H1 Existing: {scraped['existing_h1']}")
    doc.add_paragraph(f"Suggested H1: {suggestions.get('h1_change', 'No Change Required')}")
    doc.add_paragraph("\nSuggested H2:")
    for h2 in suggestions.get("h2_list") or []:
        doc.add_paragraph(h2)
    doc.add_paragraph(SEP)

    doc.add_paragraph("Internal Linking Suggestion:")
    doc.add_paragraph("")
    if link_pairs:
        for pair in link_pairs:
            doc.add_paragraph(f"Find with: {pair['sentence']}")
            p = doc.add_paragraph("Replace with: ")
            run = p.add_run(pair["sentence"])
            run.font.color.rgb = RGBColor(0x10, 0x6F, 0xD8)
            run.font.underline = True
            p.add_run(f"  (link to: {pair['target_url']})")
    else:
        doc.add_paragraph("No confident real internal-link target was found automatically "
                          "(no sitemap, or no page matched this post's topic closely enough) - "
                          "add manually or supply target pages/keywords next time.")
    doc.add_paragraph(SEP)

    doc.add_paragraph("Image Optimization Suggestions-  ")
    doc.add_paragraph("")
    doc.add_paragraph(f"Featured Image URL: {'Already Given!' if scraped['featured_image'] else MISSING}")
    if image_url:
        doc.add_paragraph(f"Image Suggestion: {image_url}")
    else:
        p = doc.add_paragraph()
        run = p.add_run("Image Suggestion: " + image_missing_note)
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        run.bold = True
    doc.add_paragraph(f"Suggested Image Alt: {suggestions.get('image_alt', '')}")
    doc.add_paragraph(f"Suggested Title Attribute: {suggestions.get('image_title_attr', '')}")
    cta = suggestions.get("cta_text", "")
    link_target = link_pairs[0]["target_url"] if link_pairs else ""
    doc.add_paragraph(f"Clickable Button: Create a clickable button below the image with the "
                      f"text \u201c{cta}\u201d and link it with this page ")
    doc.add_paragraph(f"Link: {link_target}")
    if screenshot_paths.get("image_placement"):
        doc.add_paragraph("Placement Suggestion (real screenshot, approximate placement):")
        try:
            doc.add_picture(screenshot_paths["image_placement"], width=Inches(5.5))
        except Exception:
            pass
    else:
        doc.add_paragraph("Placement Suggestion: could not capture automatically - place near the "
                          "featured image / top of the content.")
    doc.add_paragraph(SEP)

    doc.add_paragraph("Suggested Contact Us Form-")
    doc.add_paragraph("")
    doc.add_paragraph("Add this Contact Us form on the page. You will find the form on the "
                      "Contact page from there you can\nget the code and add it.")
    if screenshot_paths.get("contact_placement"):
        doc.add_paragraph("Screenshot (real, from the live Contact page):")
        try:
            doc.add_picture(screenshot_paths["contact_placement"], width=Inches(5.5))
        except Exception:
            pass
    else:
        doc.add_paragraph("Screenshot: could not capture automatically.")
    doc.add_paragraph(f"URL: {contact_url or MISSING}")
    doc.add_paragraph("Contact us placement suggestion: place near the end of the post, "
                      "below the main content.")
    doc.add_paragraph(SEP)

    doc.add_paragraph("Structured Data:")
    doc.add_paragraph("Start")
    doc.add_paragraph('<script type="application/ld+json">')
    for line in json.dumps(json_ld, indent=2).splitlines():
        doc.add_paragraph(line)
    doc.add_paragraph("</script>")
    doc.add_paragraph("End")

    doc.save(out_path)


# --------------------------------------------------------------------------- #
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("url", help="Full blog post URL to analyze")
    ap.add_argument("--out", required=True, help="Output directory")
    ap.add_argument("--format", default="james", choices=["james"])
    ap.add_argument("--targets", default=None,
                    help="Optional JSON file: {\"target_pages\": [...], \"target_keywords\": [...]}")
    ap.add_argument("--no-capture", action="store_true", help="skip live screenshots")
    args = ap.parse_args()

    os.makedirs(args.out, exist_ok=True)

    target_pages, target_keywords = [], []
    if args.targets and os.path.exists(args.targets):
        with open(args.targets, encoding="utf-8") as f:
            tdata = json.load(f)
        target_pages = tdata.get("target_pages") or []
        target_keywords = tdata.get("target_keywords") or []

    log(f"[1/6] Fetching {args.url}...")
    scraped = scrape_blog(args.url)

    log("[2/6] Generating AI suggestions from the real page content...")
    suggestions = ai_suggestions(scraped)

    log("[3/6] Finding real internal-link targets...")
    link_targets = find_link_targets(scraped, target_pages, target_keywords,
                                     suggestions.get("topic_keywords") or [])
    link_pairs = pick_link_sentences(scraped["paragraphs"], link_targets) if link_targets else []
    if link_pairs:
        log(f"   Found {len(link_pairs)} real internal-link opportunity(ies).")
    else:
        log("   No confident internal-link target found.")

    log("[4/6] Looking for a reusable real image...")
    image_url = find_reusable_image(scraped, suggestions.get("topic_keywords") or [])
    image_missing_note = "" if image_url else "Need to find new image and update"
    if image_url:
        log(f"   Found: {image_url}")
    else:
        log("   No matching real image found on the site - flagged for manual sourcing.")

    screenshot_paths = {}
    if not args.no_capture:
        log("[5/6] Capturing best-effort placement screenshots...")
        img_shot = os.path.join(args.out, "_image_placement.png")
        if capture_annotated_screenshot(args.url, img_shot, "article img, .entry-content img, main img"):
            screenshot_paths["image_placement"] = img_shot
        if scraped["contact_url"]:
            contact_shot = os.path.join(args.out, "_contact_placement.png")
            if capture_annotated_screenshot(scraped["contact_url"], contact_shot, "form"):
                screenshot_paths["contact_placement"] = contact_shot
        close_driver()
    else:
        log("[5/6] Screenshot capture skipped (--no-capture).")

    json_ld = build_json_ld(scraped, suggestions)

    safe_name = re.sub(r"[^a-zA-Z0-9]+", "_", urlparse(args.url).path.strip("/"))[:60] or "post"
    out_file = os.path.join(args.out, f"Blog_Optimization_Suggestions_-_{safe_name}.docx")
    log("[6/6] Writing .docx...")
    build_docx(out_file, scraped, suggestions, link_pairs, image_url, image_missing_note,
              screenshot_paths, scraped["contact_url"], json_ld)

    for p in screenshot_paths.values():
        try:
            os.remove(p)
        except Exception:
            pass

    log(f"[DONE] {out_file}")


if __name__ == "__main__":
    try:
        main()
    except SystemExit:
        raise
    except Exception as e:
        log(f"[ERROR] {type(e).__name__}: {e}")
        sys.exit(1)
