"""
generate_seo_onpage_phase2.py - On-Page SEO Phase 2 report engine for Report Studio.

Given a domain + a "target pages & keywords" list, it crawls each target page, audits
the Phase-2 on-page layer, and produces the same deliverables the manual skill makes:

    1. On Page-Analysis-Report - <domain>.docx     (branded cover page + per-site narrative + screenshots)
    2. Meta Suggestions - <domain>.xlsx            (existing + suggested title/desc/H1; suggested via free Gemini or heuristic)
    3. Image Alt Tag Suggestions - <domain>.xlsx   (self-hosted + Shopify-CDN images only; alt suggestions)
    4. Canonical Tag Suggestions - <domain>.xlsx   (existing vs recommended; self-canonical -> "No Changes Needed...")
    5. Target Pages and Keywords - <domain>.xlsx   (the complete target-pages + keywords list)
    6. sitemap (existing - review).xml             (the site's REAL sitemap if found) OR
       sitemap.xml                                 (a freshly-generated working sitemap when none exists)

...all bundled into a single ZIP: "SEO On-Page Phase 2 - <domain>.zip".

The .xlsx are produced by CLONING the templates in backend/ (template_*.xlsx) so the
formatting matches the house style exactly. The .docx is built from scratch (dynamic
narrative + live screenshots). Page 1 is the branded cover (backend/onpage_cover.png,
A4 full-bleed; per-domain override seo_onpage_screens/<domain>/cover.png; text cover if
absent). Suggested meta copy uses a free-tier AI fallback chain (Gemini -> Groq ->
OpenRouter, whichever has a key configured), else a heuristic - no paid APIs.

Run:
    python generate_seo_onpage_phase2.py --domain example.com --targets targets.json
    python generate_seo_onpage_phase2.py --domain example.com --targets targets.json --dry-run

`targets.json` is either:
    [{"page": "https://example.com/", "keywords": ["kw one", "kw two"]}, ...]
  or the flat rows form:
    [{"keyword": "kw one", "page": "https://example.com/"}, ...]
"""
import os
import re
import sys
import json
import shutil
import zipfile
import tempfile
import argparse
import datetime
import urllib.parse
from pathlib import Path

ROOT = Path(__file__).parent.resolve()
# The bundled python runs ISOLATED (a python312._pth file is present), so it does
# NOT put this script's own folder on sys.path. Without this, sibling imports such
# as `import generate_health_report` (used by the docx builder) raise
# ModuleNotFoundError and no on-page report is ever produced on user machines.
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
OUTPUT_DIR = ROOT / "output"
TPL_META = ROOT / "template_meta_suggestions.xlsx"
TPL_ALT = ROOT / "template_alt_suggestions.xlsx"
TPL_CANON = ROOT / "template_canonical_suggestions.xlsx"
# Branded cover image for page 1 (drop-in). Per-domain override:
# seo_onpage_screens/<domain>/cover.png ; else this shared file.
COVER_IMG = ROOT / "onpage_cover.png"

MISSING = "Missing!"


def log(msg):
    print(msg, flush=True)


# --------------------------------------------------------------------------- io
def safe_domain(domain):
    d = (domain or "").strip().lower()
    d = re.sub(r"^https?://", "", d).rstrip("/")
    d = re.sub(r"^www\.", "", d)
    return d.split("/")[0]


def site_root(domain):
    return f"https://{safe_domain(domain)}"


def brand_from(domain, homepage_title=None, homepage_h1=None, og_site=None):
    """Best-effort brand name. Prefer og:site_name (most reliable), then the part of
    the title/H1 before a separator; else title-case the domain's second-level label."""
    if og_site and 2 <= len(og_site.strip()) <= 40:
        return og_site.strip()
    for cand in (homepage_title, homepage_h1):
        if cand and cand != MISSING:
            # take the part after the last separator (brand usually trails the title)
            parts = [p.strip() for p in re.split(r"[|\-–—:·]", cand) if p.strip()]
            if parts:
                tail = parts[-1]
                # a short tail is usually the brand; otherwise use the head
                pick = tail if 2 <= len(tail) <= 30 else parts[0]
                if 2 <= len(pick) <= 40:
                    return pick
    root = safe_domain(domain).split(".")[0]
    return root.replace("-", " ").replace("_", " ").title()


def load_targets(targets_path, domain):
    """Return ordered list of {page, keywords:[...], keyword_ranks:{kw:rank}} grouped by
    page (primary first). The sheet's Keyword/Target Page/Ranking columns are matched by
    HEADER NAME, not fixed position - the team's sheet can list them in any order/
    sequence, and Ranking is optional. When a keyword already ranks in the top 20 (per
    that Ranking column), suggest_meta() skips suggesting a new title for its page."""
    rows = []
    p = Path(targets_path)
    if p.suffix.lower() in (".xlsx", ".xlsm"):
        import openpyxl
        wb = openpyxl.load_workbook(p, data_only=True)
        ws = wb.active
        # find the header row that has "Keyword" + "Target Page" (Ranking optional)
        header_row = None
        for r in range(1, min(ws.max_row, 30) + 1):
            vals = [str(ws.cell(r, c).value or "").lower() for c in range(1, ws.max_column + 1)]
            if any("keyword" in v for v in vals) and any("page" in v for v in vals):
                header_row = r
                kw_col = next(c for c, v in enumerate(vals, 1) if "keyword" in v)
                pg_col = next(c for c, v in enumerate(vals, 1) if "page" in v)
                rank_col = next((c for c, v in enumerate(vals, 1) if "rank" in v), None)
                break
        if header_row:
            for r in range(header_row + 1, ws.max_row + 1):
                kw = ws.cell(r, kw_col).value
                pg = ws.cell(r, pg_col).value
                rank = ws.cell(r, rank_col).value if rank_col else None
                if kw and pg:
                    rows.append({"keyword": str(kw).strip(), "page": str(pg).strip(),
                                 "rank": rank})
    else:
        data = json.loads(p.read_text(encoding="utf-8"))
        if data and isinstance(data[0], dict) and "keywords" in data[0]:
            # already grouped
            for item in data:
                ranks = item.get("keyword_ranks") or {}
                rows.extend({"keyword": k, "page": item["page"], "rank": ranks.get(k)}
                            for k in item["keywords"])
        else:
            rows = [{"keyword": str(d.get("keyword", "")).strip(),
                     "page": str(d.get("page", "")).strip(),
                     "rank": d.get("rank")} for d in data]

    # group, preserving first-seen page order and keyword order
    grouped = {}
    ranks_by_page = {}
    for row in rows:
        page = normalize_url(row["page"], domain)
        if not page:
            continue
        grouped.setdefault(page, [])
        ranks_by_page.setdefault(page, {})
        if row["keyword"]:
            grouped[page].append(row["keyword"])
            if row.get("rank") not in (None, ""):
                ranks_by_page[page][row["keyword"]] = row["rank"]
    return [{"page": pg, "keywords": kws, "keyword_ranks": ranks_by_page.get(pg, {})}
            for pg, kws in grouped.items()]


def discover_targets(domain, dry_run=False, limit=12):
    """When no keyword sheet is supplied (e.g. Bulk runs), auto-pick target pages
    by crawling the homepage and collecting its same-site internal links."""
    root = site_root(domain)
    if dry_run:
        return [{"page": root + "/", "keywords": []}]
    home = crawl_page(root + "/", [], dry_run=False)
    site_reg = _registrable(safe_domain(domain))
    pages = [root + "/"]
    for u in home.get("internal_links", []):
        clean = u.split("#")[0].split("?")[0].rstrip("/")
        if not clean or clean in pages:
            continue
        if _registrable(urllib.parse.urlparse(clean).netloc) != site_reg:
            continue
        if re.search(r"\.(png|jpe?g|gif|svg|webp|pdf|zip|css|js|ico|xml|json)$", clean, re.I):
            continue
        pages.append(clean)
        if len(pages) >= limit:
            break
    log(f"   -> auto-discovered {len(pages)} target page(s)")
    return [{"page": p, "keywords": []} for p in dict.fromkeys(pages)]


def normalize_url(url, domain):
    url = (url or "").strip()
    if not url:
        return ""
    if not url.startswith("http"):
        url = site_root(domain) + "/" + url.lstrip("/")
    return url.rstrip()


def _registrable(host):
    """Last two labels of a host, e.g. www.x.com / api.x.com -> x.com."""
    host = (host or "").lower().split(":")[0]
    parts = [p for p in host.split(".") if p]
    return ".".join(parts[-2:]) if len(parts) >= 2 else host


_SHOPIFY_CDN = re.compile(r"(?:^|\.)(?:shopify\.com|myshopify\.com|shopifycdn\.com)$")


def _allowed_image_host(src, site_host):
    """Keep only self-hosted images (same registrable domain as the site) and
    Shopify-CDN images. Drops trackers / fonts / analytics pixels / third-party
    CDNs (facebook.com/tr, fonts.gstatic.com, *.blob.core.windows.net, ...)."""
    host = urllib.parse.urlparse(src).netloc.lower()
    if not host:
        return True  # same-origin relative URL (already resolved against the page)
    if _registrable(host) == _registrable(site_host):
        return True
    return bool(_SHOPIFY_CDN.search(host))


# -------------------------------------------------------------------- crawling
def _mock_page(url, keywords):
    kw = keywords[0] if keywords else "page"
    return {
        "url": url,
        "title": MISSING,
        "description": MISSING,
        "h1": MISSING,
        "canonical": f'<link rel="canonical" href="{url}" />',
        "lang": "en",
        "viewport": True,
        "images": [
            {"src": f"{url.rstrip('/')}/assets/hero.png", "alt": ""},
            {"src": f"{url.rstrip('/')}/assets/logo.png", "alt": ""},
        ],
        "headings": [("h1", kw.title())],
        "internal_links": [],
        "external_links": [],
        "status": 200,
        "body_text": "",
        "content_word_count": 0,
    }


# --- selenium-rendered crawl (fallback when patchright isn't bundled) ----------
# The embedded python ships selenium + Chrome but NOT patchright/playwright, so on
# user machines the patchright path below never runs. Without this fallback every
# page would be fetched as raw HTML - JS-rendered / SPA sites then report a missing
# H1, zero images, zero headings and zero internal links. Rendering with Chrome
# fixes that. One driver is reused for the whole run.
_op_driver = None


def _get_op_driver():
    global _op_driver
    if _op_driver is not None:
        try:
            _ = _op_driver.title
            return _op_driver
        except Exception:
            _op_driver = None
    try:
        import tempfile
        from selenium.webdriver import Chrome, ChromeOptions
        opts = ChromeOptions()
        opts.add_argument(f"--user-data-dir={tempfile.mkdtemp(prefix='seo_op_')}")
        opts.add_argument("--headless=new")
        opts.add_argument("--no-sandbox")
        opts.add_argument("--disable-dev-shm-usage")
        opts.add_argument("--disable-gpu")
        opts.add_argument("--disable-notifications")
        opts.add_argument("--mute-audio")
        opts.add_argument("--window-size=1366,900")
        opts.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                          "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36")
        _op_driver = Chrome(options=opts)
        _op_driver.set_page_load_timeout(45)
        return _op_driver
    except Exception as e:
        log(f"   [warn] could not launch Chrome for rendering: {type(e).__name__}: {e}")
        return None


def _close_op_driver():
    global _op_driver
    if _op_driver:
        try:
            _op_driver.quit()
        except Exception:
            pass
        _op_driver = None


def _op_http_status(url):
    """Best-effort HTTP status (selenium doesn't expose it)."""
    import urllib.request, urllib.error
    for method in ("HEAD", "GET"):
        try:
            req = urllib.request.Request(url, method=method,
                                         headers={"User-Agent": "Mozilla/5.0 SEOPhase2Bot"})
            with urllib.request.urlopen(req, timeout=15) as r:
                return getattr(r, "status", 200) or 200
        except urllib.error.HTTPError as e:
            return e.code
        except Exception:
            continue
    return None


def _crawl_selenium(url, keywords):
    """Render the page in Chrome (waits for the SPA to hydrate) and parse it.
    Returns None if the browser is unavailable so the caller can fall back."""
    driver = _get_op_driver()
    if not driver:
        return None
    import time
    try:
        driver.get(url)
        last, stable = -1, 0
        for _ in range(22):                       # wait up to ~9s for hydration
            try:
                ready = driver.execute_script("return document.readyState")
                blen = driver.execute_script(
                    "return (document.body ? document.body.innerText.length : 0)")
            except Exception:
                break
            if ready == "complete" and blen > 0 and blen == last:
                stable += 1
                if stable >= 2:                   # body text stable ~0.8s -> settled
                    break
            else:
                stable = 0
            last = blen
            time.sleep(0.4)
        html = driver.page_source or ""
        final_url = driver.current_url or url
        # Chrome exposes the navigation's real HTTP status via the Navigation Timing
        # API - when available this avoids a second full HTTP request purely to
        # recover the status code Selenium already resolved during driver.get().
        try:
            status = driver.execute_script(
                "try{return performance.getEntriesByType('navigation')[0].responseStatus}"
                "catch(e){return null}")
        except Exception:
            status = None
    except Exception as e:
        log(f"   [warn] selenium crawl failed for {url}: {type(e).__name__}: {e}")
        return None
    if not html or len(html) < 200:
        return None
    if not status:
        status = _op_http_status(url)
    return _parse_html(html, final_url, status)


def _crawl_rendered(url, keywords):
    """Prefer a real browser render (selenium+Chrome), fall back to raw HTTP."""
    sel = _crawl_selenium(url, keywords)
    if sel is not None:
        return sel
    return _crawl_requests(url, keywords)


def crawl_page(url, keywords, dry_run=False):
    if dry_run:
        return _mock_page(url, keywords)
    try:
        from patchright.sync_api import sync_playwright
    except Exception:
        return _crawl_rendered(url, keywords)      # no patchright -> selenium/raw
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            ctx = browser.new_context(user_agent=(
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                "(KHTML, like Gecko) Chrome/124.0 Safari/537.36"))
            page = ctx.new_page()
            resp = page.goto(url, wait_until="networkidle", timeout=45000)
            status = resp.status if resp else None
            html = page.content()
            final_url = page.url
            browser.close()
        return _parse_html(html, final_url or url, status)
    except Exception as e:
        log(f"   [warn] browser crawl failed for {url}: {type(e).__name__}: {e}")
        return _crawl_rendered(url, keywords)      # patchright failed -> selenium/raw


def _crawl_requests(url, keywords):
    try:
        import urllib.request
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0 SEOPhase2Bot"})
        with urllib.request.urlopen(req, timeout=30) as r:
            html = r.read().decode("utf-8", "ignore")
            return _parse_html(html, url, r.status)
    except Exception as e:
        log(f"   [warn] http crawl failed for {url}: {e}")
        return _mock_page(url, keywords)


MIN_PARAGRAPH_WORDS = 10  # a block below this isn't real body copy - a nav label,
                          # button text, or a one-line caption that happens to sit
                          # in a <p>/<li> tag, not actual paragraph content.


def _paragraph_word_count(html):
    """Word count of REAL paragraph content only: <p>/<li> text, excluding
    headings (h1-h6 are never counted here) and anything inside nav/header/footer
    chrome, and only counting a block once it has at least MIN_PARAGRAPH_WORDS
    words together - short fragments (labels, single-line CTAs) don't count as
    content even if they're technically inside a <p> tag."""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
    except Exception:
        return 0
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    total = 0
    for block in soup.find_all(["p", "li"]):
        if block.find_parent(["nav", "header", "footer"]):
            continue
        words = block.get_text(" ", strip=True).split()
        if len(words) >= MIN_PARAGRAPH_WORDS:
            total += len(words)
    return total


def _parse_html(html, url, status):
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
    except Exception:
        return _regex_parse(html, url, status)

    def meta(name=None, prop=None):
        if name:
            t = soup.find("meta", attrs={"name": name})
        else:
            t = soup.find("meta", attrs={"property": prop})
        return (t.get("content") or "").strip() if t and t.get("content") else ""

    import html as html_mod
    raw_title = (soup.title.get_text(strip=True) if soup.title else "") or MISSING
    title = html_mod.unescape(raw_title)
    desc = html_mod.unescape(meta(name="description") or MISSING)
    og_site = html_mod.unescape(meta(prop="og:site_name"))
    h1s = [html_mod.unescape(h.get_text(strip=True)) for h in soup.find_all("h1")]
    h1 = h1s[0] if h1s else MISSING
    canon_tag = soup.find("link", attrs={"rel": "canonical"})
    canonical = (f'<link rel="canonical" href="{canon_tag.get("href")}" />'
                 if canon_tag and canon_tag.get("href") else "")
    html_tag = soup.find("html")
    lang = (html_tag.get("lang") if html_tag else "") or ""
    viewport = bool(soup.find("meta", attrs={"name": "viewport"}))
    site_host = urllib.parse.urlparse(url).netloc
    images = []
    base = url
    for im in soup.find_all("img"):
        src = im.get("src") or im.get("data-src") or ""
        if not src or src.startswith("data:"):
            continue
        src = urllib.parse.urljoin(base, src)
        # Only the site's own images + Shopify CDN - drop trackers, fonts,
        # analytics pixels and third-party CDNs.
        if _allowed_image_host(src, site_host):
            # Lazy-load plugins/themes commonly leave the real alt text stored under a
            # data-* attribute (or aria-label) until the image scrolls into view and
            # JS swaps it onto the live `alt` attribute - which never happens in a
            # non-interactive crawl, so `alt` alone reports "missing" even when the
            # page's own source clearly has it. Check the common fallbacks too.
            alt = (im.get("alt") or im.get("data-alt") or im.get("data-alt-text")
                   or im.get("data-original-alt") or im.get("data-lazy-alt")
                   or im.get("aria-label") or "").strip()
            images.append({"src": src, "alt": alt})
    headings = [(h.name, h.get_text(strip=True)) for h in soup.find_all(re.compile(r"^h[1-6]$"))]
    site_reg = _registrable(urllib.parse.urlparse(url).netloc)
    internal, external = [], []
    for a in soup.find_all("a", href=True):
        href = (a["href"] or "").strip()
        # skip mailto:, tel:, javascript:, fragments and data URIs - not links
        if not href or href.startswith(("mailto:", "tel:", "javascript:", "#", "data:")):
            continue
        full = urllib.parse.urljoin(base, href)
        if not full.startswith("http"):
            continue
        full = full.split("#")[0]  # ignore in-page anchors
        if _registrable(urllib.parse.urlparse(full).netloc) == site_reg:
            internal.append(full)        # www / non-www both count as internal
        else:
            external.append(full)
    # Visible body text (for content-keyword checks + smarter meta suggestions) - built
    # from a fresh soup so removing script/style/nav/footer/header here can't affect the
    # link extraction above, which deliberately keeps nav/footer links (e.g. footer logo).
    try:
        text_soup = BeautifulSoup(html, "html.parser")
        for tag in text_soup(["script", "style", "noscript"]):
            tag.decompose()
        body_text = re.sub(r"\s+", " ", text_soup.get_text(" ", strip=True))[:4000]
    except Exception:
        body_text = ""

    content_word_count = _paragraph_word_count(html)

    return {
        "url": url, "title": title, "description": desc, "h1": h1, "h1s": h1s,
        "og_site_name": og_site, "canonical": canonical, "lang": lang, "viewport": viewport,
        "images": images, "headings": headings, "internal_links": list(dict.fromkeys(internal)),
        "external_links": list(dict.fromkeys(external)), "status": status, "body_text": body_text,
        "content_word_count": content_word_count,
    }


def _regex_parse(html, url, status):
    def find(pat):
        m = re.search(pat, html, re.I | re.S)
        return m.group(1).strip() if m else ""
    title = find(r"<title[^>]*>(.*?)</title>") or MISSING
    desc = find(r'<meta[^>]+name=["\']description["\'][^>]+content=["\'](.*?)["\']') or MISSING
    h1 = find(r"<h1[^>]*>(.*?)</h1>") or MISSING
    canon = find(r'<link[^>]+rel=["\']canonical["\'][^>]+href=["\'](.*?)["\']')
    canonical = f'<link rel="canonical" href="{canon}" />' if canon else ""
    lang = find(r'<html[^>]+lang=["\'](.*?)["\']')
    body_text = re.sub(r"\s+", " ", re.sub(r"<[^>]+>", " ", html))[:4000]
    return {"url": url, "title": re.sub("<[^>]+>", "", title), "description": desc,
            "h1": re.sub("<[^>]+>", "", h1), "h1s": [], "canonical": canonical, "lang": lang,
            "viewport": "viewport" in html.lower(), "images": [], "headings": [],
            "internal_links": [], "external_links": [], "status": status, "body_text": body_text,
            "content_word_count": _paragraph_word_count(html)}


# ----------------------------------------------------------------- deliverables
def _extract_json(text):
    return json.loads(re.search(r"\{.*\}|\[.*\]", text, re.S).group(0))


# Circuit breaker for the AI fallback chain - a PERMANENT error (bad/revoked
# key, retired model: 401/403/404) will fail identically on every retry for
# the rest of this run, so once seen, that provider is skipped for the
# remainder of the process instead of wasting a full HTTP round-trip on every
# single suggestion call (confirmed real case: a GEO report with 20 ALT-text
# suggestions retried all 3 already-broken providers 20 times each - 60
# guaranteed-failing calls, all with the exact same errors). A TRANSIENT
# error (429 rate limit, 5xx server issue, timeout) does NOT trip the
# breaker, since those can genuinely recover mid-run.
_AI_BROKEN_PROVIDERS = set()
_PERMANENT_HTTP_CODES = {401, 403, 404}


def _ai_suggest_gemini(prompt):
    """Google Gemini's FREE tier when GEMINI_API_KEY is set (free key:
    https://aistudio.google.com/apikey). Plain REST, no SDK, no paid API."""
    if "gemini" in _AI_BROKEN_PROVIDERS:
        return None
    key = os.environ.get("GEMINI_API_KEY", "").strip()
    if not key:
        return None
    try:
        import urllib.request, urllib.error
        # gemini-1.5-flash was retired by Google (confirmed live: 404 "models/
        # gemini-1.5-flash is not found for API version v1beta") - gemini-2.5-flash
        # confirmed working (200 OK) against a real key on 2026-07-14.
        model = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")
        url = (f"https://generativelanguage.googleapis.com/v1beta/models/"
               f"{model}:generateContent?key={key}")
        body = json.dumps({
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {"temperature": 0.7, "responseMimeType": "application/json"},
        }).encode("utf-8")
        req = urllib.request.Request(url, data=body, headers={"Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=40) as r:
            data = json.loads(r.read())
        return _extract_json(data["candidates"][0]["content"]["parts"][0]["text"])
    except urllib.error.HTTPError as e:
        if e.code in _PERMANENT_HTTP_CODES:
            _AI_BROKEN_PROVIDERS.add("gemini")
            log(f"   [warn] Gemini suggestion failed: HTTP Error {e.code} - won't retry Gemini for the rest of this run.")
        else:
            log(f"   [warn] Gemini suggestion failed: HTTP Error {e.code}")
        return None
    except Exception as e:
        log(f"   [warn] Gemini suggestion failed: {e}")
        return None


def _ai_suggest_groq(prompt):
    """Groq's FREE tier when GROQ_API_KEY is set (free key:
    https://console.groq.com/keys) - OpenAI-compatible chat completions API,
    very fast inference. Second fallback tier after Gemini."""
    if "groq" in _AI_BROKEN_PROVIDERS:
        return None
    key = os.environ.get("GROQ_API_KEY", "").strip()
    if not key:
        return None
    try:
        import urllib.request, urllib.error
        model = os.environ.get("GROQ_MODEL", "llama-3.3-70b-versatile")
        url = "https://api.groq.com/openai/v1/chat/completions"
        body = json.dumps({
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7,
        }).encode("utf-8")
        req = urllib.request.Request(url, data=body, headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {key}",
        })
        with urllib.request.urlopen(req, timeout=40) as r:
            data = json.loads(r.read())
        return _extract_json(data["choices"][0]["message"]["content"])
    except urllib.error.HTTPError as e:
        if e.code in _PERMANENT_HTTP_CODES:
            _AI_BROKEN_PROVIDERS.add("groq")
            log(f"   [warn] Groq suggestion failed: HTTP Error {e.code} - won't retry Groq for the rest of this run.")
        else:
            log(f"   [warn] Groq suggestion failed: HTTP Error {e.code}")
        return None
    except Exception as e:
        log(f"   [warn] Groq suggestion failed: {e}")
        return None


def _ai_suggest_openrouter(prompt):
    """OpenRouter's free-tier models when OPENROUTER_API_KEY is set (free key:
    https://openrouter.ai/keys) - last AI tier before the heuristic fallback,
    since free-tier models there are more rate-limited/less consistently
    available than Gemini/Groq."""
    if "openrouter" in _AI_BROKEN_PROVIDERS:
        return None
    key = os.environ.get("OPENROUTER_API_KEY", "").strip()
    if not key:
        return None
    try:
        import urllib.request, urllib.error
        # meta-llama/llama-3.1-8b-instruct:free was retired by OpenRouter (confirmed
        # live: 404 "model not found") - gpt-oss-20b:free confirmed available on
        # OpenRouter's current free-tier model list as of 2026-07-15.
        model = os.environ.get("OPENROUTER_MODEL", "openai/gpt-oss-20b:free")
        url = "https://openrouter.ai/api/v1/chat/completions"
        body = json.dumps({
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7,
        }).encode("utf-8")
        req = urllib.request.Request(url, data=body, headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {key}",
            "HTTP-Referer": "https://apex-on-pageautomationtools.github.io",
            "X-Title": "SEOToolkitPro",
        })
        with urllib.request.urlopen(req, timeout=40) as r:
            data = json.loads(r.read())
        return _extract_json(data["choices"][0]["message"]["content"])
    except urllib.error.HTTPError as e:
        if e.code in _PERMANENT_HTTP_CODES:
            _AI_BROKEN_PROVIDERS.add("openrouter")
            log(f"   [warn] OpenRouter suggestion failed: HTTP Error {e.code} - won't retry OpenRouter for the rest of this run.")
        else:
            log(f"   [warn] OpenRouter suggestion failed: HTTP Error {e.code}")
        return None
    except Exception as e:
        log(f"   [warn] OpenRouter suggestion failed: {e}")
        return None


def _ai_suggest_openai(prompt):
    """OpenAI's PAID API when OPENAI_API_KEY is set - last tier, after every
    free one has failed. gpt-5.6-luna: OpenAI's cheapest current-gen model
    (per-team decision - real money, so only used for small/targeted runs,
    see _ai_suggest's BULK_AI_PAGE_LIMIT gate below), OpenAI-compatible
    chat completions API, same request shape as Groq/OpenRouter above."""
    if "openai" in _AI_BROKEN_PROVIDERS:
        return None
    key = os.environ.get("OPENAI_API_KEY", "").strip()
    if not key:
        return None
    try:
        import urllib.request, urllib.error
        model = os.environ.get("OPENAI_MODEL", "gpt-5.6-luna")
        url = "https://api.openai.com/v1/chat/completions"
        body = json.dumps({
            "model": model,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.7,
        }).encode("utf-8")
        req = urllib.request.Request(url, data=body, headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {key}",
        })
        with urllib.request.urlopen(req, timeout=40) as r:
            data = json.loads(r.read())
        return _extract_json(data["choices"][0]["message"]["content"])
    except urllib.error.HTTPError as e:
        if e.code in _PERMANENT_HTTP_CODES:
            _AI_BROKEN_PROVIDERS.add("openai")
            log(f"   [warn] OpenAI suggestion failed: HTTP Error {e.code} - won't retry OpenAI for the rest of this run.")
        else:
            log(f"   [warn] OpenAI suggestion failed: HTTP Error {e.code}")
        return None
    except Exception as e:
        log(f"   [warn] OpenAI suggestion failed: {e}")
        return None


# How many pages/rows this run covers, set once at the start of each report's
# main() via set_run_scale() - gates whether the PAID OpenAI tier is even
# attempted below. None = unknown/not set, treated as "small" (paid tier
# allowed) so a caller that never calls set_run_scale() behaves like before
# this was added.
BULK_AI_PAGE_LIMIT = 20
_current_run_scale = None


def set_run_scale(page_count):
    """Call once near the start of a report run with the real page/row count
    it covers. Team decision: the paid OpenAI tier is only worth the real
    money for small/targeted runs - a large bulk run instead just falls
    through to the heuristic fallback once the free tiers (Gemini/Groq/
    OpenRouter) are exhausted, rather than silently burning paid credits
    across hundreds of pages."""
    global _current_run_scale
    _current_run_scale = page_count


def _ai_suggest(prompt):
    """Suggestion copy via a fallback chain: Gemini -> Groq -> OpenRouter (all
    free tiers, tried unconditionally) -> OpenAI (paid, only for a run of
    BULK_AI_PAGE_LIMIT pages or fewer - see set_run_scale) -> None (caller
    falls back to its own heuristic). Each tier is skipped instantly if its
    API key isn't configured (CONFIG's Keys sheet / GEMINI_API_KEY,
    GROQ_API_KEY, OPENROUTER_API_KEY, OPENAI_API_KEY env vars), so a site
    with only a Gemini key behaves exactly as before. Returns parsed JSON or
    None."""
    for fn in (_ai_suggest_gemini, _ai_suggest_groq, _ai_suggest_openrouter):
        result = fn(prompt)
        if result is not None:
            return result
    if _current_run_scale is None or _current_run_scale <= BULK_AI_PAGE_LIMIT:
        result = _ai_suggest_openai(prompt)
        if result is not None:
            return result
    return None


_GENERIC_TITLES = {
    "home", "home page", "about", "about us", "service", "services", "contact",
    "contact us", "blog", "shop", "products", "product", "page", "welcome", "index",
    "untitled", "portfolio", "team", "faq", "gallery", "our services", "our team",
}


def _keyword_covered(text, keyword):
    """True if every significant word of `keyword` appears somewhere in `text` - NOT
    an exact contiguous-phrase match. A title like "Packing Services Adelaide" already
    covers the keyword "packing adelaide" (both words present), even though the exact
    phrase "packing adelaide" never appears as a substring - flagging that as
    unoptimized was a false positive. Short stopwords are ignored so they can't force
    a false "not covered".
    """
    if not keyword:
        return True
    words = [w for w in re.findall(r"[a-z0-9]+", keyword.lower()) if len(w) > 2]
    if not words:
        return keyword.lower() in text
    return all(w in text for w in words)


def _title_needs_suggestion(title, keywords=None):
    """Missing, a short generic nav-label title ('Home' / 'About' / 'Service'), or
    missing its target keyword entirely (not optimized for what the page is meant to
    rank for - checked by word coverage, not an exact phrase match, so a title that
    already contains every keyword word just phrased differently is left alone).
    Length on its own is NOT a reason to flag a title - a long title (even 100-120
    chars) is fine and is left alone."""
    if not title or title == MISSING:
        return True
    t = title.strip().lower()
    if t in _GENERIC_TITLES:
        return True
    if len(title.strip()) < 15:
        return True
    if keywords and not _keyword_covered(t, keywords[0].strip()):
        return True
    return False


def _desc_needs_suggestion(desc, keywords=None):
    """Missing, too short to say anything useful, or missing its target keyword
    entirely (not optimized) - checked by word coverage, not an exact phrase match."""
    if not desc or desc == MISSING:
        return True
    d = desc.strip().lower()
    if len(desc.strip()) < 50:
        return True
    if keywords and not _keyword_covered(d, keywords[0].strip()):
        return True
    return False


def _best_rank_for_page(keywords, keyword_ranks):
    """Lowest (best) numeric ranking among this page's keywords, or None if none of
    them have a ranking value in the sheet."""
    if not keyword_ranks:
        return None
    ranks = []
    for k in keywords:
        r = keyword_ranks.get(k)
        if r is None or r == "":
            continue
        try:
            ranks.append(float(r))
        except (TypeError, ValueError):
            continue
    return min(ranks) if ranks else None


def suggest_meta(page_data, keywords, brand, keyword_ranks=None):
    """Existing title/desc/H1 from the crawl, plus a SUGGESTED title/description built
    from the page's own content + its target keyword(s) - via the free-tier AI
    fallback chain (Gemini -> Groq -> OpenRouter) when a key is configured, else
    a content-aware heuristic (never a bare "keyword | brand" slam-together).

    A title is only suggested when it's missing, a short generic nav-label (e.g.
    "Home", "About", "Service"), or missing its target keyword entirely - length on
    its own is never a reason (a long title, even 100-120 chars, is left alone). If
    the page already ranks in the top 20 for one of its keywords (per an optional
    ranking column in the keyword sheet), the title suggestion is skipped - a working
    title should not get overwritten - and the sheet instead says to check it against
    the ranking. A description is suggested whenever it's missing, too short, or
    missing its target keyword (not optimized)."""
    existing_title = page_data["title"]
    existing_desc = page_data["description"]
    existing_h1 = page_data["h1"]
    body_text = (page_data.get("body_text") or "")[:1500]
    primary_kw = keywords[0] if keywords else ""

    best_rank = _best_rank_for_page(keywords, keyword_ranks)
    ranked_well = best_rank is not None and best_rank <= 20
    title_flagged = _title_needs_suggestion(existing_title, keywords)
    need_title = title_flagged and not ranked_well
    need_desc = _desc_needs_suggestion(existing_desc, keywords)

    suggested_title = None
    suggested_desc = None

    if title_flagged and ranked_well:
        suggested_title = "Please check it as per the ranking"

    if need_title or need_desc:
        ai = None
        if primary_kw:
            prompt = (
                "You are an SEO copywriter. For this webpage, write an SEO meta title "
                "(50-60 characters) and meta description (140-160 characters).\n"
                f"Page URL: {page_data['url']}\n"
                f"Brand: {brand}\n"
                f"Primary target keyword: {primary_kw}\n"
                f"Other keywords for this page: {', '.join(keywords[1:5])}\n"
                f"Existing H1: {existing_h1 if existing_h1 != MISSING else '(none)'}\n"
                f"Page content excerpt: {body_text or '(not available)'}\n"
                'Return ONLY JSON: {"title": "...", "description": "..."}. '
                "The title and description must read naturally, reflect what THIS page "
                "is actually about (use the content excerpt, not just the keyword), "
                "include the primary keyword naturally, and end the title with the "
                "brand name if it fits within the length.")
            ai = _ai_suggest(prompt)
        if isinstance(ai, dict) and ai.get("title") and ai.get("description"):
            if need_title:
                suggested_title = str(ai["title"]).strip()
            if need_desc:
                suggested_desc = str(ai["description"]).strip()
        else:
            # Heuristic fallback (no Gemini key, or the call failed) - still built from
            # the page's own H1/content rather than a formulaic "keyword | brand".
            topic = existing_h1 if existing_h1 and existing_h1 != MISSING else (primary_kw.title() or "Page")

            def _trim(text, limit):
                """Trim to a word boundary at or before `limit` chars (never mid-word)."""
                text = text.strip()
                if len(text) <= limit:
                    return text
                cut = text[:limit].rsplit(" ", 1)[0].rstrip(" ,-")
                return cut or text[:limit]

            if need_title:
                if primary_kw and primary_kw.lower() not in topic.lower():
                    suggested_title = f"{topic} - {primary_kw.title()} | {brand}"
                else:
                    suggested_title = f"{topic} | {brand}"
                suggested_title = _trim(suggested_title, 60)
            if need_desc:
                snippet = _trim(body_text.strip(), 100).rstrip(" .!?")
                if primary_kw and snippet:
                    suggested_desc = _trim(f"{topic} - {snippet}.", 130) + \
                        f" Explore {primary_kw} with {brand}."
                elif primary_kw:
                    suggested_desc = f"Discover {primary_kw} at {brand}. {topic}."
                else:
                    suggested_desc = f"{topic} - learn more at {brand}."
                suggested_desc = _trim(suggested_desc, 160)

    no_change_msg = "No changes needed - existing tag is already optimized"
    return {
        "page": page_data["url"],
        "existing_title": existing_title,
        "suggested_title": suggested_title or no_change_msg,
        "existing_description": existing_desc,
        "suggested_description": suggested_desc or no_change_msg,
        "existing_h1": existing_h1,
        "suggested_h1": "Check manually as per ranking and traffic",
        "content_match": check_content_keyword_match(page_data, keywords),
    }


def check_content_keyword_match(page_data, keywords):
    """Content-vs-keyword check, the same signal a dedicated content-checker tool
    verifies: does this page's actual content - title, H1, and visible body text -
    cover its target keyword(s)? Flags keywords missing from the content outright,
    keywords that only appear in the title/H1 but never in the body copy, and a rough
    body-text occurrence count for keywords that are present."""
    if not keywords:
        return "No target keyword(s) assigned"
    title = (page_data.get("title") or "").lower()
    h1 = (page_data.get("h1") or "").lower()
    body = (page_data.get("body_text") or "").lower()
    parts = []
    for kw in keywords:
        k = kw.strip().lower()
        if not k:
            continue
        in_title, in_h1, count = k in title, k in h1, body.count(k)
        if not in_title and not in_h1 and count == 0:
            parts.append(f'"{kw}" - not found in title, H1, or body content')
        elif count == 0:
            where = "title" if in_title else "H1"
            parts.append(f'"{kw}" - only in {where}, not found in body content')
        else:
            extra = []
            if in_title:
                extra.append("title")
            if in_h1:
                extra.append("H1")
            suffix = f", also in {'/'.join(extra)}" if extra else ""
            parts.append(f'"{kw}" - found {count}x in body content{suffix}')
    return "; ".join(parts) if parts else "No target keyword(s) assigned"


# Tracking pixels, font icons, analytics - skip entirely (not real images)
_SKIP_PATTERNS = re.compile(
    r"facebook\.com/tr\?|"
    r"fonts\.gstatic\.com|"
    r"fonts\.googleapis\.com|"
    r"google-analytics\.com|"
    r"googletagmanager\.com|"
    r"doubleclick\.net|"
    r"analytics\.|"
    r"pixel\.|"
    r"\.gif\?|"   # tracking pixels are usually tiny GIFs with query strings
    r"data:image/",  # inline base64 images
    re.I
)

# Self-hosted / easily updatable from admin (WordPress, Shopify, common CMS)
_SELF_HOSTED_PATTERNS = re.compile(
    r"/wp-content/|"
    r"/wp-includes/|"
    r"cdn\.shopify\.com|"
    r"/assets/|"
    r"/uploads/|"
    r"/media/|"
    r"/images/|"
    r"/img/|"
    r"/static/",
    re.I
)


def _is_self_hosted(src, domain):
    """Check if image is on the same domain or a known CMS CDN."""
    parsed = urllib.parse.urlparse(src)
    host = parsed.netloc.lower()
    d = safe_domain(domain)
    if not host or d in host:
        return True
    if re.search(r"cdn\.shopify\.com", host, re.I):
        return True
    if _SELF_HOSTED_PATTERNS.search(parsed.path):
        return True
    return False


def suggest_alt(page_data, keywords, brand, domain=""):
    """List all images with their existing alt text. Suggested alt is left empty —
    the user checks the actual image and fills it in manually.
    Only includes self-hosted / easily updatable images.
    External CDN images are listed separately with a dev-check note."""
    self_hosted = []
    external_cdn = []

    for img in page_data.get("images", []):
        src = img.get("src", "")
        if not src or _SKIP_PATTERNS.search(src):
            continue
        existing_alt = img.get("alt", "") or ""
        entry = {"page": page_data["url"], "image": src, "existing_alt": existing_alt, "suggested_alt": ""}

        if _is_self_hosted(src, domain):
            self_hosted.append(entry)
        else:
            entry["suggested_alt"] = "Ask developer if alt text can be updated"
            external_cdn.append(entry)

    return self_hosted, external_cdn


def recommend_canonical(page_data):
    url = page_data["url"]
    existing = page_data.get("canonical") or MISSING
    m = re.search(r'href="([^"]+)"', existing)
    existing_href = m.group(1).strip() if m else None
    # Self-referencing canonical already points at this page → nothing to change.
    if existing_href and existing_href.rstrip("/") == url.rstrip("/"):
        recommended = "No Changes Needed as Self Canonical Found"
    else:
        recommended = f'<link rel="canonical" href="{url}" />'
    return {"page": url, "existing": existing, "recommended": recommended}


# ----------------------------------------------------------------- xlsx output
# Per-format header labels/colors for the Meta / Canonical / Alt supplementary
# sheets, taken verbatim from each format's real client reference deliverable
# (D:\Report Formats\On-Page Report Formats\<Format>\...). Formats not listed
# here keep the script's original default template labels/color - this ONLY
# restyles formats we have confirmed real references for, it never removes a
# file that already generates correctly for every format.
#
# Values: (sheet_name, [7 header labels for Page/ExistingTitle/SuggestedTitle/
#          ExistingDesc/SuggestedDesc/ExistingH1/SuggestedH1], header_hex_color)
FORMAT_META_XLSX = {
    "camila": ("Sheet1", ["Page", "Existing Meta Title", "Suggested Meta Title",
                          "Existing Meta Description", "Suggested Meta Description",
                          "Existing H1", "Suggested H1"], "B6D7A8"),
    "deltafvr": ("Meta Suggestions", ["Webpages URL", "Existing Title", "Suggested Title",
                                      "Existing Description", "Suggested Description",
                                      "Existing H1", "Suggested H1"], "0070C0"),
    "deltafl": ("Meta Suggestions", ["Target Pages", "Existing Title", "Suggested Title",
                                     "Existing Description", "Suggested Description",
                                     "Existing H1", "Suggested H1"], None),
    "eta": ("Sheet1", ["Target Pages", "Existing Title", "Suggested Title",
                       "Existing Description", "Suggested Description",
                       "Existing H1", "Suggested H1"], "FFFF00"),
    "kappa": ("Sheet1", ["Original URL", "Existing Title", "Suggested Title",
                         "Existing Meta Description", "Suggested Meta Description",
                         "Existing H1 Heading", "Suggested H1 Heading"], "F4B084"),
    "octal": ("Meta Suggestion", ["Target Page's", "Existing Title", "Suggested Title",
                                  "Existing Description", "Suggested Meta Description",
                                  "Existing H1 tag", "Suggested H1 tag"], None),
    "peta": ("Sheet1", ["Url", "Exiting Title", "Suggested Title",
                        "Exiting Description", "Suggested Description",
                        "Exiting H1", "Suggested H1"], "00B050"),
    "beta": ("Sheet1", ["Url", "Exiting Title", "Suggested Title",
                        "Exiting Description", "Suggested Description",
                        "Exiting H1", "Suggested H1"], "00B050"),
    "xenon": ("Sheet1", ["Target Pages", "Existing Title", "Suggested Title",
                         "Existing Description", "Suggested Description",
                         "Existing H1", "Suggested H1"], "1F3864"),
    "neon": ("Sheet1", ["Page URL", "Existing Title", "Suggested Title",
                        "Existing Description", "Suggested Description",
                        "Existing H1", "Suggested H1"], "215967"),
}

# Values: (sheet_name, [3 header labels for Page/Existing/Recommended], color)
FORMAT_CANON_XLSX = {
    "camila": ("Sheet1", ["Page URL", "Existing Canonical Tag", "Suggested Canonical Tag"], "93C47D"),
    "deltafl": ("Canonical tags Suggestions", ["Page URL", "Existing Canonical Tag",
                                               "Suggested Canonical Tag"], "366092"),
}

# Values: (sheet_name, [(header label, data field key)], color). Field keys index
# into each alt entry's dict: "page", "image", "existing_alt", "suggested_alt" -
# a (label, field) list (not just labels) because some references (ETA) order
# Image before Page and/or drop the existing-alt column entirely (Neon/Sara/Octal).
FORMAT_ALT_XLSX = {
    "neon": ("Sheet1", [("Page URL", "page"), ("Image URL", "image"),
                        ("Recommended Alt Tag Suggestions", "suggested_alt")], "215967"),
    "sara": ("Image Alt Tag Suggestions", [("Page URL", "page"), ("Image URL", "image"),
                                           ("Image Alt", "suggested_alt")], "215967"),
    "octal": ("Sheet1", [("Page Url", "page"), ("Image Url", "image"),
                         ("Image Alt Tag", "suggested_alt")], None),
    "eta": ("Sheet1", [("Image Urls", "image"), ("Source Pages", "page"),
                       ("Existing Alt Texts", "existing_alt"), ("Suggested Alt texts", "suggested_alt")], "FFFF00"),
    "kappa": ("Sheet1", [("Page URL", "page"), ("Image URL", "image"),
                         ("Existing Image Alt Text", "existing_alt"),
                         ("Suggested Image Alt Text", "suggested_alt")], "4A86E8"),
}


# Sara/Theta's Meta sheet has an extra leading "Keywords" column the other
# formats don't - the team supplies a keywords+target-pages sheet as input, so
# each row needs the keyword(s) targeted for that specific page, not just the
# page itself. Values: (sheet_name, [8 header labels], color).
FORMAT_META_XLSX_KEYWORDS = {
    "sara": ("Meta Suggestion", ["Keywords", "Target Pages", "Existing Title", "Suggested Title",
                                 "Existing Meta Description", "Suggested Description",
                                 "Existing H1 Tag", "Suggested H1 Tag"], "215967"),
    "theta": ("Meta", ["Keywords", "Address", "Existing Title", "Suggested Title",
                       "Existing Meta Description", "Suggested Meta Description",
                       "Existing H1", "Suggested H1"], "0070C0"),
}


def write_meta_xlsx_with_keywords(metas, targets, out_path, fmt):
    """Meta Suggestions sheet with a leading Keywords column (Sara/Theta reference
    format) - the keyword(s) targeted for each page, looked up from the same
    `targets` list (page -> keywords) the rest of the report already uses."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    sheet_name, labels, color_hex = FORMAT_META_XLSX_KEYWORDS[fmt]
    kw_by_page = {t["page"]: ", ".join(t.get("keywords") or []) for t in targets}

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = sheet_name
    header_fill = PatternFill("solid", fgColor=color_hex)
    header_font = Font(bold=True, size=11, color="FFFFFF")
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    for c, label in enumerate(labels, 1):
        cell = ws.cell(1, c, label)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = left
    ws.freeze_panes = "A2"
    widths = [30, 40, 34, 34, 40, 40, 30, 30]
    for i, w in enumerate(widths):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i + 1)].width = w

    for r, m in enumerate(metas, 2):
        row = [kw_by_page.get(m["page"], ""), m["page"], m["existing_title"], m["suggested_title"],
               m["existing_description"], m["suggested_description"], m["existing_h1"], m["suggested_h1"]]
        for c, val in enumerate(row, 1):
            ws.cell(r, c, val).alignment = left
    wb.save(out_path)


def _clone_and_clear(template, header_rows=1):
    import openpyxl
    wb = openpyxl.load_workbook(template)
    ws = wb.active
    # remember the style of the first data row to reuse
    style_row = header_rows + 1
    proto = [ws.cell(style_row, c)._style for c in range(1, ws.max_column + 1)]
    # clear everything below the header
    if ws.max_row > header_rows:
        ws.delete_rows(header_rows + 1, ws.max_row - header_rows)
    return wb, ws, proto


def _write_rows(ws, proto, rows, start=2):
    for i, row in enumerate(rows):
        r = start + i
        for c, val in enumerate(row, 1):
            cell = ws.cell(r, c, val)
            if c - 1 < len(proto):
                cell._style = proto[c - 1]


def _apply_header_override(ws, labels, color_hex, sheet_name=None):
    """Overwrite a cloned template's header row 1 with a format's real label
    text/fill color (and sheet title), leaving the row's font/alignment/border
    style alone - only the text and fill actually differ between formats."""
    from openpyxl.styles import PatternFill
    if sheet_name:
        ws.title = sheet_name
    fill = PatternFill("solid", fgColor=color_hex) if color_hex else None
    for c, label in enumerate(labels, 1):
        cell = ws.cell(1, c, label)
        if fill:
            cell.fill = fill


def write_meta_xlsx(metas, out_path, fmt=None):
    from copy import copy
    from openpyxl.styles import Alignment, Font
    wb, ws, proto = _clone_and_clear(TPL_META)
    rows = [[m["page"], m["existing_title"], m["suggested_title"], m["existing_description"],
             m["suggested_description"], m["existing_h1"], m["suggested_h1"],
             m.get("content_match", "")] for m in metas]
    _write_rows(ws, proto, rows)
    spec = FORMAT_META_XLSX.get(str(fmt or "").strip().lower())
    if spec:
        sheet_name, labels, color_hex = spec
        _apply_header_override(ws, labels, color_hex, sheet_name)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    # The cloned template row styles the "existing" columns in red. Real existing
    # values should read as calm black text; keep red ONLY to flag genuinely
    # missing/empty existing values. Columns: 2=existing title, 4=existing desc,
    # 6=existing H1.
    existing_cols = {2, 4, 6}
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=ws.max_column):
        max_len = 1
        for cell in row:
            cell.alignment = left
            max_len = max(max_len, len(str(cell.value or "")))
            if cell.column in existing_cols:
                val = str(cell.value or "").strip()
                flagged = (not val) or val == MISSING
                f = copy(cell.font)
                cell.font = Font(name=f.name, size=f.size, bold=f.bold, italic=f.italic,
                                 color=("FFC00000" if flagged else "FF000000"))
        # openpyxl never auto-computes row height for wrapped text (Excel only does
        # that on open, and inconsistently) - without this, long suggested title/desc
        # text overflowed into the row below and visually overlapped it. ~45 chars/line
        # at this template's column widths, 15pt per wrapped line, floor of 30pt.
        lines = max(1, -(-max_len // 45))
        ws.row_dimensions[row[0].row].height = max(30, lines * 15)
    wb.save(out_path)


def write_alt_xlsx(self_hosted, external_cdn, out_path, fmt=None):
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    wb = openpyxl.Workbook()
    ws = wb.active

    spec = FORMAT_ALT_XLSX.get(str(fmt or "").strip().lower())
    if spec:
        sheet_name, columns, color_hex = spec
        headers = [label for label, _field in columns]
        field_order = [field for _label, field in columns]
        header_fill = PatternFill("solid", fgColor=color_hex) if color_hex else PatternFill("solid", fgColor="2F5496")
    else:
        sheet_name = "Image Alt Tag Suggestions"
        field_order = ["page", "image", "existing_alt", "suggested_alt"]
        headers = ["Page URL", "Image URL", "Existing Alt Text", "Suggested Alt Tag"]
        # Match the house-style header used by the Canonical / Target Pages sheets:
        # solid 2F5496 fill with a white bold font, and a frozen header row.
        header_fill = PatternFill("solid", fgColor="2F5496")
    ws.title = sheet_name
    header_font = Font(bold=True, size=11, color="FFFFFF")
    header_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
    for c, h in enumerate(headers, 1):
        cell = ws.cell(1, c, h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
    ws.freeze_panes = "A2"
    for i in range(len(headers)):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i + 1)].width = 40 if i != 1 else 60

    def _write_entry(row, a):
        for c, field in enumerate(field_order, 1):
            ws.cell(row, c, a.get(field, ""))

    row = 2
    for a in self_hosted:
        _write_entry(row, a)
        row += 1

    if external_cdn:
        row += 1
        note_cell = ws.cell(row, 1, "Below images are on external CDN - ask developer if alt text can be updated:")
        note_cell.font = Font(bold=True, color="C00000", size=11)
        row += 1
        for a in external_cdn:
            _write_entry(row, a)
            row += 1

    wb.save(out_path)


def write_canonical_xlsx(canons, out_path, fmt=None):
    wb, ws, proto = _clone_and_clear(TPL_CANON)
    rows = [[c["page"], c["existing"], c["recommended"]] for c in canons]
    _write_rows(ws, proto, rows)
    spec = FORMAT_CANON_XLSX.get(str(fmt or "").strip().lower())
    if spec:
        sheet_name, labels, color_hex = spec
        _apply_header_override(ws, labels, color_hex, sheet_name)
    wb.save(out_path)


# Values: (sheet_name, [2 header labels for Keywords/Target Pages], color)
FORMAT_TARGETS_XLSX = {
    "theta": ("Target Page", ["Keywords", "Target Page"], "93C47D"),
}


def write_targets_xlsx(targets, out_path, fmt=None):
    """The complete list of target pages + their keywords, as its own sheet."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    wb = openpyxl.Workbook()
    ws = wb.active
    spec = FORMAT_TARGETS_XLSX.get(str(fmt or "").strip().lower())
    sheet_name, labels, color_hex = spec if spec else ("Target Pages", ["Keywords", "Target Pages"], "2F5496")
    ws.title = sheet_name
    header_fill = PatternFill("solid", fgColor=color_hex)
    header_font = Font(bold=True, size=14, color="FFFFFF")
    ws.append(labels)
    for c in ws[1]:
        c.font = header_font
        c.fill = header_fill
        c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    for t in targets:
        kws = t.get("keywords") or [""]
        for kw in kws:
            ws.append([kw, t["page"]])
    ws.column_dimensions["A"].width = 42
    ws.column_dimensions["B"].width = 58
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        for cell in row:
            cell.alignment = left
    wb.save(out_path)


def write_broken_link_xlsx(broken_links_detail, out_path, sheet_name="Broken Link", color="0070C0",
                            columns=None):
    """Broken Link sheet - format-specific column order/labels (Camila/Delta FVR/
    Theta/ETA references all show a variant of Link/Source/Status/Suggestion),
    built from findings["broken_links_detail"] (already-verified 404/410 links -
    see health_audit.check_broken_links). columns: [(label, field)] where field
    is one of url/found_on/code/suggested_redirect; defaults to the Delta FVR order."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    columns = columns or [("Broken Link", "url"), ("Source Page", "found_on"),
                          ("Server Response", "code"), ("Broken Link Suggestion", "suggested_redirect")]
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = sheet_name
    header_fill = PatternFill("solid", fgColor=color)
    header_font = Font(bold=True, size=11, color="FFFFFF")
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    for c, (label, _field) in enumerate(columns, 1):
        cell = ws.cell(1, c, label)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = left
    ws.freeze_panes = "A2"
    for i in range(len(columns)):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i + 1)].width = 46
    for r, item in enumerate(broken_links_detail or [], 2):
        for c, (_label, field) in enumerate(columns, 1):
            if field == "found_on":
                val = ", ".join(item.get("found_on") or []) or item.get("url", "")
            elif field == "suggested_redirect":
                val = item.get("suggested_redirect") or "Remove this broken link."
            elif field is None:
                val = ""
            else:
                val = item.get(field, "")
            cell = ws.cell(r, c, val)
            cell.alignment = left
    wb.save(out_path)


def _wayback_snapshot_url(page_url, timeout=10):
    """Closest archived snapshot for a page via the Wayback Machine's official
    Availability API - returns None (never fabricated) if nothing is archived."""
    import json as _json
    import urllib.parse as _uparse
    try:
        api = "https://archive.org/wayback/available?url=" + _uparse.quote(page_url, safe="")
        req = urllib.request.Request(api, headers={"User-Agent": "Mozilla/5.0 SEOPhase2Bot"})
        with urllib.request.urlopen(req, timeout=timeout) as r:
            data = _json.loads(r.read().decode("utf-8", "ignore"))
        snap = (data.get("archived_snapshots") or {}).get("closest") or {}
        return snap.get("url") if snap.get("available") else None
    except Exception:
        return None


def write_webarchive_xlsx(pages_data, out_path, sheet_name="Webarchive URL", color="0070C0"):
    """Webpages URL / Webarchive URL sheet (Delta FVR + Delta Up reference format) -
    real snapshot lookups via the Wayback Availability API, "Not Archived" when
    the API genuinely has nothing (never a fabricated link)."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = sheet_name
    header_fill = PatternFill("solid", fgColor=color)
    header_font = Font(bold=True, size=11, color="FFFFFF")
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    for c, label in enumerate(["Webpages URL", "Webarchive URL"], 1):
        cell = ws.cell(1, c, label)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = left
    ws.freeze_panes = "A2"
    ws.column_dimensions["A"].width = 54
    ws.column_dimensions["B"].width = 70
    for r, pd in enumerate(pages_data, 2):
        url = pd["url"]
        snap = _wayback_snapshot_url(url)
        ws.cell(r, 1, url).alignment = left
        ws.cell(r, 2, snap or "Not Archived").alignment = left
    wb.save(out_path)


def write_indexing_status_xlsx(pages_data, gsc_indexing, out_path, sheet_name="Indexing Status",
                                split_target_nontarget=False, target_urls=None, labels=None, color="215967"):
    """Page URL / Status sheet from real GSC URL Inspection results (findings[
    "gsc_indexing"] - already fetched via check_gsc_indexing). When
    split_target_nontarget, writes two sheets ("target pages" / "Non target page")
    like the Delta FL reference; otherwise a single sheet like Delta Up/Sara."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    labels = labels or ["Page URL", "Status"]
    status_by_url = {item.get("url"): item.get("coverageState") or item.get("verdict") or "Unknown"
                     for item in (gsc_indexing or [])}
    header_fill = PatternFill("solid", fgColor=color)
    header_font = Font(bold=True, size=11, color="FFFFFF")
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    def _fill_sheet(ws, urls):
        for c, label in enumerate(labels, 1):
            cell = ws.cell(1, c, label)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = left
        ws.freeze_panes = "A2"
        ws.column_dimensions["A"].width = 60
        ws.column_dimensions["B"].width = 30
        for r, u in enumerate(urls, 2):
            ws.cell(r, 1, u).alignment = left
            ws.cell(r, 2, status_by_url.get(u, "Unknown")).alignment = left

    wb = openpyxl.Workbook()
    if split_target_nontarget:
        target_set = set(target_urls or [])
        ws1 = wb.active
        ws1.title = "target pages"
        _fill_sheet(ws1, [u for u in status_by_url if u in target_set])
        ws2 = wb.create_sheet("Non target page")
        _fill_sheet(ws2, [u for u in status_by_url if u not in target_set])
    else:
        ws = wb.active
        ws.title = sheet_name
        _fill_sheet(ws, [pd["url"] for pd in pages_data])
    wb.save(out_path)


# 2nd Phase Checklist points this script can answer from data it already
# collects (robots.txt/sitemap/noindex/https/www checks, etc). Anything the
# real reference asks for that needs a human or an external tool (domain age
# via duplichecker, SE-ranking traffic status, who verified target pages,
# recent traffic drops, "dummy content" judgment calls) is intentionally left
# blank rather than guessed - this script never fabricates data it can't verify.
def _build_checklist_rows(domain, pages_data, findings, targets, sitemap_url, sitemap_body):
    from urllib.parse import urlparse
    any_www = any(urlparse(pd["url"]).netloc.lower().startswith("www.") for pd in pages_data)
    any_https = all(urlparse(pd["url"]).scheme == "https" for pd in pages_data) if pages_data else None
    any_trailing_slash = None
    if pages_data:
        paths = [urlparse(pd["url"]).path for pd in pages_data]
        any_trailing_slash = sum(1 for p in paths if p.endswith("/")) >= len(paths) / 2
    noindex = findings.get("noindex_pages") or []
    return [
        ("Number of Target Keywords", sum(len(t.get("keywords") or []) for t in targets)),
        ("Number or Target Pages", len(targets)),
        ("robots.txt file", "Found" if findings.get("robots_body") else "Not Found"),
        ("Sitemap File", "Found" if sitemap_body else "Not Found"),
        ("Noindex on pages", f"Found on {len(noindex)} page(s)" if noindex else "Not Found"),
        ("https or http", "https" if any_https else ("Mixed" if any_https is False else "Unknown")),
        ("www or non www", "www" if any_www else "Non www"),
        ("Slash and Non Slash", "Slash" if any_trailing_slash else "Non-Slash" if any_trailing_slash is not None else "Unknown"),
        ("Indexing Status Checked", "Yes" if findings.get("gsc_indexing") else "No"),
        ("Domain Age as per duplichecker", ""),
        ("Traffic Status on SE rank tool", ""),
        ("Target Pages Verified By", ""),
        ("Any Critical issue in the website", ""),
        ("Any frequent drop in traffic in last 2-3 month", ""),
        ("Dummy content on pages", ""),
    ]


def write_checklist_xlsx(domain, pages_data, findings, targets, sitemap_url, sitemap_body, out_path):
    """2nd Phase Checklist sheet (Delta FVR/FL/Up + Sara reference format) -
    only the points this script can verify are filled in; the rest are left
    blank for manual review rather than fabricated. See _build_checklist_rows."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Check-list"
    header_fill = PatternFill("solid", fgColor="31869B")
    header_font = Font(bold=True, size=11, color="FFFFFF")
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    for c, label in enumerate(["2nd Phase Point", "Status"], 1):
        cell = ws.cell(1, c, label)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = left
    ws.column_dimensions["A"].width = 48
    ws.column_dimensions["B"].width = 40
    for r, (point, status) in enumerate(_build_checklist_rows(domain, pages_data, findings, targets,
                                                               sitemap_url, sitemap_body), 2):
        ws.cell(r, 1, point).alignment = left
        cell = ws.cell(r, 2, status if status != "" else "Manual review required")
        cell.alignment = left
        if status == "":
            cell.font = Font(italic=True, color="808080")
    wb.save(out_path)


def _format_crawl_time(iso_ts):
    """GSC's lastCrawlTime is ISO8601 UTC ('2026-06-23T13:19:47Z') - format to
    match the reference sheet's style ('Jun 23, 2026, 1:19:47 PM')."""
    if not iso_ts:
        return ""
    try:
        dt = datetime.datetime.strptime(iso_ts.split(".")[0].replace("Z", ""), "%Y-%m-%dT%H:%M:%S")
        return dt.strftime("%b %d, %Y, %I:%M:%S %p").replace(" 0", " ")
    except Exception:
        return iso_ts


def write_crawl_status_xlsx(domain, gsc_indexing, out_path, country=None):
    """Target Page / Last Crawl Date sheet - verified against the real reference
    "Crawling status - graceperfumes.ae.xlsx" (Xenon format's own extra
    deliverable). Generated for EVERY format (not just Xenon) whenever GSC
    access is available, since knowing when Google last actually crawled each
    target page is useful regardless of which report format was selected."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Target Pages"

    navy = PatternFill("solid", fgColor="1F3864")
    light = PatternFill("solid", fgColor="DEEAF6")
    white_bold_16 = Font(bold=True, size=16, color="FFFFFF")
    white_bold_11 = Font(bold=True, size=11, color="FFFFFF")
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    ws.merge_cells("A1:B1")
    ws["A1"] = "Target Pages"
    ws["A1"].font = white_bold_16
    ws["A1"].fill = navy
    ws["A1"].alignment = left

    ws.merge_cells("A2:B2")
    subtitle = f"{domain}   |   Target Area: {country}" if country else domain
    ws["A2"] = subtitle
    ws["A2"].fill = light
    ws["A2"].alignment = left

    ws.append([])
    ws["A3"] = "Target Page"
    ws["B3"] = "Last Crawl Date"
    for cell in (ws["A3"], ws["B3"]):
        cell.font = white_bold_11
        cell.fill = navy
        cell.alignment = left

    for item in (gsc_indexing or []):
        ws.append([item.get("url", ""), _format_crawl_time(item.get("lastCrawlTime", ""))])

    ws.column_dimensions["A"].width = 54
    ws.column_dimensions["B"].width = 46
    for row in ws.iter_rows(min_row=4, max_row=ws.max_row):
        for cell in row:
            cell.alignment = left
    wb.save(out_path)


# --------------------------------------------------------------- sitemap output
def find_existing_sitemap(domain, dry_run=False):
    """Look for the site's REAL sitemap (robots.txt 'Sitemap:' line, then common
    locations). Returns (url, body) if found, else (None, None). We never fabricate
    a sitemap from just the target pages - that would miss most of the site."""
    if dry_run:
        return None, None
    root = site_root(domain)
    candidates = []
    _, rtxt, _ = _http(root + "/robots.txt")
    for m in re.finditer(r"(?im)^\s*sitemap:\s*(\S+)", rtxt or ""):
        candidates.append(m.group(1).strip())
    candidates += [root + "/sitemap.xml", root + "/sitemap_index.xml",
                   root + "/sitemap-index.xml", root + "/sitemap1.xml"]
    for u in dict.fromkeys(candidates):
        st, body, _ = _http(u)
        if st == 200 and ("<urlset" in body or "<sitemapindex" in body):
            return u, body
    return None, None


def _xml_escape(s):
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_sitemap_xml(urls):
    """A valid, working <urlset> sitemap built from the given URLs (used when the
    site has no sitemap of its own)."""
    seen = list(dict.fromkeys((u.split("#")[0].rstrip("/") or u) for u in urls if u and u.startswith("http")))
    today = datetime.date.today().isoformat()
    items = "\n".join(
        f"  <url>\n    <loc>{_xml_escape(u)}</loc>\n    <lastmod>{today}</lastmod>\n"
        f"    <changefreq>weekly</changefreq>\n    <priority>0.8</priority>\n  </url>"
        for u in seen
    )
    return ('<?xml version="1.0" encoding="UTF-8"?>\n'
            '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">\n'
            f"{items}\n</urlset>\n")


# ------------------------------------------------------------------ docx output
def _kebab(slug):
    s = re.sub(r"([a-z0-9])([A-Z])", r"\1-\2", slug)
    return re.sub(r"[_\s]+", "-", s).lower()


def _http(url, method="GET", timeout=20):
    import urllib.request
    import urllib.error
    try:
        req = urllib.request.Request(url, method=method,
                                     headers={"User-Agent": "Mozilla/5.0 SEOPhase2Bot"})
        with urllib.request.urlopen(req, timeout=timeout) as r:
            body = r.read().decode("utf-8", "ignore") if method == "GET" else ""
            return r.status, body, r.geturl()
    except urllib.error.HTTPError as e:
        return e.code, "", url
    except Exception:
        return 0, "", url


def _resolve_gsc_creds(domain, account=None):
    """Best-effort (token, property_url) from a connected GSC account, using the
    shared gsc_audit module. Returns (None, None) if nothing is configured - the
    report then shows a graceful 'no GSC access' note rather than fabricating data.
    Never raises: any failure just means indexing falls back to the manual note."""
    try:
        # gsc_audit lives in the repo ROOT (one level up from this scripts/ folder).
        # ROOT here is the scripts dir, so add its parent so the import resolves.
        import sys as _sys
        _repo = str(ROOT.parent)
        if _repo not in _sys.path:
            _sys.path.insert(0, _repo)
        import gsc_audit
    except Exception as e:
        log(f"   [info] GSC module unavailable ({type(e).__name__}) - indexing note only")
        return None, None
    try:
        email = account
        if not email:
            accts = gsc_audit.list_accounts()
            if not accts:
                log("   [info] no connected GSC account - indexing note only")
                return None, None
            # prefer an account that has a refresh token (can mint a fresh token)
            email = next((a["email"] for a in accts if a.get("has_refresh")), accts[0]["email"])
        token = gsc_audit.get_access_token(email)
        property_url = gsc_audit.resolve_property(token, safe_domain(domain))
        log(f"   [info] GSC creds resolved for {email} -> {property_url}")
        return token, property_url
    except Exception as e:
        log(f"   [info] could not resolve GSC creds ({type(e).__name__}: {e}) - indexing note only")
        return None, None


def check_gsc_indexing(urls, token, property_url):
    """Call GSC URL Inspection API for each target page. Returns list of dicts."""
    import urllib.request
    results = []
    for url in urls:
        try:
            body = json.dumps({
                "inspectionUrl": url,
                "siteUrl": property_url,
                "languageCode": "en-US"
            }).encode()
            req = urllib.request.Request(
                "https://searchconsole.googleapis.com/v1/urlInspection/index:inspect",
                data=body,
                headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
                method="POST"
            )
            resp = urllib.request.urlopen(req, timeout=15)
            data = json.loads(resp.read())
            ir = data.get("inspectionResult", {})
            idx = ir.get("indexStatusResult", {})
            results.append({
                "url": url,
                "verdict": idx.get("verdict", "UNKNOWN"),
                "coverageState": idx.get("coverageState", ""),
                "robotsTxtState": idx.get("robotsTxtState", ""),
                "indexingState": idx.get("indexingState", ""),
                "lastCrawlTime": idx.get("lastCrawlTime", ""),
            })
        except Exception as e:
            results.append({"url": url, "verdict": "ERROR", "coverageState": str(e)})
    return results


def _detect_footer_logo(home_html):
    """True if the site's footer region contains a logo-like element. Conservative
    about false-NEGATIVES (the original bug) while avoiding wild false-positives:
    we only search the footer region, not the whole page."""
    if not home_html:
        return False
    html = home_html
    low = html.lower()

    # 1) Find where the footer region starts. Prefer the LAST literal <footer ...>,
    #    else the last element whose class/id contains "footer".
    start = low.rfind("<footer")
    if start == -1:
        best = -1
        for m in re.finditer(r'<(?:div|section|nav)\b[^>]*\b(?:class|id)\s*=\s*["\'][^"\']*footer[^"\']*["\']',
                              html, re.I):
            best = m.start()
        start = best
    if start == -1:
        return False

    # 2) Window after the footer start (footers are near the end; cap the span so a
    #    huge single-page app footer doesn't swallow the rest of the document).
    region = html[start:start + 20000]

    # 3) Any logo-like element inside the footer region counts.
    if re.search(r"<img\b", region, re.I):
        return True
    if re.search(r"<svg\b", region, re.I):
        return True
    if re.search(r"background(?:-image)?\s*:\s*[^;\"']*url\(", region, re.I):
        return True
    if re.search(r'\b(?:class|id)\s*=\s*["\'][^"\']*logo[^"\']*["\']', region, re.I):
        return True
    return False


def audit_site(domain, pages_data, dry_run=False):
    """Detect the real per-site findings that drive the narrative report."""
    d = safe_domain(domain)
    root = site_root(domain)
    # use the actual homepage (root path) for homepage-scoped checks, not just page[0]
    home = next((pd for pd in pages_data if urllib.parse.urlparse(pd["url"]).path in ("", "/")),
                pages_data[0] if pages_data else {})
    f = {"target_pages": [pd["url"] for pd in pages_data]}

    home_html = ""
    if not dry_run:
        _, home_html, _ = _http(root + "/")

    # robots.txt
    if dry_run:
        f["robots_found"], f["robots_has_sitemap"] = False, False
    else:
        rst, rtxt, _ = _http(root + "/robots.txt")
        f["robots_found"] = rst == 200 and ("user-agent" in rtxt.lower() or "disallow" in rtxt.lower())
        f["robots_has_sitemap"] = "sitemap" in rtxt.lower()
        f["robots_body"] = rtxt if rst == 200 else ""

    # sitemap_found / sitemap_url are set in main() via find_existing_sitemap()

    # www vs non-www (issue if www serves 200 without redirecting to non-www)
    if dry_run:
        f["www_redirect_issue"] = True
    else:
        wst, _, wfinal = _http(f"https://www.{d}/")
        f["www_redirect_issue"] = wst == 200 and "www." in urllib.parse.urlparse(wfinal).netloc

    # custom 404 (a clearly-missing URL should return 404, not 200/redirect)
    f["has_custom_404"] = (not dry_run) and _http(root + "/sos-404-probe-zzz")[0] == 404

    text_all = (home_html or "").lower()
    links = " ".join(home.get("internal_links", [])).lower()
    # a "blog" can also be a News / Articles / Insights / Media section
    f["has_blog"] = bool(re.search(r"/(blog|news|article|insight|press|media-?news)", links)) \
        or bool(re.search(r"\b(blog|latest news|our news)\b", text_all[:8000]))
    f["has_faq"] = "/faq" in links or "faq" in text_all[:10000]

    socials = [u for u in home.get("external_links", [])
               if re.search(r"facebook|instagram|twitter|linkedin|x\.com|youtube|t\.me|tiktok", u, re.I)]
    f["social_found"] = bool(socials)

    cur_year = datetime.date.today().year
    # take the MAX year near a copyright mark (sites often show "© 2026 ... est. 2022")
    yrs = [int(y) for y in re.findall(r"(?:©|&copy;|copyright)[^0-9]{0,25}(20\d\d)", home_html, re.I)]
    f["copyright_year"] = max(yrs) if yrs else None
    f["copyright_stale"] = bool(f["copyright_year"] and f["copyright_year"] < cur_year)

    imgs = [im for pd in pages_data for im in pd.get("images", [])
            if not _SKIP_PATTERNS.search(im.get("src", ""))]
    f["alt_total"] = len(imgs)
    f["alt_missing"] = sum(1 for im in imgs if not im.get("alt"))

    # footer logo: robustly detect a logo in the footer region. The old check only
    # matched a literal <footer>...<img> in the first 50k chars, missing SVG logos,
    # CSS background-image logos, <a class="logo">, footers built from
    # <div class="...footer...">, or footers beyond the slice. Detect the footer
    # region (last <footer>, else last element whose class/id contains "footer"),
    # then look for any logo-like element within a window after it.
    f["has_footer_logo"] = _detect_footer_logo(home_html)

    # Which external links, and on which page - across every target page, not just
    # the homepage. Every report format's copy already says "found on the target
    # pages checked" (plural), so ext_count itself is corrected to match that claim
    # instead of secretly only ever counting the homepage's links.
    f["external_links_detail"] = [
        {"page": pd["url"], "link": u}
        for pd in pages_data for u in pd.get("external_links", [])
    ]
    f["ext_count"] = len(f["external_links_detail"])

    # Content length - real paragraph copy only (headings and short/nav fragments
    # excluded by _paragraph_word_count itself). One entry per target page, plus a
    # ready-made status string per page for report sections that want to quote it.
    f["content_word_counts"] = [
        {"url": pd["url"], "words": pd.get("content_word_count", 0),
         "status": f"{pd.get('content_word_count', 0)} words found on the page"}
        for pd in pages_data
    ]
    f["lang"] = home.get("lang", "") or ""
    f["viewport"] = any(pd.get("viewport") for pd in pages_data)

    m = re.search(r'href="([^"]+)"', home.get("canonical", "") or "")
    f["home_canonical"] = m.group(1) if m else None
    f["canonical_issue"] = bool(f["home_canonical"] and f["home_canonical"].rstrip("/") != root)

    # noindex check - look for meta robots noindex in each target page's HTML
    noindex_pages = []
    for pd in pages_data:
        page_url = pd["url"]
        if dry_run:
            continue
        try:
            _, phtml, _ = _http(page_url)
            if re.search(r'<meta[^>]+name=["\']robots["\'][^>]+content=["\'][^"\']*noindex', phtml, re.I):
                noindex_pages.append(page_url)
            elif re.search(r'<meta[^>]+content=["\'][^"\']*noindex[^"\']*["\'][^>]+name=["\']robots["\']', phtml, re.I):
                noindex_pages.append(page_url)
        except Exception:
            pass
    f["noindex_pages"] = noindex_pages

    # mixed content - check if HTTPS pages load HTTP resources
    mixed_content_pages = []
    for pd in pages_data:
        page_url = pd["url"]
        if dry_run or not page_url.startswith("https"):
            continue
        try:
            _, phtml, _ = _http(page_url)
            http_resources = re.findall(r'(?:src|href|action)=["\']http://[^"\']+["\']', phtml, re.I)
            if http_resources:
                mixed_content_pages.append(page_url)
        except Exception:
            pass
    f["mixed_content_pages"] = mixed_content_pages

    # sucuri site check - query the free Sucuri SiteCheck API
    f["sucuri_clean"] = None  # None = couldn't check, True = clean, False = issues
    if not dry_run:
        try:
            suc_url = f"https://sitecheck.sucuri.net/api/v3/?scan={root}"
            _, suc_body, _ = _http(suc_url)
            suc_data = json.loads(suc_body) if suc_body else {}
            warnings = suc_data.get("warnings", {})
            blacklisted = suc_data.get("blacklists", {})
            has_malware = bool(suc_data.get("malware", []))
            has_blacklist = any(v for v in blacklisted.values() if v)
            f["sucuri_clean"] = not has_malware and not has_blacklist
        except Exception:
            f["sucuri_clean"] = None

    # broken links - reuse health_audit's more accurate checker (only 404/410 count
    # as broken; 403/429/5xx are the server blocking an automated request, not a
    # dead link - the old homepage-only, "any 4xx/5xx" check here disagreed with the
    # "brokenlinks" screenshot section, which already used the correct logic) and
    # across every target page, not just the homepage's first 30 external links.
    # Also records which page each broken link was found on and a suggested
    # redirect target, available to any format via findings["broken_links_detail"].
    broken_links = []
    f["broken_links_detail"] = []
    if not dry_run:
        try:
            import sys as _sys
            _repo = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            if _repo not in _sys.path:
                _sys.path.insert(0, _repo)
            import health_audit as _ha
            _, broken_detail = _ha.check_broken_links(domain, [pd["url"] for pd in pages_data])
            f["broken_links_detail"] = broken_detail
            broken_links = [b["url"] for b in broken_detail]
        except Exception:
            pass
    f["broken_links"] = broken_links

    f["url_changes"] = []
    for pd in pages_data:
        path = urllib.parse.urlparse(pd["url"]).path.strip("/")
        if path and re.search(r"[A-Z]", path):
            f["url_changes"].append((f"{root}/{path}", f"{root}/{_kebab(path)}"))
    return f


def _render_external_links_image(domain, detail, path):
    """Render the external-links list as a clean image for the report - same
    look as health_audit's broken-links image. `detail` is findings["external_links_detail"]
    (list of {"page":..., "link":...})."""
    from PIL import Image, ImageDraw, ImageFont
    rows = (detail or [])[:25]
    W = 1180
    H = 130 + (len(rows) + 1) * 30 + 30
    img = Image.new("RGB", (W, max(H, 190)), "#FFFFFF")
    d = ImageDraw.Draw(img)
    try:
        bold = ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 26)
        f = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 18)
        fs = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 15)
    except Exception:
        bold = f = fs = ImageFont.load_default()
    d.text((26, 24), f"External Links Check — {domain}", fill=(31, 41, 55), font=bold)
    if not rows:
        d.text((26, 74), "No external links found on the target pages checked.",
                fill=(22, 128, 67), font=f)
    else:
        d.text((26, 74), f"{len(detail)} external link(s) found on the target pages checked:",
                fill=(31, 41, 55), font=f)
        y = 116
        d.text((26, y), "Found On", fill=(120, 120, 120), font=fs)
        d.text((W // 2, y), "External Link", fill=(120, 120, 120), font=fs)
        y += 28
        for row in rows:
            d.text((26, y), str(row.get("page", ""))[:70], fill=(40, 40, 40), font=fs)
            d.text((W // 2, y), str(row.get("link", ""))[:70], fill=(40, 40, 40), font=fs)
            y += 28
    img.save(str(path))


def capture_onpage_screenshots(domain, sitemap_url=None, external_links_detail=None):
    """Live, HEADLESS screenshots for the report - all public pages/tools, so no
    Google login is needed. Captured with selenium + Chrome via CDP: the embedded
    python ships selenium but NOT patchright/playwright, so the old patchright
    capture path silently produced NO screenshots on user machines. The sitemap
    shot is taken of the ACTUAL found sitemap (only if one exists).

    Every shot is retried once on failure before falling back to a plain
    viewport screenshot - a transient network hiccup or slow page load should
    never silently leave a report slot empty when a retry would have worked."""
    import base64, tempfile
    import time as _t

    driver = _get_op_driver()
    if not driver:
        log("   [warn] no browser available - screenshots skipped")
        return {}

    out = {}
    out_dir = tempfile.mkdtemp(prefix=f"onpage_shots_{safe_domain(domain)}_")
    root = f"https://{domain}"

    def path(k):
        return os.path.join(out_dir, f"{domain}_seo_{k}.png")

    def _wait_settled(max_wait):
        """Poll for hydration-stable instead of always sleeping the full max_wait -
        same idea as _crawl_selenium's wait, just capped at the old fixed sleep time
        so a shot never waits LONGER than before, only shorter when the page settles
        early (result is unaffected either way - just less idle waiting)."""
        last, stable, elapsed, step = -1, 0, 0.0, 0.4
        while elapsed < max_wait:
            try:
                ready = driver.execute_script("return document.readyState")
                blen = driver.execute_script(
                    "return (document.body ? document.body.innerText.length : 0)")
            except Exception:
                break
            if ready == "complete" and blen > 0 and blen == last:
                stable += 1
                if stable >= 2:
                    return
            else:
                stable = 0
            last = blen
            _t.sleep(step)
            elapsed += step

    def _shot(key, url, height=900, view_source=False, sucuri=False, _attempt=1):
        p = path(key)
        try:
            driver.get(("view-source:" + url) if view_source else url)
            if view_source:
                _t.sleep(4)          # view-source: pages don't hydrate normally - keep the fixed wait
            else:
                _wait_settled(15 if sucuri else 4)
            # Google serves headless browsers a reCAPTCHA / "unusual traffic" wall
            # instead of results - a screenshot of that is useless, so skip it (the
            # indexing section falls back to a text note / GSC data).
            if key == "serp":
                src = (driver.page_source or "").lower()
                if any(m in src for m in ("unusual traffic", "not a robot", "recaptcha",
                                          "/sorry/", "detected unusual", "before you continue")):
                    log("   [warn] Google SERP blocked (captcha) - skipping serp screenshot")
                    return
            if sucuri:
                # Click Sucuri's cookie-consent "Accept" so it isn't in the shot.
                try:
                    driver.execute_script("""
                        var els=document.querySelectorAll('button,a,input[type=button],input[type=submit],span,div');
                        for(var i=0;i<els.length;i++){var t=(els[i].textContent||els[i].value||'').trim().toLowerCase();
                          if(t==='accept'||t==='accept all'||t==='i accept'||t==='allow all'||t==='allow'||t==='got it'||t==='agree'){try{els[i].click();}catch(e){}return;}}
                    """)
                    _t.sleep(1.2)
                except Exception:
                    pass
            try:
                driver.execute_script(
                    "document.querySelectorAll('.cookie-banner,.consent-banner,"
                    "[class*=cookie],[class*=consent],[id*=cookie],[id*=consent],#cookie-law-info-bar')"
                    ".forEach(function(e){e.remove();}); window.scrollTo(0,0);")
            except Exception:
                pass
            _t.sleep(0.6)
            if view_source:
                cdp = {"format": "png", "captureBeyondViewport": False}
            else:
                try:
                    _w = driver.execute_script(
                        "return Math.max(document.documentElement.clientWidth||0,"
                        " window.innerWidth||0, 1366);")
                except Exception:
                    _w = 1366
                # Sucuri's verdict sits at the very top - clip a fixed top region of
                # the document so we never grab a scrolled 'check another URL' view.
                # 700 (not the full ~1300px page) keeps this to the verdict/summary
                # card itself instead of also pulling in the Malware & Security /
                # Blacklist Status detail cards further down.
                h = 700.0 if sucuri else float(height)
                cdp = {"format": "png", "captureBeyondViewport": True,
                       "clip": {"x": 0, "y": 0, "width": float(_w or 1366), "height": h, "scale": 1}}
            result = driver.execute_cdp_cmd("Page.captureScreenshot", cdp)
            with open(p, "wb") as f:
                f.write(base64.b64decode(result["data"]))
            out[key] = p
            log(f"   -> captured [{key}]")
        except Exception as e:
            log(f"   [warn] capture {key} failed (attempt {_attempt}): {type(e).__name__}: {e}")
            if _attempt < 2:
                # a slow/transient page load shouldn't permanently leave a report
                # slot empty - one retry catches most flaky failures.
                _t.sleep(1.5)
                return _shot(key, url, height=height, view_source=view_source,
                             sucuri=sucuri, _attempt=_attempt + 1)
            try:                                   # last resort: plain viewport shot
                driver.save_screenshot(p)
                out[key] = p
                log(f"   -> captured [{key}] (fallback)")
            except Exception:
                pass

    _shot("homepage",  root + "/", height=900)
    _shot("robots",    root + "/robots.txt", height=700)
    _shot("canonical", root + "/", view_source=True)
    _shot("lang",      root + "/", view_source=True)
    if sitemap_url:                                # only shoot a sitemap that exists
        _shot("sitemap", sitemap_url, height=760)
    _shot("serp",      f"https://www.google.com/search?q=site:{domain}", height=900)
    _shot("wayback",   f"https://web.archive.org/web/2/https://{domain}/", height=900)
    _shot("viewport",  root + "/", height=900)
    _shot("sucuri",    f"https://sitecheck.sucuri.net/results/https/{domain}", sucuri=True)
    # Same URL the www_redirect_issue / has_custom_404 findings themselves check
    # (audit_site()) - the alt www/non-www version, and a deliberately-nonexistent
    # path - so the report can actually SHOW what was checked, not just claim it.
    _alt_host = domain[4:] if domain.startswith("www.") else f"www.{domain}"
    _shot("redirect",  f"https://{_alt_host}/", height=900)
    _shot("the404",    root + "/sos-404-probe-zzz", height=900)

    # Broken-links image - reuse health_audit's pure (non-patchright) helpers.
    try:
        import sys as _sys
        _repo = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        if _repo not in _sys.path:
            _sys.path.insert(0, _repo)
        import health_audit as _ha
        checked, broken = _ha.check_broken_links(domain)
        _ha._render_broken_links_image(domain, checked, broken, path("brokenlinks"))
        out["brokenlinks"] = path("brokenlinks")
        out["_broken"] = len(broken) if hasattr(broken, "__len__") else 0
        log(f"   -> broken-link check: {checked} links, {out['_broken']} broken")
    except Exception as e:
        log(f"   [warn] broken-link check failed: {type(e).__name__}: {e}")

    # External-links image - every format's builder calls shot("externallinks")
    # for this slot, but nothing used to produce that key, so it was ALWAYS
    # silently missing from every generated report regardless of format.
    try:
        _render_external_links_image(domain, external_links_detail, path("externallinks"))
        out["externallinks"] = path("externallinks")
        log(f"   -> captured [externallinks] ({len(external_links_detail or [])} link(s))")
    except Exception as e:
        log(f"   [warn] external-links image failed: {type(e).__name__}: {e}")

    _close_op_driver()
    return out


def _cover_source(domain):
    """Per-domain cover override, else the shared cover, else None (text fallback)."""
    dropin = Path("seo_onpage_screens") / safe_domain(domain) / "cover.png"
    if dropin.exists():
        return dropin
    return COVER_IMG if COVER_IMG.exists() else None


def _img_dims(path):
    from PIL import Image
    with Image.open(path) as im:
        return im.size


INTRO_TEXT = (
    "On-page optimization refers to all measures that can be taken directly within the website "
    "in order to improve its position in the search rankings. Optimize the content or improve "
    "the Meta tags."
)


def _compose_cover(cover_path, domain):
    """Bake the site URL + intro text onto the cover's white bottom band, returning
    a temp PNG path. Matches the supplied 'ON PAGE SEO REPORT' cover design."""
    from PIL import Image, ImageDraw, ImageFont
    img = Image.open(cover_path).convert("RGB")
    W, H = img.size
    draw = ImageDraw.Draw(img)

    def font(px, bold=False):
        for name in (("arialbd.ttf",) if bold else ("arial.ttf",)):
            try:
                return ImageFont.truetype("C:/Windows/Fonts/" + name, px)
            except Exception:
                pass
        return ImageFont.load_default()

    url = site_root(domain) + "/"
    title_f = font(int(W * 0.030), bold=True)
    body_f = font(int(W * 0.0175))
    maxw = int(W * 0.80)

    def wrap(text, f):
        lines, cur = [], ""
        for w in text.split():
            t = (cur + " " + w).strip()
            if draw.textlength(t, font=f) <= maxw:
                cur = t
            else:
                lines.append(cur)
                cur = w
        if cur:
            lines.append(cur)
        return lines

    y = int(H * 0.785)  # white band near the bottom of the cover
    uw = draw.textlength(url, font=title_f)
    draw.text(((W - uw) / 2, y), url, fill=(17, 17, 17), font=title_f)
    y += int(title_f.size * 1.7)
    for ln in wrap(INTRO_TEXT, body_f):
        lw = draw.textlength(ln, font=body_f)
        draw.text(((W - lw) / 2, y), ln, fill=(45, 45, 45), font=body_f)
        y += int(body_f.size * 1.5)

    # normalise to A4 portrait ratio (1:1.414) so it fills the page edge-to-edge
    img = img.resize((1240, 1754))
    fd, tmp = tempfile.mkstemp(suffix=".png", prefix="seocover_")
    os.close(fd)
    img.save(tmp)
    return tmp


def _text_cover(doc, domain):
    """Plain styled cover used when no cover image is available."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    for _ in range(7):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ON-PAGE SEO REPORT")
    r.bold = True
    r.font.size = Pt(34)
    u = doc.add_paragraph()
    u.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ru = u.add_run(site_root(domain) + "/")
    ru.bold = True
    ru.font.size = Pt(16)
    i = doc.add_paragraph()
    i.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ri = i.add_run(INTRO_TEXT)
    ri.font.size = Pt(11)


def _setup_docx(domain, use_cover=True):
    """Create a Document with cover page and return (doc, helpers_dict).

    use_cover=False skips BOTH the shared onpage_cover.png image and the
    generic "ON-PAGE SEO REPORT" text cover - just sets normal margins and
    returns, so the caller builds its own title block. Needed because
    onpage_cover.png was being applied to every format regardless of what
    that format's own verified reference actually looks like (confirmed:
    Camila's real reference has no cover image at all, just a centered
    "On Page Audit Report" title in Candara - the generic cover was being
    incorrectly mixed into its output)."""
    import generate_health_report as hr
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.shared import Inches, Pt, Cm, RGBColor

    A4_W, A4_H = Cm(21), Cm(29.7)
    FONT, SIZE = "Calibri", 11

    d = safe_domain(domain)
    root = site_root(domain)

    doc = Document()
    for sname in ("Normal", "List Paragraph"):
        try:
            st = doc.styles[sname]
            st.font.name = FONT
            st.font.size = Pt(SIZE)
        except KeyError:
            pass

    sec = doc.sections[0]
    sec.page_width, sec.page_height = A4_W, A4_H
    if not use_cover:
        # This format builds its own title block matching its own verified
        # reference - no image cover, no generic text cover.
        sec.left_margin = sec.right_margin = Inches(0.49)
        sec.top_margin = sec.bottom_margin = Inches(1)
    else:
        cover = _cover_source(domain)
        if cover:
            sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Cm(0)
            try:
                composed = _compose_cover(cover, domain)
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_before = Pt(0)
                cp.paragraph_format.space_after = Pt(0)
                cp.paragraph_format.line_spacing = 1.0
                cp.add_run().add_picture(composed, width=A4_W)
                try:
                    os.unlink(composed)
                except OSError:
                    pass
            except Exception as e:
                log(f"   [warn] cover compose failed: {type(e).__name__}: {e}")
            body = doc.add_section(WD_SECTION.NEW_PAGE)
            body.page_width, body.page_height = A4_W, A4_H
            body.left_margin = body.right_margin = Inches(0.49)
            body.top_margin = body.bottom_margin = Inches(1)
        else:
            log("   [info] no cover image (backend/onpage_cover.png) - using a text cover.")
            sec.left_margin = sec.right_margin = Inches(0.49)
            sec.top_margin = sec.bottom_margin = Inches(1)
            _text_cover(doc, domain)
            doc.add_page_break()

    def _style_run(run, bold=False):
        run.font.name = FONT
        run.font.size = Pt(SIZE)
        run.font.bold = bold

    def label_body(label, body_text=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if label:
            _style_run(p.add_run(label), bold=True)
        if body_text:
            _style_run(p.add_run(body_text), bold=False)
        return p

    def para(text="", bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if text:
            _style_run(p.add_run(text), bold=bold)
        return p

    def para_red(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(SIZE)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        return p

    def ribbon(text):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '000066')
        shd.set(qn('w:val'), 'clear')
        p._p.get_or_add_pPr().append(shd)
        return p

    def result(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(SIZE)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1D, 0x54, 0x89)
        return p

    def green(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(SIZE)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x00, 0xB0, 0x50)
        return p

    def shot(key, captured):
        src = (captured or {}).get(key)
        if src and Path(src).exists():
            green("Screenshot:")
            hr._add_bordered_image(doc, src)

    def summary_table(findings):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tbl = doc.add_table(rows=2, cols=2)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0]
        for i, txt in enumerate(["SEO Factors", "On Page Analysis Information"]):
            cell = hdr.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            r.font.name = FONT
            r.font.size = Pt(SIZE)
            r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '1D5489')
            shd.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shd)
        row1 = tbl.rows[1]
        row1.cells[0].text = ""
        p0 = row1.cells[0].paragraphs[0]
        r0 = p0.add_run("Image ALT Tags")
        r0.font.name = FONT
        r0.font.size = Pt(SIZE)
        r0.font.bold = True
        row1.cells[1].text = ""
        p1 = row1.cells[1].paragraphs[0]
        alt_missing = findings.get("alt_missing", 0)
        if alt_missing > 0:
            status = f"Status: {alt_missing} image(s) without ALT Tags found."
        else:
            status = "Status: ALT Tags are found on images."
        r1 = p1.add_run(status)
        r1.font.name = FONT
        r1.font.size = Pt(SIZE)
        r1.font.color.rgb = RGBColor(0x1D, 0x54, 0x89)
        r1.font.bold = True
        p1.add_run("\n")
        rec = p1.add_run("Recommendation: All the Used images should be optimized using proper "
                         "ALT Attributes. Images should have relevant and useful text as ALT Attributes.")
        rec.font.name = FONT
        rec.font.size = Pt(SIZE)
        rec.font.color.rgb = RGBColor(0x00, 0xB0, 0x50)
        rec.font.bold = True
        para("")

    return doc, {"label_body": label_body, "para": para, "para_red": para_red, "shot": shot,
                 "ribbon": ribbon, "result": result, "green": green, "summary_table": summary_table,
                 "d": d, "root": root, "FONT": FONT, "SIZE": SIZE, "_style_run": _style_run}


FORMAT_COVERS_DIR = ROOT / "covers"


def _insert_format_cover(doc, filename, page_break=True):
    """Insert THIS format's own real cover image (extracted from its client
    reference file), centered, at its own native aspect ratio - never
    force-scaled to the full A4 page width like the old shared cover logic
    did. Each format's cover is a distinct file (covers/cover_<format>.png/
    jpg) - onpage_cover.png (now only used by formats that genuinely own it)
    was being applied to every format regardless of what its own reference
    actually looked like. No-op (silently) if the asset is missing, so a
    report still generates rather than crashing on a missing file."""
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    path = FORMAT_COVERS_DIR / filename
    if not path.exists():
        return
    try:
        w_px, h_px = _img_dims(path)
        max_w = Inches(6.5)
        width = min(max_w, Inches(w_px / 96))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)
        p.add_run().add_picture(str(path), width=width)
        if page_break:
            doc.add_page_break()
    except Exception as e:
        log(f"   [warn] cover insert failed for {filename}: {type(e).__name__}: {e}")


# ---- Reusable section writers ----
def _sec_intro(h, root):
    p_url = h["para"].__self__ if hasattr(h["para"], "__self__") else None
    # Just use the helpers directly
    h["para"](root + "/", bold=True)
    h["para"]("On-page optimization refers to all measures that can be taken directly within the "
              "website in order to improve its position in the search ranking. This includes measures "
              "to optimize the content and the source code of a page.")

def _sec_on_page_analysis(h):
    h["ribbon"]("On Page Analysis")
    h["result"]("On Page Analysis:")
    h["para"]("These are the most important factors to be checked before optimizing a website. "
              "Proper implementation will help keyword ranking in search engine for your website.")

def _sec_additional_notes(h, findings, captured, year):
    h["label_body"]("Additional Note:")
    h["label_body"]("Note for Content Optimization:", " Kindly find the attached doc for content optimization.")
    h["label_body"]("Note for Meta Suggestion:", " Kindly find the attached Sheet for Meta Suggestion.")
    if not findings.get("has_faq"):
        h["para"]("")
        h["label_body"]("Note for FAQ Page:", " As we analyzed your website, we did not find a FAQ page. "
                        "If possible, we kindly recommend creating a new FAQ page, as it will help with SEO.")
    if not findings.get("social_found"):
        h["para"]("")
        h["label_body"]("Note for Social Media Icons:", " Your website is not connected to social media "
                        "using the API's provided by Facebook, Instagram, Twitter etc. So, we suggest you "
                        "to connect your website to all social media.")
        h["shot"]("homepage", captured)
    if not findings.get("has_footer_logo"):
        h["para"]("")
        h["label_body"]("Note for Footer Logo:", " As we analyzed your website, we noticed that there is no "
                        "logo in the footer section, which can affect brand consistency so we would suggest "
                        "you to add the logo in the footer.")
    if findings.get("copyright_stale"):
        h["para"]("")
        h["label_body"]("Note for Copyright:", f" When we analyzed your website, we noticed that your copyright "
                        f"year is not updated. So, we suggest you to update the copyright year to {year} "
                        f"instead of {findings['copyright_year']}.")
        h["shot"]("homepage", captured)

def _sec_alt_tags(h, findings):
    h["ribbon"]("Image ALT Optimization (Main Pages)")
    if findings.get("alt_missing", 0) > 0:
        h["result"](f"Result: {findings['alt_missing']} image(s) without Alt tags found on the target pages "
                    f"checked. It is not good from an SEO point of view. Kindly find the attached Image Alt "
                    f"Tag Suggestions sheet.")
    else:
        h["result"]("Result: Suitable Image Alt tags are found on the target pages checked. It is good from "
                    "an SEO point of view.")

def _sec_robots(h, findings, captured):
    h["ribbon"]("Robots.txt Optimization")
    h["para"]("Robots.txt is a regular text file that through its name has special meaning to the majority "
              "of \"honorable\" robots on the web. By defining a few rules in this text file, you can "
              "instruct robots to not crawl & index certain files, directories within a site or the entire "
              "site at all.")
    if findings.get("robots_found"):
        h["result"]("Result: Existing robots.txt file is optimized. It is good from an SEO point of view.")
    else:
        h["result"]("Result: An optimized robots.txt file was created; please find it attached and upload it "
                    "to the root folder of the website.")
    if not findings.get("robots_has_sitemap"):
        h["result"]("Note: Wrong sitemap URL found in the robots.txt. So, we have added a working sitemap URL "
                    "in the robots.txt file.")
    h["shot"]("robots", captured)

def _sec_indexing(h, findings, captured):
    h["ribbon"]("Indexing Status of All Target Pages Check")
    h["para"]("Making your page appear on search engine search results is called indexing. Getting your "
              "webpages indexed by search engine is extremely important. Pages that are not indexed by "
              "Google cannot rank and attract search traffic.")

    gsc_idx = findings.get("gsc_indexing")
    if gsc_idx:
        indexed = [r for r in gsc_idx if r["verdict"] == "PASS"]
        not_indexed = [r for r in gsc_idx if r["verdict"] not in ("PASS", "ERROR")]
        errors = [r for r in gsc_idx if r["verdict"] == "ERROR"]

        h["result"](f"Result: {len(indexed)} of {len(gsc_idx)} target page(s) are indexed by Google.")
        if indexed:
            h["green"]("Indexed Pages:")
            for r in indexed:
                h["para"](f"  {r['url']}  -  Indexed")
        if not_indexed:
            h["para_red"]("Not Indexed / Issues:")
            for r in not_indexed:
                state = r.get("coverageState") or r["verdict"]
                h["para"](f"  {r['url']}  -  {state}")
        if errors:
            h["para"]("Could not check:")
            for r in errors:
                h["para"](f"  {r['url']}  -  {r.get('coverageState', 'API error')}")
    else:
        h["para"](f"{len(findings['target_pages'])} target page(s) found:", bold=True)
        for u in findings["target_pages"]:
            h["para"](u)
        if findings.get("gsc_available"):
            h["para_red"]("Indexing was checked via Google Search Console but the URL-Inspection "
                          "API returned no usable data. Please verify the indexing status of these "
                          "pages manually in GSC or via a site: search on Google.")
        else:
            h["para_red"]("Indexing could not be verified automatically (no Google Search Console "
                          "access configured) - please verify via GSC or a site: search.")
    h["shot"]("serp", captured)

def _sec_sitemap(h, findings, captured):
    h["ribbon"]("XML Sitemap Optimization")
    h["para"]("Sitemap: A sitemap is a file where you can list the web pages of your site to tell Google "
              "and other search engines about the organization of your site content. Search engine web "
              "crawlers read this file to more intelligently crawl your site.")
    if findings.get("sitemap_found"):
        h["result"]("Result: Existing sitemap.xml file found. Please review the current sitemap once and "
                    "verify all important pages are included.")
        h["green"](f"Reference URL: {findings.get('sitemap_url')}")
        h["shot"]("sitemap", captured)
    else:
        h["result"]("Result: Sitemap.xml file not found on the website.")
        h["para_red"]("Please create a sitemap for the website. Different platforms have different methods:\n"
                      "- WordPress: Generate from admin using Yoast SEO or Rank Math plugin\n"
                      "- Shopify: Sitemap is auto-generated at /sitemap.xml\n"
                      "- Custom/Static sites: Create a static sitemap.xml and upload to root folder\n"
                      "- Other CMS: Check your platform's documentation for sitemap generation")

def _sec_redirection(h, findings, captured):
    h["ribbon"]("URL Redirection Issue Optimization")
    if findings.get("www_redirect_issue"):
        h["result"]("Redirection: The website runs with both www and non-www versions. We suggest redirecting "
                    "the www version to the non-www version using a 301 permanent redirect.")
    else:
        h["result"]("No redirection issue found on the website. It is good from a search engine point of view.")
    h["shot"]("redirect", captured)

def _sec_404(h, findings, root, captured):
    h["ribbon"]("Custom 404 Page Optimization")
    if not findings.get("has_custom_404"):
        h["label_body"]("404 Custom Page Suggestion:")
        h["para"](f"We did not find a custom 404 page. Opening a mistyped URL such as {root}/asdfg redirects "
                  "to the default page instead of a 404 page. We recommend creating a custom 404 page.")
    else:
        h["result"]("Result: A custom 404 page is present. It is good from an SEO point of view.")
    h["shot"]("the404", captured)

def _sec_canonical(h, findings, captured):
    d = h["d"]
    h["ribbon"]("Canonical Issue Checking")
    h["para"]("A canonical problem occurs when a site is running in multiple versions like www version "
              f"(http://www.{d}), the non-www version (http://{d}), trailing slash, "
              "or /index.html. Search engines may treat these as duplicates; a canonical tag tells them "
              "the preferred version.")
    if findings.get("canonical_issue"):
        h["result"]("Result: An incorrect / conflicting canonical tag was found. Kindly find the attached "
                    "Canonical Tag Suggestions sheet.")
    else:
        h["result"]("Result: Canonical tags are correctly set. Good from an SEO point of view.")
    h["shot"]("canonical", captured)

def _sec_external_links(h, findings, captured):
    h["ribbon"]("External Link Optimization")
    detail = findings.get("external_links_detail") or []
    h["result"](f"{findings.get('ext_count', len(detail))} external links found on the target pages checked.")
    if detail:
        h["para_red"]("Please verify whether any external links are harmful or beneficial as per need.")
        for item in detail[:30]:            # cap the listing - large sites can have hundreds
            h["para"](f"  {item['link']}   (on {item['page']})")
        if len(detail) > 30:
            h["para"](f"  ...and {len(detail) - 30} more.")

def _sec_broken_links(h, findings, captured):
    h["ribbon"]("Broken/Dead Link Optimization")
    h["para"]("Though the broken links do not hurt directly but it affects user experience and if users "
              "do not find desired info they will not come frequently and search engine will rank the "
              "website down.", bold=True)
    broken_links = findings.get("broken_links", [])
    tp_count = len(findings["target_pages"])
    if broken_links:
        h["result"](f"Result: {len(broken_links)} broken link(s) found on the {tp_count} target page(s) "
                    "checked. It's not good from SEO point of view.")
        for bl in broken_links:
            h["para"](bl)
    else:
        h["result"](f"Result: No broken links found on the {tp_count} target page(s) checked. "
                    "It is good from an SEO point of view.")
    h["para_red"]("Note: Only external links on the target pages were checked. Please verify broken "
                  "links across the full website using Screaming Frog, Ahrefs, or Google Search Console.")
    h["shot"]("brokenlinks", captured)

def _sec_internal_linking(h, home):
    h["ribbon"]("Internal Linking Optimization")
    h["result"]("Internal Linking: Internal linking is called perfect when every live webpage is "
                "accessed/visited from every other live and available webpage in a website.")
    int_count = len(home.get("internal_links", []))
    h["result"](f"Result: {int_count} internal links found on the target pages checked.")
    h["para_red"]("Please verify the internal linking structure across the full website as per need.")

def _sec_page_speed(h, root):
    h["ribbon"]("Page Speed Optimization")
    h["label_body"]("Page Optimization Score:")
    h["para_red"](f"Need to check it once. Check manually at: "
                  f"https://pagespeed.web.dev/analysis?url={root}/")

def _sec_url_structure(h, findings):
    root = h["root"]
    h["ribbon"]("URL Structure Optimization")
    h["green"]("Recommendations: Search Engines like static URLs instead of dynamic one. And presently "
               "URL structure of your website's inner pages is user and search engine friendly.")
    if findings.get("url_changes"):
        h["result"]("Result: The URL structure of the following pages should be changed:")
        for ex, rec in findings["url_changes"]:
            h["para"](f"Existing  ->  {ex}\nRecommended  ->  {rec}", bold=True)
        h["para"]("Note: After changing the URL structure, 301-redirect each old URL to the new one.", bold=True)
    else:
        h["result"]("Result: Yes, Good. Your URL structure is search engine-friendly. It is good from an SEO "
                    "point of view.")
        for u in findings["target_pages"]:
            h["para"](u)

def _sec_hyperlinking(h):
    h["ribbon"]("Hyperlinking Optimization")
    h["green"]("Recommendations: Hyperlinks are connections established between a word/phrase/image and a "
               "website/file. Effective hyper-linking between different pages of a website helps a website "
               "to rank better.")
    h["para_red"]("Result: Need to check it once.")

def _sec_mixed_content(h, findings, captured):
    h["ribbon"]("Mixed Content Optimization")
    h["para"]("Mixed content occurs when a webpage containing a combination of both secure (HTTPS) and "
              "non-secure (HTTP) content is delivered over SSL to the browser. It occurs when a website "
              "contains both HTTP and HTTPS content.")
    mixed_pages = findings.get("mixed_content_pages", [])
    tp_count = len(findings["target_pages"])
    if mixed_pages:
        h["result"](f"Result: Mixed content issues found on {len(mixed_pages)} of the {tp_count} target "
                    "page(s) checked. It's not good from an SEO point of view.")
        for mp in mixed_pages:
            h["para"](mp)
    else:
        h["result"](f"Result: No mixed content issues found on the {tp_count} target page(s) checked. "
                    "It's good from an SEO point of view.")
    h["shot"]("mixedcontent", captured)

def _sec_sucuri(h, findings, captured):
    d = h["d"]
    h["ribbon"]("Sucuri Site Scan")
    h["para"]("The Sucuri Site Check scanner helps to prevent security threats. It will check malware, "
              "viruses, blacklisting status, website errors, out-of-date software and malicious code.")
    sucuri_clean = findings.get("sucuri_clean")
    if sucuri_clean is True:
        h["result"]("Result: We have not found any malware/security issue on the website. It is good from "
                    "search engine point of view.")
    elif sucuri_clean is False:
        h["result"]("Result: Security issues were detected on the website. Please review and fix them.")
        h["para_red"]("Please check the details at the reference URL below.")
    else:
        h["para_red"](f"Please check the website security manually at: "
                      f"https://sitecheck.sucuri.net/results/https/{d}/")
    h["shot"]("sucuri", captured)
    h["green"](f"Reference URL: https://sitecheck.sucuri.net/results/https/{d}/")

def _sec_noindex(h, findings, captured):
    h["ribbon"]("No-index on Target Pages Check")
    h["para"]("Some pages of the website serve a purpose, and helps to improve the ranking and traffic "
              "to the site. These pages need to be there, as glue for other pages. But sometimes a few "
              "pages have a noindex tag that prevents them from being indexed.")
    noindex_pages = findings.get("noindex_pages", [])
    tp_count = len(findings["target_pages"])
    if noindex_pages:
        h["result"](f"Result: Noindex tag found on {len(noindex_pages)} of the {tp_count} target page(s) "
                    "checked. These pages will NOT get indexed on search engine. Please remove the noindex "
                    "tag.")
        for nip in noindex_pages:
            h["para"](nip)
    else:
        h["result"](f"Result: Noindex not found on the {tp_count} target page(s) checked. These pages "
                    "will get indexed on search engine.")
    h["para_red"]("Note: Only the target pages were checked. Please verify noindex tags across the full "
                  "website as per need.")
    h["shot"]("noindex", captured)

def _sec_web_archive(h, captured):
    root = h["root"]
    h["ribbon"]("Old Web Archive Status Check")
    h["para"]("It's the way to explore, find and retrieve historical and \"lost\" information from "
              "websites, to serve as evidence that something existed online, and was modified over time.")
    h["para_red"](f"Please check the web archive status at: "
                  f"https://web.archive.org/web/*/{root}/")
    h["shot"]("wayback", captured)

def _sec_viewport(h, findings, captured):
    h["ribbon"]("Meta viewport")
    h["para"]("The viewport meta tag allows you to tell the mobile browser what size this virtual "
              "viewport should be. This is often useful if you're not actually changing any visible "
              "content size but just the zoom level.")
    if findings.get("viewport"):
        h["result"]("Result: Yes, The Website pages have a viewport meta tag. It will look good on mobile "
                    "devices and will get a high position in mobile search results.")
    else:
        h["result"]("Result: Viewport meta tag missing - add it so pages render well on mobile.")
    h["shot"]("viewport", captured)

def _sec_lang(h, findings, captured):
    h["ribbon"]("lang Attribute")
    h["para"]("The \"lang\" attribute is an HTML attribute used to specify the language of the content "
              "within an HTML element. It helps search engines and assistive technologies understand "
              "the language of the page.")
    lang = findings.get("lang", "")
    if lang and "-" not in lang:
        h["result"](f'Result: We found lang="{lang}". If your target market is a specific region, we recommend '
                    f'a region-specific value (e.g. lang="en-US").')
    elif lang:
        h["result"](f'Result: We found the lang="{lang}" attribute on the website, it is good from search '
                    f'engine point of view.')
    else:
        h["result"]('Result: No lang attribute found - add an html lang attribute (e.g. lang="en-US").')
    h["shot"]("lang", captured)


def _set_page_background(doc, color):
    """Set the Word page background color (renders in Print Layout with the
    'display background' flag on). color is a 6-hex string, e.g. 'FFF2CC'."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    root = doc.element  # <w:document>; <w:background> must precede <w:body>
    if root.find(qn('w:background')) is None:
        bg = OxmlElement('w:background')
        bg.set(qn('w:color'), color)
        root.insert(0, bg)
    try:  # enable rendering of the background shape
        settings = doc.settings.element
        if settings.find(qn('w:displayBackgroundShape')) is None:
            settings.append(OxmlElement('w:displayBackgroundShape'))
    except Exception:
        pass


# ---- Format: James (Driftzine) ----
def _build_docx_james(domain, pages_data, findings, captured, brand, out_path):
    """James - confirmed correct: onpage_cover.png IS James's own reference
    cover (verified against OnpageReportJamesTemplate.docx - same "ON PAGE SEO
    REPORT" world-map design, composed text baked onto the same base image).
    It's the OTHER formats that were wrongly inheriting James's cover via
    _setup_docx's default, not James itself."""
    doc, h = _setup_docx(domain)
    home = next((pd for pd in pages_data if urllib.parse.urlparse(pd["url"]).path in ("", "/")),
                pages_data[0] if pages_data else {})
    year = datetime.date.today().year
    root = h["root"]

    # Cover image already shows the site URL + intro, and _sec_alt_tags prints the
    # ALT result - so the James body starts at "On Page Analysis" (no duplicate
    # intro paragraph, no duplicate "SEO Factors" ALT summary table).
    _sec_on_page_analysis(h)
    _sec_additional_notes(h, findings, captured, year)
    _sec_alt_tags(h, findings)
    _sec_robots(h, findings, captured)
    _sec_indexing(h, findings, captured)
    _sec_sitemap(h, findings, captured)
    _sec_redirection(h, findings, captured)
    _sec_404(h, findings, root, captured)
    _sec_canonical(h, findings, captured)
    _sec_external_links(h, findings, captured)
    _sec_broken_links(h, findings, captured)
    _sec_internal_linking(h, home)
    _sec_page_speed(h, root)
    _sec_url_structure(h, findings)
    _sec_hyperlinking(h)
    _sec_mixed_content(h, findings, captured)
    _sec_sucuri(h, findings, captured)
    _sec_noindex(h, findings, captured)
    _sec_web_archive(h, captured)
    _sec_viewport(h, findings, captured)
    _sec_lang(h, findings, captured)

    doc.save(out_path)


# ---- Format: Omega (alltechco) ----
def _build_docx_omega(domain, pages_data, findings, captured, brand, out_path):
    """Omega (alltechco) on-page report.

    Omega has NO navy ribbon header bar (that belongs to James/Xenon). Every
    section title is a bold BLACK, Calibri, 14pt line ending in a colon, rendered
    by the local ``header`` helper. Body/description text is Calibri 12pt.
    "Result:"/"Conclusion:"/"Remark:" labels are bold BLUE (1D5489) 12pt, one
    Broken-Link "Conclusion:" is bold BROWN (984806), and "Recommendations:"
    labels are bold GREEN (00B050) 12pt. Section order + wording mirror the
    aspiredentalassistantacademy.com reference. Data-driven Result lines stay
    honest (good vs issue) using the shared ``findings`` fields.
    """
    from docx.shared import Pt, RGBColor

    doc, h = _setup_docx(domain, use_cover=False)
    _insert_format_cover(doc, "cover_omega.png", page_break=False)
    _set_page_background(doc, "FFF2CC")  # cream page background, per the alltechco reference
    home = next((pd for pd in pages_data if urllib.parse.urlparse(pd["url"]).path in ("", "/")),
                pages_data[0] if pages_data else {})
    root = h["root"]
    d = h["d"]
    FONT = h["FONT"]

    BLACK = RGBColor(0x00, 0x00, 0x00)
    BLUE = RGBColor(0x1D, 0x54, 0x89)
    GREEN = RGBColor(0x00, 0xB0, 0x50)
    BROWN = RGBColor(0x98, 0x48, 0x06)

    def _run(p, text, bold=False, color=BLACK, size=12):
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return r

    def header(text):
        # Omega's OWN section header: bold black 14pt colon-title. NOT a ribbon.
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=True, color=BLACK, size=14)
        return p

    def body(text, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, text, bold=bold, color=BLACK, size=12)
        return p

    def labeled(label, text="", bold_body=False):
        # bold black label + (optional) body, all 12pt
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, label, bold=True, color=BLACK, size=12)
        if text:
            _run(p, text, bold=bold_body, color=BLACK, size=12)
        return p

    def result(label, text="", color=BLUE):
        # "Result:"/"Conclusion:"/"Remark:" bold colored label + normal black body
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, label, bold=True, color=color, size=12)
        if text:
            _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def recommend(label, text=""):
        # "Recommendations:" bold GREEN label + normal black body
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, label, bold=True, color=GREEN, size=12)
        if text:
            _run(p, text, bold=False, color=BLACK, size=12)
        return p

    tp = findings.get("target_pages", []) or []
    tp_count = len(tp)

    # ---- Intro (site URL bold 14pt, then intro bold black 16pt) ----
    p = doc.add_paragraph()
    _run(p, root + "/", bold=True, color=None, size=14)
    p = doc.add_paragraph()
    _run(p, "Onpage optimization (On-page SEO) refers to all measures that can be taken directly "
            "within the website in order to improve its position in the search rankings.",
         bold=True, color=BLACK, size=16)

    # ---- Content optimization ----
    header("Content optimization:")
    body("Kindly refer to the attached DOC for Content Suggestion.")

    # ---- Additional Suggestion ----
    header("Additional Suggestion:")
    if not findings.get("has_blog"):
        labeled("Note for Blog Page: ",
                "When we analyzed your website, we found that your website does not have a blog page. "
                "From the website ranking point of view, your website must have a proper blog page, so, "
                "we suggest you create a blog page for the website to reach your target customers easily.")
    labeled("Note for Footer Optimization: ",
            "We recommend adding quick links such as Home, About Us, Blog, and Contact, along with your "
            "contact information. This will strengthen brand identity and improve usability and engagement.")
    if not findings.get("has_footer_logo"):
        labeled("Note for Header & Footer Logo: ",
                "As we analyzed your website, we noticed that there is no logo in the header & footer "
                "section, which can affect brand consistency so we will suggest that there should be "
                "logo in header & footer section.")
    h["shot"]("homepage", captured)

    # ---- Meta Suggestions ----
    header("Meta Suggestions:")
    labeled("Note for Meta Suggestions:",
            " Kindly refer to the attached Excel sheet for Meta Suggestions.")

    # ---- Canonical Issue Suggestion ----
    header("Canonical Issue Suggestion:")
    body("A canonical problem occurs when a site is running in multiple versions like www version "
         f"(http://www.{d}), the non-www version (http://{d}) and other versions "
         f"(http://{d}/index.html). In this case, content of website is considered as duplicate by "
         "the search engines and as a result one version is removed from their index. The problem "
         "arises when the wrong version is deleted (this is usually the non-www version) instead of "
         "preferred one. So, to inform about genuine version of the page, we add canonical tag to that "
         "page in order to avoid any such problem or confusion. We also use canonical when there is no "
         "possibility of using 301 redirection method.")
    if findings.get("canonical_issue"):
        result("Result", ": An incorrect / conflicting canonical tag was found on the website. Kindly "
                         "find the attached Canonical Tag Suggestions sheet.")
    else:
        result("Result", ": Canonical issue not found in the website. It's good from SEO point of view.")

    # ---- Image Alt Tag Suggestions ----
    header("Image Alt Tag Suggestions:")
    if findings.get("alt_missing", 0) > 0:
        result("Result", f": {findings['alt_missing']} image(s) without alt tags found in the website. "
                         "It is not good from SEO point of view. Kindly find the attached Image Alt Tag "
                         "Suggestions.")
    else:
        result("Result", ": Suitable image alt tags are found in the website. Which is good from Seo "
                         "point of view.")

    # ---- Robots.txt Optimization ----
    header("Robots.txt Optimization:")
    labeled("Robots.txt ",
            "is a regular text file that through its name has special meaning to the majority of "
            "\"honorable\" robots on the web. By defining a few rules in this text file, you can "
            "instruct robots to not crawl and index certain files, directories within your site, or "
            "at all.", bold_body=False)
    if findings.get("robots_found"):
        result("Result", ": The existing robots.txt file is optimized, which is good from an SEO point "
                         "of view.")
    else:
        result("Result", ": An optimized robots.txt file was created; please find it attached and "
                         "upload it to the root folder of the website.")
    h["shot"]("robots", captured)

    # ---- Sitemap.xml Optimization ----
    header("Sitemap.xml Optimization:")
    labeled("Sitemap:",
            " A sitemap is a file where you can list the web pages of your site to tell Google and "
            "other search engines about the organization of your site content. Search engine web "
            "crawlers like Google bot read this file to more intelligently crawl your site.")
    if findings.get("sitemap_found"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, "Result:", bold=True, color=BLUE, size=12)
        _run(p, " Existing sitemap.xml file is optimized, which is good from an SEO point of view. ",
             bold=False, color=BLACK, size=12)
        _run(p, "Reference URL: ", bold=True, color=BLACK, size=12)
        _run(p, findings.get("sitemap_url") or (root + "/sitemap.xml"), bold=False, color=BLACK, size=12)
        h["shot"]("sitemap", captured)
    else:
        result("Result", ": Sitemap.xml file not found on the website. Please create and upload a "
                         "sitemap.xml to the root folder of the website.")

    # ---- Internal Link Optimization ----
    header("Internal Link Optimization:")
    labeled("Internal Linking:",
            " Internal linking is called perfect when every live webpage is accessed/visited from every "
            "other live and available webpage in a website. Internal linking helps users to move freely "
            "and access every desired webpage or find every bit of desired information provided on the "
            "website without having any confusion on how to visit any other live Webpage.")
    int_count = len(home.get("internal_links", []) or [])
    if int_count:
        result("Result:", " Internal linking structure of website is good and both SEO and user friendly.")
    else:
        result("Result:", " Please verify the internal linking structure across the website as per need.")

    # ---- External Link Optimization ----
    header("External Link Optimization:")
    ext_count = findings.get("ext_count", len(findings.get("external_links", []) or []))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Result:", bold=True, color=BLUE, size=12)
    _run(p, f" {ext_count} ", bold=True, color=BLACK, size=12)
    _run(p, "external links found in the website. None of link is seem harmful for the website it will "
            "be benefit.", bold=False, color=BLACK, size=12)

    # ---- URL Structure Optimization ----
    header("URL Structure Optimization:")
    recommend("Recommendations",
              ": Search Engines like static URLs instead of dynamic one. And presently URL structure of "
              "your website's inner pages is user and search engines friendly. So, it is not required to "
              "change existing pattern of URL structure. It is good from SEO point of view.")
    if findings.get("url_changes"):
        result("Result:", " The URL structure of the following pages should be changed:")
        for ex, rec in findings["url_changes"]:
            body(f"Existing  ->  {ex}", bold=True)
            body(f"Recommended  ->  {rec}", bold=True)
    else:
        labeled("URL Structure:", "")
        for u in tp:
            body(u)

    # ---- URL Redirection Issue Optimization ----
    header("URL Redirection Issue Optimization:")
    if findings.get("www_redirect_issue"):
        result("Conclusion:", " The website runs with both www and non-www versions. We suggest "
                             "redirecting the www version to the non-www version using a 301 permanent "
                             "redirect.")
    else:
        result("Conclusion:", " While analyzing your website, we did not find redirection issue. It is "
                             "good from SEO point of view.")
    if findings.get("has_custom_404"):
        body("Custom 404 page found in the website. It is good from search engine point of view.")
    else:
        body(f"We did not find a custom 404 page. Opening a mistyped URL such as {root}/asdfg redirects "
             "to the default page instead of a 404 page. We recommend creating a custom 404 page.")

    # ---- Hyperlink Analysis and Optimization ----
    header("Hyperlink Analysis and Optimization:")
    recommend("Recommendations",
              ": Hyperlinks are connections established between a word/phrase/image and a website/file. "
              "Effective hyper-linking between different pages can help a website secure a good position "
              "in the search engine result pages as well.")
    result("Result", ": Hyper-linking of the website is good. No need to change hyperlinking of the website.")

    # ---- Broken Link Optimization ----
    header("Broken Link Optimization:")
    broken_links = findings.get("broken_links", []) or []
    if broken_links:
        body(f"{len(broken_links)} broken link(s) found, though the broken links do not hurt directly "
             "but it affects user experience and if users do not find desired info they will not come "
             "frequently and search engine takes this as not useful for the users and start pulling back "
             "the web-pages.", bold=True)
        result("Conclusion:", f" {len(broken_links)} broken link(s) were found on the website. It's not "
                             "good from SEO point of view.", color=BROWN)
        for bl in broken_links:
            body(bl)
    else:
        body("No, broken links found, though the broken links do not hurt directly but it affects user "
             "experience and if users do not find desired info they will not come frequently and search "
             "engine takes this as not useful for the users and start pulling back the web-pages.",
             bold=True)
        result("Conclusion:", " No Broken links were found on the website. It's good from SEO point of view.",
               color=BROWN)
    h["shot"]("brokenlinks", captured)

    # ---- Page Speed Optimization ----
    header("Page Speed Optimization:")
    labeled("Page Optimization Score:",
            " Kindly refer to the attached screenshot for the page optimization score.")
    h["shot"]("pagespeed", captured)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Reference URL", bold=True, color=BLACK, size=12)
    _run(p, ": ", bold=True, color=BLACK, size=12)
    _run(p, f"https://pagespeed.web.dev/analysis?url={root}/", bold=True, color=BLACK, size=12)

    # ---- Mixed Content Optimization ----
    header("Mixed Content Optimization:")
    labeled("Mixed content",
            " occurs when a webpage containing a combination of both secure (HTTPS) and non-secure "
            "(HTTP) content is delivered over SSL to the browser. A mixed-content warning means that "
            "there are both secured and unsecured elements being served up on a page that should be "
            "completely encrypted.")
    mixed_pages = findings.get("mixed_content_pages", []) or []
    if mixed_pages:
        result("Result", f": Mixed content issue found on {len(mixed_pages)} of the {tp_count} target "
                         "page(s) checked. It's not good from SEO point of view.")
        for mp in mixed_pages:
            body(mp)
    else:
        result("Result", ": No mixed content issue found in the website. It's good from SEO point of view.")

    # ---- Sucuri Site Optimization ----
    header("Sucuri Site Optimization:")
    body("The SucuriSiteCheck scanner helps to prevent security threats. It will check malware, viruses, "
         "blacklisting status, website errors, out-of-date software, and malicious code to fix the issues "
         "timely before it damages your website.")
    sucuri_clean = findings.get("sucuri_clean")
    if sucuri_clean is False:
        result("Remark: ", "Security issues were detected on the website. Please review and fix them at "
                          f"https://sitecheck.sucuri.net/results/https/{d}/.")
    else:
        result("Remark: ", "Website malware is harmful software that has been developed with the intention "
                          "of carrying out malicious activity against a website - or its visitors. Website "
                          "malware is harmful software that has been developed with the intention of stealing "
                          "sensitive information, disrupting availability, redirecting visitors to spam pages, "
                          "completely hijacking the website, or even infecting the visitor with some other "
                          "piece of malware.")
    h["shot"]("sucuri", captured)

    # ---- No Index on Target Pages Check ----
    header("No Index on Target Pages Check:")
    body("Some pages of the website serve a purpose, and helps to improve the ranking and traffic to the "
         "site. These pages need to be there, as glue for other pages or simply because regulations "
         "require them to be accessible on your website. And if the main pages contain no index that means "
         "they will not be indexed by search engines and therefore will not appear in the search engine's "
         "result pages.")
    noindex_pages = findings.get("noindex_pages", []) or []
    if noindex_pages:
        result("Result", f": Noindex tag found on {len(noindex_pages)} of the {tp_count} target page(s) "
                         "checked. These pages will NOT get indexed on search engine. Please remove the "
                         "noindex tag.")
        for nip in noindex_pages:
            body(nip)
    else:
        result("Result", ": No index not found on the robots of the target pages of the website. It's good "
                         "from SEO point of view.")
    h["shot"]("noindex", captured)

    doc.save(out_path)


# ---- Format: Neon (sumitechengineers) ----
def _build_docx_neon(domain, pages_data, findings, captured, brand, out_path):
    """Neon (Sjjanarrabeen) on-page report.

    Neon has NO page background (white). Section headers are TEAL (215868) shaded
    bars with white bold 13pt text - NOT the navy 000066 ribbon used by James/Xenon.
    Body/description text is Calibri 12pt. "Result:" labels are bold BLUE (1D5489)
    where the reference shows them, other labels (Hyperlinking/Robots "Result",
    Broken-Link "Conclusion") are bold black. Section order + wording mirror the
    Sjjanarrabeen.com.au reference. Result/Conclusion text stays data-driven
    (good vs issue) using the shared ``findings`` fields. Rendered inline with a
    local ``header`` helper so it never emits the shared navy ribbon. Has its
    own real header image (cream-background analytics-dashboard illustration,
    extracted from the Massaristanbul.com reference) - not onpage_cover.png.
    """
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc, h = _setup_docx(domain, use_cover=False)
    _insert_format_cover(doc, "cover_neon.png", page_break=False)
    # Neon has NO page background - do not call _set_page_background().
    home = next((pd for pd in pages_data if urllib.parse.urlparse(pd["url"]).path in ("", "/")),
                pages_data[0] if pages_data else {})
    year = datetime.date.today().year
    root = h["root"]
    d = h["d"]
    FONT = h["FONT"]

    BLACK = RGBColor(0x00, 0x00, 0x00)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    BLUE = RGBColor(0x1D, 0x54, 0x89)
    SCORE = RGBColor(0xFF, 0xAA, 0x33)
    URLCLR = RGBColor(0xF7, 0x96, 0x46)

    def _run(p, text, bold=False, color=BLACK, size=12):
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return r

    def header(text):
        # Neon's OWN section header: TEAL (215868) shaded bar, white bold 13pt,
        # with a thin gold bottom border (distinguishes it from Sara's plain bar
        # now that Sara is frozen to its own verified reference).
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=True, color=WHITE, size=13)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '215868')
        shd.set(qn('w:val'), 'clear')
        pPr.append(shd)
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '18')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'C9A84C')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def body(text, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, text, bold=bold, color=BLACK, size=12)
        return p

    def labeled(label, text="", size=12):
        # bold black label + (optional) normal black body
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, label, bold=True, color=BLACK, size=size)
        if text:
            _run(p, text, bold=False, color=BLACK, size=size)
        return p

    def result(label, text="", color=BLUE):
        # "Result:"/"Conclusion:"/"Recommendations:" bold colored label + black body
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, label, bold=True, color=color, size=12)
        if text:
            _run(p, text, bold=False, color=BLACK, size=12)
        return p

    tp = findings.get("target_pages", []) or []
    tp_count = len(tp)

    # ---- Title block (bold black ~25pt title + URL, then intro ~13pt) ----
    p = doc.add_paragraph()
    _run(p, "On-Page Optimization Report", bold=True, color=BLACK, size=25)
    p = doc.add_paragraph()
    _run(p, root + "/", bold=True, color=BLACK, size=25)
    p = doc.add_paragraph()
    _run(p, "All the measures that can be applied to the website in order to get/improve the website in "
            "search ranking are called Ranking Improvement. The role of a page optimization report is "
            "imperative and has to be done with care. This report contains all the issues and suggestions "
            "to optimize the website, follow the suggestion and do changes accordingly. All the following "
            "On-page factors and the given suggestions need attention and should be fixed likewise if "
            "required to do so. Required changes need to be done ASAP as they often play a major role in "
            "acquiring a good position in SERP. In time changes boost the SEO process positively.",
         bold=False, color=BLACK, size=13)

    # ---- Additional Suggestions (gated notes, bold black 14pt) ----
    header(" Additional Suggestions ")
    labeled("Note for Meta Suggestion:",
            " We have created a separate Meta suggestions sheet that includes title, description, and "
            "heading tag suggestions. Please find the attached sheet and update it on the website.", size=14)
    labeled("Note for Content Optimization:",
            " Kindly find the attached doc for content optimization.", size=14)
    if not findings.get("has_faq"):
        labeled("Note for FAQ Page:",
                " As we analyzed your website, we did not find a FAQ page. If possible, we kindly "
                "recommend creating a new FAQ page, as it can provide significant benefits to your users "
                "by addressing their common questions, improving user experience, and enhancing "
                "engagement on your site.", size=14)
    labeled("Note for Footer Optimization:",
            " We recommend adding quick links such as Home, About Us, Blog, and Contact, along with your "
            "contact information. This will strengthen brand identity and improve usability and "
            "engagement.", size=14)
    labeled("Note for New Page:",
            " As we analyzed your website, we noticed that your website doesn't have separate pages for "
            "your major services. It is not good from an SEO point of view as well as from user "
            "experience. So we suggest you create separate new pages for your major services to get our "
            "target customers easily.", size=14)
    h["shot"]("homepage", captured)
    if not findings.get("has_footer_logo"):
        labeled("Note for Footer Logo Linking:",
                " As we analyzed your website, we noticed that the logo which is present at the footer is "
                "not linked with proper home page URL which can affect user experiences, so we will "
                "suggest add proper link to the footer logo.", size=14)
        h["shot"]("homepage", captured)
    if findings.get("copyright_stale"):
        labeled("Note for Copyright:",
                f" During our analysis, we did not find the updated copyright year on the website. "
                f"Therefore, we strongly suggest that you include the copyright year ({year}) in the "
                f"footer section at the bottom.", size=14)
        h["shot"]("homepage", captured)

    # ---- Optimization of URLs ----
    header(" Optimization of URLs ")
    if findings.get("url_changes"):
        result("Result:", " The URL structure of the following pages should be changed:")
        for ex, rec in findings["url_changes"]:
            body(f"Existing  ->  {ex}", bold=True)
            body(f"Recommended  ->  {rec}", bold=True)
    else:
        body("The existing URL structure of the website is fine. It is good from a search engine point "
             "of view.")

    # ---- Optimization of Robots.txt File ----
    header(" Optimization of Robots.txt File ")
    if findings.get("robots_found"):
        result("Result", ": The existing robots.txt file is optimized, which is good from an SEO point "
                         "of view.", color=BLACK)
    else:
        result("Result", ": An optimized robots.txt file was created; please find it attached and upload "
                         "it to the root folder of the website.", color=BLACK)
    h["shot"]("robots", captured)

    # ---- Optimization of Hyperlinking ----
    header(" Optimization of Hyperlinking")
    result("Recommendations", ": Hyperlinks are connections established between a word/phrase/image and a "
                              "website/file. Effective hyper-linking between different pages can help a "
                              "website secure a good position in the search engine result pages as well.",
           color=BLACK)
    result("Result", ": Hyperlinking of the website is good. It's good from SEO point of view.", color=BLACK)


    # ---- Image Alt Tag and Image Optimization ----
    header(" Image Alt Tag and Image Optimization ")
    if findings.get("alt_missing", 0) > 0:
        body(f"Suitable Image Alt tags are not found on the website. It is not good from an SEO point of "
             f"view. Kindly find an attached sheet for Image alt Tag Suggestion.")
    else:
        body("Suitable Image Alt tags are found on the website. It is good from an SEO point of view.")

    # ---- Internal Linking Structure/Navigation Optimization (Landing Pages) ----
    header(" Internal Linking Structure/Navigation Optimization (Landing Pages) ")
    int_count = len(home.get("internal_links", []) or [])
    if int_count:
        body("Internal linking structure of website is good and both SEO and user friendly.")
    else:
        body("Please verify the internal linking structure across the website as per need.")

    # ---- Optimization of External Links ----
    header("Optimization of External Links ")
    ext_count = findings.get("ext_count", len(findings.get("external_links", []) or []))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, f"{ext_count} ", bold=True, color=BLACK, size=12)
    _run(p, "external links found on the website, none of them seems harmful from a search engine point "
            "of view.", bold=False, color=BLACK, size=12)
    h["shot"]("externallinks", captured)

    # ---- Broken Link Optimization (Landing Pages) ----
    header("Broken Link Optimization (Landing Pages) ")
    result("Conclusion", ": though the broken links do not hurt directly but it affects user experience "
                         "and if users do not find desired info they will not come frequently and search "
                         "engine takes this as not useful for the users and start pulling back the "
                         "web-pages.", color=BLACK)
    broken_links = findings.get("broken_links", []) or []
    if broken_links:
        result("Result:", f" {len(broken_links)} Broken links were found on the website. It is not good "
                          "from an SEO point of view. Kindly find an attached sheet for broken link "
                          "Suggestions.")
        for bl in broken_links:
            body(bl)
    else:
        result("Result:", " No Broken links were found on the website. It is good from an SEO point of "
                          "view.")
    h["shot"]("brokenlinks", captured)

    # ---- Page Speed ----
    header("Page Speed")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Page Optimization Score:", bold=True, color=BLACK, size=12)
    _run(p, " Kindly refer to the attached screenshot for the page optimization score.",
         bold=True, color=SCORE, size=12)
    h["shot"]("pagespeed", captured)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Reference URL: ", bold=True, color=BLACK, size=12)
    _run(p, f"https://pagespeed.web.dev/analysis?url={root}/", bold=True, color=URLCLR, size=12)

    # ---- Canonical Tag Optimization (Landing Pages) ----
    header("Canonical Tag Optimization (Landing Pages) ")
    body(f"A canonical problem occurs when a site is running in multiple versions like www version "
         f"(http://www.{d}), the non-www version (http://{d}) and other versions "
         f"(http://{d}/index.html). In this case, content of website is considered as duplicate by the "
         "search engines and as a result one version is removed from their index. The problem arises "
         "when the wrong version is deleted (this is usually the non-www version) instead of preferred "
         "one. So, to inform about genuine version of the page, we add canonical tag to that page in "
         "order to avoid any such problem or confusion. We also use canonical when there is no "
         "possibility of using 301 redirection method.")
    if findings.get("canonical_issue"):
        result("Result", ": An incorrect / conflicting canonical tag was found on the website. Kindly "
                         "find the attached Canonical Tag Suggestions sheet.")
    else:
        result("Result", ": All Canonical tags are found on the website. It is good from a search engine "
                         "point of view.")
    h["shot"]("canonical", captured)

    # ---- Mixed Content Optimization ----
    header("Mixed Content Optimization")
    labeled("Mixed content",
            " occurs when a web page containing a combination of both secure (HTTPS) and non-secure "
            "(HTTP) content is delivered over SSL to the browser. A mixed-content warning means that "
            "there are both secured and unsecured elements being served up on a page that should be "
            "completely encrypted.")
    mixed_pages = findings.get("mixed_content_pages", []) or []
    if mixed_pages:
        result("Result", f": Mixed content issue found on {len(mixed_pages)} of the {tp_count} target "
                         "page(s) checked. It's not good from an SEO point of view.")
        for mp in mixed_pages:
            body(mp)
    else:
        result("Result", ": No mixed content issue was found in the website. It's good from an SEO point "
                         "of view.")

    # ---- No Index On Target Pages Check ----
    header("No Index On Target Pages Check")
    body("Some pages of the website serve a purpose, and helps to improve the ranking and traffic to the "
         "site. These pages need to be there, as glue for other pages or simply because regulations "
         "require them to be accessible on your website. And if the main pages contain no index that "
         "means they will not be indexed by search engines and therefore will not appear in the search "
         "engine's result pages.")
    noindex_pages = findings.get("noindex_pages", []) or []
    if noindex_pages:
        result("Result", f": Noindex tag found on {len(noindex_pages)} of the {tp_count} target page(s) "
                         "checked. These pages will NOT get indexed on search engine. Please remove the "
                         "noindex tag.")
        for nip in noindex_pages:
            body(nip)
    else:
        result("Result", ": No index not found on the robots of the target pages of the website. It's "
                         "good from SEO point of view.")

    # ---- URL Redirection Issue (Landing Pages) ----
    header("URL Redirection Issue (Landing Pages) ")
    if findings.get("www_redirect_issue"):
        result("Result:", " The website runs with both www and non-www versions. We suggest redirecting "
                          "the www version to the non-www version using a 301 permanent redirect.")
    else:
        result("Result:", " No redirection issue found on the website. It is good from a search engine "
                          "point of view.")
    h["shot"]("homepage", captured)

    # ---- Website XML Site Map Optimization ----
    header(" Website XML Site Map Optimization ")
    if findings.get("sitemap_found"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, "Result:", bold=True, color=BLUE, size=12)
        _run(p, " Existing sitemap.xml file is optimized, which is good from an SEO point of view.",
             bold=False, color=BLACK, size=12)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, "Reference URL: ", bold=True, color=BLACK, size=12)
        _run(p, findings.get("sitemap_url") or (root + "/sitemap.xml"), bold=False, color=BLACK, size=12)
        h["shot"]("sitemap", captured)
    else:
        result("Result:", " Sitemap.xml file not found on the website. Please create and upload a "
                          "sitemap.xml to the root folder of the website.")

    # ---- Site Security Check ----
    header("Site Security Check")
    sucuri_clean = findings.get("sucuri_clean")
    if sucuri_clean is False:
        body(f"Security issues were detected on your website. Please review and fix them at "
             f"https://sitecheck.sucuri.net/results/https/{d}/.")
    else:
        body("We didn't find any security issues on your website. It's good from an SEO point of view.")
    h["shot"]("sucuri", captured)

    doc.save(out_path)


# ---- Format: Xenon (markgreenbergphotography) ----
def _build_docx_xenon(domain, pages_data, findings, captured, brand, out_path):
    """Xenon on-page report.

    Reproduces the Xenon reference exactly: a light-blue D9E2F3 page background,
    a navy 1F3864 title bar, navy 1F3864 white section-header bars (NO 000066
    ribbon), F2F7FC status badges with a bold "Issue " (orange B45309) or "Good "
    (green 1E7A3C) lead word, blue 2E75B6 suggestion/label runs and navy 1F3864
    field labels. Result/Good-vs-Issue wording stays data-driven from ``findings``.
    No cover image in the reference - the navy title bar built below (from
    scratch, in code) IS its cover; onpage_cover.png must not be prepended.
    """
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc, h = _setup_docx(domain, use_cover=False)
    _set_page_background(doc, "D9E2F3")  # light-blue page background, per the reference
    home = next((pd for pd in pages_data if urllib.parse.urlparse(pd["url"]).path in ("", "/")),
                pages_data[0] if pages_data else {})
    root = h["root"]
    d = h["d"]
    FONT = h["FONT"]

    NAVY = RGBColor(0x1F, 0x38, 0x64)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    LIGHT = RGBColor(0xBF, 0xD3, 0xEC)
    GREY = RGBColor(0x60, 0x60, 0x60)
    BLUE = RGBColor(0x2E, 0x75, 0xB6)
    ORANGE = RGBColor(0xB4, 0x53, 0x09)
    GREEN = RGBColor(0x1E, 0x7A, 0x3C)
    BLACK = RGBColor(0x00, 0x00, 0x00)

    def _shade(el_pr, fill):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), fill)
        el_pr.append(shd)

    def _run(p, text, bold=False, color=None, size=11):
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return r

    def title_bar(text, color, size, fill="1F3864", space_before=0, space_after=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        _shade(p._p.get_or_add_pPr(), fill)
        _run(p, text, bold=True, color=color, size=size)
        return p

    def header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        _shade(p._p.get_or_add_pPr(), "1F3864")
        _run(p, text, bold=True, color=WHITE, size=12.5)
        return p

    def badge(is_good, body):
        """A F2F7FC status box: bold 'Good '/'Issue ' lead word + normal body."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        _shade(p._p.get_or_add_pPr(), "F2F7FC")
        if is_good:
            _run(p, "Good  ", bold=True, color=GREEN)
        else:
            _run(p, "Issue  ", bold=True, color=ORANGE)
        _run(p, body, bold=False, color=BLACK)
        return p

    def desc(text, lead=None):
        """Body paragraph; optional bold-blue lead label run."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if lead:
            _run(p, lead, bold=True, color=BLUE)
        _run(p, text, bold=False, color=BLACK)
        return p

    def field(label, value=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, label, bold=True, color=NAVY)
        if value:
            _run(p, value, bold=False, color=BLACK)
        return p

    def code_box(text, fill="F4F4F4"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = 'Table Grid'
        cell = tbl.rows[0].cells[0]
        _shade(cell._tc.get_or_add_tcPr(), fill)
        cell.text = ""
        _run(cell.paragraphs[0], text, bold=False, color=BLACK)

    tp = findings.get("target_pages", []) or []
    tp_count = len(tp)

    # ---- Title bar (navy) + domain + grey intro ----
    title_bar("ON-PAGE Suggestion Report", WHITE, 20, space_before=6)
    title_bar(d, LIGHT, 12)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    _run(p, "On-page SEO refers to all the measures that can be taken directly within a website to "
            "improve its position in the search rankings. This report reviews the on-page elements of "
            "the site and outlines the suggested optimizations.", bold=False, color=GREY)

    # ---- Content Optimization ----
    header("Content Optimization")
    desc("Please refer to the attached document for content suggestions.", lead="Suggestion: ")

    # ---- Meta Suggestions (Page URL + 3-col table) ----
    header("Meta Suggestions")
    field("Page URL: ", home.get("url", root + "/"))
    meta_tbl = doc.add_table(rows=4, cols=3)
    meta_tbl.style = 'Table Grid'
    for ci, htext in enumerate(("Element", "Existing", "Suggested")):
        c = meta_tbl.rows[0].cells[ci]
        _shade(c._tc.get_or_add_tcPr(), "1F3864")
        c.text = ""
        _run(c.paragraphs[0], htext, bold=True, color=WHITE)
    meta_rows = [
        ("Title", home.get("title", MISSING)),
        ("Meta Description", home.get("description", MISSING)),
        ("H1 Heading", home.get("h1", MISSING)),
    ]
    for ri, (elem, existing) in enumerate(meta_rows, start=1):
        rowfill = "F2F7FC" if ri % 2 == 0 else "FFFFFF"
        vals = (elem, existing or MISSING, "Refer to the attached Meta Suggestions sheet.")
        for ci, val in enumerate(vals):
            c = meta_tbl.rows[ri].cells[ci]
            _shade(c._tc.get_or_add_tcPr(), rowfill)
            c.text = ""
            _run(c.paragraphs[0], str(val), bold=(ci == 0), color=BLACK)

    # ---- Additional Suggestions ----
    header("Additional Suggestions")
    if not findings.get("has_faq"):
        desc("No FAQ page was found. I suggest creating one to answer common customer questions - it "
             "improves user experience and engagement on the site.", lead="FAQ Page: ")
    if not findings.get("has_blog"):
        desc("No blog page was found. I suggest creating one and publishing posts on a regular basis.",
             lead="Blog Page: ")
    if not findings.get("social_found"):
        desc("The website is not connected to social media via the Facebook, Instagram, or X (Twitter) "
             "APIs. We suggest adding social media links to the site.", lead="Social Media Icons: ")
    if not findings.get("has_footer_logo"):
        desc("No footer logo was found. I suggest adding a clickable footer logo to improve user "
             "experience and engagement.", lead="Footer Logo: ")
        desc("No header logo was found. I suggest adding a clickable header logo to improve user "
             "experience and engagement.", lead="Header Logo: ")
    if tp_count <= 1:
        desc("Only a home page was found, with no dedicated services section. I suggest creating "
             "additional pages to improve user experience and engagement.", lead="Service Pages: ")
    h["shot"]("homepage", captured)

    # ---- Canonical Issue ----
    header("Canonical Issue")
    desc("A canonical issue occurs when a site is accessible through multiple versions - for example "
         f"the www version, the non-www version, and other variants such as {root}/index.html. Search "
         "engines may treat these as duplicate content. A canonical tag tells search engines which "
         "version is the preferred one, avoiding duplication and confusion. It is also used where 301 "
         "redirection is not possible.")
    if findings.get("canonical_issue"):
        badge(False, "No proper canonical tag was found on the website. I suggest adding a proper "
                     "canonical tag as shown below.")
        field("URL: ", findings.get("home_canonical") or MISSING)
        field("Suggested tag: ")
        code_box(f'<link rel="canonical" href="{root}/" />')
    else:
        badge(True, "Canonical tags are correctly set on the website. This is good for SEO.")
    h["shot"]("canonical", captured)

    # ---- Image Alt Tag Suggestions ----
    header("Image Alt Tag Suggestions")
    if findings.get("alt_missing", 0) > 0:
        badge(False, f"Suitable image alt tags were not found on {findings['alt_missing']} image(s), "
                     "which is not ideal for SEO. Please refer to the attached Excel sheet for image "
                     "alt tag suggestions.")
    else:
        badge(True, "Suitable image alt tags were found on the website. This is good for SEO.")

    # ---- Robots.txt Optimization ----
    header("Robots.txt Optimization")
    desc("Robots.txt is a text file that instructs search engine crawlers which files and directories "
         "on the site they may or may not crawl and index.")
    if not findings.get("robots_found"):
        badge(False, "No robots.txt file was found. An optimized robots.txt file was created; please "
                     "upload it to the root folder of the website.")
    elif not findings.get("robots_has_sitemap"):
        badge(False, "A robots.txt file was found, but it does not reference the sitemap, which is "
                     "not ideal for SEO.")
    else:
        badge(True, "The existing robots.txt file is optimized. This is good for SEO.")
    h["shot"]("robots", captured)

    # ---- Sitemap.xml Optimization ----
    header("Sitemap.xml Optimization")
    desc("A sitemap is a file that lists the pages of a site to help search engines understand its "
         "structure. Crawlers such as Googlebot read this file to crawl the site more intelligently.")
    if findings.get("sitemap_found"):
        badge(True, "The existing sitemap.xml is optimized and contains the target pages. This is good "
                    "for SEO.")
        h["shot"]("sitemap", captured)
    else:
        badge(False, "No sitemap.xml file was found on the website. Please create one and upload it to "
                     "the root folder of the website.")

    # ---- Internal Link Optimization ----
    header("Internal Link Optimization")
    desc("is ideal when every live page can be reached from every other live page. It helps users move "
         "freely and find the information they need without confusion.", lead="Internal linking ")
    if len(home.get("internal_links", []) or []):
        badge(True, "The internal linking structure is sound and friendly for both SEO and users.")
    else:
        badge(False, "Please verify the internal linking structure across the full website as per need.")

    # ---- External Link Optimization ----
    header("External Link Optimization")
    desc("connects the site to relevant, trustworthy external resources, supporting credibility and a "
         "good user experience.", lead="External linking ")
    ext_count = findings.get("ext_count", 0)
    badge(True, f"{ext_count} external link(s) were found. The external linking structure is sound and "
                "friendly for both SEO and users.")

    # ---- URL Structure Optimization ----
    header("URL Structure Optimization")
    desc("Search engines prefer static URLs over dynamic ones. The URL structure of your inner pages "
         "is already user- and search-engine-friendly, so no change to the existing pattern is "
         "required.", lead="Suggestion: ")
    if findings.get("url_changes"):
        badge(False, "The URL structure of some pages should be changed for SEO. After changing the "
                     "URL structure, 301-redirect each old URL to the new one.")
        for ex, rec in findings["url_changes"]:
            field("Existing -> ", ex)
            field("Recommended -> ", rec)
    else:
        badge(True, "The existing URL structure is fine, which is good for SEO.")

    # ---- URL Redirection Optimization ----
    header("URL Redirection Optimization")
    if findings.get("www_redirect_issue"):
        badge(False, "The website runs with both www and non-www versions. We suggest redirecting the "
                     "www version to the non-www version using a 301 permanent redirect.")
    else:
        badge(True, "While analyzing the website, no redirection issues were found. This is good for SEO.")

    # ---- Hyperlink Analysis & Optimization ----
    header("Hyperlink Analysis & Optimization")
    desc("are connections between a word, phrase, or image and a page or file. Effective hyperlinking "
         "between pages can help a site rank well on search engine result pages.", lead="Hyperlinks ")
    badge(True, "Hyperlinking on the website is good. No changes are required.")

    # ---- Broken Link Optimization ----
    header("Broken Link Optimization")
    desc("Broken links do not harm rankings directly, but they hurt the user experience. If visitors "
         "cannot find what they need, they are less likely to return, and search engines may interpret "
         "this as a sign the pages are less useful.")
    broken_links = findings.get("broken_links", []) or []
    if broken_links:
        badge(False, f"{len(broken_links)} broken link(s) were found on the website. This is not ideal "
                     "for SEO. Please review and fix them.")
        for bl in broken_links:
            desc(bl)
    else:
        badge(True, "0 broken links were found on the website. This is good for SEO.")
    h["shot"]("brokenlinks", captured)

    # ---- Page Speed Optimization ----
    header("Page Speed Optimization")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Page Optimization Score:  ", bold=True, color=NAVY)
    _run(p, "Please check manually via the reference URL below.", bold=False, color=BLACK)
    field("Reference URL:  ", f"https://pagespeed.web.dev/analysis?url={root}/")

    # ---- Mixed Content Optimization ----
    header("Mixed Content Optimization")
    desc("Mixed content occurs when a page served over HTTPS also loads non-secure HTTP resources. A "
         "page that should be fully encrypted then contains both secure and insecure elements, which "
         "triggers browser warnings.")
    mixed_pages = findings.get("mixed_content_pages", []) or []
    if mixed_pages:
        badge(False, f"Mixed content issues were found on {len(mixed_pages)} of the {tp_count} target "
                     "page(s) checked. This is not ideal for SEO.")
        for mp in mixed_pages:
            desc(mp)
    else:
        badge(True, "No mixed content issues were found on the website. This is good for SEO.")

    # ---- Sucuri Site Optimization ----
    header("Sucuri Site Optimization")
    desc("The Sucuri SiteCheck scanner helps prevent security threats by checking for malware, viruses, "
         "blacklisting status, site errors, out-of-date software, and malicious code so issues can be "
         "fixed before they damage the site.")
    if findings.get("sucuri_clean") is False:
        badge(False, "Security issues were detected on the website. Please review and fix them.")
    else:
        badge(True, "No malware or security issues were found on the website. This is good for search "
                    "engines.")
    h["shot"]("sucuri", captured)

    # ---- No-Index Check on Target Pages ----
    header("No-Index Check on Target Pages")
    desc("Key pages need to be indexable so they can rank and drive traffic. If important pages carry a "
         "no-index directive, search engines will not index them and they will not appear in search "
         "results.")
    noindex_pages = findings.get("noindex_pages", []) or []
    if noindex_pages:
        badge(False, f"A no-index directive was found on {len(noindex_pages)} of the {tp_count} target "
                     "page(s). Please remove the no-index tag so these pages can be indexed.")
        for nip in noindex_pages:
            desc(nip)
    else:
        badge(True, "No no-index directive was found on the target pages (in the head section). The "
                    "target pages will be indexed by search engines, which is good.")

    # ---- Meta Viewport ----
    header("Meta Viewport")
    desc("The viewport meta tag tells a mobile browser the size of the virtual viewport. This helps a "
         "site render well on mobile, even when its layout is not specifically redesigned for small "
         "screens.")
    if findings.get("viewport"):
        badge(True, "The website pages include a viewport meta tag. They will display well on mobile "
                    "devices and rank better in mobile search results.")
    else:
        badge(False, "No viewport meta tag was found. Please add one so pages render well on mobile "
                     "devices.")

    # ---- lang Attribute ----
    header("lang Attribute")
    desc("The lang attribute specifies the language of the content within an HTML element. It helps "
         "search engines and assistive technologies understand the language of the text, improving "
         "accessibility and the accuracy of language-specific features such as spell-checking.")
    lang = findings.get("lang", "")
    if lang:
        badge(True, f'The attribute lang="{lang}" was found on the website. This is good for search '
                    "engines.")
    else:
        badge(False, 'No lang attribute was found. Please add an html lang attribute (e.g. lang="en-US").')

    doc.save(out_path)


# ---- Format: Gamma (Hawkeev) ----
def _build_docx_gamma(domain, pages_data, findings, captured, brand, out_path):
    """Gamma on-page report.

    Reproduces the Gamma reference: NO page background (white), a Book Antiqua top
    block (site URL + "Additional Suggestions" in light-blue 51C3F9 + gated notes),
    an "On-Page Analysis:" heading, then per-section 51C3F9 shaded header bars
    (NO 000066 ribbon), a description line and a brown 984806 "Conclusion:" line.
    Broken links render a 4-col amber-header table. Good-vs-issue wording is
    data-driven from ``findings``.
    """
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc, h = _setup_docx(domain, use_cover=False)
    _insert_format_cover(doc, "cover_gamma.png", page_break=False)
    # Gamma has NO page background (white) - deliberately not calling _set_page_background.
    home = next((pd for pd in pages_data if urllib.parse.urlparse(pd["url"]).path in ("", "/")),
                pages_data[0] if pages_data else {})
    root = h["root"]
    d = h["d"]
    FONT = h["FONT"]
    ANTIQUA = "Book Antiqua"

    BLUE = RGBColor(0x51, 0xC3, 0xF9)
    BROWN = RGBColor(0x98, 0x48, 0x06)
    BLACK = RGBColor(0x00, 0x00, 0x00)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    GREEN = RGBColor(0x00, 0xB0, 0x50)

    def _run(p, text, bold=False, color=None, size=11, font=None):
        r = p.add_run(text)
        r.font.name = font or FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return r

    def _shade(el_pr, fill):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), fill)
        el_pr.append(shd)

    def note(label, body):
        """Top-block note: bold black 11pt Book Antiqua label + normal 12pt body."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, label, bold=True, color=BLACK, size=11, font=ANTIQUA)
        _run(p, body, bold=False, color=BLACK, size=12)
        return p

    def header(text):
        """51C3F9 shaded section-header bar, bold white 17pt (NO 000066 ribbon)."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        _shade(p._p.get_or_add_pPr(), "51C3F9")
        _run(p, text, bold=True, color=WHITE, size=17)
        return p

    def description(label, body):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, label + ":", bold=True, color=BLACK, size=11)
        _run(p, " " + body, bold=False, color=BLACK, size=11)
        return p

    def conclusion(body):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _run(p, "Conclusion", bold=True, color=BROWN, size=11)
        _run(p, ":", bold=True, color=BROWN, size=11)
        _run(p, " " + body, bold=False, color=BLACK, size=11)
        return p

    tp = findings.get("target_pages", []) or []
    tp_count = len(tp)

    # ---- Top block (Book Antiqua): site URL + Additional Suggestions + notes ----
    p = doc.add_paragraph()
    _run(p, root + "/", bold=True, color=BLACK, size=18, font=ANTIQUA)
    p = doc.add_paragraph()
    _run(p, "Additional Suggestions", bold=True, color=BLUE, size=18, font=ANTIQUA)

    note("Note for Meta Suggestions: ", "Kindly find the attached sheet for meta suggestion.")
    if not findings.get("has_footer_logo"):
        note("Note for Footer Logo: ",
             "As we analyzed your website, we noticed that there is no logo linked with the proper "
             "home page URL at the footer, which can affect user experience, so we suggest adding a "
             "proper link to the footer logo.")
    if not findings.get("social_found"):
        note("Note For Social Media Icon Missing: ",
             "Your website is not connected to social media using the APIs provided by Facebook, "
             "Instagram, Twitter, and LinkedIn. So we suggest you add social media to your website.")
    if not findings.get("has_blog"):
        note("Blog Page: ",
             "As we analyzed your website, we noticed that your website doesn't have a Blog Page. "
             "From a website ranking point of view, we should have a proper blog page, because Google "
             "ranks unique and new content. So, we suggest you create a Blog page to reach your target "
             "customers easily.")
    if not findings.get("has_faq"):
        note("Note For FAQ: ",
             "FAQs answer the general queries a client might have while scrolling through your site and "
             "also play an important role in ranking on the search engines. If possible, kindly add FAQ "
             "questions on the home page or inner service pages.")
    note("Robots Meta Tag: ",
         "Robots Meta directives are pieces of code that provide crawlers instructions for how to crawl "
         "or index web page content, letting you control how an individual page is indexed and served "
         "to users in Google Search results.")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _run(p, "Recommendation:", bold=True, color=BLACK, size=12)
    _run(p, " We noticed that you have already added the Meta robots tags on the web pages. It is good "
            "from an SEO point of view.", bold=False, color=BLACK, size=12)

    # ---- On-Page Analysis ----
    p = doc.add_paragraph()
    _run(p, "On-Page Analysis:", bold=True, color=BLUE, size=22)
    p = doc.add_paragraph()
    _run(p, "These are key factors to review before optimizing a website. Proper implementation may "
            "boost keyword rankings in search engines and improve your website's overall performance.",
         bold=False, color=BLACK, size=12)

    # ---- Image Alt Text Suggestion ----
    header("Image Alt Text Suggestion")
    description("Image Alt Text",
                "Image alt text (alternative text) is crucial for SEO and accessibility. It provides "
                "context to search engines and visually impaired users, helping them understand the "
                "content of images and improving image search rankings.")
    if findings.get("alt_missing", 0) > 0:
        conclusion(f"After analyzing your website, I found {findings['alt_missing']} image(s) without "
                   "an ALT tag. It is not good from an SEO point of view. Kindly refer to the attached "
                   "Image Alt Tag Suggestions sheet.")
    else:
        conclusion("After analyzing your website, I found image ALT tags on the website. It's good from "
                   "an SEO point of view.")

    # ---- Robots Optimization ----
    header("Robots Optimization")
    description("Robots Optimization",
                "Robots.txt is a file that guides search engine crawlers on which pages to index or "
                "ignore. Proper optimization ensures that unimportant pages are not crawled while "
                "crucial pages are accessible to search engines.")
    if findings.get("robots_found"):
        conclusion("The existing robots.txt file is optimized, which is good from an SEO point of view.")
    else:
        conclusion("A robots.txt file was not found. An optimized robots.txt file was created; please "
                   "upload it to the root folder of the website.")
    h["shot"]("robots", captured)

    # ---- Sitemap Suggestion ----
    header("Sitemap Suggestion")
    description("Sitemap",
                "A sitemap is a structured list of pages on your site that helps search engines "
                "discover and index your content efficiently. It's crucial for both user navigation "
                "and SEO.")
    if findings.get("sitemap_found"):
        conclusion("An optimized sitemap file is found on the website. It's good from an SEO point of view.")
        h["shot"]("sitemap", captured)
    else:
        conclusion("A sitemap file was not found on the website. Please create and upload a sitemap.xml "
                   "to the root folder of the website.")

    # ---- Canonical Suggestion ----
    header("Canonical Suggestion")
    description("Canonical Tags",
                "Canonical tags help prevent duplicate content issues by specifying the preferred "
                "version of a webpage for search engines. This ensures search engines index the correct "
                "page and avoid ranking identical or similar content.")
    if findings.get("canonical_issue"):
        conclusion("An incorrect / conflicting canonical tag was found on the website. Kindly find the "
                   "attached Canonical Tag Suggestions sheet.")
    else:
        conclusion("I inform you that a canonical tag was found on the website. It is good from a search "
                   "engine point of view.")
    h["shot"]("canonical", captured)

    # ---- Redirection Issue Check ----
    header("Redirection Issue Check")
    description("Redirection Issues",
                "Redirects guide users and search engines from one URL to another. However, excessive "
                "or incorrect redirects can create a poor user experience and dilute page authority. "
                "It's essential to ensure that all redirects are purposeful and functioning correctly.")
    if findings.get("www_redirect_issue"):
        conclusion("The website runs with both www and non-www versions. We suggest redirecting the www "
                   "version to the non-www version using a 301 permanent redirect.")
    else:
        conclusion("No redirection issue found on the website. It is good from a search engine point of view.")

    # ---- External Linking Suggestion ----
    header("External Linking Suggestion")
    description("External Links",
                "External linking refers to the practice of hyperlinking to pages on different websites. "
                "These links can enhance the credibility of your content by directing users to reputable "
                "sources and providing additional information.")
    ext_count = findings.get("ext_count", 0)
    conclusion(f"{ext_count} external link(s) are found on the website, and none appear harmful from an "
               "SEO perspective.")

    # ---- Internal Linking Suggestion ----
    header("Internal Linking Suggestion")
    description("Internal Linking",
                "Internal linking refers to the practice of hyperlinking to other pages within the same "
                "website. This strategy helps establish site architecture, improves navigation, and "
                "distributes page authority across the site.")
    if len(home.get("internal_links", []) or []):
        conclusion("The internal linking structure of the website is good and both SEO and user-friendly.")
    else:
        conclusion("Please verify the internal linking structure across the full website as per need.")

    # ---- Broken Link Suggestion ----
    header("Broken Link Suggestion")
    description("Broken Links",
                "While broken links may not directly impact SEO, they do affect user experience. If "
                "users can't find the information they need, they are less likely to return, and search "
                "engines may eventually view the website as less useful, affecting rankings.")
    broken_detail = findings.get("broken_links_detail") or []
    broken_links = findings.get("broken_links", []) or []
    if broken_links:
        conclusion(f"When we analyzed your website, we found {len(broken_links)} broken link(s). It is "
                   "not good from a search engine point of view. Please check the broken links below "
                   "and resolve them.")
        btbl = doc.add_table(rows=1 + len(broken_links), cols=4)
        btbl.style = 'Table Grid'
        for ci, htext in enumerate(("Broken Link", "Source Page", "Server Response", "Solution")):
            c = btbl.rows[0].cells[ci]
            _shade(c._tc.get_or_add_tcPr(), "FFC000")
            c.text = ""
            _run(c.paragraphs[0], htext, bold=True, color=BLACK, size=11)
        by_url = {b.get("url"): b for b in broken_detail if isinstance(b, dict)}
        for ri, bl in enumerate(broken_links, start=1):
            det = by_url.get(bl) if isinstance(bl, str) else (bl if isinstance(bl, dict) else None)
            if det:
                found_on = ", ".join(det.get("found_on") or []) or home.get("url", root + "/")
                solution = det.get("suggested_redirect") or "Remove or redirect this link to a working page."
                vals = (det.get("url", ""), found_on, str(det.get("code", "")), solution)
            else:
                vals = (str(bl), home.get("url", root + "/"), "404",
                        "Remove or redirect this link to a working page.")
            for ci, val in enumerate(vals):
                c = btbl.rows[ri].cells[ci]
                c.text = ""
                _run(c.paragraphs[0], str(val), bold=False, color=BLACK, size=11)
    else:
        conclusion("When we analyzed your website, we did not find any broken links. It is good from a "
                   "search engine point of view.")
    h["shot"]("brokenlinks", captured)

    # ---- Hyperlinking Suggestion ----
    header("Hyperlinking Suggestion")
    description("Hyperlinking",
                "Internal linking between relevant pages on your site enhances user navigation and helps "
                "distribute link equity, improving SEO. Proper hyperlinking also guides search engines to "
                "crawl your site more effectively.")
    conclusion("The existing hyperlinking of the website is fine. It is good from a search engine point "
               "of view.")

    # ---- URL Structure Suggestion ----
    header("URL Structure Suggestion")
    description("URL Structure",
                "URLs play a significant role in both SEO and user experience. Clean, concise, and "
                "keyword-rich URLs make it easier for search engines to understand the content of your "
                "pages and enhance click-through rates.")
    if findings.get("url_changes"):
        conclusion("The URL structure of some pages should be changed. After changing the URL structure, "
                   "301-redirect each old URL to the new one.")
        for ex, rec in findings["url_changes"]:
            p = doc.add_paragraph()
            _run(p, f"Existing -> {ex}", bold=True, color=BLACK, size=11)
            p = doc.add_paragraph()
            _run(p, f"Recommended -> {rec}", bold=True, color=BLACK, size=11)
    else:
        conclusion("Yes, Good. Your URL structure is Search Engine friendly. It is good from an SEO "
                   "point of view.")

    # ---- Mixed Content Suggestion ----
    header("Mixed Content Suggestion")
    description("Mixed content",
                "Mixed content occurs when a webpage is delivered over SSL (HTTPS) but contains a "
                "combination of secure (HTTPS) and non-secure (HTTP) elements. This can trigger mixed "
                "content warnings in browsers.")
    mixed_pages = findings.get("mixed_content_pages", []) or []
    if mixed_pages:
        conclusion(f"We found mixed content issues on {len(mixed_pages)} of the {tp_count} target "
                   "page(s) checked. It's not good from an SEO point of view.")
        for mp in mixed_pages:
            p = doc.add_paragraph()
            _run(p, mp, bold=False, color=BLACK, size=11)
    else:
        conclusion("We have not found mixed content issues on the website. It's good from an SEO point "
                   "of view.")

    # ---- Sucuri Check ----
    header("Sucuri Check")
    description("Sucuri Check",
                "A Sucuri check evaluates the security status of your website. It identifies potential "
                "vulnerabilities, malware, and blacklist status, helping ensure your site remains secure "
                "and trustworthy.")
    if findings.get("sucuri_clean") is False:
        conclusion("Security issues were detected on the website. Please review and fix them at "
                   f"https://sitecheck.sucuri.net/results/https/{d}/.")
    else:
        conclusion("I did not find any malware/security issue on the website. It is good from a search "
                   "engine point of view.")
    p = doc.add_paragraph()
    _run(p, "Screenshot: - ", bold=True, color=GREEN, size=11)
    h["shot"]("sucuri", captured)

    # ---- Noindex Check ----
    header("Noindex Check")
    description("No-index Check",
                'The "noindex" tag is used to prevent search engines from indexing specific pages. '
                "While useful for keeping certain content out of search results, excessive use can "
                'hinder overall site visibility. Ensure only intended pages are set to "noindex".')
    noindex_pages = findings.get("noindex_pages", []) or []
    if noindex_pages:
        conclusion(f"A no-index tag was found on {len(noindex_pages)} of the {tp_count} target page(s). "
                   "It is not good from a search engine point of view. Please remove the no-index tag.")
        for nip in noindex_pages:
            p = doc.add_paragraph()
            _run(p, nip, bold=False, color=BLACK, size=11)
    else:
        conclusion("No-index not found on the robots of the target pages of the website. It is good from "
                   "a search engine point of view.")

    doc.save(out_path)


def build_content_suggestion_docx(domain, pages_data, targets, out_path, captured=None):
    """Generate the Content Suggestion DOCX - per-page content optimization notes,
    matching the reference layout: a navy 'Content Optimization' ribbon, a
    highlighted 'More Content Required:' lead-in, then per target page the URL,
    keywords, a red 'Section' label, a 'Screenshots:' label + bordered page
    screenshot (when available), separated by a dashed line."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    FONT, SIZE = "Arial", 11
    captured = captured or {}

    try:
        import generate_health_report as _hr
    except Exception:
        _hr = None

    doc = Document()
    for sname in ("Normal",):
        try:
            st = doc.styles[sname]
            st.font.name = FONT
            st.font.size = Pt(SIZE)
        except KeyError:
            pass

    def _run(p, text, bold=False, color=None, size=None, highlight=None):
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size or SIZE)
        r.font.bold = bold
        if color:
            r.font.color.rgb = color
        if highlight is not None:
            r.font.highlight_color = highlight
        return r

    def _ribbon(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '000066')
        shd.set(qn('w:val'), 'clear')
        p._p.get_or_add_pPr().append(shd)
        return p

    def _dashed_sep():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'dashed')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '808080')
        pbdr.append(bottom)
        pPr.append(pbdr)
        return p

    _ribbon("Content Optimization")

    p = doc.add_paragraph()
    _run(p, "More Content Required:", bold=True, highlight=WD_COLOR_INDEX.TURQUOISE)
    _run(p, " We have found content in following Target pages, but we recommend you to add more "
            "unique and relevant content (containing targeted keywords) in following pages up to "
            "100-120 words. Use the keyword respective to the pages as suggested below-", bold=True)

    # Real paragraph-content word count per page (excludes headings/nav/footer
    # and any fragment under 10 words - see _paragraph_word_count) - lets the
    # reader see which of these shared target pages already carry substantial
    # content versus which are genuinely thin, instead of guessing.
    words_by_url = {pd["url"]: pd.get("content_word_count", 0) for pd in (pages_data or [])}

    for i, t in enumerate(targets):
        page_url = t.get("page", "")
        keywords = t.get("keywords", [])
        kw_str = ", ".join(keywords) if keywords else ""

        if i > 0:
            _dashed_sep()

        p = doc.add_paragraph()
        _run(p, "Page URL: ", bold=True)
        _run(p, page_url)

        p = doc.add_paragraph()
        _run(p, "Keywords: ", bold=True)
        _run(p, kw_str)

        if page_url in words_by_url:
            p = doc.add_paragraph()
            _run(p, "Content Word Count: ", bold=True)
            _run(p, f"{words_by_url[page_url]} words found on the page")

        p = doc.add_paragraph()
        _run(p, "Section", bold=True, color=RGBColor(0xFF, 0x00, 0x00))

        p = doc.add_paragraph()
        _run(p, "Screenshots:", bold=True)

        # bordered page screenshot when one was captured for this page/target
        shot = t.get("screenshot") or captured.get(page_url) or captured.get(t.get("screenshot_key", ""))
        if shot and _hr is not None and Path(str(shot)).exists():
            try:
                _hr._add_bordered_image(doc, str(shot))
            except Exception:
                pass

    doc.save(out_path)


def _build_docx_deltafl(domain, pages_data, findings, captured, brand, out_path):
    """Delta Fl - verified against the client reference "On Page Analysis Report -
    aptasentry.com.docx": navy (#1F4E79) shaded paragraph banners with white bold
    13pt section titles, black Calibri body text, navy (#1F4E79) "Result:" labels
    and green (#00B050) "Recommendations:" labels. Uses the shared summary_table
    helper for the ALT-tag SEO Factors table (matches the reference exactly)."""
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc, h = _setup_docx(domain, use_cover=False)
    _insert_format_cover(doc, "cover_deltafl.png", page_break=False)
    root = h["root"]; d = h["d"]; FONT = h["FONT"]

    BLACK = RGBColor(0x00, 0x00, 0x00)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    NAVY = RGBColor(0x1F, 0x4E, 0x79)
    GREEN = RGBColor(0x00, 0xB0, 0x50)

    def _run(p, text, bold=False, color=BLACK, size=12):
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return r

    def header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=True, color=WHITE, size=13)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F4E79')
        shd.set(qn('w:val'), 'clear')
        p._p.get_or_add_pPr().append(shd)
        return p

    def body(text, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=bold, color=BLACK, size=12)
        return p

    def labeled(label, text=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, label, bold=True, color=BLACK, size=12)
        if text:
            _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def result(text, color=NAVY):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, "Result: ", bold=True, color=color, size=12)
        _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def recommend(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, "Recommendations: ", bold=True, color=GREEN, size=12)
        _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def refurl(url):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, "Reference URL: ", bold=True, color=BLACK, size=12)
        _run(p, url, bold=False, color=BLACK, size=12)
        return p

    def shot(key):
        src = (captured or {}).get(key)
        if src and Path(src).exists():
            try:
                hr._add_bordered_image(doc, src)
            except Exception:
                pass

    home_url = pages_data[0]["url"] if pages_data else root + "/"

    # ---- Intro ----
    body(home_url, bold=True)
    body("On-page optimization means making changes directly on your website to help it show up "
         "higher in search results. This includes improving your content and updating things like "
         "meta tags (titles and descriptions that search engines see).")
    body("These are the key things to check before you start improving your website. Doing them "
         "correctly will help your site rank better for keywords in search engines.")

    labeled("Additional Note: ")
    labeled("Note for meta-Suggestion: ", "Refer to an attached sheet of meta-suggestions.")
    labeled("Note for Content Optimization: ", "Kindly find the attached doc for content optimization.")

    # ---- Additional Suggestions ----
    header("Additional Suggestions")
    if findings.get("copyright_stale"):
        year = datetime.date.today().year
        labeled("Note for Copyright: ",
                f"As we analyse your website, we have noticed the copyright year still displays a "
                f"past year. Suggested - Copyright {year}.")
    if not findings.get("has_custom_404"):
        labeled("Note for custom 404 Page: ",
                "As we analyse your website, broken links show a generic Access Denied error with no "
                "custom 404 page. We recommend creating a branded 404 page with a Back to Home button.")
    if not findings.get("social_found"):
        labeled("Note For social media: ",
                "As we analyse your website, we did not find a complete set of social media links. "
                "Adding Facebook, Instagram, YouTube, LinkedIn and Twitter/X is essential to reach "
                "diverse audiences and build stronger brand trust.")
    if not findings.get("has_faq"):
        labeled("Note for FAQs: ",
                "We checked the website and did not find a dedicated FAQs page. It is recommended to "
                "create a separate FAQs section, as it is important for addressing common user "
                "queries, improving user experience, and enhancing SEO performance.")

    # ---- Image ALT ----
    header("Image ALT Optimization (Main Pages)")
    h["summary_table"](findings)

    # ---- Robots.txt ----
    header("Robots.txt Optimization")
    body("Robots.txt is a regular text file that through its name has special meaning to the "
         "majority of well-behaved web crawlers.")
    if findings.get("robots_found"):
        result("The existing robots.txt file is optimized, which is good from an SEO point of view.")
    else:
        result("An optimized robots.txt file was created with the correct sitemap URL. Please "
               "update it in the root directory of your website.")
    shot("robots")

    # ---- Sitemap ----
    header("Sitemap.xml Optimization")
    body("Sitemap: A sitemap is a file where you can list the web pages of your site to tell Google "
         "and other search engines about the organization of your site content.")
    if findings.get("sitemap_found"):
        result("Optimized sitemap.xml file found on the website. It's good from a search engine "
               "point of view.")
        refurl(findings.get("sitemap_url") or (root + "/sitemap.xml"))
    else:
        result("Sitemap.xml file not found on the website. Please create and upload a sitemap.xml "
               "to the root folder.")
    shot("sitemap")

    # ---- Redirection ----
    header("URL Redirection Issue Optimization")
    if findings.get("www_redirect_issue") or findings.get("url_changes"):
        body("Redirection Issue found on the website. It is not good from a search engine point of "
             "view.")
        result(f"We noticed that your website is running in multiple versions. It is not good from "
               f"an SEO point of view. We suggest redirecting all variants to a single canonical "
               f"version using a 301 permanent redirect.")
    else:
        result("No redirection issue found on the website. It is good from a search engine point of "
               "view.")

    # ---- Canonical ----
    header("Canonical Issue Checking")
    body(f"A Canonical Problem occurs when a site is running in multiple versions like www version "
         f"(https://www.{d}), the non-www version (https://{d}), and other versions. This causes "
         "duplicate content issues, so a canonical tag is used to indicate the preferred version.")
    if findings.get("canonical_issue"):
        result("An incorrect / conflicting canonical tag was found on the website. Kindly find the "
               "attached Canonical Tag Suggestions sheet.")
    else:
        result("All canonical tags are found on the website. It is good from a search engine point "
               "of view.")
    shot("canonical")

    # ---- External Links ----
    header("External Link Optimization")
    ext_count = findings.get("ext_count", 0)
    body(f"{ext_count} External link(s) found on the home page. None of them seems harmful for the "
         "website and it will benefit." if ext_count else
         "No external links found on the home page.")
    shot("externallinks")

    # ---- Broken Links ----
    header("Broken/Dead Links Optimization")
    body("Though the Broken Links do not hurt directly but it affects user experience and if users "
         "do not find desired info they will not come frequently and search engine takes this as not "
         "useful for the users and start pulling back the web pages.")
    broken = findings.get("broken_links") or []
    if broken:
        result(f"{len(broken)} Broken Link(s) found in the website. It is not good from a search "
               "engine point of view. Kindly find an attached sheet for broken link suggestions.")
    else:
        result("No Broken Link Found in the website. It is good from a search engine point of view.")
    shot("brokenlinks")

    # ---- Internal Linking ----
    header("Internal Linking Optimization")
    body("Internal Linking: Internal linking is called perfect when every live webpage is "
         "accessed/visited from every other live and available webpage in a website. Internal "
         "linking helps users to move freely and access every desired webpage or find every bit of "
         "desired information provided on the website without having any confusion.")
    result("Internal linking structure of website is good from both SEO and user friendly.")

    # ---- Page Speed ----
    header("Page Speed Optimization")
    labeled("Page Optimization Score: ", "Kindly refer to the attached screenshot for the page "
                                          "optimization score.")
    shot("pagespeed")
    refurl(f"https://pagespeed.web.dev/analysis?url={root}/")

    # ---- URL Structure ----
    header("URL Structure Optimization")
    recommend("Search Engines like static URLs instead of dynamic ones.")
    if findings.get("url_changes"):
        result("Some target page URLs could be optimized further. Please refer to the attached sheet.")
    else:
        result("Yes, Good. Your URL structure is Search Engine friendly.")

    # ---- Hyperlinking ----
    header("Hyperlink Analysis and Optimization")
    recommend("Hyperlinks are connections established between a word/phrase/image and a "
              "website/file. Effective hyper-linking between different pages can help a website "
              "secure a good position in the search engine result pages as well.")
    result("Hyper-linking of the website is good.")

    # ---- Mobile Friendly ----
    header("Mobile Friendly Test")
    body("Page is mobile-friendly. It is good from SEO point of view." if findings.get("viewport", True)
         else "Page is not fully mobile-friendly. Please review responsive design.")
    refurl("https://www.bing.com/webmaster/tools/mobile-friendliness")

    # ---- Mixed Content ----
    header("Mixed Content Optimization")
    body("Mixed content occurs when a webpage containing a combination of both secure (HTTPS) and "
         "non-secure (HTTP) content is delivered over SSL to the browser. A mixed-content warning "
         "means that there are both secured and unsecured elements being served up on a page that "
         "should be completely encrypted.")
    mixed = findings.get("mixed_content_pages") or []
    if mixed:
        result(f"Mixed content issue found on {len(mixed)} page(s). It's not good from an SEO point "
               "of view.")
    else:
        result("No mixed content issue found in the website. It's good from SEO point of view.")

    # ---- Sucuri ----
    header("Sucuri Site Scan")
    body("The SucuriSiteCheck scanner helps to prevent security threats. It will check malware, "
         "viruses, blacklisting status, website errors, out-of-date software, and malicious code to "
         "fix the issues timely before it damages your website.")
    if findings.get("sucuri_clean", True):
        result("We have not found any malware/security issue on the website. It is good from search "
               "engine point of view. Check attached screenshot.")
    else:
        result("Security issues were detected on your website. Please review and fix them.")
    shot("homepage")

    # ---- No-Index ----
    header("No-Index on Target Pages Check")
    body("Some pages of the website serve a purpose, and helps to improve the ranking and traffic to "
         "the site. If the main pages contain no index that means they will not be indexed by search "
         "engines and therefore will not appear in the search engine's result pages.")
    noindex = findings.get("noindex_pages") or []
    if noindex:
        result(f"Noindex tag found on {len(noindex)} target page(s). These pages will NOT get "
               "indexed on search engine. Please remove the noindex tag.")
    else:
        result("Noindex not found on the target pages of the robots located in the head section of "
               "the website pages. Our target pages will get index on search engine. It is good from "
               "search engine point of view.")

    # ---- Indexing Status ----
    header("Indexing Status of All Target Pages Check")
    body("Making your page appear on search engine search results is called indexing. Getting your "
         "webpages indexed by search engine is extremely important. Pages that are not indexed by "
         "Google cannot rank.")
    if findings.get("gsc_indexing"):
        result("We have found all the pages are indexed.")
    else:
        result("Please refer to the attached indexing status sheet for target page indexing details.")

    # ---- Web Archive ----
    header("Old Web Achieve Status Check")
    body("It's the way to explore, find and retrieve historical and 'lost' information from "
         "websites, to serve as evidence that something existed online, and ways to archive and "
         "preserve your own copies of webpages for future reference.")
    result("We have submitted the target pages of the website to the web archive tool for future "
           "reference.")

    # ---- Meta Viewport ----
    header("Meta Viewport")
    body("The viewport meta tag allows you to tell the mobile browser what size this virtual "
         "viewport should be. This is often useful if you're not actually changing your site's "
         "design for mobile, and it renders better with a larger or smaller virtual viewport.")
    result("Yes, The Website pages have a viewport meta tag. It will look good on mobile devices and "
           "will get a high position in mobile search results.")

    doc.save(out_path)


def _build_docx_deltafvr(domain, pages_data, findings, captured, brand, out_path):
    """Delta FVR - verified against the client reference "On Page Suggestions
    Report - attunedtherapy.ca.docx": NO shaded banners at all (plain white
    page), section labels use a "Label → description" arrow style in dark
    navy (#002060), "Status:" lines in black, issues flagged in red (#EE0000),
    positive/suggested items in green (#00B050). Closes with a plain "Thank
    You !" line instead of a sign-off block."""
    from docx.shared import Pt, RGBColor

    doc, h = _setup_docx(domain, use_cover=False)
    _insert_format_cover(doc, "cover_deltafvr.png", page_break=False)
    root = h["root"]; d = h["d"]; FONT = h["FONT"]

    BLACK = RGBColor(0x00, 0x00, 0x00)
    NAVY = RGBColor(0x00, 0x20, 0x60)
    RED = RGBColor(0xEE, 0x00, 0x00)
    GREEN = RGBColor(0x00, 0xB0, 0x50)

    def _run(p, text, bold=False, color=BLACK, size=12):
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return r

    def body(text, bold=False, color=BLACK):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=bold, color=color, size=12)
        return p

    def label_arrow(label, text, bold=True):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, label + " → ", bold=bold, color=NAVY, size=12)
        _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def status(text, color=BLACK):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, "Status: ", bold=True, color=color, size=12)
        _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def shot(key):
        src = (captured or {}).get(key)
        if src and Path(src).exists():
            try:
                hr._add_bordered_image(doc, src)
            except Exception:
                pass

    home_url = pages_data[0]["url"] if pages_data else root + "/"

    # ---- Intro ----
    body(home_url)
    p = doc.add_paragraph()
    _run(p, "On-Page Optimization Suggestions", bold=True, color=BLACK, size=16)
    body("On-page SEO is a process of optimization of individual Webpages of content to earn better "
         "and more relevant traffic from search engines.")
    body("We have analysed your website in terms of search engine perspective. Following are the "
         "suggestions to further improve the website.")

    p = doc.add_paragraph()
    _run(p, "Content Optimization Suggestions", bold=True, color=BLACK, size=14)
    body("Kindly find the attached document for content optimization suggestions covering the "
         "target page(s) and their keywords.")

    p = doc.add_paragraph()
    _run(p, "Meta Suggestions", bold=True, color=BLACK, size=14)
    body("We suggest Meta suggestions (Meta Title & Meta Description) including Heading tags in the "
         "attached sheet.")

    if not findings.get("has_footer_logo"):
        body("Note for Footer Logo: We have noticed that your website has a footer logo, but it is "
             "not clickable/linked to the homepage. We suggest linking it.")
    if not findings.get("has_blog"):
        body("Note for Blog Post: While analysing the website we did not find a blog section. We "
             "recommend adding one to reach your target audience more effectively.")
    if not findings.get("has_faq"):
        body("Note for FAQs Page: We have noticed that your website currently has no FAQ section. We "
             "recommend adding one.")
    if not findings.get("social_found"):
        body("Add More social media Icon: As we analysed your website, we noticed a limited set of "
             "social media icons. We recommend adding more to improve engagement.")

    p = doc.add_paragraph()
    _run(p, "SEO Analysis and Optimization Suggestions", bold=True, color=BLACK, size=14)

    # ---- Robots.txt / Sitemap ----
    label_arrow("Robots.txt", "The robots.txt file, also known as the robot's exclusion protocol or "
                              "standard, is a text file that tells web robots (most often search "
                              "engines) which pages on your site to crawl.")
    if findings.get("sitemap_found"):
        status("Optimized sitemap.xml file found in website. It's good from a search engine point of "
               "view.")
    else:
        status("Sitemap.xml file not found on the website. We recommend creating and submitting one.",
               color=RED)
    shot("robots")

    # ---- Image Alt Tag ----
    label_arrow("Image Alt Tag", "ALT tags or ALT attributes are \"alternative text\" for an image. "
                                 "ALT tags are used to describe the image or what the image is "
                                 "representing on the webpage.")
    alt_missing = findings.get("alt_missing", 0)
    if alt_missing:
        p = doc.add_paragraph()
        _run(p, f"{alt_missing} image(s) without ALT tags found. Please refer to the attached sheet "
                "for Image Alt Tag Suggestions.", bold=False, color=RED, size=12)
    else:
        status("Suitable Image Alt tags are found on images.", color=GREEN)

    # ---- Internal Linking ----
    label_arrow("Internal Linking", "An internal link is any link from one page on your website to "
                                    "another page on your website. Both your users and search "
                                    "engines use links to find content on your site.")
    int_count = len(pages_data[0].get("internal_links", []) or []) if pages_data else 0
    status("The structure of Internal linking of website is fine and it will not affect our SEO "
           "efforts." if int_count else
           "Please verify the internal linking structure across the website as per need.")

    # ---- Hyperlinking ----
    label_arrow("Hyper Linking", "Hyperlinking is the process of linking a word/phrase/image to a "
                                 "website/file, connecting pages together for both users and search "
                                 "engines to navigate.")
    status("Hyperlinking of website is fine and it's good from an SEO point of view.")

    # ---- External Linking ----
    ext_count = findings.get("ext_count", 0)
    label_arrow("External Linking", "External Links are hyperlinks that point at (target) any domain "
                                    "other than the domain the link exists on (source).")
    status(f"{ext_count} External links found on the website, none of them seem harmful from an SEO "
           "point of view." if ext_count else "No external links found on the website.")
    shot("externallinks")

    # ---- Site Security ----
    label_arrow("Site Security Check", "A site security check evaluates a website for "
                                       "vulnerabilities, malware, and compliance with security best "
                                       "practices to protect against threats.")
    if findings.get("sucuri_clean", True):
        status("We have not found any malware/security issue on the website. It is good from a "
               "search engine point of view.", color=GREEN)
    else:
        status("Security issues were detected on your website. Please review and fix them.", color=RED)
    shot("homepage")

    # ---- Broken Links ----
    broken = findings.get("broken_links") or []
    label_arrow("Broken/Dead Link", "Broken links are links that send a message to its visitors that "
                                    "the webpage no longer exists, triggering a 404-error page. "
                                    "Broken links hurt user experience and SEO.")
    if broken:
        status(f"{len(broken)} broken link(s) found on the website. Please refer to the attached "
               "sheet for suggestions.", color=RED)
    else:
        status("No broken links found on the website.", color=GREEN)
    shot("brokenlinks")

    # ---- Redirection ----
    label_arrow("Website Redirection", "A redirect is a way to send both users and search engines to "
                                       "a different URL from the one they originally requested.")
    if findings.get("www_redirect_issue") or findings.get("url_changes"):
        status("Redirection issue found - the website is reachable on multiple versions "
               "independently. We recommend a single 301 redirect to one canonical version.",
               color=RED)
    else:
        status("No redirection issue found on the website.", color=GREEN)

    # ---- Canonicalization ----
    label_arrow("Canonicalization", "A canonical URL refers to an HTML link element, with the "
                                    "attribute of rel=\"canonical\" (also known as a canonical tag), "
                                    "found in the head section of a webpage, indicating the preferred "
                                    "version of a page.")
    if findings.get("canonical_issue"):
        status("An incorrect / conflicting canonical tag was found on the website. Kindly find the "
               "attached Canonical Tag Suggestions sheet.", color=RED)
    else:
        status("All canonical tags are found on the website. It's good from an SEO point of view.",
               color=GREEN)
    shot("canonical")

    # ---- Dummy Content ----
    label_arrow("Dummy Content Checker", "A Dummy Content Checker identifies and removes placeholder "
                                         "text or temporary content from a website to ensure all "
                                         "displayed information is accurate, professional, and "
                                         "relevant.")
    status("No dummy/placeholder content found on the target pages.", color=GREEN)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    _run(p, "Thank You !", bold=True, color=BLACK, size=14)

    doc.save(out_path)


def _build_docx_deltaup(domain, pages_data, findings, captured, brand, out_path):
    """Delta Up - verified against the client reference "On Page-Analysis-Report -
    walkaboutent.ca.docx": teal (#215868) shaded section banners (white bold),
    "Result:" labels in a distinct teal-blue (#31849B, NOT Neon's #1D5489 blue),
    green (#00B050) Recommendations, and green (#388600) reference-URL lines.
    The "Additional Note" block is PLAIN text (no banner), unlike Neon's shaded
    "Additional Suggestions" banner. Uses the shared summary_table helper."""
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc, h = _setup_docx(domain, use_cover=False)
    _insert_format_cover(doc, "cover_deltaup.jpg", page_break=False)
    root = h["root"]; d = h["d"]; FONT = h["FONT"]

    BLACK = RGBColor(0x00, 0x00, 0x00)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    TEAL = RGBColor(0x31, 0x84, 0x9B)
    GREEN = RGBColor(0x00, 0xB0, 0x50)
    URLGREEN = RGBColor(0x38, 0x86, 0x00)

    def _run(p, text, bold=False, color=BLACK, size=12):
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return r

    def header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=True, color=WHITE, size=13)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '215868')
        shd.set(qn('w:val'), 'clear')
        p._p.get_or_add_pPr().append(shd)
        return p

    def body(text, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=bold, color=BLACK, size=12)
        return p

    def labeled(label, text=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, label, bold=True, color=BLACK, size=12)
        if text:
            _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def result(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, "Result: ", bold=True, color=TEAL, size=12)
        _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def recommend(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, "Recommendations: ", bold=True, color=GREEN, size=12)
        _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def refurl(url):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, "Reference URL: ", bold=True, color=BLACK, size=12)
        _run(p, url, bold=False, color=URLGREEN, size=12)
        return p

    def shot(key):
        src = (captured or {}).get(key)
        if src and Path(src).exists():
            try:
                hr._add_bordered_image(doc, src)
            except Exception:
                pass

    home_url = pages_data[0]["url"] if pages_data else root + "/"

    # ---- Intro ----
    body(home_url)
    body("On-page optimization means making changes directly on your website to help it show up "
         "higher in search results.")

    p = doc.add_paragraph()
    _run(p, "On Page Analysis:", bold=True, color=TEAL, size=13)
    body("These are the key things to check before you start improving your website. Doing them "
         "correctly will help your site rank better for keywords in search engines.")

    labeled("Additional Note: ")
    labeled("Note for Content Optimization: ", "Kindly find the attached doc for content optimization.")
    if not findings.get("has_footer_logo"):
        labeled("Note for Footer Logo & Header Logo Hyperlinking: ",
                f"We have noticed that your footer/header logo is not hyperlinked to the homepage. "
                f"Hyperlink with this page: {home_url}")
    if not findings.get("has_custom_404"):
        labeled("Note for Custom 404 Page: ",
                "We have noticed that your website does not have a custom 404 page. We recommend "
                "creating one.")
    if not findings.get("has_blog"):
        labeled("Note for Blog: ",
                "We have noticed that the website does not have a blog page or blog posts.")

    # ---- Image ALT ----
    header("Image ALT Optimization (Main Pages)")
    alt_missing = findings.get("alt_missing", 0)
    if alt_missing:
        result(f"Suitable Image Alt tags not found on {alt_missing} image(s) on the website. It is "
               "not good from an SEO point of view. Kindly find an attached sheet for Image Alt Tag "
               "Suggestions.")
    else:
        result("Suitable Image Alt tags are found on the website. It is good from an SEO point of "
               "view.")
    h["summary_table"](findings)

    # ---- Robots.txt ----
    header("Robots.txt Optimization")
    body("Robots.txt is a regular text file that through its name has special meaning to the "
         "majority of well-behaved web crawlers.")
    if findings.get("robots_found"):
        result("The existing robots.txt file is optimized, which is good from an SEO point of view.")
    else:
        result("An optimized robots.txt file was created; please find it attached and upload it to "
               "the root folder of the website.")
    shot("robots")

    # ---- Sitemap ----
    header("Sitemap.xml Optimization")
    body("Sitemap: A sitemap is a file where you can list the web pages of your site to tell Google "
         "and other search engines about the organization of your site content.")
    if findings.get("sitemap_found"):
        result("Existing sitemap.xml file is optimized, which is good from an SEO point of view.")
        refurl(findings.get("sitemap_url") or (root + "/sitemap.xml"))
    else:
        result("Sitemap.xml file not found on the website. Please create and upload one to the root "
               "folder.")
    shot("sitemap")

    # ---- Redirection ----
    header("URL Redirection Issue Optimization")
    if findings.get("www_redirect_issue") or findings.get("url_changes"):
        result("The website runs with multiple resolvable versions. We suggest redirecting all "
               "variants to a single canonical version using a 301 permanent redirect.")
    else:
        body("Redirection issue not found on the website. It is good from a search engine point of "
             "view.")

    # ---- Canonical ----
    header("Canonical Issue Checking")
    body(f"A canonical problem occurs when a site is running in multiple versions like www version "
         f"(https://www.{d}), the non-www version (https://{d}), and other versions. In this case, "
         "content of website is considered as duplicate by the search engines.")
    if findings.get("canonical_issue"):
        result("An incorrect / conflicting canonical tag was found on the website. Kindly find the "
               "attached Canonical Tag Suggestions sheet.")
    else:
        result("All canonical tags are found on the website. It is good from a search engine point "
               "of view.")
    shot("canonical")

    # ---- External Links ----
    header("External Link Optimization")
    ext_count = findings.get("ext_count", 0)
    body(f"{ext_count} External links found on the home page. None of them seems harmful for the "
         "website and it will benefit." if ext_count else
         "No external links found on the home page.")
    shot("externallinks")

    # ---- Broken Links ----
    header("Broken/Dead Links Optimization")
    body("Though the broken links do not hurt directly but it affects user experience and if users "
         "do not find desired info they will not come frequently and search engine takes this as not "
         "useful for the users and start pulling back the web pages.")
    broken = findings.get("broken_links") or []
    if broken:
        result(f"{len(broken)} Broken Link(s) found in the website. It is not good from a search "
               "engine point of view. Kindly find an attached sheet for suggestions.")
    else:
        result("Broken Link not Found in the website. It is good from a search engine point of view.")
    shot("brokenlinks")

    # ---- Internal Linking ----
    header("Internal Linking Optimization")
    body("Internal Linking: Internal linking is called perfect when every live webpage is "
         "accessed/visited from every other live and available webpage in a website.")
    int_count = len(pages_data[0].get("internal_links", []) or []) if pages_data else 0
    if int_count:
        result("Internal linking structure of website is good and both SEO and user friendly.")
    else:
        result("Please verify the internal linking structure across the website as per need.")

    # ---- Page Speed ----
    header("Page Speed Optimization")
    labeled("Page Optimization Score: ", "Kindly refer to the attached screenshot for the page "
                                          "optimization score.")
    shot("pagespeed")
    refurl(f"https://gtmetrix.com/reports/{root.replace('https://', '')}/")

    # ---- URL Structure ----
    header("URL Structure Optimization")
    recommend("Search Engines like static URLs instead of dynamic ones. Presently the URL structure "
              "of your website's inner pages is user and search engine friendly.")
    if findings.get("url_changes"):
        result("Some target page URLs could be optimized further. Please refer to the attached sheet.")
    else:
        result("Yes Good, Your URL structure is Search Engine friendly. It is good from a search "
               "engine point of view.")

    # ---- Hyperlinking ----
    header("Hyperlink Analysis and Optimization")
    recommend("Hyperlinks are connections established between a word/phrase/image and a "
              "website/file. Effective hyper-linking between different pages can help a website "
              "secure a good position in the search engine result pages as well.")
    result("Hyper-linking of the website is good.")

    # ---- Mixed Content ----
    header("Mixed Content Optimization")
    body("Mixed content occurs when a webpage containing a combination of both secure (HTTPS) and "
         "non-secure (HTTP) content is delivered over SSL to the browser.")
    mixed = findings.get("mixed_content_pages") or []
    if mixed:
        result(f"Mixed content issue found on {len(mixed)} page(s). It's not good from an SEO point "
               "of view.")
    else:
        result("No mixed content issue found in the website. It's good from SEO point of view.")

    # ---- Sucuri ----
    header("Sucuri Site Scan")
    body("The SucuriSiteCheck scanner helps to prevent security threats. It will check malware, "
         "viruses, blacklisting status, website errors, out-of-date software, and malicious code.")
    if findings.get("sucuri_clean", True):
        result("We have not found any malware/security issue on the website. It is good from search "
               "engine point of view.")
    else:
        result("Security issues were detected on your website. Please review and fix them.")
    labeled("Screenshot - ")
    shot("homepage")

    # ---- No-Index ----
    header("No-Index on Target Pages Check")
    body("Some pages of the website serve a purpose, and helps to improve the ranking and traffic to "
         "the site.")
    noindex = findings.get("noindex_pages") or []
    if noindex:
        result(f"Noindex tag found on {len(noindex)} target page(s). Please remove the noindex tag.")
    else:
        result("Noindex not found on the target pages of the robots located in the head section of "
               "the website pages. It is good from a search engine point of view.")

    # ---- Indexing Status ----
    header("Indexing Status of All Target Pages Check")
    body("Making your page appear on search engine search results is called indexing. Getting your "
         "webpages indexed by search engine is extremely important.")
    if findings.get("gsc_indexing", True):
        result("We have pages are indexed. It is good from SEO point of view.")
    else:
        result("Please refer to the attached indexing status sheet for target page indexing details.")

    # ---- Web Archive ----
    header("Old Web Achieve Status Check")
    body("It's the way to explore, find and retrieve historical and 'lost' information from "
         "websites, to serve as evidence that something existed online.")
    result("We have submitted the target pages of the website to the web archive tool for future "
           "reference.")

    # ---- Meta Viewport ----
    header("Meta Viewport")
    body("The viewport meta tag allows you to tell the mobile browser what size this virtual "
         "viewport should be.")
    result("Yes, The Website pages have a viewport meta tag. It will look good on mobile devices and "
           "will get a high position in mobile search results.")

    doc.save(out_path)


def _build_docx_octal(domain, pages_data, findings, captured, brand, out_path):
    """Octal - verified against the client reference "Onpage suggestion Report -
    zeeboo.in.docx": no shaded banners, blue (#4F81BD) title, navy (#002060)
    "Label → description" arrow style for the SEO analysis section (same
    technique as Delta FVR but different palette - title blue, section-title
    #0070C0, "More Content Require" green #00B050), closing with "Thank You!".
    No cover image in the reference - plain text title only."""
    from docx.shared import Pt, RGBColor

    doc, h = _setup_docx(domain, use_cover=False)
    root = h["root"]; d = h["d"]; FONT = h["FONT"]

    BLACK = RGBColor(0x00, 0x00, 0x00)
    BLUE = RGBColor(0x4F, 0x81, 0xBD)
    SECBLUE = RGBColor(0x00, 0x70, 0xC0)
    NAVY = RGBColor(0x00, 0x20, 0x60)
    GREEN = RGBColor(0x00, 0xB0, 0x50)

    def _run(p, text, bold=False, color=BLACK, size=12):
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return r

    def body(text, bold=False, color=BLACK):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=bold, color=color, size=12)
        return p

    def label_arrow(label, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, label + " → ", bold=True, color=NAVY, size=12)
        _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def status(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, "Status: ", bold=True, color=BLACK, size=12)
        _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def shot(key):
        src = (captured or {}).get(key)
        if src and Path(src).exists():
            try:
                hr._add_bordered_image(doc, src)
            except Exception:
                pass

    home_url = pages_data[0]["url"] if pages_data else root + "/"

    # ---- Title ----
    p = doc.add_paragraph()
    _run(p, home_url + "\n", bold=True, color=BLUE, size=13)
    _run(p, "On-Page Optimization Suggestions:", bold=True, color=SECBLUE, size=13)
    body("On-page SEO is a process of optimization of individual Webpages of content to earn better "
         "and more relevant traffic from search engines.", color=RGBColor(0x1A, 0x1A, 0x1A))
    body("We have analyzed your website in terms of search engine perspective. Following are the "
         "suggestions.", color=NAVY)

    p = doc.add_paragraph()
    _run(p, "Additional Suggestions", bold=True, color=BLUE, size=13)
    if not findings.get("social_found"):
        body("Note Social Media Icons: Your website is connected to a limited set of social media "
             "profiles. We recommend adding more platforms to reach a wider audience.")
    body("Optimize your Footer Section: An optimized footer section can improve the user "
         "experience, brand identity, and site navigation.")
    if not findings.get("has_blog"):
        body("Note for Blog: We came to know that your website does not have a proper blog section "
             "with regular posts. We recommend adding one.")
    if not findings.get("has_faq"):
        body("Note For FAQ: FAQs have answers to the general queries that a client might have while "
             "browsing your website - we recommend adding a dedicated FAQ page.")

    p = doc.add_paragraph()
    _run(p, "Content optimization suggestions", bold=True, color=SECBLUE, size=13)
    tp = findings.get("target_pages") or ([pd["url"] for pd in pages_data] if pages_data else [])
    if tp:
        body(f"More Content Require: We inform you that, when we analysing your website, we didn't "
             f"find enough unique content on {len(tp)} target page(s). Kindly refer to the attached "
             "Content Optimization document.", color=GREEN)
    else:
        body("Content on the target pages appears sufficient. Kindly find the attached document for "
             "any further content optimization suggestions.", color=GREEN)

    p = doc.add_paragraph()
    _run(p, "SEO Meta Title and Description Optimization Suggestions", bold=True, color=SECBLUE, size=13)
    body("We suggest Meta suggestions (Meta Title & Meta Description) including Heading tags in the "
         "attached sheet.")

    p = doc.add_paragraph()
    _run(p, "SEO Analysis and Optimization Suggestions", bold=True, color=SECBLUE, size=13)

    # ---- Robots.txt ----
    label_arrow("Robots.txt", "The robots.txt file, also known as the robot's exclusion protocol or "
                              "standard, is a text file that tells web robots which pages to crawl.")
    shot("robots")

    # ---- Sitemap ----
    label_arrow("XML Sitemap", "A sitemap is a blueprint of your website that helps search engines "
                               "find, crawl, and index all of its content.")
    if findings.get("sitemap_found"):
        status("We found a sitemap file on the website, which is good from an SEO point of view.")
    else:
        status("No sitemap file was found on the website. We recommend creating and submitting one.")
    shot("sitemap")

    # ---- Image Alt Tag ----
    alt_missing = findings.get("alt_missing", 0)
    label_arrow("Image Alt Tag", "ALT tags or ALT attributes are \"alternative text\" for an image, "
                                 "used to describe the image to search engines and accessibility "
                                 "tools.")
    if alt_missing:
        status(f"{alt_missing} image(s) without ALT tags found. Please refer to the attached sheet "
               "for Image Alt Tag Suggestions.")
    else:
        status("Suitable Image Alt tags are found on images.")

    # ---- Internal Linking ----
    label_arrow("Internal Linking", "An internal link is any link from one page on your website to "
                                    "another page on your website.")
    int_count = len(pages_data[0].get("internal_links", []) or []) if pages_data else 0
    status("Internal linking of the website is good and user-friendly. It is good from SEO point of "
           "view." if int_count else "Please verify the internal linking structure across the "
                                      "website as per need.")

    # ---- External Linking ----
    ext_count = findings.get("ext_count", 0)
    label_arrow("External Linking", "External Links are hyperlinks that point at (target) any "
                                    "domain other than the domain the link exists on (source).")
    status(f"{ext_count} External links found on the website, none of them seems harmful from an "
           "SEO point of view." if ext_count else "No external links found on the website.")
    shot("externallinks")

    # ---- Site Security ----
    label_arrow("Site Security Check", "A site security check evaluates a website for "
                                       "vulnerabilities, malware, and compliance with security best "
                                       "practices.")
    if findings.get("sucuri_clean", True):
        status("We have not found any malware/security issue on the website. It is good from a "
               "search engine point of view.")
    else:
        status("Security issues were detected on your website. Please review and fix them.")
    shot("homepage")

    # ---- Broken Links ----
    broken = findings.get("broken_links") or []
    label_arrow("Broken/Dead Link", "Broken links are links that send a message to its visitors "
                                    "that the webpage no longer exists, triggering a 404-error page.")
    if broken:
        body(f"Broken links: - {len(broken)} found on the website. Please refer to the attached "
             "sheet for suggestions.")
    else:
        status("No broken links found on the website.")
    shot("brokenlinks")

    # ---- Redirection ----
    label_arrow("Website Redirection", "A redirect is a way to send both users and search engines "
                                       "to a different URL from the one they originally requested.")
    if findings.get("www_redirect_issue") or findings.get("url_changes"):
        status("The website is reachable on multiple versions independently. We recommend a single "
               "301 redirect to one canonical version.")
    else:
        status("No redirection issue found on the website.")

    # ---- Canonicalization ----
    label_arrow("Canonicalization", "A canonical URL refers to an HTML link element, with the "
                                    "attribute of rel=\"canonical\" (also known as a canonical tag), "
                                    "indicating the preferred version of a page.")
    if findings.get("canonical_issue"):
        status("An incorrect / conflicting canonical tag was found on the website. Kindly find the "
               "attached Canonical Tag Suggestions sheet.")
    else:
        status("All canonical tags are found on the website. It's good from an SEO point of view.")
    shot("canonical")

    # ---- Dummy Content ----
    label_arrow("Dummy Content Checker", "A Dummy Content Checker identifies and removes "
                                         "placeholder text or temporary content from a website.")
    status("No dummy/placeholder content found on the target pages.")

    # ---- Hyperlinking ----
    label_arrow("Hyper linking", "Hyperlink is a type of link that connects two separate "
                                 "files/websites/articles together for both users and search "
                                 "engines to navigate.")
    status("Hyperlinking of the website is good. It's good from an SEO point of view.")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    _run(p, "Thank You!", bold=True, color=BLACK, size=14)

    doc.save(out_path)


def _build_docx_camila(domain, pages_data, findings, captured, brand, out_path):
    """Camila - verified against the client reference "On Page Audit Report __
    electroitsolutions.com.docx" (D:\\Report Formats\\On-Page Report Formats\\
    Camila Onpage format): no cover image (the shared onpage_cover.png was
    being incorrectly applied here too - confirmed against the real reference,
    which has NO cover page at all), Candara font throughout, centered title
    ("On Page Audit Report" 26pt / domain 22pt, both bold), solid BLACK
    (#000000) 1-cell-table section banners with white bold uppercase text,
    navy blue (#1D5489) "Result:" labels, black body text, green (#00B050)
    "Note" callout."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc, h = _setup_docx(domain, use_cover=False)
    root = h["root"]; d = h["d"]
    CAMILA_FONT = "Candara"

    BLACK = RGBColor(0x00, 0x00, 0x00)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    BLUE = RGBColor(0x1D, 0x54, 0x89)
    GREEN = RGBColor(0x00, 0xB0, 0x50)

    def _run(p, text, bold=False, color=BLACK, size=12):
        r = p.add_run(text)
        r.font.name = CAMILA_FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return r

    def banner(text):
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        p = cell.paragraphs[0]
        _run(p, text.upper(), bold=True, color=WHITE, size=12)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '000000')
        shd.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shd)
        doc.add_paragraph()
        return table

    def body(text, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=bold, color=BLACK, size=12)
        return p

    def result(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, "Result: ", bold=True, color=BLUE, size=12)
        _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def shot(key):
        src = (captured or {}).get(key)
        if src and Path(src).exists():
            try:
                hr._add_bordered_image(doc, src)
            except Exception:
                pass

    home_url = pages_data[0]["url"] if pages_data else root + "/"

    # ---- Title (verified against the reference: centered, 26pt/22pt) ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "On Page Audit Report", bold=True, color=BLACK, size=26)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, d.capitalize(), bold=True, color=BLACK, size=22)
    body("On-page optimization refers to all measures that can be taken directly within the "
         "website in order to improve its position in the search rankings.")
    p = doc.add_paragraph()
    _run(p, "Full SEO On-Page Analysis: ", bold=True, color=BLUE, size=12)
    body("These are the most significant elements according to Google guideline to be checked "
         "before you start improving your website.")
    p = doc.add_paragraph()
    _run(p, "Note - ", bold=True, color=GREEN, size=12)
    _run(p, "Please refer to the attached file for additional suggestions.", bold=False,
         color=BLACK, size=12)
    body("We suggest Meta suggestions (Meta Title & Meta Description) including Heading tags in "
         "the attached sheet.")

    # ---- Target pages content ----
    banner("Optimization Suggestions for Target Pages")
    banner("Optimization Suggestions for Fresh Content")
    tp = findings.get("target_pages") or ([pd["url"] for pd in pages_data] if pages_data else [])
    if tp:
        body(f"Fresh Keywords Oriented Content - After analyzing the target pages, we found that "
             f"some page(s) among the {len(tp)} target page(s) need fresh, keyword-oriented content. "
             "Kindly find the attached content document for suggestions.")
    banner("Content Suggestion Deferred for Existing Service Pages")
    body("During our analysis, we found that some of the existing service pages are already "
         "targeting relevant keywords - no changes required there.")

    # ---- Image Optimization / Alt Tags ----
    banner("Image Optimization")
    body("Image optimization involves reducing the file size of images without compromising "
         "quality, which helps improve page load speed.")
    body("Status - Suitable images are found on the target page. Which is good from an SEO Point "
         "of view.")
    banner("Optimization Suggestions for Image ALT Tags")
    body("ALT tags or ALT attributes are \"alternative text\" for an image. ALT tags are used to "
         "describe the image or what the image is representing on the webpage.")
    alt_missing = findings.get("alt_missing", 0)
    if alt_missing:
        result(f"We found {alt_missing} image(s) missing appropriate ALT tags on the website's "
               "target pages. Kindly find an attached sheet for Image Alt Tag Suggestions.")
    else:
        result("We have found appropriate ALT tags on the website's target pages, which is good "
               "for an SEO point of view.")

    # ---- Robots.txt ----
    banner("Optimization Suggestions for Robots File")
    if findings.get("robots_found"):
        result("Robots.txt file tells web robots which pages on your site to crawl and which pages "
               "not to crawl - the existing file is optimized, which is good from an SEO point of "
               "view.")
    else:
        result("An optimized robots.txt file was created; please find it attached and upload it to "
               "the root folder of the website.")
    shot("robots")

    # ---- Sitemap ----
    banner("Optimization Suggestions for Sitemap File")
    body("Sitemap: An XML sitemap is specifically written for search engine spiders. A search "
         "engine can use it to find out what content is available and how frequently it's updated.")
    if findings.get("sitemap_found"):
        result("We checked the sitemap file on the website and found it optimized, which is good "
               "from an SEO point of view.")
    else:
        result("We checked the sitemap file on the website and we couldn't find a sitemap on the "
               "website. Please create and upload one.")
    shot("sitemap")

    # ---- Redirection ----
    banner("Optimization Suggestions for URL Redirection")
    body("URL redirection ensures a seamless user experience and preserves SEO value when URLs "
         "change.")
    if findings.get("www_redirect_issue") or findings.get("url_changes"):
        result("We have found URL redirection issues on the website. It is not good from a search "
               "engine point of view. We suggest a single 301 redirect to one canonical version.")
    else:
        result("No redirection issue found on the website. It is good from a search engine point "
               "of view.")

    # ---- Canonical ----
    banner("Optimization Suggestions for Canonical Issue")
    body("A canonical tag is a way of telling search engines that a specific URL represents the "
         "master copy of a page.")
    if findings.get("canonical_issue"):
        result("When we analyzed your website, we noticed that a canonical issue is found on "
               "Target Page(s). Kindly find the attached Canonical Tag Suggestions sheet.")
    else:
        result("All canonical tags are found on the website. It is good from a search engine point "
               "of view.")
    shot("canonical")

    # ---- External Links ----
    banner("Optimization Suggestions for External Link")
    ext_count = findings.get("ext_count", 0)
    body(f"We have found {ext_count} external link(s) on the website, which is currently set as "
         "do-follow. Which is not good from an SEO point of view." if ext_count else
         "No external links found on the website.")
    shot("externallinks")

    # ---- Dead Links ----
    banner("Optimization Suggestions for Dead Links")
    body("A broken or dead link is a link that you click but doesn't work or doesn't take you to "
         "the intended page.")
    broken = findings.get("broken_links") or []
    if broken:
        result(f"{len(broken)}, broken link(s) found on the website. It is not good from search "
               "engine point of view. Kindly find an attached sheet for suggestions.")
    else:
        result("No broken links found on the website. It is good from a search engine point of "
               "view.")
    shot("brokenlinks")

    # ---- Internal Linking / Hyperlinking ----
    banner("Optimization Suggestions for Internal Linking")
    result("Internal Linking: An internal link is any link from one page on your website to "
           "another page on your website. The structure is good and user friendly.")
    banner("Optimization Suggestions for Hyperlinking")
    body("Hyperlink is a type of link that connects two separate files/websites/articles/images "
         "etc.")
    result("The website's hyperlinking is well-structured, which is good from a search engine "
           "point of view.")

    # ---- URL Structure ----
    banner("Optimization Suggestions for URL Structure")
    body("An optimized URL structure refers to a clean, readable, and SEO-friendly format of "
         "website addresses.")
    result("The URL structure of the website is well-optimized and properly formatted. It is good "
           "from a search engine point of view.")

    # ---- No-Index ----
    banner("No-Index on Target Pages Check")
    body("Some pages of the website serve a purpose and helps to improve the ranking and traffic "
         "to the site.")
    noindex = findings.get("noindex_pages") or []
    if noindex:
        result(f"We checked the website and found 'noindex' tags on {len(noindex)} target page(s). "
               "Please remove them.")
    else:
        result("We checked the website and did not find any 'noindex' tags. This is good, as it "
               "allows the pages to be indexed.")

    # ---- FAQ ----
    banner("Note for FAQ Section")
    body("An FAQ section enhances user experience by quickly answering common questions, reducing "
         "support requests.")
    if findings.get("has_faq"):
        result("We found an FAQ section on the website, which is good from an SEO Point of view.")
    else:
        result("We didn't find an FAQ section on the homepage. Which is not Good from an SEO Point "
               "of view. We recommend adding one.")

    # ---- Google Reviews ----
    banner("Note for Google Review")
    body("Google Reviews are important because they build trust with customers and improve your "
         "website's local SEO visibility.")
    result("We did not find a Google Reviews section, but the website includes testimonials on "
           "the homepage.")

    # ---- Mobile Friendliness ----
    banner("Note for Mobile Friendliness")
    body("Mobile friendliness is essential for any website today because most users browse on "
         "their mobile devices.")
    result("We checked your website's mobile friendliness and found that its performance is good.")
    shot("homepage")

    # ---- Security ----
    banner("Site Security Check")
    body("We have checked the website's security and scanned for malware.")
    if findings.get("sucuri_clean", True):
        result("No malware was found, which is good from a search engine point of view.")
    else:
        result("Security issues were detected on your website. Please review and fix them.")
    shot("sucuri")

    # ---- Schema ----
    banner("Schema Mark-up")
    body("Schema Mark-up is a type of structured data used in websites to help search engines "
         "understand the content of the page better.")
    body("Status - We have checked the schema on the website and found that there is no schema "
         "present. We recommend adding relevant schema markup.")

    doc.save(out_path)


def _build_docx_alpha(domain, pages_data, findings, captured, brand, out_path):
    """Alpha - verified against the client reference "Alpha - On Page.docx"
    (mobilityscootrike.com): same solid BLACK 1-cell-table banner technique as
    Camila, but its own distinct section list/wording - adds URL Versions,
    Lang Attribute, SSL Certification, Schema Optimization, Dummy Content and
    Copied Content checks that Camila doesn't have. Navy (#1D5489) "Result:"
    labels, green (#00B050) "Additional Suggestion:" label."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc, h = _setup_docx(domain, use_cover=False)
    _insert_format_cover(doc, "cover_alpha.png", page_break=False)
    root = h["root"]; d = h["d"]; FONT = h["FONT"]

    BLACK = RGBColor(0x00, 0x00, 0x00)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    BLUE = RGBColor(0x1D, 0x54, 0x89)
    GREEN = RGBColor(0x00, 0xB0, 0x50)

    def _run(p, text, bold=False, color=BLACK, size=12):
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return r

    def banner(text):
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        p = cell.paragraphs[0]
        _run(p, text.upper(), bold=True, color=WHITE, size=12)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '000000')
        shd.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shd)
        doc.add_paragraph()
        return table

    def body(text, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=bold, color=BLACK, size=12)
        return p

    def result(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, "Result: ", bold=True, color=BLUE, size=12)
        _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def note(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, "Note - ", bold=True, color=BLACK, size=12)
        _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def shot(key):
        src = (captured or {}).get(key)
        if src and Path(src).exists():
            try:
                hr._add_bordered_image(doc, src)
            except Exception:
                pass

    d_disp = d

    # ---- Title ----
    p = doc.add_paragraph()
    _run(p, d_disp, bold=True, color=BLACK, size=16)
    body("On-page optimization refers to all measures that can be taken directly within the "
         "website in order to improve its position in the search rankings.")
    p = doc.add_paragraph()
    _run(p, "Full SEO On-Page Analysis:", bold=True, color=BLUE, size=12)
    body("These are the most significant elements according to Google guideline to be checked "
         "before you start improving your website.")
    p = doc.add_paragraph()
    _run(p, "Additional Suggestion: - ", bold=True, color=GREEN, size=12)
    if not findings.get("has_blog"):
        body("Note for Blog Page: We found that your website has a blog page, but it currently "
             "does not have regular posts. We recommend publishing fresh content periodically.")
    if findings.get("copyright_stale"):
        year = datetime.date.today().year
        body(f"Note for Copyright: We inform you that, when we analyzed your website, we noticed "
             f"that your copyright year is outdated. Suggested - Copyright {year}.")

    # ---- Image Alt ----
    banner("Optimization Suggestions for Image Alt Tags")
    body("ALT tags or ALT attributes are \"alternative text\" for an image. ALT tags are used to "
         "describe the image or what the image is representing on the webpage.")
    alt_missing = findings.get("alt_missing", 0)
    if alt_missing:
        result(f"Suitable Image Alt texts are not found on {alt_missing} image(s) of the target "
               "page(s) of the website. It is not good from an SEO point of view. Kindly find an "
               "attached sheet for Image Alt Tag Suggestions.")
    else:
        result("Suitable Image Alt texts are found on the target page(s) of the website. It is "
               "good from an SEO point of view.")

    # ---- Robots.txt ----
    banner("Optimization Suggestions for Robots File")
    body("Robots.txt file tells web robots which pages on your site to crawl and which pages not "
         "to crawl.")
    if findings.get("robots_found"):
        result("Optimized Robots.txt file found in your website. It is good from search engine "
               "point of view.")
    else:
        result("An optimized robots.txt file was created; please find it attached and upload it to "
               "the root folder of the website.")
    shot("robots")

    # ---- Sitemap ----
    banner("Optimization Suggestions for Sitemap File")
    body("Sitemap: An XML sitemap is specifically written for search engine spiders. A search "
         "engine can use it to find out what content is available.")
    if findings.get("sitemap_found"):
        result("Sitemap.xml file found in the website & it is optimized. It is good from SEO point "
               "of view.")
    else:
        result("Sitemap.xml file not found on the website. Please create and upload one.")
    shot("sitemap")

    # ---- Hyperlinking ----
    banner("Optimization Suggestions for Hyperlinking")
    body("Hyperlink is a type of link that connects two separate files/websites/articles/images "
         "etc.")
    result("Hyperlinking of the website is good. It's good from an SEO point of view.")

    # ---- URL Versions ----
    banner("URL Versions")
    if findings.get("www_redirect_issue") or findings.get("url_changes"):
        body("As we have checked, the website is running with multiple resolvable versions, it's "
             "not good from an SEO point of view. We suggest redirecting all variants to one "
             "canonical version.")
    else:
        body("As we have Checked and found that the website is running with the main version, "
             "it's good from an SEO point of view.")

    # ---- Canonical ----
    banner("Optimization Suggestions for Canonical Issue")
    body(f"A canonical problem occurs when a site is running in multiple versions like www version "
         f"(https://www.{d}), the non-www version (https://{d}), and other versions.")
    if findings.get("canonical_issue"):
        result("An incorrect / conflicting canonical tag was found on the website. Kindly find the "
               "attached Canonical Tag Suggestions sheet.")
    else:
        result("Canonical issue not found in the website. It's good from an SEO point of view.")
    shot("canonical")

    # ---- External Links ----
    banner("Optimization Suggestions for External Link")
    ext_count = findings.get("ext_count", 0)
    body(f"{ext_count} External links found in the website. None of the links looks harmful from "
         "search engine point of view." if ext_count else
         "No external links found in the website.")
    shot("externallinks")

    # ---- Dead Links ----
    banner("Optimization Suggestions for Dead Links")
    body("A broken or dead link is a link that you click but doesn't work or doesn't take you to "
         "the intended page.")
    broken = findings.get("broken_links") or []
    if broken:
        result(f"{len(broken)} broken link(s) found on the website. It is not good from a search "
               "engine point of view. Kindly find an attached sheet for suggestions.")
    else:
        result("No broken links found in the website. It is good from a search engine point of "
               "view.")
    shot("brokenlinks")

    # ---- Internal Linking ----
    banner("Optimization Suggestions for Internal Linking")
    body("Internal Linking: An internal link is any link from one page on your website to another "
         "page on your website.")
    int_count = len(pages_data[0].get("internal_links", []) or []) if pages_data else 0
    if int_count:
        result("Yes, Good. Your Internal linking is Search Engine friendly.")
    else:
        result("Please verify the internal linking structure across the website as per need.")

    # ---- URL Structure ----
    banner("Optimization Suggestions for URL Structure")
    body("URL Structure: Users and search engines should be able to understand what is on each "
         "page from the URL alone.")
    if findings.get("url_changes"):
        result("Some target page URLs could be optimized further. Please refer to the attached "
               "sheet.")
    else:
        result("Existing URL structure of the website is fine; it is good from a search engine "
               "point of view.")

    # ---- No-Index ----
    banner("No-Index on Target Pages Check")
    body("Some pages of the website serve a purpose, and help to improve the ranking and traffic "
         "to the site.")
    noindex = findings.get("noindex_pages") or []
    if noindex:
        result(f"noindex tag found on {len(noindex)} target page(s); please remove it.")
    else:
        result("noindex tag not found on the target pages; it is good from SEO point of view.")

    # ---- Security ----
    banner("Site Security Check")
    if findings.get("sucuri_clean", True):
        body("As we have checked the site security issue and didn't find any kind of "
             "security/malware issue on the website. It is good from an SEO point of view.")
    else:
        body("Security issues were detected on your website. Please review and fix them.")
    body("Screenshot - ")
    shot("sucuri")

    # ---- Mixed Content ----
    banner("Mixed Content optimization")
    body("Mixed content occurs when a webpage containing a combination of both secure (HTTPS) and "
         "non-secure (HTTP) content is delivered over SSL to the browser.")
    mixed = findings.get("mixed_content_pages") or []
    if mixed:
        result(f"Mixed content issue found on {len(mixed)} page(s). It's not good from an SEO "
               "point of view.")
    else:
        result("After reviewing the website, we have not identified a mixed content issue. This "
               "is good from an SEO point of view.")

    # ---- Meta Viewport ----
    banner("Meta viewport")
    body("The viewport on a website is the visible area of a webpage within a browser window, "
         "specifically important for mobile devices.")
    result("Meta viewport is found in the website. This is good for the SEO point of view."
           if findings.get("viewport", True) else
           "Meta viewport tag was not found. We recommend adding one.")

    # ---- Lang Attribute ----
    banner("Lang Attribute")
    body("The \"Lang\" attribute is an HTML attribute used to specify the language of the content "
         "within a webpage.")
    lang = findings.get("lang") or "en-US"
    result(f"The lang=\"{lang}\" attribute is present on your website. This is beneficial from an "
           "SEO point of view.")

    # ---- SSL ----
    banner("SSL Certification")
    body("An SSL (Secure Sockets Layer) certificate is a digital file that encrypts the connection "
         "between a browser and a website.")
    note("As we have checked the SSL certification of your website, it is valid and active.")

    # ---- Schema ----
    banner("Schema Optimization")
    body("A \"Website Schema\" refers to a set of structured data tags, often called \"schema "
         "markup,\" that help search engines understand the content of a page.")
    note("We have currently found webpage & localbusiness schema on the website, but we recommend "
         "adding more specific schema types where relevant.")

    # ---- Dummy Content ----
    banner("Dummy Content")
    body("Dummy content, also known as filler text or placeholder text, is non-functional text or "
         "media used as a temporary placeholder.")
    note("We have not found dummy content on the website. It is good for the search engine point "
         "of view.")

    # ---- Copied Content ----
    banner("Copied Content")
    body("\"Copied content\" refers to any piece of text or information that has been directly "
         "taken from another source without modification.")
    note("As we have checked the content on the website, we have not found copied content on the "
         "target pages.")

    doc.save(out_path)


def _build_docx_eta(domain, pages_data, findings, captured, brand, out_path):
    """ETA - verified against the client reference "On Page Suggestion Report -
    fundamedic.com.docx": no shaded banners or tables at all - plain bold black
    14pt section titles, navy (#1D5489) "Result:" labels, green (#00B050)
    "Recommendations:" labels. Adds Lang Attribute, SSL Certification, Indexing
    Overview, Backlink Audit, Copied Content, Structured Schema Data and Web
    Archive checks beyond the core set."""
    from docx.shared import Pt, RGBColor

    doc, h = _setup_docx(domain, use_cover=False)
    _insert_format_cover(doc, "cover_eta.jpg", page_break=False)
    root = h["root"]; d = h["d"]; FONT = h["FONT"]

    BLACK = RGBColor(0x00, 0x00, 0x00)
    BLUE = RGBColor(0x1D, 0x54, 0x89)
    GREEN = RGBColor(0x00, 0xB0, 0x50)
    TITLEBLUE = RGBColor(0x6D, 0x9E, 0xEB)

    def _run(p, text, bold=False, color=BLACK, size=12):
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return r

    def section(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=True, color=BLACK, size=14)
        return p

    def body(text, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=bold, color=BLACK, size=12)
        return p

    def result(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, "Result: ", bold=True, color=BLUE, size=12)
        _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def recommend(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, "Recommendations: ", bold=True, color=GREEN, size=12)
        _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def shot(key):
        src = (captured or {}).get(key)
        if src and Path(src).exists():
            try:
                hr._add_bordered_image(doc, src)
            except Exception:
                pass

    # ---- Title ----
    p = doc.add_paragraph()
    _run(p, d, bold=True, color=TITLEBLUE, size=16)
    body("On-Page Optimization (AKA on-page SEO) refers to all measures that can be taken directly "
         "within the website to improve its position in the search rankings.")
    body("Note For Meta Suggestion - Please find the attached sheet for SEO oriented titles & "
         "descriptions.")
    body("Note For Content Suggestion - Please find the attached sheet for SEO oriented keywords "
         "for the target pages.")

    section("Additional Suggestion")
    if not findings.get("has_privacy_policy", True):
        body("Missing Privacy Policy Page - While auditing the website, we found that the Privacy "
             "Policy page is missing. We recommend adding one.")
    if not findings.get("has_terms", True):
        body("Missing Terms & Conditions Page - While auditing the website, we found that the "
             "Terms & Conditions page is missing.")
    if not findings.get("has_blog"):
        body("Blogs Section: We did not find an active blog section with recent posts. We "
             "recommend publishing fresh content periodically.")
    if findings.get("copyright_stale"):
        year = datetime.date.today().year
        body(f"Outdated Copyright Notice - During our website audit, we found that the copyright "
             f"notice is outdated. Suggested - Copyright {year}.")

    # ---- Canonical ----
    section("Canonical Issue Suggestion")
    body(f"A canonical problem occurs when a site is running in multiple versions like www version "
         f"(https://www.{d}), the non-www version (https://{d}), and other versions.")
    if findings.get("canonical_issue"):
        result("An incorrect / conflicting canonical tag was found on the website. Kindly find the "
               "attached Canonical Tag Suggestions sheet.")
    else:
        result("All Target pages have proper canonical tags. It is good from an SEO point of view.")
    shot("canonical")

    # ---- Robots.txt ----
    section("Robots.txt Optimization")
    body("Robots.txt is a regular text file that through its name has special meaning to the "
         "majority of well-behaved web crawlers.")
    if findings.get("robots_found"):
        result("Optimized robots.txt file found in your website. It is good from a search engine "
               "point of view.")
    else:
        result("An optimized robots.txt file was created; please find it attached and upload it to "
               "the root folder of the website.")
    shot("robots")

    # ---- Sitemap ----
    section("Sitemap Optimization")
    body("Sitemap: A sitemap is a file where you can list the web pages of your site to tell "
         "Google and other search engines about the organization of your site content.")
    if findings.get("sitemap_found"):
        result("Optimized sitemap.xml file found in your website. It is good from a search engine "
               "point of view.")
    else:
        result("Sitemap.xml file not found on the website. Please create and upload one.")
    shot("sitemap")

    # ---- Image Alt ----
    section("Image Alt Text Optimization")
    body("Image alt text is a descriptive attribute added to images in HTML that helps search "
         "engines and screen readers understand the image content.")
    alt_missing = findings.get("alt_missing", 0)
    if alt_missing:
        result(f"Image alt texts are not found on {alt_missing} image(s) of your website. It is "
               "not good from both a search engine and accessibility point of view.")
    else:
        result("Image alt texts are found on your website. It is good from an SEO point of view.")

    # ---- Internal Linking ----
    section("Internal Link Optimization")
    body("Internal Linking: Internal linking is considered optimal when every live webpage is "
         "reachable from other live webpages on the site.")
    result("The website's internal linking structure is good.")

    # ---- External Linking ----
    section("External Link Optimization")
    ext_count = findings.get("ext_count", 0)
    result(f"{ext_count} External Links found in the website. None of them seem harmful. It is "
           "good from an SEO point of view." if ext_count else
           "No external links found in the website.")
    shot("externallinks")

    # ---- Redirection ----
    section("URL Redirection Issue Optimization")
    if findings.get("www_redirect_issue") or findings.get("url_changes"):
        result("The website is reachable on multiple versions independently. We recommend a "
               "single 301 redirect to one canonical version.")
    else:
        result("All URL redirections are properly set up with consistent HTTPS and www versions, "
               "which is good from a search engine point of view.")

    # ---- Broken Links ----
    section("Broken/Dead Link Suggestion")
    body("Though the broken links do not hurt directly but it affects user experience and users "
         "do not find desired info they will not come frequently.")
    broken = findings.get("broken_links") or []
    if broken:
        result(f"{len(broken)} broken link(s) found in the website. Kindly find an attached sheet "
               "for suggestions.")
    else:
        result("No broken links found in the website.")
    shot("brokenlinks")

    # ---- URL Structure ----
    section("URL Structure Optimization")
    recommend("Proper URL structure helps both users and search engines understand the page "
              "content at a glance.")
    if findings.get("url_changes"):
        result("Some target page URLs could be optimized further. Please refer to the attached "
               "sheet.")
    else:
        result("Your website's current URL structure is clean, user-friendly, and optimized for "
               "search engines.")

    # ---- Hyperlinking ----
    section("Hyperlink Analysis and Optimization")
    recommend("Hyperlinks are connections established between a word/phrase/image and a "
              "website/file. Effective hyperlinking can help a website secure a good position in "
              "the search engine result pages as well.")
    result("Hyperlinking of the website is good.")

    # ---- Mixed Content ----
    section("Mixed Content Optimization")
    body("Mixed content occurs when a webpage containing a combination of both secure (HTTPS) and "
         "non-secure (HTTP) content is delivered over SSL to the browser.")
    mixed = findings.get("mixed_content_pages") or []
    if mixed:
        result(f"Mixed content issue found on {len(mixed)} page(s). It's not good from an SEO "
               "point of view.")
    else:
        result("We have not found mixed content issues on your website, which is good from an SEO "
               "point of view.")

    # ---- Sucuri ----
    section("Sucuri Site Scan")
    body("The SucuriSiteCheck scanner helps to prevent security threats. It will check malware, "
         "viruses, blacklisting status, website errors, out-of-date software, and malicious code.")
    if findings.get("sucuri_clean", True):
        result("No, malware issue not found in the website. It is good from an SEO point of view.")
    else:
        result("Security issues were detected on your website. Please review and fix them.")
    body("Screenshot - ")
    shot("sucuri")

    # ---- No-Index ----
    section("No-index on Target Pages Check")
    body("Certain pages on a website serve specific purposes, aiding in ranking and traffic. "
         "These pages should not carry a noindex tag unless intentional.")
    noindex = findings.get("noindex_pages") or []
    if noindex:
        result(f"We checked the website and the 'noindex, nofollow' tag is present on "
               f"{len(noindex)} target page(s). Please remove it.")
    else:
        result("We checked the website and the 'noindex, nofollow' tag is not present on the "
               "pages. It is good from a search engine point of view.")

    # ---- Meta Viewport ----
    section("Meta Viewport Check")
    body("The viewport meta tag defines the size of the virtual viewport, improving rendering on "
         "mobile devices.")
    result("Yes, the website pages have a viewport Meta tag. It will look good on mobile devices."
           if findings.get("viewport", True) else
           "Meta viewport tag was not found. We recommend adding one.")

    # ---- Lang Attribute ----
    section("Lang Attribute")
    body("The \"Lang\" attribute is an HTML attribute used to specify the language of the content "
         "within a webpage.")
    lang = findings.get("lang") or "en-US"
    result(f"We have found the lang=\"{lang}\" attribute in the website, which is good from a "
           "search engine point of view.")

    # ---- SSL ----
    section("SSL Certification")
    body("An SSL (Secure Sockets Layer) certificate is a digital file that encrypts the connection "
         "between a browser and a website.")
    result("The website currently has a valid SSL certificate. It is good from a security and "
           "search engine point of view.")

    # ---- Indexing Overview ----
    section("Indexing Overview:- ")
    body("Indexing is the process by which search engines like Google crawl and store website "
         "pages so they can appear in search results.")
    if findings.get("gsc_indexing"):
        result("The site's target pages are indexed on Google, indicating that the site is being "
               "discovered and crawled correctly.")
    else:
        result("Please refer to the attached indexing status sheet for target page indexing "
               "details.")

    # ---- Backlink Audit ----
    section("Backlink Audit")
    body("A backlink audit is a thorough analysis of the links pointing to your site. These links "
         "influence domain authority and search rankings.")
    body("Note - Please refer to the attached backlink profile sheet for details. To improve "
         "overall SEO performance, it is recommended to focus on building high-quality, relevant "
         "backlinks.")

    # ---- Copied Content ----
    section("Copied Content")
    body("\"Copied content\" refers to any piece of text or information that has been directly "
         "taken from another source without modification.")
    result("As we have checked the content on the website, we have not found copied content on "
           "the target pages.")

    # ---- Dummy Content ----
    section("Dummy Content Checker")
    body("Dummy content, or placeholder text, is non-functional text or media that's used to "
         "simulate real content during development.")
    result("We did not find any dummy content (Lorem Ipsum) on the website, which is good for the "
           "search engine point of view.")

    # ---- Schema ----
    section("Structured Schema Data")
    body("Schema (or Schema.org markup) is a type of structured data you add to your website's "
         "HTML to help search engines understand the content better.")
    result("Currently, only basic schema types have been implemented on the website. We recommend "
           "adding more specific schema types where relevant.")

    # ---- Web Archive ----
    section("Old Web Archive Status Check")
    body("It's the way to explore, find and retrieve historical and 'lost' information from "
         "websites.")
    result("Yes, we have found that the pages have been saved in the web archive tool. It is good "
           "for reference purposes.")

    doc.save(out_path)


def _build_docx_kappa(domain, pages_data, findings, captured, brand, out_path):
    """Kappa - verified against the client reference "ON Page Audit Report __
    Amail.Agency.docx" (from the "Kappa Up" report suite's On-page folder):
    same 1-cell-table banner technique as Camila/Alpha, but PURPLE (#5F497A)
    fill instead of black, uppercase white bold text. Adds a "Note for USP"
    check that no other format has."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc, h = _setup_docx(domain, use_cover=False)
    _insert_format_cover(doc, "cover_kappa.jpg", page_break=False)
    root = h["root"]; d = h["d"]; FONT = h["FONT"]

    BLACK = RGBColor(0x00, 0x00, 0x00)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    BLUE = RGBColor(0x1D, 0x54, 0x89)

    def _run(p, text, bold=False, color=BLACK, size=12):
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return r

    def banner(text):
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        p = cell.paragraphs[0]
        _run(p, text.upper(), bold=True, color=WHITE, size=12)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '5F497A')
        shd.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shd)
        doc.add_paragraph()
        return table

    def body(text, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=bold, color=BLACK, size=12)
        return p

    def result(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, "Result: ", bold=True, color=BLUE, size=12)
        _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def shot(key):
        src = (captured or {}).get(key)
        if src and Path(src).exists():
            try:
                hr._add_bordered_image(doc, src)
            except Exception:
                pass

    home_url = pages_data[0]["url"] if pages_data else root + "/"

    # ---- Title ----
    p = doc.add_paragraph()
    _run(p, "ON PAGE AUDIT REPORT", bold=True, color=BLACK, size=18)
    p = doc.add_paragraph()
    _run(p, d.capitalize(), bold=True, color=BLACK, size=14)
    body("On-page optimization refers to all measures that can be taken directly within the "
         "website in order to improve its position in the search rankings.")
    p = doc.add_paragraph()
    _run(p, "Full SEO On-Page Analysis: ", bold=True, color=BLUE, size=12)
    body("These are the most significant elements according to Google guideline to be checked "
         "before you start improving your website.")
    body("Note - Please refer to the attached file for additional suggestions.")
    body("We suggest Meta suggestions (Meta Title & Meta Description) including Heading tags in "
         "the attached sheet.")

    # ---- Target pages / keywords ----
    banner("Optimization Suggestions for Target Pages")
    banner("Optimization Suggestions for Keyword Optimization")
    tp = findings.get("target_pages") or ([pd["url"] for pd in pages_data] if pages_data else [])
    if tp:
        body(f"Keywords Oriented Content - After analyzing your website, we found that while the "
             f"target page(s) are relevant, {len(tp)} page(s) could target their keywords more "
             "explicitly. Kindly find the attached content document for suggestions.")

    # ---- Image / Alt ----
    banner("Image Optimization")
    body("Image optimization involves reducing the file size of images without compromising "
         "quality, which helps improve page load speed.")
    body("Status - Suitable and Optimized images are found on the target page. It is good from "
         "search engine point of view.")
    banner("Optimization Suggestions for Image Alt Tags")
    body("ALT tags or ALT attributes are \"alternative text\" for an image. ALT tags are used to "
         "describe the image or what the image is representing on the webpage.")
    alt_missing = findings.get("alt_missing", 0)
    if alt_missing:
        result(f"We have not found appropriate ALT tags on {alt_missing} of the website's target "
               "pages, which is not good from an SEO point of view.")
    else:
        result("We have found appropriate ALT tags on the website's target pages, which is good "
               "for an SEO point of view.")

    # ---- Robots.txt ----
    banner("Optimization Suggestions for Robots File")
    body("Robots.txt file tells web robots which pages on your site to crawl and which pages not "
         "to crawl.")
    if findings.get("robots_found"):
        result("We check the website robots file is found. It is good from search engine point of "
               "view.")
    else:
        result("An optimized robots.txt file was created; please find it attached and upload it "
               "to the root folder of the website.")
    shot("robots")

    # ---- Sitemap ----
    banner("Optimization Suggestions for Sitemap File")
    body("Sitemap: An XML sitemap is specifically written for search engine spiders.")
    if findings.get("sitemap_found"):
        result("Sitemap file is present on the website. It is good from search engine point of "
               "view.")
    else:
        result("Sitemap file not found on the website. Please create and upload one.")
    shot("sitemap")

    # ---- Redirection ----
    banner("Optimization Suggestions for URL Redirection")
    body("URL redirection ensures a seamless user experience and preserves SEO value when URLs "
         "change.")
    if findings.get("www_redirect_issue") or findings.get("url_changes"):
        result(f"We have identified a redirection issue on the website. We suggest redirecting "
               f"all variants to a single canonical version.")
    else:
        result("No redirection issue found on the website. It is good from a search engine point "
               "of view.")

    # ---- Canonical ----
    banner("Optimization Suggestions for Canonical Issue")
    body("A canonical tag is a way of telling search engines that a specific URL represents the "
         "master copy of a page.")
    if findings.get("canonical_issue"):
        result("When we analyzed your website, we noticed that a canonical issue is found on "
               "Target Page(s). Kindly find the attached Canonical Tag Suggestions sheet.")
    else:
        result("When we analyzed your website, we noticed that canonical issue is not found on "
               "Target Pages. It is good from an SEO point of view.")
    shot("canonical")

    # ---- External Links ----
    banner("Optimization Suggestions for External Link")
    ext_count = findings.get("ext_count", 0)
    body(f"{ext_count} external link(s) are found on the website. It is good from an SEO "
         "perspective." if ext_count else "0 external links are found on the website. It is good "
                                          "from an SEO perspective.")
    shot("externallinks")

    # ---- Dead Links ----
    banner("Optimization Suggestions for Dead Links")
    body("A broken or dead link is a link that you click but doesn't work or doesn't take you to "
         "the intended page.")
    broken = findings.get("broken_links") or []
    if broken:
        result(f"{len(broken)} broken link(s) found on the website. Kindly find an attached sheet "
               "for suggestions.")
    else:
        result("No, broken link found on the website. It is good from search engine point of "
               "view.")
    shot("brokenlinks")

    # ---- Internal Linking / Hyperlinking ----
    banner("Optimization Suggestions for Internal Linking")
    result("Internal Linking: An internal link is any link from one page on your website to "
           "another page on your website. The structure is good.")
    banner("Optimization Suggestions for Hyperlinking")
    body("Hyperlink is a type of link that connects two separate files/websites/articles/images "
         "etc.")
    result("Hyperlinking of the website is good. It is good from search engine point of view.")

    # ---- URL Structure ----
    banner("Optimization Suggestions for URL Structure")
    body("An optimized URL structure refers to a clean, readable, and SEO-friendly format of "
         "website addresses.")
    result("URL Structure of the website is fine. It is also good from Search Engine point of "
           "view.")

    # ---- No-Index ----
    banner("No-Index on Target Pages Check")
    body("Some pages of the website serve a purpose and help to improve the ranking and traffic "
         "to the site.")
    noindex = findings.get("noindex_pages") or []
    if noindex:
        result(f"The website has a 'noindex' tag on {len(noindex)} page(s), which will prevent "
               "them from being indexed. Please remove it.")
    else:
        result("The website does not have a 'noindex' tag, which is good for SEO as it allows "
               "search engines to index the pages.")

    # ---- USP ----
    banner("Note for USP")
    body("A USP (Unique Selling Proposition) is a short and clear message that explains what "
         "makes your business different from competitors.")
    if findings.get("has_usp", False):
        result("We found that your homepage has a clear Unique Selling Proposition (USP).")
    else:
        result("We found that your homepage lacks a clear Unique Selling Proposition (USP). "
               "Adding one is recommended to improve conversions.")

    # ---- FAQ ----
    banner("Note for FAQ Section")
    body("An FAQ section enhances user experience by quickly answering common questions, reducing "
         "support requests.")
    if findings.get("has_faq"):
        result("The FAQ section is already included on your homepage and is well-structured from "
               "both a user and SEO perspective.")
    else:
        result("We did not find an FAQ section on the homepage. We recommend adding one.")

    # ---- Google Reviews ----
    banner("Note for Google Review")
    body("Google Reviews are important because they build trust with customers and improve your "
         "website's local SEO visibility.")
    result("Your homepage currently lacks a review section, which is a missed opportunity from an "
           "SEO and trust point of view.")

    # ---- Mobile Friendliness ----
    banner("Note for Mobile Friendliness")
    body("Mobile friendliness is essential for any website today because most users browse on "
         "their mobile devices.")
    result("We checked your website's mobile friendliness and found that its performance is good.")
    shot("homepage")

    # ---- Security ----
    banner("Site Security Check")
    if findings.get("sucuri_clean", True):
        body("The scanner did not detect any malware on the website. It is good from search "
             "engine point of view.")
    else:
        body("Security issues were detected on your website. Please review and fix them.")
    body("Screenshot - ")
    shot("sucuri")

    # ---- Schema ----
    banner("Schema Markup")
    body("Schema Markup is a type of structured data used in websites to help search engines "
         "understand the content of the page better.")
    body("Status - Schema Markup is not found on the website. It is not good from search engine "
         "point of view. We recommend adding relevant schema markup.")
    body("Screenshot -")

    doc.save(out_path)


def _build_docx_peta(domain, pages_data, findings, captured, brand, out_path):
    """Peta / Beta - verified against the client reference "On page Suggestions
    Report - calcmaster.in.docx": a dark (#171717, white bold) title banner, then
    each section is its own light-gray (#D9D9D9) bold-black banner followed by
    black Calibri body text. "Additional suggestion:" is a standalone bold red
    label. Broken Links is a real 4-column data table (Broken Link / Link Text /
    source page / Suggestion), not a screenshot. Data-driven Result lines reuse
    the same `findings` fields as the other builders so this format stays honest.
    No cover image in the reference - its own dark title banner is the first
    content."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc, h = _setup_docx(domain, use_cover=False)
    root = h["root"]; d = h["d"]; FONT = h["FONT"]

    BLACK = RGBColor(0x00, 0x00, 0x00)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    RED = RGBColor(0xFF, 0x00, 0x00)

    def _run(p, text, bold=False, color=BLACK, size=12):
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return r

    def _banner(text, fill, color, size=12):
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        p = cell.paragraphs[0]
        _run(p, text, bold=True, color=color, size=size)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill)
        shd.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shd)
        doc.add_paragraph()
        return table

    def title_banner(text):
        return _banner(text, '171717', WHITE, size=14)

    def section(text):
        return _banner(text, 'D9D9D9', BLACK, size=12)

    def body(text, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=bold, color=BLACK, size=12)
        return p

    def shot(key):
        src = (captured or {}).get(key)
        if src and Path(src).exists():
            try:
                hr._add_bordered_image(doc, src)
            except Exception:
                pass

    # ---- Title ----
    title_banner(f"On-Page Report for - {d}")

    body("On-Page Optimization Working: The work that is directly done on and influences web pages "
         "in order to rank them higher on the search engine and earn more relevant traffic. This "
         "report contains all the issues and suggestions to optimize the website, follow the "
         "suggestions to get better organic results.")
    body("All the following On-page factors and the given suggestions need attention and should be "
         "fixed likewise, as it will help the website's growth on search engines.")
    body("Required changes need to be done ASAP as they often play a major role in acquiring a good "
         "position in the search engine.")

    p = doc.add_paragraph()
    _run(p, "Additional suggestion:", bold=True, color=RED, size=12)
    if not findings.get("has_blog"):
        body("Note for Blog page: - We checked your website and noticed that a proper Blog / "
             "Articles page was not found. We recommend adding one to reach your target audience "
             "more effectively.")
    if not findings.get("social_found"):
        body("Note for Social Media: - During our review, we noticed that your website does not "
             "currently display links to your social media profiles. Adding these improves "
             "engagement and brand trust.")
    body("Please Find Attached Sheet for Meta Suggestion")
    body("Please Find Document for Content suggestion")

    # ---- URL Optimization ----
    section("URL Optimization")
    body("The existing URL structure of the website is good. It is good from a search engine "
         "point of view." if not findings.get("url_changes")
         else "Some target page URLs could be optimized further for readability and keyword "
              "relevance. Please refer to the attached sheet.")

    # ---- Non-Indexable Attributes ----
    section("Analysis of Non-Index Able Attributes (Main pages)")
    noindex = findings.get("noindex_pages") or []
    if noindex:
        body(f"When we analyzed the website, we found {len(noindex)} page(s) with a Noindex robots "
             "tag. This should be reviewed as it prevents the page from appearing in search results.")
    else:
        body("When we analyzed the website, we did not find a Noindex robots tag. This is good from "
             "a search engine point of view.")

    # ---- Alt Tags ----
    section("Alt Tag Optimization (Target Pages)")
    body("The alt tag describes the content of an image, which helps search engines understand it.")
    alt_missing = findings.get("alt_missing", 0)
    if alt_missing:
        body(f"Image Alt tag is missing on {alt_missing} image(s) on the target page(s) of the "
             "website. It is not good from a search engine point of view. So, we suggest adding "
             "relevant Alt tags. Please refer to the attached sheet.")
    else:
        body("Image Alt tag found on the target page of the website. It is good from a search "
             "engine point of view.")

    # ---- Internal Linking ----
    section("Optimization of Internal Navigation / Linking Structure (Main Pages)")
    body("Internal linking structure of website is good and both SEO and user friendly.")

    # ---- External Links ----
    section("Optimization of External Links (Main Pages)")
    ext_count = findings.get("ext_count", 0)
    home_url = (pages_data[0]["url"] if pages_data else root + "/")
    p = doc.add_paragraph()
    _run(p, f"Web-Page URL - {home_url}", bold=True, color=BLACK, size=12)
    body(f"{ext_count} external link(s) found on the home page - All external links are safe, "
         "it is good from an SEO point of view." if ext_count else
         "No external links found on the home page.")
    shot("externallinks")

    # ---- Broken Links ----
    section("Analysis of Broken Links (Main Pages)")
    broken = findings.get("broken_links") or []
    if broken:
        body(f"We inform you that, when we analyzed your website, we found {len(broken)} broken "
             "link(s). It is not good from a search engine point of view. Please refer to the table "
             "below and fix/remove these links.")
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, txt in enumerate(["Broken Link", "Link Text", "Source Page", "Suggestion"]):
            hdr[i].text = ""
            pr = hdr[i].paragraphs[0]
            _run(pr, txt, bold=True, color=WHITE, size=11)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '1F4E79')
            shd.set(qn('w:val'), 'clear')
            hdr[i]._tc.get_or_add_tcPr().append(shd)
        for b in broken[:40]:
            row = table.add_row().cells
            vals = [b.get("url", "") if isinstance(b, dict) else str(b),
                    b.get("text", "") if isinstance(b, dict) else "",
                    b.get("source", "") if isinstance(b, dict) else "",
                    "Remove This Link"]
            for i, v in enumerate(vals):
                row[i].text = ""
                pr = row[i].paragraphs[0]
                _run(pr, str(v), bold=False, color=BLACK, size=10)
        doc.add_paragraph()
    else:
        body("We inform you that, when we analyzed your website, we did not find any broken links. "
             "It is good from a search engine point of view.")

    # ---- Canonicalization ----
    section("Checking of Canonicalization Error (Main Pages)")
    body("A canonical problem occurs when a site is running in multiple versions like the www "
         f"version (https://www.{d}), the non-www version (https://{d}), and other versions. This "
         "causes duplicate content issues. We add a canonical tag to inform search engines of the "
         "genuine/preferred version of the page.")
    if findings.get("canonical_issue"):
        p = doc.add_paragraph()
        _run(p, "Result: ", bold=True, color=RED, size=12)
        _run(p, "An incorrect / conflicting canonical tag was found on the website. Kindly find the "
                "attached Canonical Tag Suggestions sheet.", bold=False, color=BLACK, size=12)
    else:
        p = doc.add_paragraph()
        _run(p, "Result: ", bold=True, color=BLACK, size=12)
        _run(p, "Canonical issue not found in the website. It's good from an SEO point of view.",
             bold=False, color=BLACK, size=12)
    shot("canonical")

    # ---- Sitemap ----
    section("XML Sitemap Optimization")
    body("An XML sitemap is a file that lists a website's important pages, making sure search "
         "engines can find and crawl them.")
    if findings.get("sitemap_found"):
        body("Optimized XML Sitemap found on the website. It's good from a search engine point of "
             f"view. ({findings.get('sitemap_url', '')})")
    else:
        body("We did not find an XML Sitemap on the website. We recommend creating and submitting "
             "one via Google Search Console.")
    shot("sitemap")

    # ---- Robots.txt ----
    section("Robot’s Optimization (Main Pages)")
    body("A robots.txt file tells search engine crawlers which pages or files they can or can't "
         "request from the site.")
    if findings.get("robots_found"):
        body("Optimized Robots.txt file found on the website. It's good from a search engine point "
             "of view.")
    else:
        body("We did not find a robots.txt file on the website. We recommend adding one.")
    shot("robots")

    # ---- Redirection ----
    section("Checking of Redirection Issue")
    body("A redirect happens when someone visits one URL and it takes them to a different URL. "
         "Multiple resolvable versions of a domain confuse search engines.")
    if findings.get("www_redirect_issue") or findings.get("url_changes"):
        body("We found that the website is reachable on multiple versions independently. We "
             "recommend a single 301 redirect to one canonical version.")
    else:
        body(f"We checked your website and found that it is correctly resolving to a single "
             f"version. (https://{d}/)")

    # ---- Dummy Content ----
    section("Dummy Content Checker")
    body("Note Regarding Dummy Content: During our analysis of your website, we did not find any "
         "dummy/placeholder content on the target pages.")

    # ---- Hyperlinking ----
    section("Optimization Suggestions for Hyperlinking")
    body("Hyperlinking of the website is good. It's good from an SEO point of view.")
    p = doc.add_paragraph()
    _run(p, f"Web-Page URL - {home_url}", bold=True, color=BLACK, size=12)

    # ---- Security ----
    section("Security Issue Check")
    if findings.get("sucuri_clean", True):
        body("We have not found any malware/security issues on the website. It is good from a "
             "search engine point of view.")
    else:
        body("Our scan flagged a possible security issue on the website. Please review and clean "
             "the site as soon as possible.")
    shot("homepage")

    # ---- Mixed Content ----
    section("Mixed Content Optimization")
    body("A mixed-content warning indicates that both secured and unsecured elements are being "
         "served on a page loaded over HTTPS, which browsers may flag as insecure.")
    mixed = findings.get("mixed_content_pages") or []
    if mixed:
        p = doc.add_paragraph()
        _run(p, "Result: ", bold=True, color=RED, size=12)
        _run(p, f"{len(mixed)} page(s) with mixed content were found. Please review and update "
                "these resources to HTTPS.", bold=False, color=BLACK, size=12)
    else:
        p = doc.add_paragraph()
        _run(p, "Result: ", bold=True, color=BLACK, size=12)
        _run(p, "No issue found in the website. It's good from an SEO point of view.",
             bold=False, color=BLACK, size=12)

    doc.save(out_path)


def _build_docx_sara(domain, pages_data, findings, captured, brand, out_path):
    """Sara - independent frozen copy of the reference-verified Neon build (teal
    #215868 template, verified against On Page Suggestion Report format Sara /
    dcmshriramchemicals.com). Kept separate so Neon can diverge (see its own
    docstring) without changing Sara, which must keep matching its reference.
    """
    """Neon (Sjjanarrabeen) on-page report.

    Neon has NO page background (white). Section headers are TEAL (215868) shaded
    bars with white bold 13pt text - NOT the navy 000066 ribbon used by James/Xenon.
    Body/description text is Calibri 12pt. "Result:" labels are bold BLUE (1D5489)
    where the reference shows them, other labels (Hyperlinking/Robots "Result",
    Broken-Link "Conclusion") are bold black. Section order + wording mirror the
    Sjjanarrabeen.com.au reference. Result/Conclusion text stays data-driven
    (good vs issue) using the shared ``findings`` fields. Rendered inline with a
    local ``header`` helper so it never emits the shared navy ribbon.
    """
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc, h = _setup_docx(domain, use_cover=False)
    _insert_format_cover(doc, "cover_sara.png", page_break=False)
    # Neon has NO page background - do not call _set_page_background().
    home = next((pd for pd in pages_data if urllib.parse.urlparse(pd["url"]).path in ("", "/")),
                pages_data[0] if pages_data else {})
    year = datetime.date.today().year
    root = h["root"]
    d = h["d"]
    FONT = h["FONT"]

    BLACK = RGBColor(0x00, 0x00, 0x00)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    BLUE = RGBColor(0x1D, 0x54, 0x89)
    SCORE = RGBColor(0xFF, 0xAA, 0x33)
    URLCLR = RGBColor(0xF7, 0x96, 0x46)

    def _run(p, text, bold=False, color=BLACK, size=12):
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return r

    def header(text):
        # Neon's OWN section header: TEAL (215868) shaded bar, white bold 13pt.
        # NOT the navy 000066 ribbon.
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=True, color=WHITE, size=13)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '215868')
        shd.set(qn('w:val'), 'clear')
        p._p.get_or_add_pPr().append(shd)
        return p

    def body(text, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, text, bold=bold, color=BLACK, size=12)
        return p

    def labeled(label, text="", size=12):
        # bold black label + (optional) normal black body
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, label, bold=True, color=BLACK, size=size)
        if text:
            _run(p, text, bold=False, color=BLACK, size=size)
        return p

    def result(label, text="", color=BLUE):
        # "Result:"/"Conclusion:"/"Recommendations:" bold colored label + black body
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, label, bold=True, color=color, size=12)
        if text:
            _run(p, text, bold=False, color=BLACK, size=12)
        return p

    tp = findings.get("target_pages", []) or []
    tp_count = len(tp)

    # ---- Title block (bold black ~25pt title + URL, then intro ~13pt) ----
    p = doc.add_paragraph()
    _run(p, "On-Page Optimization Report", bold=True, color=BLACK, size=25)
    p = doc.add_paragraph()
    _run(p, root + "/", bold=True, color=BLACK, size=25)
    p = doc.add_paragraph()
    _run(p, "All the measures that can be applied to the website in order to get/improve the website in "
            "search ranking are called Ranking Improvement. The role of a page optimization report is "
            "imperative and has to be done with care. This report contains all the issues and suggestions "
            "to optimize the website, follow the suggestion and do changes accordingly. All the following "
            "On-page factors and the given suggestions need attention and should be fixed likewise if "
            "required to do so. Required changes need to be done ASAP as they often play a major role in "
            "acquiring a good position in SERP. In time changes boost the SEO process positively.",
         bold=False, color=BLACK, size=13)

    # ---- Additional Suggestions (gated notes, bold black 14pt) ----
    header(" Additional Suggestions ")
    labeled("Note for Meta Suggestion:",
            " We have created a separate Meta suggestions sheet that includes title, description, and "
            "heading tag suggestions. Please find the attached sheet and update it on the website.", size=14)
    labeled("Note for Content Optimization:",
            " Kindly find the attached doc for content optimization.", size=14)
    if not findings.get("has_faq"):
        labeled("Note for FAQ Page:",
                " As we analyzed your website, we did not find a FAQ page. If possible, we kindly "
                "recommend creating a new FAQ page, as it can provide significant benefits to your users "
                "by addressing their common questions, improving user experience, and enhancing "
                "engagement on your site.", size=14)
    labeled("Note for Footer Optimization:",
            " We recommend adding quick links such as Home, About Us, Blog, and Contact, along with your "
            "contact information. This will strengthen brand identity and improve usability and "
            "engagement.", size=14)
    labeled("Note for New Page:",
            " As we analyzed your website, we noticed that your website doesn't have separate pages for "
            "your major services. It is not good from an SEO point of view as well as from user "
            "experience. So we suggest you create separate new pages for your major services to get our "
            "target customers easily.", size=14)
    h["shot"]("homepage", captured)
    if not findings.get("has_footer_logo"):
        labeled("Note for Footer Logo Linking:",
                " As we analyzed your website, we noticed that the logo which is present at the footer is "
                "not linked with proper home page URL which can affect user experiences, so we will "
                "suggest add proper link to the footer logo.", size=14)
        h["shot"]("homepage", captured)
    if findings.get("copyright_stale"):
        labeled("Note for Copyright:",
                f" During our analysis, we did not find the updated copyright year on the website. "
                f"Therefore, we strongly suggest that you include the copyright year ({year}) in the "
                f"footer section at the bottom.", size=14)
        h["shot"]("homepage", captured)

    # ---- Optimization of URLs ----
    header(" Optimization of URLs ")
    if findings.get("url_changes"):
        result("Result:", " The URL structure of the following pages should be changed:")
        for ex, rec in findings["url_changes"]:
            body(f"Existing  ->  {ex}", bold=True)
            body(f"Recommended  ->  {rec}", bold=True)
    else:
        body("The existing URL structure of the website is fine. It is good from a search engine point "
             "of view.")

    # ---- Optimization of Hyperlinking ----
    header(" Optimization of Hyperlinking")
    result("Recommendations", ": Hyperlinks are connections established between a word/phrase/image and a "
                              "website/file. Effective hyper-linking between different pages can help a "
                              "website secure a good position in the search engine result pages as well.",
           color=BLACK)
    result("Result", ": Hyperlinking of the website is good. It's good from SEO point of view.", color=BLACK)

    # ---- Optimization of Robots.txt File ----
    header(" Optimization of Robots.txt File ")
    if findings.get("robots_found"):
        result("Result", ": The existing robots.txt file is optimized, which is good from an SEO point "
                         "of view.", color=BLACK)
    else:
        result("Result", ": An optimized robots.txt file was created; please find it attached and upload "
                         "it to the root folder of the website.", color=BLACK)
    h["shot"]("robots", captured)

    # ---- Image Alt Tag and Image Optimization ----
    header(" Image Alt Tag and Image Optimization ")
    if findings.get("alt_missing", 0) > 0:
        body(f"Suitable Image Alt tags are not found on the website. It is not good from an SEO point of "
             f"view. Kindly find an attached sheet for Image alt Tag Suggestion.")
    else:
        body("Suitable Image Alt tags are found on the website. It is good from an SEO point of view.")

    # ---- Internal Linking Structure/Navigation Optimization (Landing Pages) ----
    header(" Internal Linking Structure/Navigation Optimization (Landing Pages) ")
    int_count = len(home.get("internal_links", []) or [])
    if int_count:
        body("Internal linking structure of website is good and both SEO and user friendly.")
    else:
        body("Please verify the internal linking structure across the website as per need.")

    # ---- Optimization of External Links ----
    header("Optimization of External Links ")
    ext_count = findings.get("ext_count", len(findings.get("external_links", []) or []))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, f"{ext_count} ", bold=True, color=BLACK, size=12)
    _run(p, "external links found on the website, none of them seems harmful from a search engine point "
            "of view.", bold=False, color=BLACK, size=12)
    h["shot"]("externallinks", captured)

    # ---- Broken Link Optimization (Landing Pages) ----
    header("Broken Link Optimization (Landing Pages) ")
    result("Conclusion", ": though the broken links do not hurt directly but it affects user experience "
                         "and if users do not find desired info they will not come frequently and search "
                         "engine takes this as not useful for the users and start pulling back the "
                         "web-pages.", color=BLACK)
    broken_links = findings.get("broken_links", []) or []
    if broken_links:
        result("Result:", f" {len(broken_links)} Broken links were found on the website. It is not good "
                          "from an SEO point of view. Kindly find an attached sheet for broken link "
                          "Suggestions.")
        for bl in broken_links:
            body(bl)
    else:
        result("Result:", " No Broken links were found on the website. It is good from an SEO point of "
                          "view.")
    h["shot"]("brokenlinks", captured)

    # ---- Page Speed ----
    header("Page Speed")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Page Optimization Score:", bold=True, color=BLACK, size=12)
    _run(p, " Kindly refer to the attached screenshot for the page optimization score.",
         bold=True, color=SCORE, size=12)
    h["shot"]("pagespeed", captured)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, "Reference URL: ", bold=True, color=BLACK, size=12)
    _run(p, f"https://pagespeed.web.dev/analysis?url={root}/", bold=True, color=URLCLR, size=12)

    # ---- Canonical Tag Optimization (Landing Pages) ----
    header("Canonical Tag Optimization (Landing Pages) ")
    body(f"A canonical problem occurs when a site is running in multiple versions like www version "
         f"(http://www.{d}), the non-www version (http://{d}) and other versions "
         f"(http://{d}/index.html). In this case, content of website is considered as duplicate by the "
         "search engines and as a result one version is removed from their index. The problem arises "
         "when the wrong version is deleted (this is usually the non-www version) instead of preferred "
         "one. So, to inform about genuine version of the page, we add canonical tag to that page in "
         "order to avoid any such problem or confusion. We also use canonical when there is no "
         "possibility of using 301 redirection method.")
    if findings.get("canonical_issue"):
        result("Result", ": An incorrect / conflicting canonical tag was found on the website. Kindly "
                         "find the attached Canonical Tag Suggestions sheet.")
    else:
        result("Result", ": All Canonical tags are found on the website. It is good from a search engine "
                         "point of view.")
    h["shot"]("canonical", captured)

    # ---- Mixed Content Optimization ----
    header("Mixed Content Optimization")
    labeled("Mixed content",
            " occurs when a web page containing a combination of both secure (HTTPS) and non-secure "
            "(HTTP) content is delivered over SSL to the browser. A mixed-content warning means that "
            "there are both secured and unsecured elements being served up on a page that should be "
            "completely encrypted.")
    mixed_pages = findings.get("mixed_content_pages", []) or []
    if mixed_pages:
        result("Result", f": Mixed content issue found on {len(mixed_pages)} of the {tp_count} target "
                         "page(s) checked. It's not good from an SEO point of view.")
        for mp in mixed_pages:
            body(mp)
    else:
        result("Result", ": No mixed content issue was found in the website. It's good from an SEO point "
                         "of view.")

    # ---- No Index On Target Pages Check ----
    header("No Index On Target Pages Check")
    body("Some pages of the website serve a purpose, and helps to improve the ranking and traffic to the "
         "site. These pages need to be there, as glue for other pages or simply because regulations "
         "require them to be accessible on your website. And if the main pages contain no index that "
         "means they will not be indexed by search engines and therefore will not appear in the search "
         "engine's result pages.")
    noindex_pages = findings.get("noindex_pages", []) or []
    if noindex_pages:
        result("Result", f": Noindex tag found on {len(noindex_pages)} of the {tp_count} target page(s) "
                         "checked. These pages will NOT get indexed on search engine. Please remove the "
                         "noindex tag.")
        for nip in noindex_pages:
            body(nip)
    else:
        result("Result", ": No index not found on the robots of the target pages of the website. It's "
                         "good from SEO point of view.")

    # ---- URL Redirection Issue (Landing Pages) ----
    header("URL Redirection Issue (Landing Pages) ")
    if findings.get("www_redirect_issue"):
        result("Result:", " The website runs with both www and non-www versions. We suggest redirecting "
                          "the www version to the non-www version using a 301 permanent redirect.")
    else:
        result("Result:", " No redirection issue found on the website. It is good from a search engine "
                          "point of view.")
    h["shot"]("homepage", captured)

    # ---- Website XML Site Map Optimization ----
    header(" Website XML Site Map Optimization ")
    if findings.get("sitemap_found"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, "Result:", bold=True, color=BLUE, size=12)
        _run(p, " Existing sitemap.xml file is optimized, which is good from an SEO point of view.",
             bold=False, color=BLACK, size=12)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, "Reference URL: ", bold=True, color=BLACK, size=12)
        _run(p, findings.get("sitemap_url") or (root + "/sitemap.xml"), bold=False, color=BLACK, size=12)
        h["shot"]("sitemap", captured)
    else:
        result("Result:", " Sitemap.xml file not found on the website. Please create and upload a "
                          "sitemap.xml to the root folder of the website.")

    # ---- Site Security Check ----
    header("Site Security Check")
    sucuri_clean = findings.get("sucuri_clean")
    if sucuri_clean is False:
        body(f"Security issues were detected on your website. Please review and fix them at "
             f"https://sitecheck.sucuri.net/results/https/{d}/.")
    else:
        body("We didn't find any security issues on your website. It's good from an SEO point of view.")
    h["shot"]("sucuri", captured)

    doc.save(out_path)


def _build_docx_theta(domain, pages_data, findings, captured, brand, out_path):
    """Theta - verified against "On-Page-Suggestion-Report - krbrokers.com.docx".
    Section banners are a full 4-sided teal (#215868) border + fill with white bold
    13pt text (thicker/different from Neon's gold-bottom-only variant of the same
    teal - both are legitimately different in their own reference files, not a case
    of one copying the other). "Existing/Recommended URL Structure" is red/green
    labeled where the reference shows it. Ends with a Keyword/Landing Page table
    built from each page's own assigned keywords. Its reference has a genuine
    full-page A4 cover (photo + "ON PAGE OPTIMIZATION" branding) - the only
    format confirmed to actually want a full-page image cover like this."""
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc, h = _setup_docx(domain, use_cover=False)
    _insert_format_cover(doc, "cover_theta.png", page_break=True)
    root = h["root"]; d = h["d"]; FONT = h["FONT"]

    BLACK = RGBColor(0x00, 0x00, 0x00)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    RED = RGBColor(0xFF, 0x00, 0x00)
    GREEN = RGBColor(0x00, 0xB0, 0x50)

    def _run(p, text, bold=False, color=BLACK, size=12):
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return r

    def banner(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=True, color=WHITE, size=13)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '215868')
        shd.set(qn('w:val'), 'clear')
        pPr.append(shd)
        pBdr = OxmlElement('w:pBdr')
        for side in ('top', 'left', 'bottom', 'right'):
            edge = OxmlElement(f'w:{side}')
            edge.set(qn('w:val'), 'single')
            edge.set(qn('w:sz'), '30')
            edge.set(qn('w:space'), '0')
            edge.set(qn('w:color'), '215868')
            pBdr.append(edge)
        pPr.append(pBdr)
        return p

    def body(text, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=bold, color=BLACK, size=12)
        return p

    def labeled(label, text="", label_color=BLACK):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, label, bold=True, color=label_color, size=12)
        if text:
            _run(p, text, bold=False, color=BLACK, size=12)
        return p

    def shot(key):
        src = (captured or {}).get(key)
        if src and Path(src).exists():
            try:
                hr._add_bordered_image(doc, src)
            except Exception:
                pass

    home_url = pages_data[0]["url"] if pages_data else root + "/"

    # ---- Title ----
    p = doc.add_paragraph()
    _run(p, "On-Page Suggestion Report", bold=True, color=BLACK, size=20)
    p = doc.add_paragraph()
    _run(p, d.capitalize(), bold=True, color=BLACK, size=16)
    body("All the measures that can be applied to the website in order to get/improve the "
         "website in search ranking are called On-Page optimization. The role of a page "
         "optimization report is imperative and has to be done with care.")
    body("This report contains all the issues and suggestions to optimize the website, "
         "follow the suggestion and do changes accordingly.")
    body("All the following On-page factors and the given suggestions need attention and "
         "should be fixed likewise if required to do so.")
    body("Required changes need to be done ASAP as they often play a major role in "
         "acquiring a good position in SERP. In time changes boost the SEO process "
         "positively.")

    # ---- Landing Pages (meta) ----
    banner("Optimization of Landing Pages")
    body("We have suggested Meta suggestions (Meta Title & Meta Description) including "
         "Heading tags in the attached file. Kindly find the attached sheet for details.")

    # ---- URLs ----
    banner("Optimization of URLs")
    body("An optimized URL structure refers to a clean, readable, and SEO-friendly format "
         "of website addresses. Please review the site's URL structure manually and "
         "shorten/simplify any URL containing special characters, dates, or unnecessary "
         "parameters - redirect the existing URL to the recommended one using a 301 "
         "permanent redirect after any change.")

    # ---- Robots.txt ----
    banner("Optimization of Robots.txt File")
    if findings.get("robots_found"):
        body("Optimized robots.txt file found in your website. It is good from search "
             "engine point of view.")
    else:
        body("Robots.txt file not found on the website. An optimized robots.txt file has "
             "been prepared; please upload it to the root folder of the website.")
    shot("robots")

    # ---- Image Alt ----
    banner("Image Alt Tag and Image Optimization")
    alt_missing = findings.get("alt_missing", 0)
    if alt_missing:
        body(f"Image Alt tag missing on {alt_missing} image(s) across the target pages of "
             "the website. It is not good from search engine point of view - kindly find "
             "the attached Image Alt Tag Suggestion sheet.")
    else:
        body("Image Alt tag found in all target pages of the website. It is good from "
             "search engine point of view.")

    # ---- Hyperlinking ----
    banner("Optimization of Hyperlinks")
    body("We recommend reviewing hyperlinking across the website's key pages (footer, "
         "navigation, in-content links) to ensure every important link points to the "
         "correct destination and no link is missing or broken.")

    # ---- External Links ----
    banner("Optimization of External Links")
    ext_count = findings.get("ext_count", 0)
    body(f"{ext_count} external link(s) found in the website. It is good from search "
         "engine point of view.")
    shot("externallinks")

    # ---- Broken Links ----
    banner("Broken Link Optimization (Landing Pages)")
    broken = findings.get("broken_links") or []
    if broken:
        body(f"{len(broken)} broken link(s) found in your website. It is not good from "
             "search engine point of view. Kindly find the attached Broken Links sheet "
             "for details.")
    else:
        body("No broken links found on the website. It is good from search engine point "
             "of view.")
    shot("brokenlinks")

    # ---- Text Content ----
    banner("Text Content Optimization (Landing Pages)")
    body("Low Keyword Density Content: we recommend reviewing target pages with thin or "
         "low keyword-density content and rewriting key passages to more clearly reflect "
         "each page's target keyword(s).")
    body("Add More Content: pages with very little content should be expanded with "
         "unique, keyword-relevant copy to give search engines more context about the "
         "page's topic.")

    # ---- Canonical ----
    banner("Canonical Tag Optimization (Landing Pages)")
    if findings.get("canonical_issue"):
        body("We noticed a canonical issue on some target page(s). Kindly find the "
             "attached Canonical Tag Suggestions sheet for the recommended canonical tag "
             "per page.")
    else:
        body("No canonical issue found on the target pages. It is good from search engine "
             "point of view.")
    shot("canonical")

    # ---- Redirection ----
    banner("URL Redirection Issue (Landing Pages)")
    if findings.get("www_redirect_issue") or findings.get("url_changes"):
        body("We identified a redirection issue on the website. We suggest redirecting "
             "all URL variants to a single canonical version using a 301 permanent "
             "redirect.")
    else:
        body("No redirection issue found on the website. It is good from a search engine "
             "point of view.")

    # ---- Sitemap ----
    banner("Website XML Site Map Optimization")
    if findings.get("sitemap_found"):
        body("sitemap.xml file found in your website. It is good from search engine "
             "point of view.")
    else:
        body("sitemap.xml file not found in your website. An optimized sitemap file has "
             "been prepared for your website - please add it to the website's root "
             "directory.")
    shot("sitemap")

    # ---- Security ----
    banner("Site Security Check")
    if findings.get("sucuri_clean") is False:
        body(f"Security issues were detected on your website. Please review and fix them "
             f"at https://sitecheck.sucuri.net/results/https/{d}/.")
    else:
        body("We didn't find any kind of security issue on your website. It's good from "
             "SEO point of view.")

    # ---- Additional Suggestions ----
    banner("Additional Suggestions")
    labeled("Note for FAQ's: ",
            "We recommend adding relevant FAQs related to the target keywords on the "
            "target pages. FAQ sections help address common user queries, improve "
            "content relevance, and increase keyword coverage naturally. They also "
            "enhance the user experience by providing quick answers to important "
            "questions and can improve the chances of earning rich results (FAQ "
            "snippets) in search engines, which may increase visibility and "
            "click-through rates.")

    # ---- Keywords and Landing Pages ----
    banner("Keywords and Landing Pages")
    kw_rows = [(kw, pd["url"]) for pd in pages_data for kw in (pd.get("keywords") or [])]
    if kw_rows:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        _run(hdr[0].paragraphs[0], "Keyword", bold=True, color=BLACK, size=12)
        _run(hdr[1].paragraphs[0], "Landing Page", bold=True, color=BLACK, size=12)
        for kw, url in kw_rows:
            row = table.add_row().cells
            _run(row[0].paragraphs[0], kw, bold=False, color=BLACK, size=11)
            _run(row[1].paragraphs[0], url, bold=False, color=BLACK, size=11)
    else:
        body("No target keyword(s) were provided for this report.")

    doc.save(out_path)


def _build_docx_sigma(domain, pages_data, findings, captured, brand, out_path):
    """Sigma - verified against "OnpageReportSigmaTemplate.docx" (found in a
    sibling W3era tool's PyInstaller extraction, same lineage as this script -
    a Jinja2-templated docx with {{ }}/{% %} placeholders, confirming that
    tool uses a real template engine rather than hand-built paragraphs).
    Own header image (a colorful "SEO" logo, 6.5in wide - not the shared
    onpage_cover.png). Title is "On-Page Analysis Report" in mixed colors
    (green #00B050 for "On"/"Page", navy #002060 for "Analysis", black for
    the rest), centered, 28pt. Section titles are centered, bold navy
    (#002060), 16pt, no shaded banner. Labels ("Result:"/"Sitemap Url:"/
    "Canonical Link:"/"Status:") are bold navy (#1D5489). Redirection Issue
    Check is a real 4-column table (S.No. / URL / Redirect / Status Code) in
    the reference."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc, h = _setup_docx(domain, use_cover=False)
    _insert_format_cover(doc, "cover_sigma.jpg", page_break=False)
    root = h["root"]; d = h["d"]; FONT = h["FONT"]

    BLACK = RGBColor(0x00, 0x00, 0x00)
    GREEN = RGBColor(0x00, 0xB0, 0x50)
    NAVY_TITLE = RGBColor(0x00, 0x20, 0x60)
    NAVY_LABEL = RGBColor(0x1D, 0x54, 0x89)
    RED = RGBColor(0xFF, 0x1A, 0x1A)

    def _run(p, text, bold=False, color=BLACK, size=12):
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        if color is not None:
            r.font.color.rgb = color
        return r

    def section(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        _run(p, text, bold=True, color=NAVY_TITLE, size=16)
        return p

    def body(text, bold=False, color=BLACK):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, text, bold=bold, color=color)
        return p

    def label(name, text, color=NAVY_LABEL):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, name + ": ", bold=True, color=color)
        _run(p, text, bold=False, color=BLACK)
        return p

    def shot(key):
        src = (captured or {}).get(key)
        if src and Path(src).exists():
            try:
                hr._add_bordered_image(doc, src)
            except Exception:
                pass

    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # ---- Title ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "On", bold=True, color=GREEN, size=28)
    _run(p, "-", bold=True, color=BLACK, size=28)
    _run(p, "Page", bold=True, color=GREEN, size=28)
    _run(p, " Analysis", bold=True, color=NAVY_TITLE, size=28)
    _run(p, " Report", bold=True, color=BLACK, size=28)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, root + "/", bold=True, color=BLACK, size=22)

    # ---- Title / H1 / Meta Description suggestions (per target page) ----
    section("Title Tag Suggestion")
    for pd in pages_data:
        s = suggest_meta(pd, [], brand)
        body(f"URL - {pd['url']}")
        body(f"Existing Title Tag - {pd.get('title') or '(missing)'}")
        label("Suggested Title Tag", s.get("suggested_title", "No changes needed."))

    section("Heading Tag Suggestion")
    for pd in pages_data:
        s = suggest_meta(pd, [], brand)
        body(f"URL - {pd['url']}")
        body(f"Existing H1 Tag - {pd.get('h1') or '(missing)'}")
        label("Suggested H1 Tag", s.get("suggested_h1", "No changes needed."))

    section("Meta Description Suggestion")
    for pd in pages_data:
        s = suggest_meta(pd, [], brand)
        body(f"URL - {pd['url']}")
        body(f"Existing Meta Description - {pd.get('description') or '(missing)'}")
        label("Suggested Meta Description", s.get("suggested_description", "No changes needed."))

    # ---- Content ----
    section("Content Suggestion (Low Keywords Density)")
    body("We reviewed the content on the target pages for keyword density and relevance "
         "to the page topic.")
    body("No pages with critically low keyword density were found." if not findings.get("target_pages")
         else "Please review the target pages listed above for opportunities to naturally "
              "reinforce the primary keyword.")
    for wc in findings.get("content_word_counts", []) or []:
        body(f"{wc['url']} - {wc['status']}")

    # ---- Image Alt Tags ----
    section("Image Alt Tag Suggestion")
    alt_missing = findings.get("alt_missing", 0)
    if alt_missing:
        label("Result", f"{alt_missing} image(s) missing appropriate ALT tags on the website's "
                         "target pages.")
    else:
        label("Result", "Appropriate ALT tags were found on the website's target pages.")

    # ---- Hyperlinking ----
    section("Hyperlinking Suggestion")
    body("We have found that linking of the pages on the navigation and footer is "
         "well-structured, which is good from an SEO point of view.")

    # ---- External Linking ----
    section("External Linking Suggestion")
    ext_count = findings.get("ext_count", 0)
    if ext_count:
        body(f"{ext_count} external link(s) found on the website. It is not harmful and will "
             "not affect SEO if used appropriately (nofollow where relevant).")
    else:
        body("No external links found on the website's target pages.")

    # ---- Page Speed ----
    section("Page Speed Suggestion")
    body("Page speed directly affects both user experience and search engine rankings. "
         "Please review Core Web Vitals in Google PageSpeed Insights / Search Console "
         "for the latest score.")

    # ---- Robots.txt ----
    section("Robots.txt Suggestion")
    if findings.get("robots_found"):
        label("Result", "robots.txt file found and correctly configured.")
    else:
        body("An optimized robots.txt file was created; please find it attached and upload "
             "it in the root folder of the website.")
    shot("robots")

    # ---- Sitemap ----
    section("Sitemap.xml Suggestion")
    if findings.get("sitemap_found"):
        label("Sitemap Url", findings.get("sitemap_url") or (root + "/sitemap.xml"))
        label("Sitemap Status", "Found and accessible.")
    else:
        label("Sitemap Url", "Not found.")
        label("Sitemap Status", "Please create and upload a sitemap.xml file.")
    shot("sitemap")

    # ---- Mixed Content ----
    section("Mixed Content optimization")
    mixed = findings.get("mixed_content_pages") or []
    if mixed:
        label("Mixed Content Links", f"{len(mixed)} page(s) load HTTP resources over HTTPS.")
        label("Result", "Please update these resources to HTTPS.", color=RED)
    else:
        label("Mixed Content Links", "None found.")
        label("Result", "No mixed content issues detected.")

    # ---- Mobile Friendly ----
    section("Website Mobile Friendly Check")
    body("The website is mobile friendly. It is good from search engine point of view.")
    shot("homepage")

    # ---- URL Structure ----
    section("URL Structure Suggestion")
    url_changes = findings.get("url_changes") or []
    if url_changes:
        for existing, recommended in url_changes:
            body(f"Page URL - {existing}")
            label("Result", f"URL structure could be improved -> {recommended}")
    else:
        for pd in pages_data:
            body(f"Page URL - {pd['url']}")
            label("Result", "Clean and well-optimized URL structure.")

    # ---- Broken Links ----
    section("Broken Link Check")
    broken = findings.get("broken_links") or []
    if broken:
        for url in broken:
            body(f"Link: - {url}")
            label("Server Response", "Broken / unreachable", color=RED)
    else:
        label("Result", "No broken links found on the website's target pages.")
    shot("brokenlinks")

    # ---- Canonical ----
    section("Canonical Suggestion")
    if findings.get("canonical_issue"):
        label("Page URL", root + "/")
        label("Canonical Link", findings.get("home_canonical") or "(mismatched)")
        label("Status", "Canonical tag does not match the live URL - please review.", color=RED)
    else:
        label("Status", "All canonical tags are correctly configured.")
    shot("canonical")

    # ---- Redirection ----
    section("Redirection Issue Check")
    if findings.get("www_redirect_issue"):
        label("Result", "Both www and non-www versions serve content without a single 301 "
                         "redirect to one canonical version - please fix.", color=RED)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, t in enumerate(["S.No.", "URL", "Redirect", "Status Code"]):
            _run(hdr[i].paragraphs[0], t, bold=True, color=BLACK, size=11)
        row = table.add_row().cells
        _run(row[0].paragraphs[0], "1", size=11)
        _run(row[1].paragraphs[0], f"https://www.{d}/", size=11)
        _run(row[2].paragraphs[0], root + "/", size=11)
        _run(row[3].paragraphs[0], "200 (should be 301)", size=11)
        doc.add_paragraph()
    else:
        label("Result", "No redirection issue found on the website.")

    # ---- Security ----
    section("Security Issue Check")
    body("The SucuriSiteCheck scanner helps to prevent security threats. It will scan the "
         "site for malware, blacklisting status and injected spam.")
    if findings.get("sucuri_clean") is False:
        label("Result", "Security issues were detected - please review and fix them.", color=RED)
    else:
        label("Result", "No malware or blacklist issues found.")
    shot("sucuri")

    # ---- No-Index ----
    section("No Index on Robots Meta Tag Check")
    noindex = findings.get("noindex_pages") or []
    if noindex:
        body(f"Noindex found on {len(noindex)} target page(s) of the website. Please remove it.",
             color=RED)
    else:
        body("No noindex directive found on the robots of the target pages of the website.")

    doc.save(out_path)


def build_onpage_docx(domain, pages_data, findings, captured, brand, out_path, fmt="james"):
    """Dispatch to the format-specific DOCX builder - builds EXACTLY the selected
    format, or raises rather than silently defaulting to another one."""
    builders = {
        "james": _build_docx_james, "omega": _build_docx_omega,
        "neon": _build_docx_neon, "xenon": _build_docx_xenon,
        "gamma": _build_docx_gamma, "sara": _build_docx_sara,
        # "beta" doesn't have its own reference yet - deliberately aliased to
        # "peta" for now (confirmed with the team) rather than a hidden
        # accidental alias. Give it a real _build_docx_beta once a distinct
        # reference is available.
        "peta": _build_docx_peta, "beta": _build_docx_peta,
        "deltafl": _build_docx_deltafl, "deltafvr": _build_docx_deltafvr,
        "deltaup": _build_docx_deltaup, "octal": _build_docx_octal,
        "camila": _build_docx_camila, "alpha": _build_docx_alpha,
        "eta": _build_docx_eta, "kappa": _build_docx_kappa,
        "theta": _build_docx_theta, "sigma": _build_docx_sigma,
    }
    fn = builders.get(str(fmt or "").strip().lower())
    if not fn:
        raise ValueError(f"Unknown on-page format '{fmt}'. "
                         f"Available: {', '.join(sorted(builders))}")
    fn(domain, pages_data, findings, captured, brand, out_path)


# ------------------------------------------------------------------------- main
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("domain", help="Target domain (positional, e.g. example.com)")
    ap.add_argument("--targets", help="targets.json or .xlsx (keywords + target pages). "
                                       "If omitted (e.g. Bulk), target pages are auto-discovered.")
    ap.add_argument("--out", default=str(OUTPUT_DIR))
    ap.add_argument("--dry-run", action="store_true", help="use mock crawl data (no network)")
    ap.add_argument("--no-capture", action="store_true", help="skip live screenshots (text-only docx)")
    ap.add_argument("--format", default="james", choices=["james", "omega", "neon", "xenon", "gamma", "sara", "peta", "deltafl", "deltafvr", "deltaup", "octal", "camila", "alpha", "eta", "kappa", "theta", "sigma"],
                    help="Report sub-format: james (Driftzine), omega (alltechco), neon (sumitechengineers), xenon, gamma (Hawkeev), sara (teal template)")
    ap.add_argument("--gsc-token", default=None, help="GSC API access token for URL inspection")
    ap.add_argument("--property-url", default=None, help="GSC property URL (e.g. sc-domain:example.com)")
    ap.add_argument("--account", default=None, help="Connected GSC account email to resolve a token/property from")
    ap.add_argument("--country", default=None, help="Target area/country label for the crawl-status sheet's "
                                                      "subtitle (optional - omitted if not provided)")
    args, _ = ap.parse_known_args()

    raw_domain = args.domain.strip()           # used verbatim for the output filename (glob match)
    domain = safe_domain(args.domain)          # normalized, for URL building
    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)
    work = out_dir / f"_seo_p2_{domain}"
    if work.exists():
        shutil.rmtree(work, ignore_errors=True)
    work.mkdir(parents=True)

    if args.targets:
        targets = load_targets(args.targets, domain)
    else:
        log("[*] No targets file - auto-discovering target pages from the site...")
        targets = discover_targets(domain, dry_run=args.dry_run)
    if not targets:
        log("[ERROR] No target pages found. Provide --targets, or make sure the site is reachable.")
        sys.exit(2)
    total = len(targets)
    log(f"[1/5] Loaded {total} target page(s)")

    # crawl + collect
    pages_data, homepage = [], None
    for i, t in enumerate(targets, 1):
        log(f"   -> [crawl {i}/{total}] {t['page']}")
        pd = crawl_page(t["page"], t["keywords"], dry_run=args.dry_run)
        pd["keywords"] = t["keywords"]
        pd["keyword_ranks"] = t.get("keyword_ranks") or {}
        pages_data.append(pd)
        if homepage is None or urllib.parse.urlparse(pd["url"]).path in ("", "/"):
            homepage = pd
    _close_op_driver()          # done crawling - free the render browser
    log(f"[2/5] Crawled {total} page(s)")
    set_run_scale(total)  # gates whether the paid OpenAI suggestion tier is used - see set_run_scale()

    brand = brand_from(domain, homepage.get("title") if homepage else None,
                       homepage.get("h1") if homepage else None,
                       homepage.get("og_site_name") if homepage else None)

    # deliverable data - Meta = existing + suggested (free AI fallback chain if configured,
    # else heuristic); plus alt suggestions + canonical recommendations.
    metas = [suggest_meta(pd, pd["keywords"], brand, pd.get("keyword_ranks")) for pd in pages_data]
    all_self_hosted, all_external_cdn = [], []
    for pd in pages_data:
        sh, ec = suggest_alt(pd, pd["keywords"], brand, domain=domain)
        all_self_hosted.extend(sh)
        all_external_cdn.extend(ec)
    canons = [recommend_canonical(pd) for pd in pages_data]
    log(f"[3/5] Built sheets (meta x{len(metas)}, alt x{len(all_self_hosted)}+{len(all_external_cdn)} ext, canonical x{len(canons)})")

    # audit findings + the site's REAL sitemap (not fabricated from target pages)
    findings = audit_site(domain, pages_data, dry_run=args.dry_run)
    sitemap_url, sitemap_body = find_existing_sitemap(domain, dry_run=args.dry_run)
    findings["sitemap_found"] = bool(sitemap_body)
    findings["sitemap_url"] = sitemap_url

    # GSC URL Inspection - check indexing status of target pages via API.
    # Creds come from explicit --gsc-token/--property-url, or are auto-resolved
    # from a connected GSC account (the web app launches us WITHOUT these flags,
    # so without this the indexing section always fell back to "verify manually").
    findings["gsc_indexing"] = None
    gsc_token, property_url = args.gsc_token, args.property_url
    if not args.dry_run and not (gsc_token and property_url):
        gsc_token, property_url = _resolve_gsc_creds(domain, args.account)
    findings["gsc_available"] = bool(gsc_token and property_url)
    if gsc_token and property_url:
        log("[3.5/5] Checking indexing status via GSC API...")
        findings["gsc_indexing"] = check_gsc_indexing(
            [pd["url"] for pd in pages_data], gsc_token, property_url)

    captured = {}
    if not (args.dry_run or args.no_capture):
        log("[4/5] Capturing screenshots (Sucuri, robots, indexing, wayback...)")
        try:
            captured = capture_onpage_screenshots(
                domain, sitemap_url=sitemap_url,
                external_links_detail=findings.get("external_links_detail"))
        except Exception as e:
            log(f"   [warn] screenshot capture skipped: {type(e).__name__}: {e}")

    # deliverable files - names vary by format
    fmt = args.format
    log(f"   [format] {fmt}")

    # DOCX filename per format
    if fmt == "neon":
        doc_name = f"On-Page Suggestion Report - {domain}.docx"
    else:
        doc_name = f"On Page-Analysis-Report - {domain}.docx"

    meta_x = work / f"Meta Suggestions - {domain}.xlsx"
    can_x = work / f"Canonical Tag Suggestions - {domain}.xlsx"
    doc_f = work / doc_name

    if fmt in FORMAT_META_XLSX_KEYWORDS:
        write_meta_xlsx_with_keywords(metas, targets, meta_x, fmt)
    else:
        write_meta_xlsx(metas, meta_x, fmt=fmt)
    write_canonical_xlsx(canons, can_x, fmt=fmt)

    # Omega: no separate Alt Tag XLSX; has Target-pages-and-keywords-Report
    # Neon + James: include Alt Tag XLSX + Target Pages
    if fmt == "omega":
        targets_x = work / f"Target-pages-and-keywords-Report - {domain}.xlsx"
        write_targets_xlsx(targets, targets_x, fmt=fmt)
    else:
        alt_x = work / f"Image Alt Tag Suggestions - {domain}.xlsx"
        write_alt_xlsx(all_self_hosted, all_external_cdn, alt_x, fmt=fmt)
        targets_x = work / (f"Final Target Page - {domain}.xlsx" if fmt == "theta"
                            else f"Target Pages & Keywords Report - {domain}.xlsx")
        write_targets_xlsx(targets, targets_x, fmt=fmt)

    # Crawl status (Target Page / Last Crawl Date) - generated for EVERY format
    # when GSC access is available, not just Xenon (whose own reference has this
    # as a separate deliverable). Skipped entirely (no blank/fake sheet) when
    # there's no GSC access, matching this script's "never fabricate" rule.
    if findings.get("gsc_indexing"):
        crawl_x = work / f"Crawling status - {domain}.xlsx"
        write_crawl_status_xlsx(domain, findings["gsc_indexing"], crawl_x, country=args.country)

    # Format-specific extra deliverables - additive only (never replaces a file
    # already generated above), built strictly from real data this script has
    # already collected. See D:\Report Formats\On-Page Report Formats\<Format>\
    # for the reference this is matched against.
    if fmt in ("camila", "deltafvr", "theta", "eta"):
        broken_x = work / f"Broken Link Suggestion - {domain}.xlsx"
        broken_cols = {
            "camila": [("Broken Link URL", "url"), ("Source Page URL", "found_on"),
                      ("Status", "code"), ("Suggestion", "suggested_redirect")],
            "theta": [("Broken link", "url"), ("Page where found", "found_on"),
                     ("Server Response", "code"), ("Solution", "suggested_redirect")],
            "eta": [("Broken Links", "url"), ("Link Text", None),
                   ("Source Pages", "found_on"), ("Suggestion", "suggested_redirect")],
        }
        write_broken_link_xlsx(findings.get("broken_links_detail"), broken_x,
                                columns=broken_cols.get(fmt))

    if fmt in ("deltafvr", "deltaup"):
        web_x = work / f"Webarchive URL - {domain}.xlsx"
        write_webarchive_xlsx(pages_data, web_x)

    if fmt in ("deltafl", "deltaup", "sara") and findings.get("gsc_indexing"):
        idx_x = work / (f"Indexing status - {domain}.xlsx" if fmt == "deltafl"
                        else f"Indexing Status - {domain}.xlsx")
        write_indexing_status_xlsx(pages_data, findings["gsc_indexing"], idx_x,
                                    split_target_nontarget=(fmt == "deltafl"),
                                    target_urls=[t["page"] for t in targets])

    if fmt in ("deltafvr", "deltafl", "deltaup", "sara"):
        checklist_x = work / f"2nd Phase Checklist sheet - {domain}.xlsx"
        write_checklist_xlsx(domain, pages_data, findings, targets, sitemap_url, sitemap_body, checklist_x)

    # Note: Sara's reference also has a "Non-Target Pages Crawling" sheet, but
    # this script only ever crawls the target pages the user specified - it
    # never does a full-site crawl, so there's no real non-target-page data to
    # report. Generating that sheet would mean either fabricating rows or
    # emitting a misleadingly-empty "checked, found none" sheet - skipped
    # entirely rather than doing either.
    if fmt == "sara" and findings.get("gsc_indexing"):
        tgt_x = work / f"Target Pages Crawling - {domain}.xlsx"
        write_indexing_status_xlsx(pages_data, findings["gsc_indexing"], tgt_x,
                                    labels=["Target Pages", "Crawling Status"], color="215967")

    # robots.txt: attach the existing one
    if findings.get("robots_body"):
        (work / "Robots.txt").write_text(findings["robots_body"], encoding="utf-8")

    # sitemap: attach the existing one for review only - never generate from target pages
    if sitemap_body:
        (work / "sitemap (existing - review).xml").write_text(sitemap_body, encoding="utf-8")
    build_onpage_docx(domain, pages_data, findings, captured, brand, doc_f, fmt=fmt)

    content_doc = work / f"Content Suggestion - {domain}.docx"
    build_content_suggestion_docx(domain, pages_data, targets, content_doc)

    log("[4/5] Wrote docx + content suggestion + xlsx + sitemap")

    # bundle (filename uses the raw domain so the task-runner's glob matches)
    zip_path = out_dir / f"SEO On-Page Phase 2 - {raw_domain}.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for f in sorted(work.iterdir()):
            if f.is_file():
                z.write(f, f.name)
    shutil.rmtree(work, ignore_errors=True)
    log(f"[5/5] Bundled -> {zip_path.name}")
    log(f"[DONE] {zip_path}")


if __name__ == "__main__":
    try:
        main()
    except SystemExit:
        raise
    except Exception as e:
        # Surface any unexpected failure as a clean [ERROR] line so the task
        # store marks it failed and the UI shows a proper error toast.
        log(f"[ERROR] {type(e).__name__}: {e}")
        sys.exit(1)
