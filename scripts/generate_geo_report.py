"""
GEO / AI Optimization report generator - the "apart from performance" core of the
monthly GEO deliverable set (see the real reference: onlinesafetytraining.ca month 7).

Given a domain + target pages, this:
  1. Crawls each target page's real content (reusing generate_seo_onpage_phase2's crawler)
  2. Generates AI-searchable keywords + long-tail/NLP FAQ queries + answers, grounded in
     that page's actual content (via Gemini, same free-tier pattern as suggest_meta())
  3. For each generated query, checks - WITHOUT login, WITHOUT any paid API - whether the
     site is already being cited in real AI-answer results for that query:
       - Perplexity's public search (no login required) - reliable, real answer text +
         cited source domains extracted from the live page.
       - Google's AI Overview - best-effort only. Headless automated requests get a
         captcha/"unusual traffic" wall from Google almost every time (confirmed live,
         same limitation the existing SERP screenshot in generate_seo_onpage_phase2.py
         already has to skip past) - so this is attempted but not relied on.
       - ChatGPT (chatgpt.com) - best-effort only. Its real chat interface requires a
         logged-in account to get an actual answer (confirmed live) - anonymous access
         is not reliably available, so this is attempted and skipped gracefully rather
         than faked.
  4. Writes a "Keywords and Content (FAQs) Work for GEO" xlsx matching the real
     reference's columns, plus extra columns showing which terms the site is ALREADY
     being cited for right now (an "opportunity" vs "already appearing" signal).

This intentionally does NOT build the Performance Report (rank tracking / Ahrefs /
GA4 screenshots) - that is separate, bigger scope (Ahrefs + GSC + rank-tracker
integration) to be scoped once this core piece is tested.
"""
import argparse
import datetime
import json
import os
import re
import sys
import time
from pathlib import Path

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT))
import generate_seo_onpage_phase2 as op  # reuse crawler, AI helper, browser driver

OUTPUT_DIR = ROOT.parent / "output"


def log(msg):
    print(msg, flush=True)


# --------------------------------------------------------------------------- #
# Step 1: content-grounded AI-searchable keyword + FAQ generation
# --------------------------------------------------------------------------- #
def generate_ai_faqs(page_data, brand, n=5, seed_keywords=None):
    """AI-searchable keyword + long-tail/NLP FAQ query + answer, per page - grounded
    in that page's own crawled content (title/H1/body text), matching the real
    reference xlsx's columns exactly: AI-Searchable Keyword | Long-Tail/NLP Query
    (FAQ) | Answers. Falls back to a simple heuristic (from title/H1) if no Gemini
    key is configured or the call fails - same fallback philosophy as suggest_meta().

    seed_keywords: optional list of keywords the team already targets for this
    site (given up front instead of only auto-discovered from content) - when
    given, generation is steered to build FAQs around THOSE terms first, so the
    team doesn't have to separately go find AI-searchable phrasing themselves."""
    title = page_data.get("title") or ""
    h1 = page_data.get("h1") or ""
    body = (page_data.get("body_text") or "")[:2500]
    url = page_data["url"]

    result = None
    if body or title or h1:
        seed_line = ""
        if seed_keywords:
            seed_line = (
                "\nThe team already targets these existing SEO keywords for this site - "
                "ground the entries around these terms where relevant to this page, rather "
                f"than only inventing new topics: {', '.join(seed_keywords)}\n"
            )
        prompt = (
            "You are an AI-search (GEO/answer-engine) optimization specialist writing "
            f"content for {brand}'s page at {url}.\n"
            f"Page title: {title}\nPage H1: {h1}\n"
            f"Page content excerpt: {body or '(not available)'}\n"
            f"{seed_line}\n"
            f"Generate exactly {n} distinct entries. Each entry needs:\n"
            '- "keyword": a short AI-searchable keyword phrase someone would type into '
            "ChatGPT/Perplexity/Google AI Overview to find this page's topic\n"
            '- "query": a natural-language, long-tail question phrased the way a real '
            "person would ask an AI assistant (not a bare keyword)\n"
            '- "answer": a 2-4 sentence answer to that question, written from THIS '
            "page's actual content (not generic filler), naturally mentioning the "
            f"brand ({brand}) once\n\n"
            "IMPORTANT: write the keyword, query, and answer in the SAME language as "
            "the page title/H1/content excerpt above - never translate to English if "
            "the page itself isn't in English.\n\n"
            'Return ONLY a JSON array: [{"keyword": "...", "query": "...", "answer": "..."}, ...]'
        )
        result = op._ai_suggest(prompt)

    if isinstance(result, list) and result:
        out = []
        for item in result[:n]:
            if isinstance(item, dict) and item.get("query") and item.get("answer"):
                out.append({
                    "keyword": str(item.get("keyword", "")).strip(),
                    "query": str(item["query"]).strip(),
                    "answer": str(item["answer"]).strip(),
                })
        if out:
            return out

    # Heuristic fallback - no AI key configured, or every tier of the free-tier
    # chain (Gemini/Groq/OpenRouter) failed/returned nothing usable. This is a
    # single generic entry, not `n` real ones - log it loudly so a run that
    # silently degraded to this (confirmed real case: kidscover.eu got only 1
    # templated FAQ per page instead of the requested 5) is obvious from the
    # log instead of looking like a legitimate low-FAQ result.
    log(f"   [!] No working AI key found for {url} - add a Gemini/Groq/OpenRouter API "
        f"key in Admin -> Sync API Keys, then re-run this report. Using 1 generic "
        f"placeholder FAQ for now instead of {n} real ones.")
    topic = h1 if h1 and h1 != op.MISSING else (title if title and title != op.MISSING else "this page")
    return [{
        "keyword": topic.lower(),
        "query": f"What should I know about {topic}?",
        "answer": f"{brand} covers {topic} on this page - see {url} for full details.",
    }]


# --------------------------------------------------------------------------- #
# Step 2: no-login AI-visibility checks
# --------------------------------------------------------------------------- #
def _domain_cited(text, domain):
    if not text:
        return False
    root = domain[4:] if domain.startswith("www.") else domain
    brand_token = root.split(".")[0]
    return root.lower() in text.lower() or (len(brand_token) > 3 and brand_token.lower() in text.lower())


def _check_perplexity(driver, query, domain):
    """Perplexity's public search - no login required, confirmed live and reliable.
    Returns {checked, cited, snippet}."""
    try:
        driver.get(f"https://www.perplexity.ai/search?q={query.replace(' ', '+')}")
        # Poll for the answer to actually render instead of a fixed sleep - same
        # hydration-wait idea as generate_seo_onpage_phase2's _wait_settled().
        text = ""
        for _ in range(20):
            time.sleep(1)
            try:
                text = driver.execute_script("return document.body.innerText;") or ""
            except Exception:
                continue
            if len(text) > 800 and "Searching the web" not in text[-200:]:
                break
        cited = _domain_cited(text, domain)
        return {"engine": "perplexity", "checked": True, "cited": cited,
                "snippet": text[:400]}
    except Exception as e:
        return {"engine": "perplexity", "checked": False, "cited": False,
                "snippet": f"error: {type(e).__name__}"}


def _check_google_ai_overview(driver, query, domain):
    """Google AI Overview - BEST EFFORT ONLY. Headless requests get a captcha/
    'unusual traffic' wall from Google almost every time (confirmed live) - this is
    the exact same limitation the existing SERP screenshot already has to skip past
    in generate_seo_onpage_phase2.capture_onpage_screenshots(). Reports checked=False
    (not a false negative) whenever blocked, rather than pretending it ran."""
    try:
        url = f"https://www.google.com/search?q={query.replace(' ', '+')}&hl=en&gl=us"
        driver.get(url)
        time.sleep(3)
        src = (driver.page_source or "").lower()
        if any(m in src for m in ("unusual traffic", "not a robot", "recaptcha",
                                   "/sorry/", "detected unusual", "before you continue")):
            return {"engine": "google_ai_overview", "checked": False, "cited": False,
                    "snippet": "blocked (captcha)"}
        has_overview = any(m in src for m in ("ai overview", "ai-generated", "generative ai is experimental"))
        if not has_overview:
            return {"engine": "google_ai_overview", "checked": True, "cited": False,
                    "snippet": "no AI Overview shown for this query"}
        cited = _domain_cited(src, domain)
        return {"engine": "google_ai_overview", "checked": True, "cited": cited, "snippet": ""}
    except Exception as e:
        return {"engine": "google_ai_overview", "checked": False, "cited": False,
                "snippet": f"error: {type(e).__name__}"}


def _check_chatgpt_anon(driver, query, domain):
    """ChatGPT (chatgpt.com) - BEST EFFORT ONLY. Its real chat interface requires a
    logged-in account to get an actual answer (confirmed live) - this detects the
    login wall and reports checked=False rather than faking a result."""
    try:
        driver.get("https://chatgpt.com/")
        time.sleep(3)
        src = (driver.page_source or "").lower()
        if any(m in src for m in ("log in", "sign up", "stay logged out")):
            return {"engine": "chatgpt", "checked": False, "cited": False,
                    "snippet": "login required - not available anonymously"}
        # If no login wall is detected, this deployment allows anonymous chat -
        # not currently reachable to verify further without a live account, so
        # still report unchecked rather than guessing at a selector.
        return {"engine": "chatgpt", "checked": False, "cited": False,
                "snippet": "anonymous chat UI present but not yet wired up"}
    except Exception as e:
        return {"engine": "chatgpt", "checked": False, "cited": False,
                "snippet": f"error: {type(e).__name__}"}


def check_ai_visibility(driver, query, domain):
    """Runs all three engine checks for one query. Perplexity is the reliable one;
    Google/ChatGPT are best-effort and commonly report checked=False."""
    return [
        _check_perplexity(driver, query, domain),
        _check_google_ai_overview(driver, query, domain),
        _check_chatgpt_anon(driver, query, domain),
    ]


VISIBILITY_WORKERS = 2  # deliberately small - real headless Chrome instances hitting
                        # Perplexity/Google concurrently from the SAME IP; kept low to
                        # stay conservative on captcha/rate-limit risk rather than
                        # maximize throughput.


def _new_visibility_driver():
    """A genuinely NEW, independent Chrome instance - NOT generate_seo_onpage_phase2's
    shared singleton (_get_op_driver() always returns the SAME driver object), which
    is fine for that module's sequential screenshot capture but useless here: running
    checks concurrently needs one real browser process per worker."""
    import tempfile
    from selenium.webdriver import Chrome, ChromeOptions
    opts = ChromeOptions()
    opts.add_argument(f"--user-data-dir={tempfile.mkdtemp(prefix='seo_geo_vis_')}")
    opts.add_argument("--headless=new")
    opts.add_argument("--no-sandbox")
    opts.add_argument("--disable-dev-shm-usage")
    opts.add_argument("--disable-gpu")
    opts.add_argument("--disable-notifications")
    opts.add_argument("--mute-audio")
    opts.add_argument("--window-size=1366,900")
    opts.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                      "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36")
    driver = Chrome(options=opts)
    driver.set_page_load_timeout(45)
    return driver


def run_visibility_checks_batch(queries, domain):
    """queries: list of query strings (duplicates are checked only once).
    Runs up to VISIBILITY_WORKERS independent Chrome instances concurrently -
    genuine parallelism, not tabs (a single Selenium driver only ever controls
    one tab at a time regardless of how many are open, so multiple tabs on one
    driver would NOT actually run concurrently). Falls back to the existing
    shared single-driver path if a second browser instance can't be launched
    (e.g. a memory-constrained machine), rather than failing the whole run.
    Returns {query: [perplexity_cell, google_cell, chatgpt_cell]}."""
    unique_queries = list(dict.fromkeys(queries))
    if not unique_queries:
        return {}
    results = {}

    n_workers = min(VISIBILITY_WORKERS, len(unique_queries))
    drivers = []
    try:
        drivers = [_new_visibility_driver() for _ in range(n_workers)]
    except Exception as e:
        log(f"   [warn] could not launch {n_workers} parallel browser instance(s) ({e}) - "
            f"falling back to a single shared browser.")
        for d in drivers:
            try:
                d.quit()
            except Exception:
                pass
        drivers = [op._get_op_driver()]

    def _cells_for(driver, query):
        log(f"   -> checking AI visibility: {query[:70]}")
        engine_results = check_ai_visibility(driver, query, domain)
        by_engine = {r["engine"]: r for r in engine_results}
        return [
            "Yes" if by_engine["perplexity"]["cited"] else ("Not checked" if not by_engine["perplexity"]["checked"] else "No"),
            "Yes" if by_engine["google_ai_overview"]["cited"] else ("Not checked" if not by_engine["google_ai_overview"]["checked"] else "No"),
            "Yes" if by_engine["chatgpt"]["cited"] else ("Not checked" if not by_engine["chatgpt"]["checked"] else "No"),
        ]

    def _worker(worker_idx):
        driver = drivers[worker_idx]
        for i in range(worker_idx, len(unique_queries), len(drivers)):
            q = unique_queries[i]
            try:
                results[q] = _cells_for(driver, q)
            except Exception as e:
                log(f"   [warn] visibility check failed for {q[:50]!r}: {e}")
                results[q] = ["Not checked", "Not checked", "Not checked"]

    import concurrent.futures
    try:
        with concurrent.futures.ThreadPoolExecutor(max_workers=len(drivers)) as pool:
            list(pool.map(_worker, range(len(drivers))))
    finally:
        # Only quit drivers WE created here - never quit the shared singleton
        # from _get_op_driver() (other code may still need it this run).
        for d in drivers:
            if d is not None and d is not getattr(op, "_op_driver", None):
                try:
                    d.quit()
                except Exception:
                    pass
    return results


# --------------------------------------------------------------------------- #
# Step 3: combine into the reference-matching xlsx
# --------------------------------------------------------------------------- #
def build_all_faqs(pages_data, brand, seed_keywords=None, faqs_per_page=5):
    """Generate the FAQ/keyword set ONCE per page and reuse it everywhere (xlsx,
    Schema, llms-full.txt) - matching the real reference, where every deliverable
    draws from the exact same FAQ data instead of each file inventing its own.
    Returns {url: [{"keyword","query","answer"}, ...]}."""
    faqs_by_page = {}
    for pd in pages_data:
        log(f"   -> generating {faqs_per_page} FAQ(s): {pd['url']}")
        faqs_by_page[pd["url"]] = generate_ai_faqs(pd, brand, n=faqs_per_page, seed_keywords=seed_keywords)
    return faqs_by_page


def build_geo_keywords_xlsx(domain, pages_data, faqs_by_page, out_path, check_visibility=True, seed_keywords=None):
    """seed_keywords: optional list of SEO keywords the team already has for this
    domain (given up front, e.g. from their existing keyword research) - these
    are also checked directly for AI-citation as their own rows, so the team can
    see right away whether a keyword they already have is already showing up in
    AI answers, instead of only getting brand-new AI-generated terms."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = Workbook()
    ws = wb.active
    ws.title = "Keywords"
    headers = ["URL", "AI-Searchable Keyword", "Long-Tail/NLP Query (FAQ)", "Answers", "Source"]
    if check_visibility:
        headers += ["Cited on Perplexity", "Cited on Google AI Overview", "Cited on ChatGPT"]
    ws.append(headers)
    for c in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=c)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F3864")
        cell.alignment = Alignment(wrap_text=True, vertical="center")

    # Build every row's data FIRST (no visibility cells yet), then run all the
    # needed visibility checks together as one batch through run_visibility_
    # checks_batch()'s parallel-browser pool, instead of checking one query at
    # a time interleaved with row-building - lets 2 independent Chrome
    # instances work through the query list concurrently.
    pending_rows = []
    for pd in pages_data:
        for item in faqs_by_page.get(pd["url"], []):
            pending_rows.append([pd["url"], item["keyword"], item["query"], item["answer"],
                                  "Auto-generated from page content", item["query"]])
    # Direct check of the team's own given keywords, as-typed - not just the
    # AI-rephrased question derived from them, so they can see their EXACT
    # existing keyword's current AI-citation status.
    if seed_keywords:
        home_url = pages_data[0]["url"] if pages_data else f"https://{domain}/"
        for kw in seed_keywords:
            pending_rows.append([home_url, kw, kw, "", "User-provided keyword (checked as-is)", kw])

    visibility_by_query = {}
    if check_visibility and pending_rows:
        visibility_by_query = run_visibility_checks_batch([r[5] for r in pending_rows], domain)

    total_rows = 0
    for url, keyword, query, answer, source, check_query in pending_rows:
        row = [url, keyword, query, answer, source]
        if check_visibility:
            row += visibility_by_query.get(check_query, ["Not checked", "Not checked", "Not checked"])
        ws.append(row)
        total_rows += 1

    ws.column_dimensions["A"].width = 45
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 45
    ws.column_dimensions["D"].width = 60
    ws.column_dimensions["E"].width = 30
    for col in "FGH":
        ws.column_dimensions[col].width = 22
    for row in ws.iter_rows(min_row=2, max_col=4):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    wb.save(out_path)
    log(f"   -> {total_rows} keyword/FAQ row(s) written to {out_path}")
    return total_rows


# --------------------------------------------------------------------------- #
# llms-full.txt - the whole site as markdown (AI-crawler-facing file), ending
# with the same FAQ set as the xlsx/Schema doc, matching the real reference's
# structure ("# https://domain/ llms-full.txt", per-page content, "## FAQs").
# Best-effort rendering: built from what's actually crawled (headings/paragraph
# text/images), not a pixel-identical clone of whatever dedicated llms.txt
# plugin/tool produced the reference - that tool has access to the fully
# rendered page; this reuses the same crawler as every other report here.
# --------------------------------------------------------------------------- #
def _fetch_html(url):
    import urllib.request
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0 SEOPhase2Bot"})
        with urllib.request.urlopen(req, timeout=30) as r:
            return r.read().decode("utf-8", "ignore")
    except Exception as e:
        log(f"   [warn] could not fetch {url} for markdown: {e}")
        return ""


def _inline_markdown(el, base_url):
    import urllib.parse as _up
    parts = []
    for node in el.children:
        if isinstance(node, str):
            parts.append(str(node))
        elif getattr(node, "name", None) == "a":
            href = node.get("href", "")
            href = _up.urljoin(base_url, href) if href else ""
            text = node.get_text(" ", strip=True)
            parts.append(f"[{text}]({href})" if text and href else text)
        elif getattr(node, "name", None) == "img":
            src = node.get("src") or node.get("data-src") or ""
            src = _up.urljoin(base_url, src) if src else ""
            if src:
                parts.append(f"![{node.get('alt', '')}]({src})")
        elif getattr(node, "name", None) in ("strong", "b"):
            parts.append(f"**{node.get_text(' ', strip=True)}**")
        elif getattr(node, "name", None) in ("em", "i"):
            parts.append(f"*{node.get_text(' ', strip=True)}*")
        elif hasattr(node, "get_text"):
            parts.append(node.get_text(" ", strip=True))
    return re.sub(r"\s+", " ", " ".join(parts)).strip()


def _html_to_markdown(html, base_url):
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    body = soup.body or soup
    out = []
    for el in body.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "img"]):
        if el.find_parent(["nav", "header", "footer"]):
            continue
        name = el.name
        if name and len(name) == 2 and name[0] == "h" and name[1].isdigit():
            text = el.get_text(" ", strip=True)
            if text:
                out.append("#" * int(name[1]) + " " + text)
        elif name == "img":
            if el.find_parent(["p", "li"]):
                continue  # already covered inline by its parent block
            import urllib.parse as _up
            src = el.get("src") or el.get("data-src") or ""
            if src:
                out.append(f"![{el.get('alt', '')}]({_up.urljoin(base_url, src)})")
        else:
            text = _inline_markdown(el, base_url)
            if text and len(text.split()) >= 3:
                out.append(text)
    return "\n\n".join(out)


def build_llms_full_txt(domain, pages_data, faqs_by_page, out_path):
    lines = [f"# https://{domain}/ llms-full.txt", ""]
    for pd in pages_data:
        html = _fetch_html(pd["url"])
        title = pd.get("title") or pd.get("h1") or pd["url"]
        if title and title != op.MISSING:
            lines.append(f"# {title}")
            lines.append("")
        md = _html_to_markdown(html, pd["url"]) if html else (pd.get("body_text") or "")
        if md:
            lines.append(md)
            lines.append("")

    lines.append("## FAQs")
    lines.append("")
    for pd in pages_data:
        for item in faqs_by_page.get(pd["url"], []):
            lines.append(f"[{item['query']}]({pd['url']})")
            lines.append("")
            lines.append(item["answer"])
            lines.append("")

    out_path.write_text("\n".join(lines), encoding="utf-8")
    log(f"   -> llms-full.txt written to {out_path} ({len(lines)} lines)")


# --------------------------------------------------------------------------- #
# Schema Suggestions - FAQPage JSON-LD per page, built directly from the SAME
# FAQ data as the Keywords xlsx (no new AI calls) - matches the real reference:
# title, yellow-highlighted warning note, then per-URL "FAQs Schema" (Heading 2)
# blocks with a literal <script type="application/ld+json"> block.
# --------------------------------------------------------------------------- #
def build_schema_docx(domain, faqs_by_page, out_path):
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run(f"Schema - {domain}")
    r.bold = True
    r.font.size = Pt(22)

    note_p = doc.add_paragraph()
    note_r = note_p.add_run(
        "Note – Before adding the new suggested schema, remove the existing FAQs "
        "schema from the page and then update the following schema to prevent any conflicts.")
    note_r.bold = True
    rpr = note_r._element.get_or_add_rPr()
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), "yellow")
    rpr.append(highlight)

    blocks = 0
    for url, faqs in faqs_by_page.items():
        if not faqs:
            continue
        doc.add_heading("FAQs Schema", level=2)
        p2 = doc.add_paragraph()
        r2 = p2.add_run(f"Page URL - {url}")
        r2.bold = True

        schema_obj = {
            "@context": "https://schema.org",
            "@type": "FAQPage",
            "mainEntity": [
                {"@type": "Question", "name": item["query"],
                 "acceptedAnswer": {"@type": "Answer", "text": item["answer"]}}
                for item in faqs
            ],
        }
        doc.add_paragraph('<script type="application/ld+json">')
        for line in json.dumps(schema_obj, indent=2, ensure_ascii=False).split("\n"):
            doc.add_paragraph(line)
        doc.add_paragraph("</script>")
        blocks += 1

    doc.save(out_path)
    log(f"   -> Schema Suggestions written to {out_path} ({blocks} page block(s))")


# --------------------------------------------------------------------------- #
# Technical Optimization Report - image ALT-tag audit, matching the real
# reference: title, intro, per-image "Page URL:/Image Source:/Image ALT
# Suggestion:" bullets, closing green bold "Suggestion:" call-to-action.
# --------------------------------------------------------------------------- #
def _suggest_alt_text(page_url, img_src, brand, page_title=""):
    prompt = (
        f"Write a concise, descriptive image ALT text (under 125 characters) for an image on "
        f"{brand}'s page at {page_url} (page topic: {page_title or 'unknown'}). "
        f"The image file is: {img_src}. Infer what the image likely shows from the filename/page "
        "topic. Return ONLY the ALT text, no quotes, no extra commentary."
    )
    result = op._ai_suggest(f'Return ONLY JSON: {{"alt": "..."}}\n\n{prompt}')
    if isinstance(result, dict) and result.get("alt"):
        return str(result["alt"]).strip()
    # Heuristic fallback - filename-derived
    name = img_src.rsplit("/", 1)[-1].rsplit(".", 1)[0]
    name = re.sub(r"[-_]+", " ", name).strip() or "image"
    return f"{name} - {brand}".strip(" -")


def build_technical_optimization_docx(domain, pages_data, brand, out_path, max_images=20):
    from docx import Document
    from docx.shared import Pt, RGBColor

    missing = []
    for pd in pages_data:
        for im in pd.get("images", []):
            if not (im.get("alt") or "").strip():
                missing.append((pd["url"], im.get("src", ""), pd.get("title") or pd.get("h1") or ""))

    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Technical Optimization Suggestions")
    r.bold = True
    r.font.size = Pt(22)
    p = doc.add_paragraph()
    r = p.add_run("Technical Optimization Analysis:")
    r.bold = True
    doc.add_paragraph(
        "We reviewed the target pages for technical on-page issues that can affect how AI "
        "crawlers and search engines read and cite this content.")
    p = doc.add_paragraph()
    r = p.add_run("Image ALT Optimization:")
    r.bold = True

    if missing:
        doc.add_paragraph(
            f"{len(missing)} image(s) are missing ALT tags across the target pages checked - "
            "ALT text helps both accessibility and AI/search engines understand image content.")
        to_process = missing[:max_images]
        # Each ALT suggestion is an independent AI call (a full Gemini/Groq/
        # OpenRouter round-trip) - running them one at a time made a page with
        # many missing-ALT images take a very long time. Same bounded-
        # concurrency pattern already used for broken-link checking elsewhere
        # in this codebase. Order in the output doc is preserved even though
        # completion order isn't.
        import concurrent.futures
        alts = [None] * len(to_process)
        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as pool:
            futures = {pool.submit(_suggest_alt_text, url, src, brand, title): i
                       for i, (url, src, title) in enumerate(to_process)}
            done = 0
            for fut in concurrent.futures.as_completed(futures):
                alts[futures[fut]] = fut.result()
                done += 1
                log(f"   -> ALT suggestion {done}/{len(to_process)} done")
        for (url, src, title), alt in zip(to_process, alts):
            doc.add_paragraph(f"Page URL: - {url}", style="List Paragraph")
            doc.add_paragraph(f"Image Source: - {src}", style="List Paragraph")
            doc.add_paragraph(f"Image ALT Suggestion: - {alt}", style="List Paragraph")
        if len(missing) > max_images:
            doc.add_paragraph(f"...and {len(missing) - max_images} more image(s) not shown here.")
    else:
        doc.add_paragraph("No missing ALT tags were found on the target pages checked.")

    p = doc.add_paragraph()
    r = p.add_run("Suggestion: - Please review and approve the above ALT text suggestions for implementation.")
    r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0xB0, 0x50)
    doc.save(out_path)
    log(f"   -> Technical Optimization Report written to {out_path} ({len(missing)} missing ALT(s))")


# --------------------------------------------------------------------------- #
# Existing Content Modification Suggestions - AI-grounded Find/Replace pairs,
# each verified to be a REAL exact substring of the page's actual crawled
# content before being kept (a suggestion whose "find" text doesn't really
# exist on the page can't be applied, which would defeat the point).
# --------------------------------------------------------------------------- #
def generate_content_modifications(page_data, brand, n=2):
    body = (page_data.get("body_text") or "")[:2500]
    if not body:
        return []
    prompt = (
        "You are editing existing website copy for AI-search optimization. Below is real text "
        f"from {brand}'s page at {page_data['url']}.\nContent: {body}\n\n"
        f"Find up to {n} short passages (1-2 sentences each) in this content that could be "
        "improved - fixing grammar/typos and/or turning a plain statement into an FAQ-style "
        "lead-in question naturally answered by the surrounding text. Each passage you pick MUST "
        "be an EXACT quote copied from the content above (so it can be found with a find/replace) "
        "- do not paraphrase the 'find' text.\n"
        'Return ONLY JSON: [{"find": "<exact original passage>", "replace": "<edited passage>"}, ...]'
    )
    result = op._ai_suggest(prompt)
    out = []
    if isinstance(result, list):
        for item in result[:n]:
            if isinstance(item, dict) and item.get("find") and item.get("replace"):
                find_text = str(item["find"]).strip()
                if find_text and find_text in body:
                    out.append({"find": find_text, "replace": str(item["replace"]).strip()})
    return out


def build_content_modification_docx(domain, pages_data, brand, out_path):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Existing Content Modification Suggestions")
    r.bold = True
    r.font.size = Pt(22)

    any_content = False
    for pd in pages_data:
        log(f"   -> content modifications: {pd['url']}")
        mods = generate_content_modifications(pd, brand)
        if not mods:
            continue
        any_content = True
        p = doc.add_paragraph()
        r = p.add_run(f"Page URL: - {pd['url']}")
        r.bold = True
        for m in mods:
            p1 = doc.add_paragraph()
            r1 = p1.add_run(f"Find: - {m['find']}")
            r1.bold = True
            p2 = doc.add_paragraph()
            r2 = p2.add_run(f"Replace With: - {m['replace']}")
            r2.bold = True
    if not any_content:
        doc.add_paragraph("No content modification suggestions were generated for the pages checked.")
    doc.save(out_path)
    log(f"   -> Content Modification Suggestions written to {out_path}")


# --------------------------------------------------------------------------- #
# Internal Linking Suggestions - AI-suggested anchor text + target page, again
# verified against the real crawled content before being kept, with a REAL
# w:hyperlink run in the "Replace With" paragraph (not just blue text). The
# candidate target pool is the sitemap (not just the handful of target pages
# in this run) so a suggestion can point to any real page on the site, not
# just whichever few pages happen to be in the current batch - and every
# suggested target_url is verified against that real set before being kept,
# so a hallucinated/unrelated URL never makes it into the output.
# --------------------------------------------------------------------------- #
def get_sitemap_urls(domain, cap=300):
    """Real page URLs from the site's own sitemap (following one level of
    sitemap-index nesting, capped to a handful of sub-sitemaps so a huge site
    doesn't take forever). Returns [] if no sitemap is found."""
    import xml.etree.ElementTree as ET
    sitemap_url, body = op.find_existing_sitemap(domain)
    if not body:
        return []

    def _locs(xml_body):
        try:
            root = ET.fromstring(xml_body.encode("utf-8", "ignore"))
        except Exception:
            return [], []
        ns = "{http://www.sitemaps.org/schemas/sitemap/0.9}"
        urls = [el.text.strip() for el in root.iter(f"{ns}loc") if el.text]
        if not urls:
            urls = [el.text.strip() for el in root.iter("loc") if el.text]
        is_index = root.tag.endswith("sitemapindex")
        return urls, is_index

    urls, is_index = _locs(body)
    if not is_index:
        return urls[:cap]

    # sitemap index - fetch a handful of sub-sitemaps for their real page URLs
    page_urls = []
    for sub_url in urls[:5]:
        _, sub_body, _ = op._http(sub_url)
        if not sub_body:
            continue
        sub_urls, _ = _locs(sub_body)
        page_urls.extend(sub_urls)
        if len(page_urls) >= cap:
            break
    return page_urls[:cap]


def generate_internal_links(page_data, other_pages, sitemap_urls, brand, n=2):
    body = (page_data.get("body_text") or "")[:2000]
    candidates = list(other_pages)
    crawled_urls = {p["url"] for p in other_pages} | {page_data["url"]}
    extra = [u for u in sitemap_urls if u not in crawled_urls][:20]
    candidates_display = [f"- {p['url']} ({p.get('h1') or p.get('title') or ''})" for p in other_pages[:8]]
    candidates_display += [f"- {u}" for u in extra]
    valid_targets = crawled_urls | set(extra)
    if not body or not candidates_display:
        return []
    targets = "\n".join(candidates_display)
    prompt = (
        f"You are adding internal links to existing website copy for {brand}.\n"
        f"Current page: {page_data['url']}\nContent: {body}\n\n"
        f"Other REAL pages on the site you could link to (from the sitemap):\n{targets}\n\n"
        f"Find up to {n} short sentences (EXACT quotes copied from the content above - do not "
        "paraphrase) where inserting an internal link to ONE OF THE PAGES LISTED ABOVE would make "
        "sense - the target must be genuinely relevant to what the sentence is about, not just any "
        "page. Use a natural anchor text that is a substring of that same sentence.\n"
        'Return ONLY JSON: [{"find": "<exact original sentence>", "anchor_text": "<words within '
        'find that become the link>", "target_url": "<one of the exact page URLs listed above>"}, ...]'
    )
    result = op._ai_suggest(prompt)
    out = []
    if isinstance(result, list):
        for item in result[:n]:
            if not isinstance(item, dict):
                continue
            find_text = str(item.get("find", "")).strip()
            anchor = str(item.get("anchor_text", "")).strip()
            target = str(item.get("target_url", "")).strip()
            # target_url must be a REAL candidate page - not the current page
            # itself, and not something the AI invented - otherwise the link
            # goes nowhere relevant, which is worse than no suggestion at all.
            if (find_text and anchor and target and find_text in body and anchor in find_text
                    and target in valid_targets and target != page_data["url"]):
                out.append({"find": find_text, "anchor_text": anchor, "target_url": target})
    return out


def _add_hyperlink(paragraph, text, url):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                           is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build_internal_linking_docx(domain, pages_data, brand, out_path):
    from docx import Document
    from docx.shared import Pt

    log("   -> fetching sitemap for internal-linking candidates...")
    sitemap_urls = get_sitemap_urls(domain)
    log(f"   -> {len(sitemap_urls)} URL(s) found in sitemap")

    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Internal Linking Suggestions")
    r.bold = True
    r.font.size = Pt(22)

    any_content = False
    for pd in pages_data:
        others = [o for o in pages_data if o["url"] != pd["url"]]
        log(f"   -> internal linking: {pd['url']}")
        links = generate_internal_links(pd, others, sitemap_urls, brand)
        if not links:
            continue
        any_content = True
        p = doc.add_paragraph()
        r = p.add_run(f"Page URL: - {pd['url']}")
        r.bold = True
        for l in links:
            p1 = doc.add_paragraph()
            r1 = p1.add_run("Find With: - ")
            r1.bold = True
            p1.add_run(l["find"])
            p2 = doc.add_paragraph()
            r2 = p2.add_run("Replace With: - ")
            r2.bold = True
            before, _, after = l["find"].partition(l["anchor_text"])
            if before:
                p2.add_run(before)
            _add_hyperlink(p2, l["anchor_text"], l["target_url"])
            if after:
                p2.add_run(after)
    if not any_content:
        doc.add_paragraph("No internal linking suggestions were generated for the pages checked.")
    doc.save(out_path)
    log(f"   -> Internal Linking Suggestions written to {out_path}")


# --------------------------------------------------------------------------- #
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("domain")
    ap.add_argument("--targets", help="targets.json or .xlsx (target pages).")
    ap.add_argument("--pages", default=None,
                     help="One or more target page URLs (newline or comma separated) - the team's "
                          "GEO work is done per-page, so at least one page (the homepage if that's "
                          "all there is) is required unless --targets is given instead.")
    ap.add_argument("--out", default=str(OUTPUT_DIR))
    ap.add_argument("--dry-run", action="store_true")
    ap.add_argument("--no-visibility-check", action="store_true",
                     help="skip the live no-login AI-visibility checks (faster, keyword/FAQ generation only)")
    ap.add_argument("--keywords", default=None,
                     help="Comma or newline separated SEO keywords the team already has for this domain - "
                          "used to steer FAQ generation and checked directly for AI-citation as their own rows.")
    ap.add_argument("--keywords-only", action="store_true",
                     help="only generate the Keywords/FAQ xlsx - skip llms-full.txt, Schema, Technical, "
                          "Content Modification, and Internal Linking (faster, fewer AI calls)")
    ap.add_argument("--faqs-per-page", type=int, default=5,
                     help="Number of AI-searchable keyword/FAQ entries to generate per page (default 5).")
    args = ap.parse_args()

    faqs_per_page = max(1, min(20, args.faqs_per_page))

    seed_keywords = None
    if args.keywords:
        seed_keywords = [k.strip() for k in re.split(r"[,\n]", args.keywords) if k.strip()]

    domain = op.safe_domain(args.domain)
    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)
    work = out_dir / f"_geo_{domain}"
    if work.exists():
        import shutil
        shutil.rmtree(work, ignore_errors=True)
    work.mkdir(parents=True)

    if not args.targets and not args.pages:
        log("[ERROR] At least one target page is required - the team's GEO work is done per-page, "
            "not just per-domain. Provide --pages (e.g. just the homepage URL if that's all there "
            "is) or --targets.")
        sys.exit(2)

    if args.targets:
        targets = op.load_targets(args.targets, domain)
    else:
        page_urls = [op.normalize_url(u.strip(), domain)
                     for u in re.split(r"[,\n]", args.pages) if u.strip()]
        page_urls = list(dict.fromkeys(u for u in page_urls if u))
        targets = [{"page": u, "keywords": []} for u in page_urls]
    if not targets:
        log("[ERROR] No target pages found.")
        sys.exit(2)
    log(f"[1/4] Loaded {len(targets)} target page(s)")

    pages_data = []
    for i, t in enumerate(targets, 1):
        log(f"   -> [crawl {i}/{len(targets)}] {t['page']}")
        pages_data.append(op.crawl_page(t["page"], t.get("keywords", []), dry_run=args.dry_run))
    op._close_op_driver()
    log(f"[2/4] Crawled {len(pages_data)} page(s)")

    homepage = pages_data[0] if pages_data else None
    brand = op.brand_from(domain, homepage.get("title") if homepage else None,
                           homepage.get("h1") if homepage else None,
                           homepage.get("og_site_name") if homepage else None)

    log(f"[3/4] Generating AI-searchable keywords/FAQs ({faqs_per_page} per page)...")
    faqs_by_page = build_all_faqs(pages_data, brand, seed_keywords=seed_keywords, faqs_per_page=faqs_per_page)

    log("[4/4] Building deliverables" +
        ("" if args.no_visibility_check else " + checking live AI visibility") + "...")
    xlsx_path = work / f"{domain} - Keywords and Content (FAQs) Work for GEO.xlsx"
    build_geo_keywords_xlsx(domain, pages_data, faqs_by_page, xlsx_path,
                             check_visibility=not args.no_visibility_check,
                             seed_keywords=seed_keywords)

    if not args.keywords_only:
        build_llms_full_txt(domain, pages_data, faqs_by_page, work / "llms-full.txt")
        build_schema_docx(domain, faqs_by_page,
                           work / f"{domain} - (GEO Work) Schema Suggestions.docx")
        build_technical_optimization_docx(domain, pages_data, brand,
                                           work / f"{domain} - (GEO Work) Technical Optimization Report.docx")
        build_content_modification_docx(domain, pages_data, brand,
                                         work / f"{domain} - (GEO Work) Existing Content Modification Suggestions.docx")
        build_internal_linking_docx(domain, pages_data, brand,
                                     work / f"{domain} - (GEO Work) Internal Linking Suggestions.docx")

    # bundle (filename uses the domain so the task-runner's glob matches, same
    # pattern as generate_seo_onpage_phase2.py's ZIP)
    import shutil, zipfile
    zip_path = out_dir / f"GEO Report - {domain}.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for f in sorted(work.iterdir()):
            if f.is_file():
                z.write(f, f.name)
    shutil.rmtree(work, ignore_errors=True)
    log(f"[DONE] Bundled -> {zip_path}")


if __name__ == "__main__":
    main()
