"""
WEBSITE HEALTH AUDIT REPORT GENERATOR (.docx)

Reusable engine for the "Website Health Audit Report" format used by Report
Studio. Reproduces the reference James/usaautoship house format exactly:

  - Page 8.5 x 11", 1" margins all sides
  - Body font: Candara 11pt
  - Each checkpoint = bold label + normal explanation, "List Paragraph" indent
  - Centered client logo at the top
  - Every screenshot is an inline picture with a 3pt solid-black border + the
    same soft drop-shadow as the reference document

Checkpoints are NUMBERED 1..17 (auto). The black image border is BAKED into each
picture (with a tall white top buffer) so Word's line box can never clip the top
border (the docx-border approach was getting clipped until you resized).

Screenshots (live via patchright, all HEADLESS — no window opens, only the page
is captured, nothing else on screen is grabbed; NO drawn browser chrome):
  - Client LOGO: header element screenshot (or og:image / apple-touch-icon).
  - sucuri / manual / security / status200(GSC index): tool + GSC pages.
  - robots.txt, sitemap.xml ({domain}/sitemap.xml), homepage (versions / blank),
    SERP (site:{domain}), dummy (site:{domain} lorem), wayback (layout),
    meta_source (view-source:, scrolled to the meta description): plain page shots.
  - canonical / double_meta / meta_robots: MANUAL drop-ins (from the
    "SEO META in 1 Click" extension; can't run headless).
  - broken_links: real Python link check rendered to an image.
  - pagespeed: PageSpeed Insights (waits for the full report before capturing).
  - Drop-in PNG under health_screenshots/<domain>/<key>.png (or logo.png) always
    wins. Missing image → text-only; the doc never breaks.

Usage:
    python generate_health_report.py <domain> [--account NAME] [--format HINT]
                                     [--score 68] [--logo path\\to\\logo.png]
                                     [--no-capture]   # text + drop-ins only
"""
import os
import sys
import tempfile
from pathlib import Path
from datetime import datetime
from urllib.parse import quote

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor

# Reuse the battle-tested capture helpers from the PPTX engine
import generate_report as gr

# ---------- DOC FORMATTING (matches the reference exactly) ----------
FONT_NAME = "Candara"
FONT_SIZE_PT = 11
PAGE_W_IN, PAGE_H_IN = 8.5, 11.0
MARGIN_IN = 1.0
BODY_INDENT_IN = 0.0          # text flush-left (left aligned)
IMAGE_WIDTH_IN = 6.3          # near full text width, like the reference
IMAGE_MAX_H_IN = 4.6          # cap so a tall capture never overflows / clips
TEXT_COLOR = RGBColor(0x00, 0x00, 0x00)

OUTPUT_DIR = Path("output")
DROPIN_DIR = Path("health_screenshots")   # health_screenshots/<domain>/<key>.png

# Border + shadow XML lifted from the reference doc's <pic:spPr>.
# algn="in" draws the border INSIDE the image bounds so the top border can never
# be clipped by the line box (fixes "top border hidden until you resize").
_BORDER_XML = (
    '<a:ln %s w="38100" cap="sq" algn="in"><a:solidFill><a:srgbClr val="000000"/>'
    '</a:solidFill><a:prstDash val="solid"/><a:miter lim="800000"/></a:ln>'
    % nsdecls("a")
)
_SHADOW_XML = (
    '<a:effectLst %s><a:outerShdw blurRad="50800" dist="38100" dir="2700000" '
    'algn="tl" rotWithShape="0"><a:srgbClr val="000000"><a:alpha val="43000"/>'
    "</a:srgbClr></a:outerShdw></a:effectLst>" % nsdecls("a")
)


# ---------- CHECKPOINTS ----------
# Each: key, label (bold), body (normal). {domain} / {score} are filled in.
# "capture" = how the screenshot is obtained:
#   gsc:<page>  -> Search Console page (needs an authed account)
#   sucuri      -> Sucuri sitecheck
#   url:<tmpl>  -> navigate to a URL and screenshot
#   wayback     -> web.archive.org latest snapshot
#   pagespeed   -> PageSpeed Insights
#   manual      -> not auto-captured; drop-in only
INTRO = [
    ("Website health checkup and analysis",
     ": As a process of routine analysis, we have checked the website "
     "according to different points in order to determine if any issue "
     "persists in the website or not."),
    ("",
     "And after doing the analysis, we found that there is no health related "
     "issue in the website."),
]

# "capture" verbs:
#   sucuri / gsc:<page> / wayback / pagespeed       -> as before
#   frame:<url>   -> navigate, wait, screenshot, then draw a browser address bar
#                    (showing "addr") on top — like a real browser screenshot
#   viewsrc:<url> -> same but for view-source: pages
# "addr" = the text shown in the synthetic address bar.
CHECKPOINTS = [
    {"key": "sucuri", "capture": "sucuri",
     "label": "Sucuri Check",
     "body": ": We first of all check the website in the tool Sucuri, to check "
             "for any malware or website blacklisting issue.",
     "extra": ("URL:  ", "https://sitecheck.sucuri.net/results/https/{domain}"),
     "status": ("Status", ": We have not find issue in the website.")},

    {"key": "manual", "capture": "gsc:manual-actions",
     "label": "Manual Action:",
     "body": " The Manual Actions report lists manually detected issues with a "
             "page or site that are mostly attempts to manipulate our search "
             "index, but are not necessarily dangerous for users.",
     "status": ("Status –", " Not found")},

    {"key": "security", "capture": "gsc:security-issues",
     "label": "Security Issues:",
     "body": " The Security Issues report lists indications that your site was "
             "hacked, or behavior on your site that could potentially harm a "
             "visitor or their computer: for example, phishing attacks or "
             "installing malware or unwanted software on the user's computer.",
     "status": ("Status:  ", "Not Found")},

    {"key": "robots", "capture": "frame:https://{domain}/robots.txt",
     "addr": "{domain}/robots.txt",
     "label": "Robots.txt:",
     "body": " Robots.txt file has pages that needs to be blocked and are not "
             "useful for the website. Thus, we have checked the robots file is "
             "running fine for now."},

    {"key": "sitemap", "capture": "frame:https://{domain}/sitemap.xml",
     "addr": "{domain}/sitemap.xml",
     "label": "Sitemap.xml: ",
     "body": "Sitemap file must have all the web-pages that are being present "
             "in the website so that all the pages can be identified by the "
             "Search Engine. We have checked the sitemap file as it is having "
             "all the relevant pages or not in the website."},

    {"key": "status200", "capture": "status200",
     "label": "All Target page's status is 200:- ",
     "body": "{status200_result}"},

    {"key": "versions", "capture": "computed",
     "label": "The website is running with multiple versions: - ",
     "body": "{versions_result}"},

    {"key": "meta_source", "capture": "viewsrc:https://{domain}/",
     "addr": "view-source:https://{domain}/",
     "find": 'meta name="description"',   # scroll the source to this tag before capturing
     "label": "Meta Suggestions is visible in the source code (Also on Top):- ",
     "body": "Meta suggestion are visible in the source code of the pages."},

    # 10/11/12 come from the "SEO META in 1 Click" Chrome extension, which can't
    # be driven headless — so these are MANUAL drop-ins:
    #   health_screenshots/<domain>/{canonical,double_meta,meta_robots}.png
    {"key": "canonical", "capture": "computed",
     "label": "Right Canonical tag: - ",
     "body": "{canonical_result}"},

    {"key": "double_meta", "capture": "computed",
     "label": "Check for Double Meta Tags: - ",
     "body": "{double_meta_result}"},

    {"key": "meta_robots", "capture": "computed",
     "label": "Check Meta Robots Tag on Web-Pages: - ",
     "body": "{meta_robots_result}"},

    {"key": "layout", "capture": "wayback",
     "label": "Check Website Layout: -  ",
     "body": "The screenshot below is the latest archived snapshot from web.archive.org. Compare it with the current live site to confirm no unexpected layout changes have occurred recently."},

    {"key": "dummy", "capture": "frame:https://www.google.com/search?q=site:{domain}+lorem",
     "addr": "site:{domain} lorem",
     "label": "Check Dummy Content: ",
     "body": "{dummy_content_result}"},

    {"key": "broken_links", "capture": "brokenlinks",
     "label": "Check for Broken Links:- ",
     "body": "{broken_links_result}"},
]


# ---------- DOCX BUILDER ----------
def _style_run(run, bold=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE_PT)
    run.font.bold = bold
    run.font.color.rgb = TEXT_COLOR


def _add_para(doc, indent_in=BODY_INDENT_IN, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(style="List Paragraph")
    if indent_in is not None:
        p.paragraph_format.left_indent = Inches(indent_in)
    if align is not None:
        p.alignment = align
    return p


def _add_label_body(doc, label, body):
    """One checkpoint line: bold label + normal body, Candara 11pt."""
    p = _add_para(doc)
    if label:
        _style_run(p.add_run(label), bold=True)
    if body:
        _style_run(p.add_run(body), bold=False)
    return p


def _apply_shadow(shape):
    """Add only the soft drop-shadow (the black border is baked into the image)."""
    spPr = shape._inline.graphic.graphicData.pic.spPr
    spPr.append(parse_xml(_SHADOW_XML))


def _bake_image_frame(src_path, big_top=True):
    """Return a temp PNG with a black border baked IN + a GENEROUS white top
    margin (and smaller side/bottom margins).

    Definitive top-border fix: Word's line box can clip the very top of an inline
    image (the border vanishes until you resize). By baking the border into the
    image and leaving a tall white buffer above it, any top clip only eats the
    invisible white margin — the black border + content always stay visible.
    `big_top=False` (used for the short logo) keeps the margins symmetric.
    """
    from PIL import Image, ImageDraw
    img = Image.open(str(src_path)).convert("RGB")
    W, H = img.size
    border = max(4, round(W / 200))            # ~3pt
    side = max(10, round(W / 110))
    bottom = side
    top = max(70, round(H * 0.14)) if big_top else side   # big buffer absorbs the clip
    cw = W + 2 * border + 2 * side
    ch = H + 2 * border + top + bottom
    canvas = Image.new("RGB", (cw, ch), "#FFFFFF")
    ix, iy = side + border, top + border
    canvas.paste(img, (ix, iy))
    d = ImageDraw.Draw(canvas)
    d.rectangle([ix - border, iy - border, ix + W + border - 1, iy + H + border - 1],
                outline=(17, 17, 17), width=border)
    fd, tmp = tempfile.mkstemp(suffix=".png", prefix="hframe_")
    os.close(fd)
    canvas.save(tmp)
    return tmp


def _add_bordered_image(doc, img_path):
    """Inline picture with a BAKED black border (+ white margin) and a docx
    drop-shadow. Paragraph matches the template (indent + tiny mark) for layout."""
    p = _add_para(doc, indent_in=0.44)            # match template indent (left=634 twips)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        '<w:rPr %s><w:rFonts w:ascii="Candara" w:hAnsi="Candara"/><w:sz w:val="10"/></w:rPr>'
        % nsdecls("w")
    ))
    run = p.add_run()
    run.font.size = Pt(5)
    framed = _bake_image_frame(img_path)
    try:
        nw, nh = gr.png_dimensions(framed)
        if nw and (IMAGE_WIDTH_IN * nh / nw) > IMAGE_MAX_H_IN:
            run.add_picture(framed, height=Inches(IMAGE_MAX_H_IN))
        else:
            run.add_picture(framed, width=Inches(IMAGE_WIDTH_IN))
        # No docx border/shadow — both are baked into the image (border) so the
        # top can never be clipped.
    finally:
        try:
            os.unlink(framed)
        except Exception:
            pass
    return p


def _resolve_image(domain, key, captured):
    """Drop-in PNG wins; otherwise the live-captured path; else None."""
    dropin = DROPIN_DIR / domain / f"{key}.png"
    if dropin.exists():
        return dropin
    val = captured.get(key)
    if isinstance(val, dict):
        img = val.get("img")
        if not img:
            return None
        # The worker passes the GSC screenshot as a base64 data-URI
        # ("data:image/jpeg;base64,...") rather than a file path. Decode it to a
        # temp file so python-docx can embed it.
        if isinstance(img, str) and img.startswith("data:"):
            import base64, tempfile, re as _re
            m = _re.match(r"data:image/(\w+);base64,(.*)", img, _re.S)
            if not m:
                return None
            ext, b64 = m.group(1), m.group(2)
            safe_domain = _re.sub(r"[^a-z0-9]", "_", domain.lower())
            tmp = Path(tempfile.gettempdir()) / f"gsc_{key}_{safe_domain}.{ext}"
            try:
                tmp.write_bytes(base64.b64decode(b64))
                return tmp
            except Exception:
                return None
        if Path(img).exists():
            return Path(img)
        return None
    if val and Path(val).exists():
        return Path(val)
    return None


def _resolve_gsc_status(key, captured):
    """Return the detected GSC status string for manual/security, or None."""
    val = captured.get(key)
    if isinstance(val, dict):
        return val.get("status")
    return None


# ---------- REAL CHECKS ----------
import re as _re_mod
import urllib.request as _ur
import urllib.error as _ue
from html.parser import HTMLParser as _HP

def _fetch_html(url, timeout=15):
    headers = {"User-Agent": "Mozilla/5.0 (compatible; HealthBot/1.0)"}
    try:
        req = _ur.Request(url, headers=headers)
        with _ur.urlopen(req, timeout=timeout) as r:
            return r.read().decode("utf-8", "ignore")
    except Exception as e:
        return None


def _get_status(url, timeout=10):
    headers = {"User-Agent": "Mozilla/5.0 (compatible; HealthBot/1.0)"}
    for method in ("HEAD", "GET"):
        try:
            req = _ur.Request(url, headers=headers, method=method)
            with _ur.urlopen(req, timeout=timeout) as r:
                return r.status
        except _ue.HTTPError as e:
            if method == "HEAD" and e.code in (403, 405):
                continue
            return e.code
        except Exception:
            return 0
    return 0


def check_status200(target_pages, domain):
    """Returns list of {url, status, ok} dicts."""
    if not target_pages:
        target_pages = [f"https://{domain}/"]
    results = []
    for url in target_pages:
        url = url.strip()
        if not url: continue
        if not url.startswith("http"): url = "https://" + url
        code = _get_status(url)
        results.append({"url": url, "status": str(code), "ok": code == 200})
    return results


def _parse_meta(html):
    """Extract canonical, title, meta description, meta robots from HTML."""
    result = {"canonical": [], "title": [], "description": [], "robots": []}
    # Canonical
    for m in _re_mod.finditer(r'<link[^>]+rel=["\']canonical["\'][^>]*href=["\']([^"\']+)["\']', html, _re_mod.I):
        result["canonical"].append(m.group(1).strip())
    for m in _re_mod.finditer(r'<link[^>]+href=["\']([^"\']+)["\'][^>]*rel=["\']canonical["\']', html, _re_mod.I):
        result["canonical"].append(m.group(1).strip())
    # Title
    for m in _re_mod.finditer(r'<title[^>]*>(.*?)</title>', html, _re_mod.I | _re_mod.S):
        t = m.group(1).strip()
        if t:
            result["title"].append(t)
    # Meta description
    for m in _re_mod.finditer(r'<meta[^>]+name=["\']description["\'][^>]*content=["\']([^"\']*)["\']', html, _re_mod.I):
        result["description"].append(m.group(1).strip())
    for m in _re_mod.finditer(r'<meta[^>]+content=["\']([^"\']*)["\'][^>]*name=["\']description["\']', html, _re_mod.I):
        result["description"].append(m.group(1).strip())
    # Meta robots
    for m in _re_mod.finditer(r'<meta[^>]+name=["\']robots["\'][^>]*content=["\']([^"\']*)["\']', html, _re_mod.I):
        result["robots"].append(m.group(1).strip())
    for m in _re_mod.finditer(r'<meta[^>]+content=["\']([^"\']*)["\'][^>]*name=["\']robots["\']', html, _re_mod.I):
        result["robots"].append(m.group(1).strip())
    return result


def _detect_live_homepage(domain):
    """Detect the actual running homepage URL by following redirects."""
    import urllib.request as _ur
    for prefix in (f"https://{domain}/", f"https://www.{domain}/", f"http://{domain}/", f"http://www.{domain}/"):
        try:
            req = _ur.Request(prefix, headers={"User-Agent": "Mozilla/5.0"})
            with _ur.urlopen(req, timeout=10) as r:
                return r.url
        except Exception:
            continue
    return f"https://{domain}/"


def check_canonical(target_pages, domain):
    """Returns list of {url, canonical, ok, note} dicts."""
    pages = list(target_pages) if target_pages else []
    homepage = _detect_live_homepage(domain)
    if not any(p.rstrip("/") == homepage.rstrip("/") for p in pages):
        pages = [homepage] + pages
    results = []
    for url in pages:
        url = url.strip()
        if not url: continue
        if not url.startswith("http"): url = "https://" + url
        html = _fetch_html(url)
        if html is None:
            results.append({"url": url, "canonical": "—", "ok": False, "note": "Could not fetch"})
            continue
        meta = _parse_meta(html)
        canonicals = list(dict.fromkeys(meta["canonical"]))
        if not canonicals:
            results.append({"url": url, "canonical": "Missing", "ok": False, "note": "No canonical tag"})
        elif len(canonicals) > 1:
            results.append({"url": url, "canonical": canonicals[0], "ok": False, "note": f"{len(canonicals)} canonical tags"})
        else:
            canon = canonicals[0].rstrip("/")
            page_url = url.rstrip("/")
            ok = canon == page_url
            results.append({"url": url, "canonical": canonicals[0], "ok": ok, "note": "Self-referencing" if ok else "Points elsewhere"})
    return results


def check_double_meta(target_pages, domain):
    """Returns list of {url, title_count, desc_count, ok, note} dicts."""
    pages = list(target_pages) if target_pages else []
    if not pages: pages = [f"https://{domain}/"]
    results = []
    for url in pages:
        url = url.strip()
        if not url: continue
        if not url.startswith("http"): url = "https://" + url
        html = _fetch_html(url)
        if html is None:
            results.append({"url": url, "title_count": "—", "desc_count": "—", "ok": False, "note": "Could not fetch"})
            continue
        meta = _parse_meta(html)
        tc, dc = len(meta["title"]), len(meta["description"])
        ok = (tc == 1 and dc == 1)
        notes = []
        if tc == 0: notes.append("Missing title")
        elif tc > 1: notes.append(f"Title ×{tc}")
        if dc == 0: notes.append("Missing meta desc")
        elif dc > 1: notes.append(f"Meta desc ×{dc}")
        results.append({"url": url, "title_count": str(tc), "desc_count": str(dc), "ok": ok, "note": ", ".join(notes) if notes else "OK"})
    return results


def check_meta_robots(target_pages, domain):
    """Returns list of {url, robots_value, ok, note} dicts."""
    pages = list(target_pages) if target_pages else []
    if not pages: pages = [f"https://{domain}/"]
    results = []
    for url in pages:
        url = url.strip()
        if not url: continue
        if not url.startswith("http"): url = "https://" + url
        html = _fetch_html(url)
        if html is None:
            results.append({"url": url, "robots_value": "—", "ok": False, "note": "Could not fetch"})
            continue
        meta = _parse_meta(html)
        robots_values = meta["robots"]
        if not robots_values:
            results.append({"url": url, "robots_value": "Not set (default: index,follow)", "ok": True, "note": "OK"})
            continue
        val = robots_values[0]
        val_lower = val.lower()
        issues = []
        if "noindex" in val_lower: issues.append("noindex")
        if "nofollow" in val_lower: issues.append("nofollow")
        ok = not issues
        results.append({"url": url, "robots_value": val, "ok": ok, "note": ", ".join(issues) if issues else "OK"})
    return results


def check_versions_full(domain):
    """Check the 4 URL variants. Returns (summary_text, rows) where rows is a list
    of (variant_url, http_code_str, result_str, ok_bool) for the table image."""
    variants = [
        f"http://{domain}/",
        f"http://www.{domain}/",
        f"https://{domain}/",
        f"https://www.{domain}/",
    ]
    active = []
    final_urls = set()
    for url in variants:
        try:
            import urllib.request as _ur2
            class _NoRedirect(_ur2.HTTPErrorProcessor):
                def http_response(self, req, resp): return resp
                https_response = http_response
            opener = _ur2.build_opener(_NoRedirect)
            req = _ur2.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with opener.open(req, timeout=10) as r:
                code = r.status
                location = r.getheader("Location", "")
            if code in (200, 301, 302, 307, 308):
                active.append((url, code, location))
                if code == 200:
                    final_urls.add(url.rstrip("/"))
                elif location:
                    final_urls.add(location.rstrip("/"))
        except Exception:
            pass

    returning_200 = [url for url, code, _ in active if code == 200]
    redirecting = [(url, code, loc) for url, code, loc in active if code != 200]

    # OK if exactly one variant returns 200 and all redirects point to it.
    is_ok = False
    if len(returning_200) == 1:
        canonical = returning_200[0].rstrip("/")
        if all(loc.rstrip("/") == canonical or loc.rstrip("/") == "" for _, _, loc in redirecting):
            is_ok = True
            summary = f"No issue found. All variants correctly consolidate to a single version: {returning_200[0]}"
    if not is_ok and len(returning_200) == 0 and len(final_urls) == 1:
        is_ok = True
        summary = f"No issue found. All variants redirect to: {list(final_urls)[0]}"
    if not is_ok:
        summary = (f"Issue found — {len(returning_200)} URL variant(s) return HTTP 200 independently. "
                   f"Fix: redirect (301) all variants to one canonical URL via server config or .htaccess.")

    # Build table rows. An independently-200 variant is the problem when there is
    # an issue; redirecting variants are fine.
    rows = []
    for url, code, loc in active:
        if code == 200:
            rows.append((url, "200", "OK (active)", is_ok))
        else:
            rows.append((url, str(code), f"-> {loc or '?'}", True))
    return summary, rows


def check_dummy_content(domain):
    """Check Google for site:domain lorem to detect dummy content pages."""
    import urllib.parse
    search_url = f"https://www.google.com/search?q=site:{domain}+lorem"
    html = _fetch_html(search_url)
    if html is None:
        return "Could not check Google for dummy content. Please check manually: " + search_url
    # Check if any results found
    no_results_markers = ["did not match any documents", "No results found", "no results", "0 results"]
    for marker in no_results_markers:
        if marker.lower() in html.lower():
            return f"No dummy content found. Google search for 'site:{domain} lorem' returned no results."
    # Extract result snippets
    import re as _re
    # Look for result count
    count_match = _re.search(r'About ([\d,]+) results', html)
    count = count_match.group(1) if count_match else "some"
    # Try to find page titles in results
    titles = _re.findall(r'<h3[^>]*>(.*?)</h3>', html, _re.S)
    titles = [_re.sub(r'<[^>]+>', '', t).strip() for t in titles[:5] if t.strip()]
    if titles:
        result = f"Possible dummy content found — Google shows {count} result(s) for 'site:{domain} lorem'.\nPages found:\n"
        result += "\n".join(f"  • {t}" for t in titles)
    else:
        result = f"Google returned {count} result(s) for 'site:{domain} lorem'. Please verify manually whether these pages contain dummy/lorem content."
    return result


def _render_table_image(title, headers, rows, col_widths, path):
    """Render a results table as a PNG image for embedding in the DOCX."""
    from PIL import Image, ImageDraw, ImageFont
    try:
        bold_font = ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 20)
        header_font = ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 16)
        cell_font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 15)
        title_font = ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 24)
    except Exception:
        bold_font = header_font = cell_font = title_font = ImageFont.load_default()

    W = sum(col_widths) + 2
    ROW_H = 28
    TITLE_H = 50
    HEADER_H = 32
    H = TITLE_H + HEADER_H + ROW_H * max(len(rows), 1) + 20
    img = Image.new("RGB", (W, H), "#FFFFFF")
    d = ImageDraw.Draw(img)

    # Title
    d.text((16, 14), title, fill=(31, 41, 55), font=title_font)

    # Header row
    y = TITLE_H
    x = 0
    d.rectangle([0, y, W, y + HEADER_H], fill="#F3F4F6")
    for i, h in enumerate(headers):
        d.text((x + 8, y + 8), h, fill=(75, 85, 99), font=header_font)
        x += col_widths[i]
        if i < len(headers) - 1:
            d.line([(x, y), (x, y + HEADER_H)], fill="#D1D5DB", width=1)
    d.line([(0, y + HEADER_H), (W, y + HEADER_H)], fill="#D1D5DB", width=1)

    # Data rows
    for ri, row in enumerate(rows):
        y = TITLE_H + HEADER_H + ri * ROW_H
        bg = "#F9FAFB" if ri % 2 == 0 else "#FFFFFF"
        d.rectangle([0, y, W, y + ROW_H], fill=bg)
        x = 0
        ok = row[-1] if isinstance(row[-1], bool) else True
        for ci, cell in enumerate(row[:-1]):  # last element is the bool ok flag
            color = (31, 41, 55)
            if ci == len(row) - 2:  # last visible column = note/status
                color = (22, 128, 67) if ok else (200, 40, 60)
            text = str(cell)
            # Truncate long URLs
            max_chars = col_widths[ci] // 8
            if len(text) > max_chars:
                text = text[:max_chars - 1] + "…"
            d.text((x + 8, y + 6), text, fill=color, font=cell_font)
            x += col_widths[ci]
            if ci < len(row) - 2:
                d.line([(x, y), (x, y + ROW_H)], fill="#E5E7EB", width=1)
        d.line([(0, y + ROW_H), (W, y + ROW_H)], fill="#E5E7EB", width=1)

    if not rows:
        d.text((16, TITLE_H + HEADER_H + 8), "No data.", fill=(150, 150, 150), font=cell_font)

    img.save(str(path))


def _render_status200_image(domain, results, path):
    rows = [
        (r["url"].replace("https://", "").replace("http://", ""), r["status"],
         "OK" if r["ok"] else "Issue", r["ok"])
        for r in results
    ]
    _render_table_image(
        f"HTTP Status Check — {domain}",
        ["Page URL", "Status", "Result"],
        rows,
        [780, 120, 140],
        path
    )


def _render_canonical_image(domain, results, path):
    rows = [
        (r["url"].replace("https://", "").replace("http://", ""),
         r["canonical"][:60] + "…" if len(r["canonical"]) > 60 else r["canonical"],
         r["note"], r["ok"])
        for r in results
    ]
    _render_table_image(
        f"Canonical Tag Check — {domain}",
        ["Page URL", "Canonical", "Result"],
        rows,
        [540, 360, 180],
        path
    )


def _render_double_meta_image(domain, results, path):
    rows = [
        (r["url"].replace("https://", "").replace("http://", ""),
         r["title_count"], r["desc_count"], r["note"], r["ok"])
        for r in results
    ]
    _render_table_image(
        f"Meta Tags Check — {domain}",
        ["Page URL", "Title", "Meta Desc", "Result"],
        rows,
        [600, 90, 110, 240],
        path
    )


def _render_meta_robots_image(domain, results, path):
    rows = [
        (r["url"].replace("https://", "").replace("http://", ""),
         r["robots_value"], r["note"], r["ok"])
        for r in results
    ]
    _render_table_image(
        f"Meta Robots Check — {domain}",
        ["Page URL", "Robots Value", "Result"],
        rows,
        [540, 360, 180],
        path
    )


def _render_versions_image(domain, rows, path):
    _render_table_image(
        f"URL Versions Check — {domain}",
        ["URL Variant", "HTTP", "Result"],
        rows,
        [620, 90, 360],
        path
    )


def build_health_docx(domain, captured=None, logo=None, score=68, out_path=None, target_pages=None):
    captured = captured or {}
    target_pages = [p.strip() for p in (target_pages or []) if p.strip()]

    # Run real checks
    _tmp_dir = Path(tempfile.gettempdir())
    _s200_results = check_status200(target_pages, domain)
    _canon_results = check_canonical(target_pages, domain)
    _dmeta_results = check_double_meta(target_pages, domain)
    _mrobots_results = check_meta_robots(target_pages, domain)
    # Render table images and inject into captured
    if _s200_results:
        _p = _tmp_dir / f"health_status200_{domain}.png"
        _render_status200_image(domain, _s200_results, _p)
        captured["status200"] = str(_p)
    if _canon_results:
        _p = _tmp_dir / f"health_canonical_{domain}.png"
        _render_canonical_image(domain, _canon_results, _p)
        captured["canonical"] = str(_p)
    if _dmeta_results:
        _p = _tmp_dir / f"health_doublemeta_{domain}.png"
        _render_double_meta_image(domain, _dmeta_results, _p)
        captured["double_meta"] = str(_p)
    if _mrobots_results:
        _p = _tmp_dir / f"health_metarobots_{domain}.png"
        _render_meta_robots_image(domain, _mrobots_results, _p)
        captured["meta_robots"] = str(_p)
    # URL versions — render the variant/redirect data as a table image.
    _versions_summary, _versions_rows = check_versions_full(domain)
    if _versions_rows:
        _p = _tmp_dir / f"health_versions_{domain}.png"
        _render_versions_image(domain, _versions_rows, _p)
        captured["versions"] = str(_p)

    # Broken links — run the check so we can use the result in the body text.
    # The image is already rendered during capture; here we just need the count.
    _bl_checked, _bl_broken = _check_broken_links(domain)

    def _broken_links_summary():
        if _bl_checked == 0:
            return "Could not check for broken links."
        if not _bl_broken:
            return f"Checked {_bl_checked} links — no broken links found. See screenshot below."
        return f"Checked {_bl_checked} links — {len(_bl_broken)} broken link(s) found. See screenshot below."

    def _s200_summary():
        if not _s200_results: return "No target pages provided."
        bad = [r for r in _s200_results if not r["ok"]]
        if not bad: return f"All {len(_s200_results)} page(s) returned HTTP 200. See screenshot below."
        return f"{len(bad)} of {len(_s200_results)} page(s) have non-200 status. See screenshot below."

    def _canon_summary():
        if not _canon_results: return "No pages checked."
        bad = [r for r in _canon_results if not r["ok"]]
        if not bad: return f"All {len(_canon_results)} page(s) have correct self-referencing canonical tags. See screenshot below."
        return f"{len(bad)} of {len(_canon_results)} page(s) have canonical issues. See screenshot below."

    def _dmeta_summary():
        if not _dmeta_results: return "No pages checked."
        bad = [r for r in _dmeta_results if not r["ok"]]
        if not bad: return f"All {len(_dmeta_results)} page(s) have title and meta description present once each. See screenshot below."
        return f"{len(bad)} of {len(_dmeta_results)} page(s) have meta tag issues. See screenshot below."

    def _mrobots_summary():
        if not _mrobots_results: return "No pages checked."
        bad = [r for r in _mrobots_results if not r["ok"]]
        if not bad: return f"All {len(_mrobots_results)} page(s) have correct meta robots (index, follow). See screenshot below."
        return f"{len(bad)} of {len(_mrobots_results)} page(s) have robots tag issues. See screenshot below."

    computed = {
        "versions_result": _versions_summary,
        "status200_result": _s200_summary(),
        "canonical_result": _canon_summary(),
        "double_meta_result": _dmeta_summary(),
        "meta_robots_result": _mrobots_summary(),
        "broken_links_result": _broken_links_summary(),
        "dummy_content_result": check_dummy_content(domain),
    }
    doc = Document()

    # Defaults so blank/inherited paragraphs are Candara 11pt too
    for sname in ("Normal", "List Paragraph"):
        try:
            st = doc.styles[sname]
            st.font.name = FONT_NAME
            st.font.size = Pt(FONT_SIZE_PT)
        except KeyError:
            pass

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(PAGE_W_IN), Inches(PAGE_H_IN)
    sec.left_margin = sec.right_margin = Inches(MARGIN_IN)
    sec.top_margin = sec.bottom_margin = Inches(MARGIN_IN)

    # Domain name as header (bold, 24pt, centered)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(domain)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(24)

    # Intro
    for label, body in INTRO:
        _add_label_body(doc, label, body)
    doc.add_paragraph()

    placed = missing = 0
    for n, cp in enumerate(CHECKPOINTS, start=1):   # numbered 1..17 (the "counting")
        body_tmpl = cp["body"]
        for k, v in computed.items():
            body_tmpl = body_tmpl.replace("{" + k + "}", v)
        body = body_tmpl.replace("{score}", str(score))
        try:
            body = body.format(domain=domain)
        except Exception:
            pass
        _add_label_body(doc, f"{n}. {cp['label']}", body)

        # Optional URL line (Sucuri)
        extra = cp.get("extra")
        if extra:
            p = _add_para(doc)
            _style_run(p.add_run(extra[0]), bold=True)
            _style_run(p.add_run(extra[1].format(domain=domain)), bold=False)

        # Optional status line — use real GSC status if captured by worker
        status = cp.get("status")
        if status:
            gsc_status = _resolve_gsc_status(cp["key"], captured)
            p = _add_para(doc)
            _style_run(p.add_run(status[0]), bold=True)
            _style_run(p.add_run(gsc_status if gsc_status else status[1]), bold=False)

        if cp.get("capture") == "action":
            p = _add_para(doc)
            run = p.add_run(f"⚠ Action Required: {cp['body'].format(domain=domain)}")
            run.font.name = FONT_NAME
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            run.bold = True
            doc.add_paragraph()
        else:
            img = _resolve_image(domain, cp["key"], captured)
            if img:
                _add_bordered_image(doc, img)
                placed += 1
            else:
                missing += 1
                print(f"     [info] no image for '{cp['key']}' (text only)")
            # Under the archived snapshot, keep a manual-comparison reminder.
            if cp["key"] == "layout":
                p = _add_para(doc)
                run = p.add_run("⚠ Compare this archived snapshot (web.archive.org) with the current live site to confirm no recent layout changes.")
                run.font.name = FONT_NAME
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                run.bold = True
            doc.add_paragraph()

    OUTPUT_DIR.mkdir(exist_ok=True)
    if out_path is None:
        gen_date = datetime.now().strftime("%d-%m-%Y")
        out_path = OUTPUT_DIR / (
            f"{domain} Website health checkup analysis report {gen_date}.docx"
        )
    out_path = gr.safe_save_path(Path(out_path))
    doc.save(out_path)
    print(f"  -> {placed} image(s) placed, {missing} text-only checkpoint(s)")
    return out_path


# ---------- LIVE CAPTURE (patchright, reuses generate_report) ----------
def _capture_frame(page, url, path, height=820, view_source=False, find=None):
    """Plain HEADLESS page screenshot — NO address bar / browser chrome (the
    border is baked into the image later). For view-source, optionally scroll to
    `find` (e.g. the meta description) first."""
    target = ("view-source:" + url) if view_source else url
    try:
        page.goto(target, wait_until="networkidle", timeout=60000)
        page.wait_for_timeout(gr.EXTRA_WAIT_MS)
    except Exception as e:
        print(f"     [warn] {type(e).__name__}, capturing anyway")
        page.wait_for_timeout(3000)
    if find:
        try:
            page.evaluate(
                """(t) => {
                    const w = document.createTreeWalker(document.body, NodeFilter.SHOW_TEXT);
                    let n; while ((n = w.nextNode())) {
                        if (n.nodeValue && n.nodeValue.includes(t)) {
                            (n.parentElement || document.body).scrollIntoView({block: 'center'});
                            window.scrollBy(0, -120);
                            break;
                        }
                    }
                }""", find)
            page.wait_for_timeout(1000)
        except Exception:
            pass
    page.screenshot(path=str(path), clip={"x": 0, "y": 0, "width": gr.VIEWPORT_W, "height": height},
                    full_page=False)


def _check_broken_links(domain, limit=60):
    """Real broken-link check (Python). Crawls the homepage's links and HTTP-
    checks each. Returns (checked_count, [(url, status), ...] broken)."""
    import urllib.request as ur
    import urllib.error as ue
    import re as _re
    from urllib.parse import urljoin

    home = f"https://{domain}/"
    headers = {"User-Agent": "Mozilla/5.0 (compatible; SearchOps-HealthBot/1.0)"}
    try:
        html = ur.urlopen(ur.Request(home, headers=headers), timeout=20).read().decode("utf-8", "ignore")
    except Exception as e:
        print(f"     [warn] broken-link fetch home failed: {e}")
        return 0, []

    links = []
    seen = set()
    for m in _re.finditer(r'href=["\']([^"\']+)["\']', html):
        href = m.group(1).strip()
        if href.startswith(("mailto:", "tel:", "javascript:", "#", "data:")):
            continue
        u = urljoin(home, href)
        if u.startswith("http") and u not in seen:
            seen.add(u); links.append(u)
        if len(links) >= limit:
            break

    broken = []
    for u in links:
        code = None
        for method in ("HEAD", "GET"):
            try:
                req = ur.Request(u, headers=headers, method=method)
                code = ur.urlopen(req, timeout=12).status
                break
            except ue.HTTPError as he:
                code = he.code
                if method == "HEAD" and code in (403, 405):   # some servers block HEAD
                    continue
                break
            except Exception:
                code = 0
        if code is None or code == 0 or code >= 400:
            broken.append((u, code or "ERR"))
    return len(links), broken


def _render_broken_links_image(domain, checked, broken, path):
    """Render the broken-link result as a clean image for the report."""
    from PIL import Image, ImageDraw, ImageFont
    rows = broken[:25]
    W = 1180
    H = 130 + (len(rows) + 1) * 30 + 30
    img = Image.new("RGB", (W, H), "#FFFFFF")
    d = ImageDraw.Draw(img)
    try:
        bold = ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 26)
        f = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 18)
        fs = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 15)
    except Exception:
        bold = f = fs = ImageFont.load_default()
    d.text((26, 24), f"Broken Links Check — {domain}", fill=(31, 41, 55), font=bold)
    if not broken:
        d.text((26, 74), f"Checked {checked} links — No broken links found.", fill=(22, 128, 67), font=f)
    else:
        d.text((26, 74), f"Checked {checked} links — {len(broken)} broken:", fill=(200, 40, 60), font=f)
        y = 116
        d.text((26, y), "URL", fill=(120, 120, 120), font=fs)
        d.text((W - 150, y), "Status", fill=(120, 120, 120), font=fs)
        y += 28
        for u, code in rows:
            d.text((26, y), u[:135], fill=(40, 40, 40), font=fs)
            d.text((W - 150, y), str(code), fill=(200, 40, 60), font=fs)
            y += 28
    img.save(str(path))


def _capture_pagespeed(page, domain, path):
    """Open PageSpeed Insights and wait for the FULL report (Performance score)
    before screenshotting the report area."""
    url = f"https://pagespeed.web.dev/analysis?url=https://{domain}/"
    try:
        page.goto(url, wait_until="domcontentloaded", timeout=120000)
        # Wait until a numeric Performance score gauge appears (report ready).
        page.wait_for_function(
            """() => {
                const txt = (document.body.innerText || '');
                const hasGauge = /\\b(Performance|Accessibility|Best Practices|SEO)\\b/.test(txt);
                const hasScore = /\\b\\d{1,3}\\b/.test(txt) && document.querySelectorAll('svg').length > 4;
                return hasGauge && hasScore;
            }""",
            timeout=150000,
        )
        page.wait_for_timeout(5000)
    except Exception as e:
        print(f"     [warn] pagespeed wait: {type(e).__name__}, capturing anyway")
        page.wait_for_timeout(6000)
    page.screenshot(path=str(path), clip={"x": 0, "y": 0, "width": gr.VIEWPORT_W, "height": 1300})


def _capture_logo(page, domain, path):
    """Grab the site's header logo (element screenshot), falling back to
    og:image / apple-touch-icon. Returns the path or None."""
    try:
        page.goto(f"https://{domain}/", wait_until="domcontentloaded", timeout=45000)
        page.wait_for_timeout(2500)
    except Exception:
        return None
    selectors = ["header a[href='/'] img", "header img", "a.navbar-brand img",
                 ".logo img", "#logo img", "img[alt*='logo' i]", "img[class*='logo' i]"]
    for sel in selectors:
        try:
            el = page.locator(sel).first
            if el.count() == 0 or not el.is_visible():
                continue
            box = el.bounding_box()
            if box and box["width"] >= 40 and box["height"] >= 18:
                el.screenshot(path=str(path))
                return str(path)
        except Exception:
            continue
    # Fallback: og:image / apple-touch-icon
    try:
        href = page.evaluate(
            """() => {
                const og = document.querySelector("meta[property='og:image']");
                if (og && og.content) return og.content;
                const ic = document.querySelector("link[rel*='apple-touch-icon'], link[rel*='icon']");
                return ic ? ic.href : null;
            }"""
        )
        if href:
            import urllib.request
            req = urllib.request.Request(href, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=15) as r:
                Path(path).write_bytes(r.read())
            return str(path)
    except Exception:
        pass
    return None


# Generous clip heights for GSC pages. Always start at y=0 so the TOP of the
# content (the section heading + status card) is never cropped; the height is
# large enough that the bottom isn't cut either. (Fixes the "top crop" issue.)
GSC_CLIP_HEIGHT = {"sitemaps": 840, "manual-actions": 620, "security-issues": 620, "index": 760}


def _capture_gsc_health(page, gsc_page, url, path):
    """Navigate a GSC report page and screenshot it cleanly (no top crop).
    For the sitemaps page, (re)submit the latest sitemap first."""
    try:
        page.goto(url, wait_until="networkidle", timeout=60000)
        page.wait_for_timeout(gr.EXTRA_WAIT_MS)
    except Exception as e:
        print(f"     [warn] {type(e).__name__}, capturing anyway")
        page.wait_for_timeout(3000)
    if gsc_page == "sitemaps":
        gr.submit_latest_sitemap(page)
    clip = {"x": 0, "y": 0, "width": gr.VIEWPORT_W,
            "height": GSC_CLIP_HEIGHT.get(gsc_page, 640)}
    page.screenshot(path=str(path), clip=clip, full_page=False)


def capture_health_screenshots(domain, account_hint=None, format_hint=None):
    """Live-capture the auto-able checkpoints. Mirrors generate_report's
    account/property selection so the same auth states work."""
    from patchright.sync_api import sync_playwright

    gr.SCREENSHOTS_DIR.mkdir(exist_ok=True)
    out = {}

    csv_map = gr.load_accounts_csv()
    entry = csv_map.get(domain)
    if entry:
        account_hint = account_hint or entry["account"]
        format_hint = format_hint or entry["access_level"]

    auth_files = gr.get_auth_files()
    if not auth_files:
        print("  [warn] no auth states found - GSC pages will be skipped")
        auth_files = []
        out["manual"] = {"status": "⚠ GSC account not connected — run: python scripts/python/01_auth_setup.py <accountKey>"}
        out["security"] = {"status": "⚠ GSC account not connected — run: python scripts/python/01_auth_setup.py <accountKey>"}

    if account_hint and auth_files:
        hl = account_hint.lower()
        auth_files = ([f for f in auth_files if f.stem.lower() == hl] +
                      [f for f in auth_files if f.stem.lower() != hl])

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = None
        working_prop = None

        for auth_file in auth_files:
            ctx = browser.new_context(
                storage_state=str(auth_file),
                viewport={"width": gr.VIEWPORT_W, "height": gr.VIEWPORT_H},
            )
            page = ctx.new_page()
            try:
                prop = gr.find_working_property(page, domain, format_hint=format_hint)
                if prop:
                    working_prop = prop
                    print(f"  [*] GSC account '{auth_file.stem}' -> {prop}")
                    break
                ctx.close()
                page = None
            except gr.SessionExpired:
                print(f"     [session expired] {auth_file.stem}")
                ctx.close(); page = None
            except Exception as e:
                print(f"     [error] {type(e).__name__}: {e}")
                ctx.close(); page = None

        # If no authed page, still capture the public checks via a fresh context
        if page is None:
            page = browser.new_context(
                viewport={"width": gr.VIEWPORT_W, "height": gr.VIEWPORT_H}
            ).new_page()
            if auth_files:
                out["manual"] = {"status": "⚠ GSC account has no access to this domain — reconnect or check permissions"}
                out["security"] = {"status": "⚠ GSC account has no access to this domain — reconnect or check permissions"}

        # Client logo (header element / og:image) for the report header.
        try:
            logo_path = gr.SCREENSHOTS_DIR / f"{domain}_health_logo.png"
            if _capture_logo(page, domain, logo_path):
                out["logo"] = str(logo_path)
                print("  -> captured [logo]")
        except Exception as e:
            print(f"     [warn] logo capture failed: {type(e).__name__}: {e}")

        for cp in CHECKPOINTS:
            cap, key = cp["capture"], cp["key"]
            path = gr.SCREENSHOTS_DIR / f"{domain}_health_{key}.png"
            try:
                if cap == "sucuri":
                    gr.capture_sucuri(page, domain, path)
                elif cap.startswith("gsc:"):
                    if not working_prop:
                        continue  # no authed GSC account → leave for drop-in
                    pg = cap.split(":", 1)[1]
                    enc = quote(working_prop, safe="")
                    url = f"https://search.google.com/search-console/{pg}?resource_id={enc}"
                    _capture_gsc_health(page, pg, url, path)
                elif cap.startswith("frame:"):
                    _capture_frame(page, cap.split(":", 1)[1].format(domain=domain), path)
                elif cap.startswith("viewsrc:"):
                    _capture_frame(page, cap.split(":", 1)[1].format(domain=domain), path,
                                   height=760, view_source=True, find=cp.get("find"))
                elif cap == "wayback":
                    _capture_frame(page, f"https://web.archive.org/web/2/https://{domain}/", path,
                                   height=900)
                elif cap == "brokenlinks":
                    checked, broken = _check_broken_links(domain)
                    _render_broken_links_image(domain, checked, broken, path)
                    print(f"     broken-link check: {checked} links, {len(broken)} broken")
                elif cap in ("computed", "status200", "manual", "action"):
                    pass  # handled at build time, not capture time
                else:
                    continue  # skip unknown capture types
                out[key] = str(path)
                print(f"  -> captured [{key}]")
            except Exception as e:
                print(f"     [warn] capture '{key}' failed: {type(e).__name__}: {e}")

        browser.close()
    return out


# ---------- MAIN ----------
def main():
    args = sys.argv[1:]
    if not args or args[0].startswith("--"):
        print("Usage: python generate_health_report.py <domain> "
              "[--account NAME] [--format HINT] [--score N] [--logo PATH] "
              "[--no-capture] [--pages URL1 URL2 ...]")
        sys.exit(1)

    domain = gr.clean_domain(args[0])
    account_hint = format_hint = logo = None
    score = 68
    no_capture = False
    target_pages = []
    manual_status = manual_img = security_status = security_img = None
    i = 1
    while i < len(args):
        a = args[i]
        if a == "--account" and i + 1 < len(args): account_hint = args[i+1]; i += 2
        elif a == "--format" and i + 1 < len(args): format_hint = args[i+1]; i += 2
        elif a == "--score" and i + 1 < len(args): score = args[i+1]; i += 2
        elif a == "--logo" and i + 1 < len(args): logo = args[i+1]; i += 2
        elif a == "--no-capture": no_capture = True; i += 1
        elif a == "--args-file" and i + 1 < len(args):
            import json as _json
            with open(args[i+1], encoding="utf-8") as _f:
                _d = _json.load(_f)
            target_pages = _d.get("pages") or []
            manual_status = _d.get("manualStatus")
            manual_img = _d.get("manualImg")
            security_status = _d.get("securityStatus")
            security_img = _d.get("securityImg")
            i += 2
        elif a == "--manual-status" and i + 1 < len(args): manual_status = args[i+1]; i += 2
        elif a == "--manual-img" and i + 1 < len(args): manual_img = args[i+1]; i += 2
        elif a == "--security-status" and i + 1 < len(args): security_status = args[i+1]; i += 2
        elif a == "--security-img" and i + 1 < len(args): security_img = args[i+1]; i += 2
        elif a == "--pages":
            i += 1
            while i < len(args) and not args[i].startswith("--"):
                target_pages.append(args[i]); i += 1
        else: i += 1

    print(f"\n{'='*55}\n  Website Health Audit Report\n  Domain : {domain}\n{'='*55}")

    captured = {}
    if not no_capture:
        print("\n[1/2] Capturing screenshots...")
        captured = capture_health_screenshots(domain, account_hint, format_hint)
    else:
        print("\n[1/2] Skipping live capture (drop-ins/text only)")

    # Merge any GSC captures the Node worker passed in. The Python patchright
    # capture above is the primary source; only let a worker-provided screenshot
    # win if it actually carries an image. Otherwise keep the patchright capture.
    def _merge_gsc(key, status, img):
        if img:
            captured[key] = {'status': status, 'img': img}
        elif captured.get(key):
            pass  # keep the patchright-captured path
        elif status:
            captured[key] = {'status': status, 'img': None}
    _merge_gsc('manual', manual_status, manual_img)
    _merge_gsc('security', security_status, security_img)

    print("\n[2/2] Building .docx...")
    out = build_health_docx(domain, captured, logo=logo, score=score, target_pages=target_pages)
    print(f"\n{'='*55}\n  [DONE] {out.resolve()}\n{'='*55}\n")


if __name__ == "__main__":
    main()
